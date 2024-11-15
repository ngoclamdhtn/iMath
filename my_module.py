import math
import sympy as sp
import numpy as np
from sympy import *
import random
from fractions import Fraction
import os
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os, shutil,re, sys, subprocess
from datetime import datetime
from PIL import Image
#Xuất ra màn hình desktop file word chứa câu hỏi
def Xuatfile(cauhoi):
    # Đường dẫn đến màn hình desktop (tuỳ thuộc vào hệ điều hành)
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

    # Tạo một đối tượng Document
    doc = Document()
    doc.add_paragraph(cauhoi)

    # Lưu file với tên "cauhoi.docx" trên màn hình desktop
    file_path = os.path.join(desktop_path, "cauhoi.docx")
    doc.save(file_path)

    print(f"File 'cauhoi.docx' đã được tạo trên màn hình desktop.")

    # Mở file Word
    subprocess.Popen(["start", "winword", file_path], shell=True)
    return

def xu_li_heso_1(a):
    if a==1:
        heso=""
    elif a==-1:
        heso="-" 
    else:
        heso=a 
    return heso

# Xử lí trả về đáp án đúng để ghi câu Chọn đáp án
def tra_ve_dap_an(list_PA):    
    if "*" in list_PA[0]: 
            dap_an="A"
    elif "*" in list_PA[1]: 
            dap_an="B"
    elif "*" in list_PA[2]: 
            dap_an="C"
    else:
            dap_an="D"   
    return dap_an

# Xử lí trả về đáp án đúng để ghi câu Chọn đáp án
def tra_ve_TF(list_PA):
    list_TF=[]
    for phan_tu in list_PA:
        if "*" in phan_tu:
            list_TF.append("đúng")
        else:
            list_TF.append("sai")
    return list_TF
#Hiện số a 
def show_num(a):
    if a==0:
        t=""
    else:
        t=a
    return t

#Hiện số b theo ẩn st

def show_num_with_variable(b,st):
    if b>0:
        t=f"+{b}{st}"
    if  b==0:
        t=f""
    if b<0:
        t=f"{b}{st}"            
    if b==1:
        t=f"+{st}"
    if b==-1:
        t=f"-{st}"
    return t

#Trả về dấu của một số để đưa vào latex
def tao_dau(a):
    dau="+"
    if a<0 or a==0:
        dau=""
    return dau

#Trả về phép cộng với ngoặc đơn
def show_tong(a,b):
    if b>0:
        ketqua=f"{a}+{b}"
    else:
        ketqua=f"{a}+({b})"
    return ketqua

#Trả về phép trừ với ngoặc đơn
def show_hieu(a,b):
    if b>=0:
        ketqua=f"{a}-{b}"
    else:
        ketqua=f"{a}-({b})"
    return ketqua
    
def thay_dau_congtru(st):
    ketqua=st.replace("-+","-").replace("--","+").replace("+-","-").replace("++","+").replace("+ +","+").replace("+ -","-").replace("- -","+").replace("- +","-")
    return ketqua

def thay_hinh_hoc(st):
    ketqua=st.replace("1I","I").replace("1E","E").replace("1M","M")
    ketqua=ketqua.replace("=1a","=a")
    return ketqua

def frac_to_dfrac(st):
    ketqua=st.replace(f"\\frac",f"\\dfrac")
    return ketqua

#Trả về dạng phân số 
def hien_phan_so(t):
    m=Rational(t).limit_denominator(100000000000)
    return m
#Tạo dấu cho một số

#Trả về giá trị cos,sin
def hien_sin_cos(goc,ham_so):
    if ham_so=="sin":
        if goc==0:
            gia_tri=0
        if goc==30 or goc ==150:
            gia_tri=1/2
        if goc==45 or goc ==135:
            gia_tri=sqrt(2)/2
        if goc==60 or goc ==120:
            gia_tri=sqrt(3)/2
        if goc==90:
            gia_tri=1
        if goc==180:
            gia_tri=0
                    
    if ham_so=="cos":
        if goc==0:
            gia_tri=1
        if goc==30:
            gia_tri=sqrt(3)/2
        if goc==45:
            gia_tri=sqrt(2)/2
        if goc==60:
            gia_tri=1/2
        if goc==90:
            gia_tri=0
        if goc==120:
            gia_tri=-1/2
        if goc==135:
            gia_tri=-sqrt(2)/2
        if goc==150:
            gia_tri=-sqrt(3)/2
        if goc==180:
            gia_tri=-1
    return gia_tri

def hien_tan(goc):
    if goc==30:
        gia_tri=sqrt(3)/3
    if goc==45:
        gia_tri=1
    if goc==60:
        gia_tri=sqrt(3)
    return gia_tri

def tinh_tan(a):
    goc_rad = rad(a)       
    kq = tan(goc_rad)        
    kq = simplify(kq.rewrite(sqrt))
    return kq

def tinh_sin(a):
    goc_rad = rad(a)       
    kq = sin(goc_rad)        
    kq = simplify(kq.rewrite(sqrt))
    return kq
    
def tinh_cos(a):
    goc_rad = rad(a)       
    kq = cos(goc_rad)        
    kq = simplify(kq.rewrite(sqrt))
    return kq

#Hàm rút gọn căn thức
def rutgon_can(x):
    from sympy import nsimplify
    if x.is_number:
        ns = nsimplify(x)
        if ns != x and x.equals(ns):
            return ns
    return x
def check_so_nguyen(so):
    ketqua=False
    if isinstance(so, int):
        ketqua=True
    return ketqua

#Kiểm tra số nguyên
# def check_so_nguyen(so):
#     ketqua=False
#     if so==int(so):
#         ketqua=True
#     return ketqua

def tra_ve_so_nguyen(danh_sach):
    so_phantu=len(danh_sach)
    for i in range(so_phantu):
        if danh_sach[i]==int(danh_sach[i]):
            danh_sach[i]=int(danh_sach[i])
    for i in range(so_phantu):
        danh_sach[i]=round(danh_sach[i],2)
        danh_sach[i]=str(danh_sach[i])
        danh_sach[i]=danh_sach[i].replace(".",",")
    return danh_sach
       

#Hàm rút gọn căn thức
def canh_thu_3_tg(a,b):
    c = random.randint(abs(a - b) + 1, a + b - 1)
    return c

def tim_boi_của_n_nho_hon_m(n, m):
    multiples = []
    for i in range(0, m):
        if i % n == 0:
            multiples.append(i)
    return multiples

# Tạo ngẫu nhiên tập hợp 
def tao_taphop(so_phantu,so_batdau,so_ketthuc):
    my_set = set(random.sample(range(so_batdau, so_ketthuc), so_phantu))
    return my_set

def thay_kihieu_taphop(string):
    st = string.replace(f"\\{{set()\\}}",f"\\emptyset") 
    return st
def xoa_ngoac_vuong(string):
    find_what="["
    replace_with=""
    st = string.replace(find_what,replace_with)

    find_what="]"
    replace_with=""
    st = st.replace(find_what,replace_with)
    return st

# Tạo 2 phân số ngẫu nhiên có tổng bình phương bằng 1
def tao_2_phanso():
    a = random.randint(1,10)
    b = random.randint(11,20)
    dau = (-1)**random.randint(1,2)

    # Tính phân số thứ nhất (ngẫu nhiên nhỏ hơn 1)
    fraction1 = Rational(dau*a/b).limit_denominator(100)

    # Tính phân số thứ hai sao cho tổng bằng 1
    fraction2 = rutgon_can(sp.sqrt(b**2-a**2)/b)

    return fraction1,fraction2
def khong_trung_so(a,b,c,d):
    list_so=[a,b,c,d]    
    if list_so[1]==list_so[0]:
         list_so[1]=max(list_so)+1
    if list_so[1]==list_so[2]:
         list_so[2]=max(list_so)+1
    if list_so[1]==list_so[3]:
         list_so[3]=max(list_so)+1        
    if list_so[2]==list_so[0] or list_so[2]==list_so[3] :
         list_so[2]=max(list_so)+1
    if list_so[3]==list_so[0] or list_so[3]==list_so[1] or list_so[3]==list_so[2]:
         list_so[3]=max(list_so)+1  

    return list_so
    
#Tạo đa thức biến st và bậc k. Kết quả: Đa thức, hệ số của bậc, hệ số tự do.
def random_polynomial(st,k):
    if st== "x":
        # Định nghĩa biến
        x = sp.symbols('x')
    if st== "n":
        # Định nghĩa biến
        x = sp.symbols('n')  

    # Tạo đa thức với các hệ số
    coefficients = [random.randint(-15, 15) for _ in range(k)]
    if coefficients[k-1]==0:
        coefficients[k-1] += 1
        
    if coefficients[k-1]==0:
        coefficients[k-1] -= 1    
        
    polynomial = sum(coefficients[i] * x**i for i in range(k))

    # Lấy tất cả các hệ số của đa thức
    all_coefficients = [polynomial.coeff(x, i) for i in range(k)]

    #Lấy hệ số của bậc của đa thức
    coeff_bac= all_coefficients[k-1]

    #Lấy hệ số tự do của đa thức
    coeff_tudo= all_coefficients[0]

    return polynomial,coeff_bac,coeff_tudo

#Tạo đa thức luôn dương.

def random_dathuc_bac2_luon_duong(st):
    if st== "x":
        # Định nghĩa biến
        x = sp.symbols('x')
    if st== "n":
        # Định nghĩa biến
        x = sp.symbols('n')  

    # Tạo đa thức với các hệ số
    coefficients =[ ]
    coefficients.append(random.randint(-5,5))    
    coefficients.append (random.randint(-5,5))
    coefficients.append ( random.randint(1,10))
    
        
    polynomial = sum(coefficients[i] * x**i for i in range(3))

    # Lấy tất cả các hệ số của đa thức
    all_coefficients = [polynomial.coeff(x, i) for i in range(3)]

    #Lấy hệ số của bậc của đa thức
    coeff_bac= all_coefficients[2]

    #Lấy hệ số tự do của đa thức
    coeff_tudo= all_coefficients[0]

    return polynomial,coeff_bac,coeff_tudo

#Giải bất phương trình bậc hai
def solve_bpt_bac2(a,b,c,dau_bpt,bien_x):
    delta,x_1,x_2=tinh_va_dau_delta(a,b,c)
    #ax^2 + bx + c>0
    if dau_bpt==">0" and a>0:    
            if delta == "<0":
                nghiem=f"$\\forall {bien_x} \\in \\mathbb{{R}}$"
            elif delta=="=0":
                nghiem=f"${bien_x}\\ne {latex(hien_phan_so(x_1))}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${bien_x} < {x_1}$ hoặc ${bien_x} > {x_2}$"
    if dau_bpt==">0" and a<0:
            if delta == "<0" or delta == "=0" :
                nghiem=f"không tồn tại ${bien_x}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${x_1} < {bien_x}< {x_2}$"

    #ax^2 + bx + c<0
    if dau_bpt=="<0" and a>0:

            if delta == "<0" or delta == "=0" :
                nghiem=f"không tồn tại ${bien_x}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${x_1} < {bien_x} <{x_2}$"            
    if dau_bpt=="<0" and a<0:       
            if delta == "<0":
                nghiem=f"$\\forall {bien_x} \\in \\mathbb{{R}}$"
            elif delta=="=0":
                nghiem=f"${bien_x}\\ne {latex(hien_phan_so(x_1))}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${bien_x} < {x_1}$ hoặc ${bien_x} > {x_2}$"

    #ax^2 + bx + c>0
    if dau_bpt==">=0" and a>0:
        
            if delta == "<0" or delta== "=0":
                nghiem=f"$\\forall {bien_x} \\in \\mathbb{{R}}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${bien_x} \\le {x_1}$ hoặc  ${bien_x}\\ge {x_2}$"
    if dau_bpt==">=0" and a<0:
            if delta == "<0":
                nghiem=f"không tồn tại ${bien_x}$"
            elif delta == "=0":
                nghiem=f"${bien_x} = {latex(hien_phan_so(x_1))}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${x_1} \\le {bien_x} \\le {x_2}$"
    #ax^2 + bx + c<=0
    if dau_bpt=="<=0" and a>0:
            if delta == "<0":
                nghiem=f"không tồn tại ${bien_x}$"
            elif delta == "=0":
                nghiem=f"${bien_x} = {latex(hien_phan_so(x_1))}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${x_1} \\le {bien_x} \\le {x_2}$"
            
    if dau_bpt=="<=0" and a<0:           

            if delta == "<0" or delta== "=0":
                nghiem=f"$\\forall {bien_x} \\in \\mathbb{{R}}$"
            else:
                if check_so_nguyen(x_1):            
                    x_1=latex(hien_phan_so(x_1))
                    x_2=latex(hien_phan_so(x_2))
                else:
                    x_1=latex(x_1)
                    x_2=latex(x_2)
                nghiem=f"${bien_x} \\le {x_1}$ hoặc ${bien_x} \\ge {x_2}$"
    return nghiem




#Lấy thư mục hiện tại
def get_folder_name():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))    
    return current_directory

#Lấy thư mục hình
def get_folder_hinh():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'HINH VE') 
    return hinh_folder_path

#Lấy thư mục doc
def get_folder_doc():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'DOC') 
    return hinh_folder_path

#Lấy thư mục LATEX
def get_folder_latex():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'LATEX') 
    return hinh_folder_path

#Lấy thư mục UPDATE
def get_folder_update():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'UPDATE') 
    return hinh_folder_path

#Lấy thư mục icon
def get_folder_icon():
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'ICON') 
    return hinh_folder_path

#Biên dịch lệnh Latex
def run_latex(content):
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    latex_folder_path = os.path.join(current_directory, 'LATEX')

    # Kiểm tra xem thư mục con có tên "HINH VE" đã tồn tại chưa
    if not os.path.exists(latex_folder_path):         
        os.makedirs(latex_folder_path) 

    name = datetime.now().strftime("%d-%m_%H-%M-%S")

    # Tạo đường dẫn đến file trong thư mục hiện tại
    file_path = os.path.join(latex_folder_path, f'{name}.tex')

    # Mở file và ghi nội dung vào đó
    with open(file_path, 'w',encoding="utf-8") as file:
        file.write(content)

    #Di chuyển đến thư mục chứa file tex
    os.chdir(os.path.dirname(file_path))

    #Biên dịch tệp LaTeX thành PDF
    latex_command = f"pdflatex -jobname={name} {name}.tex"
    subprocess.run(latex_command, shell=True)

    #Lấy danh sách tất cả các tệp trong thư mục trừ các tệp có đuôi là ".PDF" ".TEX"
    #files_to_delete = [file for file in glob.glob("*") if not file.lower().endswith([".pdf",".tex"])]
    

    # Xóa các tệp linh tinh
    # for file_to_delete in files_to_delete:
    #     os.remove(file_to_delete)
    return

#Chuyển PDF thành ảnh PNG    
def pdftoimage(content):
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'HINH VE')

    # Kiểm tra xem thư mục con có tên "HINH VE" đã tồn tại chưa
    if not os.path.exists(hinh_folder_path):         
        os.makedirs(hinh_folder_path) 


    # Tạo đường dẫn đến file trong thư mục hiện tại
    file_path = os.path.join(hinh_folder_path, 'hinhve.tex')

    # Mở file và ghi nội dung vào đó
    with open(file_path, 'w',encoding="utf-8") as file:
        file.write(content)

    #Di chuyển đến thư mục chứa file tex
    os.chdir(os.path.dirname(file_path))

    #Biên dịch tệp LaTeX thành PDF
    latex_command = "pdflatex -jobname=hinhve hinhve.tex"
    subprocess.run(latex_command, shell=True)

    #Chuyển đổi PDF thành hình ảnh PNG
    convert_command = "convert -density 300 -alpha on hinhve.pdf hinhve.png"
    subprocess.run(convert_command, shell=True)

    #Lấy danh sách tất cả các tệp trong thư mục "HINH" trừ các tệp có đuôi là ".PNG"
    files_to_delete = [file for file in glob.glob("*") if not file.lower().endswith(".png")]


    # Xóa các tệp linh tinh
    for file_to_delete in files_to_delete:
        os.remove(file_to_delete)
    return

#Chuyển PDF thành PNG có đặt tên theo thời gian
def pdftoimage_timename(content):
    #Lấy đường dẫn thư mục hiện tại của script Python
    current_directory = os.path.dirname(os.path.abspath(__file__))
    hinh_folder_path = os.path.join(current_directory, 'HINH VE')
    latex_folder_path = os.path.join(current_directory, 'LATEX')

    # Kiểm tra xem thư mục con có tên "HINH VE" đã tồn tại chưa
    if not os.path.exists(hinh_folder_path):         
        os.makedirs(hinh_folder_path)

    #Copy file ex_test.sty
    source_file_path = f"{latex_folder_path}\\ex_test.sty"
    if os.path.exists(source_file_path):                             
            destination_file_path = os.path.join(hinh_folder_path, os.path.basename(source_file_path))                                
            shutil.copy2(source_file_path, destination_file_path)
            
    
    #Lấy ngày giờ hiện tại
    name = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    # Tạo đường dẫn đến file trong thư mục hiện tại
    file_path = os.path.join(hinh_folder_path, f"{name}.tex")

    # Mở file và ghi nội dung vào đó
    with open(file_path, 'w', encoding="utf-8") as file:
        file.write(content)

    
    #Di chuyển đến thư mục chứa file tex
    os.chdir(os.path.dirname(file_path))

    #Biên dịch tệp LaTeX thành PDF
    latex_command = f"pdflatex -jobname={name} {name}.tex"
    subprocess.run(latex_command, shell=True)

    #Chuyển đổi PDF thành hình ảnh PNG
    convert_command = f"convert -density 300 -alpha on {name}.pdf {name}.png"
    subprocess.run(convert_command, shell=True)

    #Lấy danh sách tất cả các tệp trong thư mục "HINH" trừ các tệp có đuôi là ".PNG"
    files_to_delete = [file for file in glob.glob("*") if not file.lower().endswith(".png")]

    # Xóa các tệp linh tinh
    for file_to_delete in files_to_delete:
        os.remove(file_to_delete)
    return name

def get_image_dimensions(file_path):
    try:
        # Mở ảnh
        with Image.open(file_path) as img:
            # Lấy kích thước ảnh trong pixels
            width_px, height_px = img.size

            # Lấy độ phân giải (DPI) của ảnh
            dpi = img.info.get('dpi', (72, 72))

            # Tính kích thước ảnh trong inches
            width_in = width_px / dpi[0]
            height_in = height_px / dpi[1]

            chieu_rong =f"{round(width_in,2)}"
            chieu_cao=f"{round(height_in,2)}"

            return chieu_rong, chieu_cao
        
    except Exception as e:
        print(f"Error: {e}")
        return None


#Tìm kiếm tên ảnh và chèn ảnh từ thư mục HINHVE
def find_and_insert_image(doc):
    folder_hinh = get_folder_hinh()
    # Lấy danh sách tất cả các tệp trong thư mục HINHVE
    if os.path.exists(folder_hinh):
        # Lấy danh sách tất cả các tệp trong thư mục
        file_list = [file.split('.')[0] for file in os.listdir(folder_hinh) if os.path.isfile(os.path.join(folder_hinh, file))]
    if file_list != []:
        for file_name in file_list:
            for i in range(len(doc.paragraphs)):
                if file_name in doc.paragraphs[i].text:
                    image_path = os.path.join(folder_hinh,f"{file_name}.PNG")
                    chieu_rong= get_image_dimensions(image_path)[1] 
                    #print(chieu_rong)
                     # Thêm đoạn văn trống trước đoạn văn chứa ảnh
                    empty_paragraph = doc.paragraphs[i].insert_paragraph_before()               
                    run = doc.paragraphs[i].add_run()                           
                    
                    #Chỉnh kích thước bảng mẫu số liệu ghép nhóm
                    if chieu_rong=="1.74":
                        run.add_picture(image_path, width=Inches(6.5),height=Inches(0.6))

                    elif chieu_rong=="1.71":
                        run.add_picture(image_path, width=Inches(6.5),height=Inches(0.6))

                    #Chỉnh kích thước bảng biến thiên bậc 2
                    elif chieu_rong =="5.18":
                        run.add_picture(image_path, width=Inches(3.0),height=Inches(1.5))

                    #Chỉnh kích thước xét dấu đạo hàm
                    elif chieu_rong =="3.56":
                        run.add_picture(image_path, width=Inches(4.3),height=Inches(1))
                    elif chieu_rong =="2.9":
                        run.add_picture(image_path, width=Inches(5),height=Inches(1))

                    #Chỉnh kích thước bảng xét dấu bậc 2 vô nghiệm
                    elif chieu_rong =="2.56":
                        run.add_picture(image_path, width=Inches(3.0),height=Inches(0.65))

                    #Chỉnh kích thước bảng biến thiên hàm bậc 3 có 2 nghiệm
                    elif chieu_rong=="5.33":
                        run.add_picture(image_path, width=Inches(3.8),height=Inches(2))

                    #Chỉnh kích thước bảng biến thiên hàm phân thức bậc nhất
                    elif chieu_rong=="5.53":
                        run.add_picture(image_path, width=Inches(2.7),height=Inches(1.4))

                    elif chieu_rong =="7.22":
                        run.add_picture(image_path, width=Inches(3),height=Inches(3.48))

                    elif chieu_rong =="7.25":
                        run.add_picture(image_path, width=Inches(3),height=Inches(2.34))

                    #Chỉnh kích thước bảng biến thiên tìm đường tiệm cận
                    elif chieu_rong =="6.71":
                        run.add_picture(image_path, width=Inches(4.45),height=Inches(1.56))

                    #Chỉnh kích thước hình chóp hình thang
                    elif chieu_rong =="9.32":
                        run.add_picture(image_path, width=Inches(2.8),height=Inches(2.0))

                    #Chỉnh kích thước hình lăng trụ xiên tam giác 
                    elif chieu_rong =="9.68":
                        run.add_picture(image_path, width=Inches(2.0),height=Inches(2.0))

                    #Chỉnh đồ thị bậc 2
                    elif chieu_rong =="13.07":
                        run.add_picture(image_path, width=Inches(2.5),height=Inches(2.0))

                    #Chỉnh vòng tròn lượng giác
                    elif chieu_rong =="13.44":
                        run.add_picture(image_path, width=Inches(3),height=Inches(3))

                    #Chỉnh kích thước bảng xét dấu đạo hàm
                    elif chieu_rong =="14.19" and chieu_cao=="3.56":
                        run.add_picture(image_path, width=Inches(3.0),height=Inches(0.8))

                    #Chỉnh kích thước bảng biến thiên tiệm cận
                    elif chieu_rong =="15.83" and chieu_cao=="5.85":
                        run.add_picture(image_path, width=Inches(5.5),height=Inches(2.3))

                    elif chieu_rong =="19.32":
                        run.add_picture(image_path, width=Inches(3.45),height=Inches(0.65))                    

                    #Chỉnh đồ thị bậc 4
                    elif chieu_rong =="10.12":
                        run.add_picture(image_path, width=Inches(3.0),height=Inches(3.0)) 
                    else:
                        run.add_picture(image_path, width=Inches(3.0),height=Inches(2.0))
                    # Căn giữa đoạn văn chứa hình ảnh
                    doc.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    doc.paragraphs[i+1].text =""              
                    break
    return
def check_dau(a):
    if a>0:
        dau=">0"
    if a<0:
        dau="<0"
    if a==0:
        dau="=0"
    return dau

#Code vẽ điểm biểu diễn
def code_ve_mot_diem(ten_diem,a,b):
    if a>0 and b>0: 
        vitri="above right"
        vitri_a="below"
        vitri_b="left"
    elif a<0 and b>0: 
        vitri="above left"
        vitri_a="below"
        vitri_b="right"
    elif a>0 and b<0: 
        vitri="below right"
        vitri_a="above"
        vitri_b="left"
    elif a<0 and b<0: 
        vitri="below left"
        vitri_a="above"
        vitri_b="right"
    else:
        vitri="below left"
        vitri_a="above"
        vitri_b="right"
    truc_x, truc_y=abs(a), abs(b)
    code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20]({-truc_x-1},{-truc_y-1}) grid({truc_x+1},{truc_y+1});\n\
\\draw[->] ({-truc_x-1},0)--({truc_x+1},0) node[below left] {{$x$}};\n\
\\draw[->] (0,{-truc_y-1})--(0,{truc_y+1}) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
 \\foreach \\x in {{{a}}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [{vitri_a}] {{\\footnotesize$\\x$}};\n\
\\foreach \\y in {{{b}}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [{vitri_b}] {{\\footnotesize$\\y$}};\n\
\\begin{{scope}}\n\
 \\clip ({-truc_x-1},{-truc_y-1}) rectangle ({truc_x+1},{truc_y+1});\n\
\\draw[fill=black] ({a},{b})  node[{vitri}]{{${ten_diem}$}} circle (1pt);\n\
\\draw[dashed] ({a},{b})--({a},0) ;\n\
\\draw[dashed] ({a},{b})--(0,{b}) ;\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n"
    return code

#Code vẽ đồ thị hàm số bậc 1   
def codelatex_dothi_bac_1(a,b):
    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20](-6,-6.5)grid(6,6.5);\n\
\\draw[->] (-6,0)--(6,0) node[below left] {{$x$}};\n\
\\draw[->] (0,-6.5)--(0,6.5) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
 \\foreach \\x in {{-4,-3,-2,-1,1,2,3,4}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
\\foreach \\y in {{-4,-3,-2,-1,1,2,3,4}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
\\begin{{scope}}\n\
 \\clip (-6.5,-6.5) rectangle (6.5,6.5);\n\
\\draw[samples=200,domain=-6.5:6.5,smooth,variable=\\x] plot (\\x,{{{a}*(\\x)+{b}}});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}} "
    return  code
def tao_chuoi_so_nguyen(a, b):
    # Tạo danh sách các số nguyên từ a đến b
    danh_sach_so = list(range(a, b + 1))
    danh_sach_so = [x for x in danh_sach_so if x != 0 and x%2==0]
    
    # Chuyển đổi các số thành chuỗi
    chuoi_so = [str(x) for x in danh_sach_so]
    
    # Nối chuỗi bằng dấu ,
    chuoi_ket_qua = ",".join(chuoi_so)
    
    return chuoi_ket_qua

#Code vẽ đồ thị hàm số bậc 2   
def codelatex_dothi_bac_2(a,b,c):
    x=sp.symbols("x")
    f=a*x**2 + b*x + c
    x_0=round(-b/(2*a),1)
    y_0=round(f.subs(x,x_0),1)
    x_min, x_max= x_0-3, x_0+3
    if x_min>0: 
        x_min=-1
        x_max=2*(x_0+1)
    if x_max<0: 
        x_max=1
        x_min=2*(1-x_0)
    d=tinh_va_dau_delta(a,b,c)[0]
    #Vô nghiệm, nghiệm kép
    if d=="<0" or d=="=0":
        if a>0:        
            y_min, y_max= y_0 - 2, y_0 + 3
            if y_min > 0: 
                y_min=-2            
        else:
            y_min, y_max= y_0 - 3, y_0 + 0.3
            if y_max < 0: 
                y_max=2
    #2 nghiệm
            
    else:
        if a>0:
            y_min= y_0 - 2                  
            y_max=2-y_min
        else:
            y_max= y_0+2      
            y_min=-2-y_max

    #Tạo số cho trục x
    so_truc_x=tao_chuoi_so_nguyen(int(x_min), int(x_max))
    so_truc_y=tao_chuoi_so_nguyen(int(y_min), int(y_max))

    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.5]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20]({x_min},{y_min-1}) grid ({x_max+0.5},{y_max+1});\n\
\\draw[->] ({x_min},0)--({x_max+0.5},0) node[below right] {{$x$}};\n\
\\draw[->] (0,{y_min-1})--(0,{y_max+1}) node[right] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
 \\foreach \\x in {{-1,1,{so_truc_x}}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
\\foreach \\y in {{-1,1,{so_truc_y}}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
\\begin{{scope}}\n\
 \\clip ({x_min},{y_min-1}) rectangle ({x_max},{y_max+1});\n\
\\draw[samples=200,domain={x_min}:{x_max},smooth,variable=\\x, color=blue] plot (\\x,{{{a}*(\\x)^2+{b}*(\\x)+{c}}});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}} "
    return  code

def codelatex_dothi_bac_2_no_header(a,b,c):
    x=sp.symbols("x")
    f=a*x**2 + b*x + c
    x_0=round(-b/(2*a),1)
    y_0=round(f.subs(x,x_0),1)
    x_min, x_max= x_0-3, x_0+3
    if x_min>0: 
        x_min=-1
        x_max=2*(x_0+1)
    if x_max<0: 
        x_max=1
        x_min=2*(1-x_0)
    d=tinh_va_dau_delta(a,b,c)[0]
    #Vô nghiệm, nghiệm kép
    if d=="<0" or d=="=0":
        if a>0:        
            y_min, y_max= y_0 - 2, y_0 + 3
            if y_min > 0: 
                y_min=-2            
        else:
            y_min, y_max= y_0 - 3, y_0 + 0.3
            if y_max < 0: 
                y_max=2
    #2 nghiệm
            
    else:
        if a>0:
            y_min= y_0 - 2           
            y_max=2-y_min
        else:
            y_max= y_0+2          
            y_min=-2-y_max

    #Tạo số cho trục x
    so_truc_x=tao_chuoi_so_nguyen(int(x_min), int(x_max))
    so_truc_y=tao_chuoi_so_nguyen(int(y_min), int(y_max) )

    code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.5]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20]({x_min},{y_min-1}) grid ({x_max+0.5},{y_max+1});\n\
\\draw[->] ({x_min},0)--({x_max+0.5},0) node[below right] {{$x$}};\n\
\\draw[->] (0,{y_min-1})--(0,{y_max+1}) node[right] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
 \\foreach \\x in {{-1,1,{so_truc_x}}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
\\foreach \\y in {{-1,1,{so_truc_y}}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
\\begin{{scope}}\n\
 \\clip ({x_min},{y_min-1}) rectangle ({x_max},{y_max+1});\n\
\\draw[samples=200,domain={x_min}:{x_max},smooth,variable=\\x, color=blue] plot (\\x,{{{a}*(\\x)^2+{b}*(\\x)+{c}}});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n"
    return  code

#Code vẽ đồ thị hàm số bậc 3   
def codelatex_dothi_bac_3(a,b,c,d):
    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20](-6,-7)grid(6,7);\n\
\\draw[->] (-6,0)--(6,0) node[below left] {{$x$}};\n\
\\draw[->] (0,-7)--(0,7) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{-3,-2,-1,1,2,3}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize $\\x$}};\n\
\\foreach \\y in {{-3,-2,-1,1,2,3}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize $\\y$}};\n\
\\begin{{scope}}\n\
\\clip (-6.5,-7) rectangle (6.5,7);\n\
\\draw[domain=-5.5:5.5,smooth,variable=\\x]\n\
plot (\\x,{{{a}*(\\x)^3+{b}*(\\x)^2+{c}*(\\x)+{d}}});\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}"
    return  code

#Code vẽ đồ thị hàm số bậc 3   
def code_dothi_bac_3(a,b,c,d):
    x=sp.symbols("x")
    f=a*x*3+b*x**2+c*x+d
    a1,b1,c1=3*a,2*b,c
    dau,x_1,x_2=tinh_va_dau_delta(a1,b1,c1)
    if dau==">0":
        y_1, y_2=f.subs(x,x_1), f.subs(x,x_2)
        y_min=min(y_1,y_2)-2.5
        y_max=max(y_1,y_2)+2.5
        if y_max <=0: y_max=2
        if y_min>=0: y_min=-2
        y_min=f"{round(y_min,1):.1f}"
        y_max=f"{round(y_max,1):.1f}"
    else:
        y_min,y_max=-5,5

    code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20](-5,{y_min})grid(5,{y_max});\n\
\\draw[->] (-5,0)--(5,0) node[below left] {{$x$}};\n\
\\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{-3,-2,-1,1,2,3}}\n\
\\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize $\\x$}};\n\
\\foreach \\y in {{-1,1}}\n\
\\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize $\\y$}};\n\
\\begin{{scope}}\n\
\\clip (-5,{y_min}) rectangle (5,{y_max});\n\
\\draw[samples=200,domain=-5:5,smooth,variable=\\x]\n\
plot (\\x,{{{a}*(\\x)^3+{b}*(\\x)^2+{c}*(\\x)+{d}}});\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n"
    return  code

#Code vẽ đồ thị hàm số bậc 4   
def codelatex_dothi_bac_4(a,b,c):
    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20](-4,-5)grid(4,5);\n\
\\draw[->] (-4,0)--(4,0) node[below left] {{$x$}};\n\
\\draw[->] (0,-5)--(0,5) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{-3,-2,-1,1,2,3}}\n\
    \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
\\foreach \\y in {{-4,-3,-2,-1,1,2,3,4}}\n\
    \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
\\begin{{scope}}\n\
\\clip (-4,-5) rectangle (4,5);\n\
\\draw[samples=200,domain=-3.5:3.5,smooth,variable=\\x]\n\
plot (\\x,{{{a}*(\\x)^4+{b}*(\\x)^2+{c}}});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"
    return  code

def code_dothi_bac_4(a,b,c):
    if a>0 and a*b<0: 
        x_0=sqrt(-b/(2*a))
        y_0=a*(-b/(2*a))**2+b*(-b/(2*a))+c
        x_min, x_max=round(-x_0-1.5, 2), round(x_0+1.5, 2)
        if c>=0:
            y_min, y_max =round(y_0-1.5,2), c+2
        if c<0:
            y_min, y_max =round(y_0-1.5,2),2
        day_x=",".join(f"{i}" for i in range(int(x_min),int(x_max)+1))
        day_x=day_x[:len(day_x)-2]
        day_x=day_x.replace(",0","")

        day_y=",".join(f"{i}" for i in range(int(y_min),int(y_max)+1))
        day_y=day_y[:len(day_y)-2]
        day_y=day_y.replace(",0","")

        code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
    \\tikzset{{every node/.style={{scale=0.9}}}}\n\
    \\draw[gray!20]({x_min},{y_min})grid({x_max},{y_max});\n\
    \\draw[->] ({x_min},0)--({x_max},0) node[below left] {{$x$}};\n\
    \\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{$y$}};\n\
    \\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
    \\foreach \\x in {{{day_x}}}\n\
        \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
    \\foreach \\y in {{{day_y}}}\n\
        \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
    \\begin{{scope}}\n\
    \\clip ({x_min},{y_min}) rectangle ({x_max},{y_max});\n\
    \\draw[samples=200,domain={x_min}:{x_max},smooth,variable=\\x]\n\
    plot (\\x,{{{a}*(\\x)^4+{b}*(\\x)^2+{c}}});\n\
    \\end{{scope}}\n\
    \\end{{tikzpicture}}\n"
    elif a<0 and a*b<0:
        x_0=sqrt(-b/(2*a))
        y_0=a*(-b/(2*a))**2+b*(-b/(2*a))+c
        x_min, x_max=round(-x_0-1.5, 2), round(x_0+1.5, 2)
        if c>=0:
            y_min, y_max =-2, round(y_0+1.5,2)
        if c<0:
            y_min, y_max =c-2,max(2,round(y_0+1.5,2))
        day_x=",".join(f"{i}" for i in range(int(x_min),int(x_max)+1))
        day_x=day_x[:len(day_x)-2]
        day_x=day_x.replace(",0","")

        day_y=",".join(f"{i}" for i in range(int(y_min),int(y_max)+1))
        day_y=day_y[:len(day_y)-2]
        day_y=day_y.replace(",0","")

        code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
    \\tikzset{{every node/.style={{scale=0.9}}}}\n\
    \\draw[gray!20]({x_min},{y_min})grid({x_max},{y_max});\n\
    \\draw[->] ({x_min},0)--({x_max},0) node[below left] {{$x$}};\n\
    \\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{$y$}};\n\
    \\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
    \\foreach \\x in {{{day_x}}}\n\
        \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
    \\foreach \\y in {{{day_y}}}\n\
        \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
    \\begin{{scope}}\n\
    \\clip ({x_min},{y_min}) rectangle ({x_max},{y_max});\n\
    \\draw[samples=200,domain={x_min}:{x_max},smooth,variable=\\x]\n\
    plot (\\x,{{{a}*(\\x)^4+{b}*(\\x)^2+{c}}});\n\
    \\end{{scope}}\n\
    \\end{{tikzpicture}}\n"

    else:
        if c>=0:
            y_min, y_max=-2, c+2
        if c<0:
            y_min, y_max=c-2, 3

        day_y=",".join(f"{i}" for i in range(int(y_min)+1,int(y_max)+1))
        day_y=day_y[:len(day_y)-2]
        day_y=day_y.replace(",0","")

        code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
    \\tikzset{{every node/.style={{scale=0.9}}}}\n\
    \\draw[gray!20](-4,{y_min}) grid(4,{y_max});\n\
    \\draw[->] (-4,0)--(4,0) node[below left] {{$x$}};\n\
    \\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{$y$}};\n\
    \\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
    \\foreach \\x in {{-3,-2,-1,1,2,3}}\n\
        \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize$\\x$}};\n\
    \\foreach \\y in {{{day_y}}}\n\
        \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize$\\y$}};\n\
    \\begin{{scope}}\n\
    \\clip (-4,{y_min}) rectangle (4,{y_max});\n\
    \\draw[samples=200,domain=-3.5:3.5,smooth,variable=\\x]\n\
    plot (\\x,{{{a}*(\\x)^4+{b}*(\\x)^2+{c}}});\n\
    \\end{{scope}}\n\
    \\end{{tikzpicture}}\n"

    return  code



#Code vẽ đồ thị hàm số y=(ax^2+bx+c)/(dx+e)   
def codelatex_dothi_phanthuc_bac2(a,b,c,d,e):
    t_1=-e/d-0.1
    t_2=-e/d+0.1    
    code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20](-7.5,-7)grid(7.5,7);\n\
\\draw[->] (-7.5,0)--(7.5,0) node[below left] {{$x$}};\n\
\\draw[->] (0,-7)--(0,7) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{-6,-5,-4,-3,-2,-1,1,2,3,4,5,6}}\n\
    \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize $\\x$}};\n\
\\foreach \\y in {{-6,-5,-4,-3,-2,-1,1,2,3,4,5,6}}\n\
    \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize $\\y$}};\n\
\\begin{{scope}}\n\
\\clip (-7.5,-7) rectangle (7.5,7);\n\
\\draw[samples=200,domain=-7.5:{t_1},smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)^2+{b}*(\\x)+{c})/({d}*(\\x)+{e})}});\n\
\\draw[samples=200,domain={t_2}:7.5,smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)^2+{b}*(\\x)+{c})/({d}*(\\x)+{e})}});\n\
\\draw[samples=200,domain=-7:7,smooth,dashed,variable=\\x]\n\
plot (\\x,{{{a/d}*(\\x) + {(-a*e + b*d)/d**2}}});\n\
\\draw[dashed] ({-e/d},-7)--({-e/d},7);\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n"
    return code

def code_dothi_phanthuc_bac2(a,b,c,d,e):
    x=sp.symbols("x")
    f=(a/d)*x+(-a*e+b*d)/d**2
    y_0=f.subs(x,-e/d)

    y_min,y_max=int(y_0)-15,int(y_0)+15    
    numbers = [f'{i}' for i in range(y_min+1, y_max) if (i!=0 and i%2==0 )]    
    chuoi_so_y = ','.join(numbers)

    x_min, x_max=int(-e/d)-8, int(-e/d)+8
    numbers = [f'{i}' for i in range(x_min+1, x_max) if (i!=0 and i%2==0)]    
    chuoi_so_x = ','.join(numbers)

    t_1=-e/d-0.1
    t_2=-e/d+0.1    
    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.4]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20]({x_min},{y_min})grid({x_max},{y_max});\n\
\\draw[->] ({x_min},0)--({x_max},0) node[below left] {{\\footnotesize $x$}};\n\
\\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{\\footnotesize $y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{{chuoi_so_x}}}\n\
    \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize $\\x$}};\n\
\\foreach \\y in {{{chuoi_so_y}}}\n\
    \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize $\\y$}};\n\
\\begin{{scope}}\n\
\\clip ({x_min},{y_min}) rectangle ({x_max},{y_max});\n\
\\draw[samples=100,domain={x_min}:{t_1},smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)^2+{b}*(\\x)+{c})/({d}*(\\x)+{e})}});\n\
\\draw[samples=100,domain={t_2}:{x_max},smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)^2+{b}*(\\x)+{c})/({d}*(\\x)+{e})}});\n\
\\draw[samples=100,domain={x_min}:{x_max},smooth,dashed,variable=\\x]\n\
plot (\\x,{{{a/d}*(\\x) + {(-a*e + b*d)/d**2}}});\n\
\\draw[dashed] ({-e/d},{y_min})--({-e/d},{y_max});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"
    return code

#Code vẽ đồ thị hàm số y=(ax+b)/(cx+d)   
def code_dothi_phanthuc_bac1(a,b,c,d):
    x=sp.symbols("x")
    t_1=-d/c-0.1
    t_2=-d/c+0.1
    f=(a*x+b)/(c*x+d)
    x_0, y_0=-d/c, a/c
    x_min,x_max=int(x_0)-7, int(x_0)+7
    y_min, y_max=int(y_0)-7,int(y_0)+7
    if x_min>0: 
        x_min=-2.5
    if x_max<0: 
        x_max=2.5
    if y_min>0: 
        y_min=-2.5
    if y_max<0: 
        y_max=2.5

    numbers = [f'{i}' for i in range(x_min+1, x_max) if (i!=0 and i%2==0)]    
    chuoi_so_x = ','.join(numbers)
    numbers = [f'{i}' for i in range(y_min+1, y_max) if (i!=0 and i%2==0 )]    
    chuoi_so_y = ','.join(numbers)

    code = f"\\begin{{tikzpicture}}[line join=round, line cap=round,>=stealth,thick,scale=0.6]\n\
\\tikzset{{every node/.style={{scale=0.9}}}}\n\
\\draw[gray!20]({x_min},{y_min})grid({x_max},{y_max});\n\
\\draw[->] ({x_min},0)--({x_max},0) node[below left] {{$x$}};\n\
\\draw[->] (0,{y_min})--(0,{y_max}) node[below left] {{$y$}};\n\
\\draw (0,0) node [below left] {{\\footnotesize $O$}};\n\
\\foreach \\x in {{{chuoi_so_x}}}\n\
    \\draw[thin] (\\x,1pt)--(\\x,-1pt) node [below] {{\\footnotesize $\\x$}};\n\
\\foreach \\y in {{{chuoi_so_y}}}\n\
    \\draw[thin] (1pt,\\y)--(-1pt,\\y) node [left] {{\\footnotesize $\\y$}};\n\
\\begin{{scope}}\n\
\\clip ({x_min},{y_min}) rectangle ({x_max},{y_max});\n\
\\draw[samples=200,domain={x_min}:{t_1},smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)+{b})/({c}*(\\x)+{d})}});\n\
\\draw[samples=200,domain={t_2}:{x_max},smooth,variable=\\x]\n\
plot (\\x,{{({a}*(\\x)+{b})/({c}*(\\x)+{d})}});\n\
\\draw[dashed,samples=200,domain={x_min}:{x_max},smooth,variable=\\x]\n\
plot (\\x,{{({a/c}}});\n\
\\draw[dashed] ({-d/c},{y_min})--({-d/c},{y_max});\n\
\\end{{scope}}\n\
\\end{{tikzpicture}}\n"
    return  code

#Code vẽ đồ thị hàm số y=(ax+b)/(cx+d)
def codelatex_dothi_phanthuc_bac1(a,b,c,d):  
    code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\
{code_dothi_phanthuc_bac1(a,b,c,d)}\
\\end{{document}}"
    return code



# Code latex bảng biến thiên bậc hai
def codelatex_bbt_bac2(a,b,c):
    x_0 = latex(hien_phan_so(-b/(2*a)))
    y_0 = latex(hien_phan_so((-b**2+4*a*c)/(4*a)))
    if a>0:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$, ${x_0}$, $+\\infty$}} \n\
                \\tkzTabVar{{+/$+\\infty$ , -/ ${y_0}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
    else:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}}\n\
                \\begin{{document}}\n \
                \\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$,${x_0}$,$+\\infty$}}\n\
                \\tkzTabVar{{-/$-\\infty$ , +/ ${y_0}$ /, -/$-\\infty$ /}}\n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
    return  code
def code_bbt_bac2(a,b,c):
    x_0 = latex(hien_phan_so(-b/(2*a)))
    y_0 = latex(hien_phan_so((-b**2+4*a*c)/(4*a)))
    if a>0:
        code =f" \\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$, ${x_0}$, $+\\infty$}} \n\
                \\tkzTabVar{{+/$+\\infty$ , -/ ${y_0}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n"

    else:
        code =f"\\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$,${x_0}$,$+\\infty$}}\n\
                \\tkzTabVar{{-/$-\\infty$ , +/ ${y_0}$ /, -/$-\\infty$ /}}\n\
                \\end{{tikzpicture}}\n"
    return  code

# Code latex bảng biến thiên phân thức bậc 1/1
def codelatex_bbt_phanthucbac1(a,b,c,d):
    x_0 = latex(hien_phan_so(-d/c))
    y_0 = latex(hien_phan_so(a/c))
    if a*d-b*c<0:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.8]\n\
\\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
{{$x$ /1,$y'$ /1,$y$ /2}}\n\
{{$-\\infty$,${x_0}$,$+\\infty$}}\n\
\\tkzTabLine{{,-, d ,-,}}\n\
\\tkzTabVar{{+/ ${y_0}$ / , -D+/ $-\\infty$ / $+\\infty$ , -/ ${y_0}$ /}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n "
    else:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.8]\n\
\\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
{{$x$ /1,$y'$ /1,$y$ /2}}\n\
{{$-\\infty$,${x_0}$,$+\\infty$}}\n\
\\tkzTabLine{{,+, d ,+,}}\n\
\\tkzTabVar{{-/${y_0}$ / , +D-/ $+\\infty$ /$-\\infty$  , +/ ${y_0}$/}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}"
    return code

# Code latex bảng biến thiên phân thức bậc 1/1
def codelatex_bbt_phanthucbac2(a,b,c,d,e):
    x=sp.symbols("x")
    x_0 = latex(hien_phan_so(-e/d))
    f=(a*x**2+b*x+c)/(d*x+e)
    g=diff(f,"x")
    equation=Eq(g,0)
    tap_nghiem=solve(equation,x)
    if "I" in str(tap_nghiem[0]):  
        if a>0:
            code =f"\\documentclass[border=2pt]{{standalone}}\n\
    \\usepackage{{tkz-tab,tikz}}\n\
    \\usetikzlibrary{{calc,intersections,patterns}}\n\
    \\begin{{document}}\n\
    \\begin{{tikzpicture}}[scale=0.8]\n\
    \\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
    {{$x$ /1,$y'$ /1,$y$ /2}}\n\
    {{$-\\infty$,${x_0}$,$+\\infty$}}\n\
    \\tkzTabLine{{,+, d ,+,}}\n\
    \\tkzTabVar{{-/$-\\infty$ / , +D-/ $+\\infty$ /$-\\infty$  , +/ $+\\infty$/}}\n\
    \\end{{tikzpicture}}\n\
    \\end{{document}}"
        else:
            code =f"\\documentclass[border=2pt]{{standalone}}\n\
    \\usepackage{{tkz-tab,tikz}}\n\
    \\usetikzlibrary{{calc,intersections,patterns}}\n\
    \\begin{{document}}\n\
    \\begin{{tikzpicture}}[scale=0.8]\n\
    \\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
    {{$x$ /1,$y'$ /1,$y$ /2}}\n\
    {{$-\\infty$,${x_0}$,$+\\infty$}}\n\
    \\tkzTabLine{{,-, d ,-,}}\n\
    \\tkzTabVar{{+/ $+\\infty$ / , -D+/ $-\\infty$ / $+\\infty$ , -/ $-\\infty$ /}}\n\
    \\end{{tikzpicture}}\n\
    \\end{{document}}\n "
    else:
        x_1,x_2=tap_nghiem[0],tap_nghiem[1]
        y_1=f.subs(x,x_1)
        y_2=f.subs(x,x_2)
        if check_so_nguyen(x_1):
            x_1=latex(hien_phan_so(x_1))
            x_2=latex(hien_phan_so(x_2))
        else:
            x_1=latex(x_1)
            x_2=latex(x_2)

        if check_so_nguyen(y_1):
            y_1=latex(hien_phan_so(y_1))
            y_2=latex(hien_phan_so(y_2))
        else:
            y_1=latex(y_1)
            y_2=latex(y_2)

        if a>0:
            code =f"\\documentclass[border=2pt]{{standalone}}\n\
    \\usepackage{{tkz-tab,tikz}}\n\
    \\usetikzlibrary{{calc,intersections,patterns}}\n\
    \\begin{{document}}\n\
    \\begin{{tikzpicture}}[scale=0.8]\n\
    \\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
    {{$x$ /1,$y'$ /1,$y$ /2}}\n\
    {{$-\\infty$,${x_1}$, ${x_0}$, ${x_2}$, $+\\infty$}}\n\
\\tkzTabLine{{ ,+,z,-,d,-,z,+, }}\n\
\\tkzTabVar{{-/$-\\infty$,+/${y_1}$,-D+/$-\\infty$/$+\\infty$,-/${y_2}$,+/$+\\infty$}}\n\
    \\end{{tikzpicture}}\n\
    \\end{{document}}"
        else:
            code =f"\\documentclass[border=2pt]{{standalone}}\n\
    \\usepackage{{tkz-tab,tikz}}\n\
    \\usetikzlibrary{{calc,intersections,patterns}}\n\
    \\begin{{document}}\n\
    \\begin{{tikzpicture}}[scale=0.8]\n\
    \\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
    {{$x$ /1,$y'$ /1,$y$ /2}}\n\
    {{$-\\infty$,${x_1}$, ${x_0}$, ${x_2}$, $+\\infty$}}\n\
\\tkzTabLine{{ ,-,z,+,d,+,z,-, }}\n\
\\tkzTabVar{{+/$+\\infty$,-/${y_1}$,+D-/$+\\infty$/$-\\infty$,+/${y_2}$,-/$-\\infty$}}\n\
    \\end{{tikzpicture}}\n\
    \\end{{document}}"
    return code

def code_bbt_phanthucbac1(a,b,c,d):
    x_0 = latex(hien_phan_so(-d/c))
    y_0 = latex(hien_phan_so(a/c))
    if a*d-b*c<0:
        code =f"\\begin{{tikzpicture}}[scale=1]\n\
\\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
{{$x$ /1,$y'$ /1,$y$ /2}}\n\
{{$-\\infty$,${x_0}$,$+\\infty$}}\n\
\\tkzTabLine{{,-, d ,-,}}\n\
\\tkzTabVar{{+/ ${y_0}$ / , -D+/ $-\\infty$ / $+\\infty$ , -/ ${y_0}$ /}}\n\
\\end{{tikzpicture}}\n"
    else:
        code =f"\\begin{{tikzpicture}}[scale=1]\n\
\\tkzTabInit[nocadre=false,lgt=1.5,espcl=3]\n\
{{$x$ /1,$y'$ /1,$y$ /2}}\n\
{{$-\\infty$,${x_0}$,$+\\infty$}}\n\
\\tkzTabLine{{,+, d ,+,}}\n\
\\tkzTabVar{{-/${y_0}$ / , +D-/ $+\\infty$ /$-\\infty$  , +/ ${y_0}$/}}\n\
\\end{{tikzpicture}}\n"
    return code 



# Code latex bảng biến thiên bậc 3 hai nghiệm
def codelatex_bbt_bac3_2nghiem(a,x_1,x_2,y_1,y_2):
    if a>0:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
            \\usepackage{{tkz-tab,tikz}} \n\
            \\usetikzlibrary{{calc,intersections,patterns}}\n\
            \\begin{{document}} \n\
            \\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n \
            \\tkzTabLine{{,+,0,-,0,+,}} \n \
            \\tkzTabVar{{-/$-\\infty$ ,+/ ${y_1}$, -/ ${y_2}$ /, +/$+\\infty$ /}} \n\
            \\end{{tikzpicture}}\n \
            \\end{{document}}\n"
    else:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
            \\usepackage{{tkz-tab,tikz}} \n\
            \\usetikzlibrary{{calc,intersections,patterns}}\n\
            \\begin{{document}} \n\
            \\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n \
            \\tkzTabLine{{,-,0,+,0,-,}} \n \
            \\tkzTabVar{{+/$+\\infty$ ,-/ ${y_1}$, +/ ${y_2}$ /, -/$-\\infty$ /}} \n\
            \\end{{tikzpicture}}\n \
            \\end{{document}}\n"
    return code

def code_bbt_bac3_2nghiem(a,x_1,x_2,y_1,y_2):
    if a>0:
            code = f"\\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n \
            \\tkzTabLine{{,+,0,-,0,+,}} \n \
            \\tkzTabVar{{-/$-\\infty$ ,+/ ${y_1}$, -/ ${y_2}$ /, +/$+\\infty$ /}} \n\
            \\end{{tikzpicture}}\n"         
    else:
            code = f"\\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n \
            \\tkzTabLine{{,-,0,+,0,-,}} \n \
            \\tkzTabVar{{+/$+\\infty$ ,-/ ${y_1}$, +/ ${y_2}$ /, -/$-\\infty$ /}} \n\
            \\end{{tikzpicture}}\n"

    return code
#Tạo hàm bậc 2
def tao_ham_bac_2():
    x=symbols("x")
    a = random.choice([random.randint(-7, -1), random.randint(1, 7)])
    b = random.choice([random.randint(-10, -1), random.randint(1, 10)])
    c = random.randint(-8, 8)
    ham=a*x**2+b*x+c
    return ham
    
#Tính nghiệm và xét dấu delta
def tinh_va_dau_delta(a,b,c):
    d=b**2-4*a*c
    if d<0:
        dau="<0"
        x_1=""
        x_2=""
    if d==0:
        dau="=0"
        x_1=-b/(2*a)
        x_2=-b/(2*a)
    if d>0:
        dau=">0"
        x_1=(-b-sp.sqrt(d))/(2*a)
        x_2=(-b+sp.sqrt(d))/(2*a)
        if x_1>x_2:
            t=x_2
            x_2=x_1
            x_1=t
    return dau, x_1, x_2

#Khởi đầu file latex
def start_latex_file():
    content = r"""
\documentclass[border=2pt]{standalone}
\usepackage[T1]{fontenc}
\usepackage[utf8]{inputenc}
\usepackage{tkz-euclide}
\usepackage{tikz,tkz-tab,tkz-linknodes}
%\usetikzlibrary{snakes}
\usepackage{tabvar}
%\usetkzobj{all}
%\usepackage{tikz-3dplot}
\usetikzlibrary{shapes.geometric,arrows,calc,intersections,angles,quotes,patterns,snakes,positioning}

"""
    return content

def moi_truong_anh_latex(code):
    code_header=r"""\documentclass[border=2pt]{standalone}
    \usepackage[utf8]{inputenc}
    \usepackage[vietnamese]{babel}
    \usepackage{tkz-tab,tikz}
    \usepackage[dethi]{ex_test}   
    \usetikzlibrary{shapes.geometric,arrows,calc,intersections,angles,quotes,patterns,snakes,positioning}
    \begin{document} """   

    code_footer=r"""\end{document} """
    
    code_latex=f"{code_header}\n\n{code}\n\n{code_footer}"            
    return code_latex
    

def moi_truong_latex(code):
    code_header=r"""\documentclass[12pt,a4paper]{article}
\usepackage[top=1.5cm, bottom=1.5cm, left=2.0cm, right=1.5cm] {geometry}
\usepackage{amsmath,amssymb,txfonts}
\usepackage{tkz-euclide}
\usepackage{setspace}
\usepackage{lastpage}

\usepackage{tikz,tkz-tab}
%\usepackage[solcolor]{ex_test}
\usepackage[dethi]{ex_test} % Chỉ hiển thị đề thi
%\usepackage[loigiai]{ex_test} % Hiển thị lời giải
%\usepackage[color]{ex_test} % Khoanh các đáp án
\usetikzlibrary{shapes.geometric,arrows,calc,intersections,angles,quotes,patterns,snakes,positioning}
\everymath{\displaystyle}

\def\colorEX{\color{purple}}
%\def\colorEX{}%Không tô màu đáp án đúng trong tùy chọn loigiai
\renewtheorem{ex}{\color{violet}Câu}
\renewcommand{\FalseEX}{\stepcounter{dapan}{{\bf \textcolor{blue}{\Alph{dapan}.}}}}
\renewcommand{\TrueEX}{\stepcounter{dapan}{{\bf \textcolor{blue}{\Alph{dapan}.}}}}

%---------- Khai báo viết tắt, in đáp án
\newcommand{\hoac}[1]{ %hệ hoặc
    \left[\begin{aligned}#1\end{aligned}\right.}
\newcommand{\heva}[1]{ %hệ và
    \left\{\begin{aligned}#1\end{aligned}\right.}

%Tiêu đề
\newcommand{\tenso}{}
\newcommand{\tentruong}{}
\newcommand{\tenkythi}{}
\newcommand{\tenmonthi}{}
\newcommand{\thoigian}{}
\newcommand{\tieude}[1]{
    \noindent
     \begin{minipage}[b]{6cm}
    \centerline{\textbf{\fontsize{11}{0}\selectfont \tenso}}
    \centerline{\fontsize{11}{0}\selectfont \tentruong}  
  \end{minipage}\hspace{1cm}
  \begin{minipage}[b]{11cm}
    \centerline{\textbf{\fontsize{11}{0}\selectfont \tenkythi}}
    \centerline{\textbf{\fontsize{11}{0}\selectfont \tenmonthi}}
    \centerline{\textit{\fontsize{11}{0}\selectfont Thời \underline{gian làm bài: \thoigian  } phút }}
  \end{minipage}
  \vspace*{3mm}
  \noindent
  \begin{minipage}[t]{12cm}
    \textbf{Họ, tên thí sinh:}\dotfill\\
    \textbf{Số báo danh:}\dotfill
  \end{minipage}\hfill
  \begin{minipage}[b]{3cm}
    \setlength\fboxrule{1pt}
    \setlength\fboxsep{3pt}
    \vspace*{3mm}\fbox{\bf Mã đề thi #1}
  \end{minipage}\\
}

\newcommand{\chantrang}[2]{\rfoot{Trang \thepage $-$ Mã đề #2}}
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0pt} 
\renewcommand{\footrulewidth}{0pt}

\begin{document}
%Thiết lập giãn dọng 1.5cm 
%\setlength{\lineskip}{1.5em}



%Nội dung trắc nghiệm bắt đầu ở đây
"""            
    code_footer=r"""

\end{document}"""
    code_latex=f"{code_header}\n\n{code}\n\n{code_footer}"            
    return code_latex


def moi_truong_latex_loigiai(code):
    code_header=r"""\documentclass[12pt,a4paper]{article}
\usepackage[top=1.5cm, bottom=1.5cm, left=2.0cm, right=1.5cm] {geometry}
\usepackage{amsmath,amssymb,txfonts}
\usepackage{tkz-euclide}
\usepackage{setspace}
\usepackage{lastpage}

\usepackage{tikz,tkz-tab}
%\usepackage[solcolor]{ex_test}
%\usepackage[dethi]{ex_test} % Chỉ hiển thị đề thi
\usepackage[loigiai]{ex_test} % Hiển thị lời giải
%\usepackage[color]{ex_test} % Khoanh các đáp án
\everymath{\displaystyle}

\def\colorEX{\color{purple}}
%\def\colorEX{}%Không tô màu đáp án đúng trong tùy chọn loigiai
\renewtheorem{ex}{\color{violet}Câu}
\renewcommand{\FalseEX}{\stepcounter{dapan}{{\bf \textcolor{blue}{\Alph{dapan}.}}}}
\renewcommand{\TrueEX}{\stepcounter{dapan}{{\bf \textcolor{blue}{\Alph{dapan}.}}}}

%---------- Khai báo viết tắt, in đáp án
\newcommand{\hoac}[1]{ %hệ hoặc
    \left[\begin{aligned}#1\end{aligned}\right.}
\newcommand{\heva}[1]{ %hệ và
    \left\{\begin{aligned}#1\end{aligned}\right.}

%Tiêu đề
\newcommand{\tenso}{}
\newcommand{\tentruong}{}
\newcommand{\tenkythi}{}
\newcommand{\tenmonthi}{}
\newcommand{\thoigian}{}
\newcommand{\tieude}[1]{
   \begin{tabular}{cm{3cm}cm{3cm}cm{3cm}}
    {\bf \tenso} & & {\bf \tenkythi} \\
    {\bf \tentruong} & & {\bf \tenmonthi}\\
    && {\bf Thời gian: \bf \thoigian \, phút}\\
    && { \fbox{\bf Mã đề: #1}}
   \end{tabular}\\\\
    
   {Họ tên HS: \dotfill Số báo danh \dotfill}\\
}
\newcommand{\chantrang}[2]{\rfoot{Trang \thepage $-$ Mã đề #2}}
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0pt} 
\renewcommand{\footrulewidth}{0pt}
\usetikzlibrary{shapes.geometric,arrows,calc,intersections,angles,quotes,patterns,snakes,positioning}

\begin{document}
%Thiết lập giãn dọng 1.5cm 
%\setlength{\lineskip}{1.5em}
%Nội dung trắc nghiệm bắt đầu ở đây
"""            
    code_footer=r"""
\end{document}"""
    code_latex=f"{code_header}\n\n{code}\n\n{code_footer}"            
    return code_latex



def codelatex_bxd_bac2(a,b,c):
    dau,x_1,x_2=tinh_va_dau_delta(a,b,c)[0:3]    
    if dau==">0":   
        x_1=latex(hien_phan_so(x_1))
        x_2=latex(hien_phan_so(x_2))
        if a>0:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=1, espcl=1.3] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n\
            \\tkzTabLine{{,+,0,-,0,+,}}\n\
            \\end{{tikzpicture}}\n"
        else:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=1, espcl=1.3] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,${x_1}$,${x_2}$,$+\\infty$}}\n\
            \\tkzTabLine{{,-,0,+,0,-,}}\n\
            \\end{{tikzpicture}}\n"
    if dau=="=0":
        x_1=latex(hien_phan_so(x_1))
        if a>0:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=2, espcl=4] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,${x_1}$,$+\\infty$}}\n\
            \\tkzTabLine{{,+,0,+,}}\n\
            \\end{{tikzpicture}}\n"
        else:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=2, espcl=4] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,${x_1}$,$+\\infty$}}\n\
            \\tkzTabLine{{,-,0,-,}}\n\
            \\end{{tikzpicture}}\n"
    if dau=="<0":
        if a>0:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=2, espcl=4] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,$+\\infty$}}\n\
            \\tkzTabLine{{,+,}}\n\
            \\end{{tikzpicture}}\n"
        else:
            code =f"\\begin{{tikzpicture}}\n\
            \\tkzTabInit[nocadre=false, lgt=2, espcl=4] \n\
            {{$x$ /1,$f(x)$ /1}}\n\
            {{$-\\infty$,$+\\infty$}}\n\
            \\tkzTabLine{{,-,}}\n\
            \\end{{tikzpicture}}\n"
    return code

def kiem_tra_chinh_phuong(n):
    ketqua=False
    if n < 0:
        return False  # Số âm không thể là số chính phương
    if sqrt(n)==int(sqrt(n)):
        return True

def codelatex_bbt_bac3(a,b,c,d):
    x=sp.symbols("x")
    f=a*x**3 + b*x**2 + c*x + d
    a1,b1,c1 = 3*a, 2*b, c
    dau,x_1,x_2=tinh_va_dau_delta(a1,b1,c1)[0:3]

    if dau==">0": 
        y_1=f.subs(x,x_1)
        y_2=f.subs(x,x_2)
        delta=b1**2-4*a1*c1
        sqrt_delta=sp.sqrt(delta)
        if a>0:
            if kiem_tra_chinh_phuong(b1**2-4*a1*c1):            
                code = f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}}\n \
                \\tkzTabInit[nocadre=false, lgt=0.8, espcl=1.3] \n \
                {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
                {{$-\\infty$,${latex(hien_phan_so(x_1))}$,${latex(hien_phan_so(x_2))}$,$+\\infty$}}\n \
                \\tkzTabLine{{,+,0,-,0,+,}} \n \
                \\tkzTabVar{{-/$-\\infty$ ,+/ ${latex(hien_phan_so(y_1))}$, -/ ${latex(hien_phan_so(y_2))}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n \
                \\end{{document}}\n"
            else:
                code = f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}}\n \
                \\tkzTabInit[nocadre=false, lgt=0.8, espcl=1.3] \n \
                {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
                {{$-\\infty$,${latex(x_1)}$,${latex(x_2)}$,$+\\infty$}}\n \
                \\tkzTabLine{{,+,0,-,0,+,}} \n \
                \\tkzTabVar{{-/$-\\infty$ ,+/ ${latex(y_1)}$, -/ ${latex(y_2)}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n \
                \\end{{document}}\n"

        else:
            if kiem_tra_chinh_phuong(b1**2-4*a1*c1):
                code = f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}}\n \
                \\tkzTabInit[nocadre=false, lgt=0.8, espcl=1.3] \n \
                {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
                {{$-\\infty$,${latex(x_1)}$,${latex(x_2)}$,$+\\infty$}}\n \
                \\tkzTabLine{{,-,0,+,0,-,}} \n \
                \\tkzTabVar{{+/$+\\infty$ ,-/ ${latex(y_1)}$, +/ ${latex(y_2)}$ /, -/$-\\infty$ /}} \n\
                \\end{{tikzpicture}}\n \
                \\end{{document}}\n"
            else:
                code = f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}}\n \
                \\tkzTabInit[nocadre=false, lgt=0.8, espcl=1.3] \n \
                {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
                {{$-\\infty$,${latex(hien_phan_so(x_1))}$,${latex(hien_phan_so(x_2))}$,$+\\infty$}}\n \
                \\tkzTabLine{{,-,0,+,0,-,}} \n \
                \\tkzTabVar{{+/$+\\infty$ ,-/ ${latex(hien_phan_so(y_1))}$, +/ ${latex(hien_phan_so(y_2))}$ /, -/$-\\infty$ /}} \n\
                \\end{{tikzpicture}}\n \
                \\end{{document}}\n"
    if dau=="=0": 
        if a>0:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=0.5,espcl=1.2] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ , ${latex(hien_phan_so(x_1))}$, $+\\infty$}}\n\
\\tkzTabLine{{,+,0,+,}} \n\
\\tkzTabVar{{-/$-\\infty$,R,+/$+\\infty$}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"
        else:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=0.5,espcl=1.2] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ , ${latex(hien_phan_so(x_1))}$, $+\\infty$}}\n\
\\tkzTabLine{{,-,0,-,}} \n\
\\tkzTabVar{{+/$+\\infty$,R,-/$-\\infty$}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"
    if dau=="<0":
        if a>0:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=0.5,espcl=1.6] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ ,$+\\infty$}}\n\
\\tkzTabLine{{,+,}}\n\
\\tkzTabVar{{-/$-\\infty$,+/$+\\infty$}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"
        else:
            code = f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tkz-tab,tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=0.5,espcl=1.6] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ ,$+\\infty$}}\n\
\\tkzTabLine{{,-,}}\n\
\\tkzTabVar{{+/$+\\infty$,-/$-\\infty$}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}\n"       
    return code

def code_bbt_bac3(a,b,c,d):
    x=sp.symbols("x")
    f=a*x**3 + b*x**2 + c*x + d
    a1,b1,c1 = 3*a, 2*b, c
    dau,x_1,x_2=tinh_va_dau_delta(a1,b1,c1)[0:3]

    if dau==">0": 
        y_1=f.subs(x,x_1)
        y_2=f.subs(x,x_2)
        if a>0:
            code = f"\\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${latex(hien_phan_so(x_1))}$,${latex(hien_phan_so(x_2))}$,$+\\infty$}}\n \
            \\tkzTabLine{{,+,0,-,0,+,}} \n \
            \\tkzTabVar{{-/$-\\infty$ ,+/ ${latex(hien_phan_so(y_1))}$, -/ ${latex(hien_phan_so(y_2))}$ /, +/$+\\infty$ /}} \n\
            \\end{{tikzpicture}}\n"
        else:
            code = f" \\begin{{tikzpicture}}\n \
            \\tkzTabInit[nocadre=false, lgt=1, espcl=2] \n \
            {{$x$ /0.7,$y'$ /0.7,$y$ /1.7}}\n \
            {{$-\\infty$,${latex(hien_phan_so(x_1))}$,${latex(hien_phan_so(x_2))}$,$+\\infty$}}\n \
            \\tkzTabLine{{,-,0,+,0,-,}} \n \
            \\tkzTabVar{{+/$+\\infty$ ,-/ ${latex(hien_phan_so(y_1))}$, +/ ${latex(hien_phan_so(y_2))}$ /, -/$-\\infty$ /}} \n\
            \\end{{tikzpicture}}\n"
    if dau=="=0": 
        if a>0:
            code = f"\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=1,espcl=3] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ , ${latex(hien_phan_so(x_1))}$, $+\\infty$}}\n\
\\tkzTabLine{{,+,0,+,}} \n\
\\tkzTabVar{{-/$-\\infty$,R,+/$+\\infty$}}\n\
\\end{{tikzpicture}}\n"
        else:
            code = f"\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=1,espcl=3] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ , ${latex(hien_phan_so(x_1))}$, $+\\infty$}}\n\
\\tkzTabLine{{,-,0,-,}} \n\
\\tkzTabVar{{+/$+\\infty$,R,-/$-\\infty$}}\n\
\\end{{tikzpicture}}\n"
    if dau=="<0":
        if a>0:
            code = f"\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=1,espcl=3] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ ,$+\\infty$}}\n\
\\tkzTabLine{{,+,}}\n\
\\tkzTabVar{{-/$-\\infty$,+/$+\\infty$}}\n\
\\end{{tikzpicture}}\n"
        else:
            code = f"\\begin{{tikzpicture}}\n\
\\tkzTabInit[lgt=1,espcl=3] \n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.2}}\n\
{{$-\\infty$ ,$+\\infty$}}\n\
\\tkzTabLine{{,-,}}\n\
\\tkzTabVar{{+/$+\\infty$,-/$-\\infty$}}\n\
\\end{{tikzpicture}}\n"       
    return code

# Code latex bảng biến thiên bậc bốn trùng phương
def code_bbt_bac4(a,b,c):
    x=sp.symbols("x")
    f=a*x**4+b*x**2+c
    if a>0 and b>=0:
        code =f"\\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$, $0$, $+\\infty$}} \n\
                \\tkzTabVar{{+/$+\\infty$ , -/ ${latex(hien_phan_so(c))}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n"
    if a<0 and b<=0:
        code =f"\\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y$ /2}}\n\
                {{$-\\infty$,$0$,$+\\infty$}}\n\
                \\tkzTabVar{{-/$-\\infty$ , +/ ${latex(hien_phan_so(c))}$ /, -/$-\\infty$ /}}\n\
                \\end{{tikzpicture}}\n"
    if a>0 and b<0:

        x_0=(sqrt(-b*2*a))/(2*a)        
        y_0=f.subs(x,x_0)
        if check_so_nguyen(x_0):
            x_1=latex(hien_phan_so(rutgon_can(sqrt(-b*2*a)/(2*a))))
        else:
            x_1=f"{latex((sqrt(-b*2*a))/(2*a))}"
        code =f"\\begin{{tikzpicture}} \n\
                \\tkzTabInit[lgt=1,espcl=1.5]\n\
                {{$x$  /1, $y'$ /1,$y$  /2}}\n\
                {{$-\\infty$ , $-{x_1}$,$0$,${x_1}$, $+\\infty$}}\n\
                \\tkzTabLine{{,-,0,+,0,-,0,+}} \n\
                \\tkzTabVar{{+/$+\\infty$ ,-/${latex(hien_phan_so(y_0))}$,+/${latex(hien_phan_so(c))}$,-/${latex(hien_phan_so(y_0))}$,+/$+\\infty$}}\n\
                \\end{{tikzpicture}}\n"
    if a<0 and b>0: 
        x_0=(sqrt(-b*2*a))/(-2*a) 
        y_0=f.subs(x,x_0)
        if check_so_nguyen(x_0):
            x_1=latex(hien_phan_so(rutgon_can((sqrt(-b*2*a))/(-2*a))))
        else:
            x_1=latex(rutgon_can(sqrt(-b*2*a)/(-2*a)))
        code =f"\\begin{{tikzpicture}} \n\
\\tkzTabInit[lgt=1,espcl=1.5]\n\
{{$x$  /1, $y'$ /1,$y$  /2}}\n\
{{$-\\infty$ , $-{x_1}$,$0$,${x_1}$, $+\\infty$}}\n\
\\tkzTabLine{{,+,0,-,0,+,0,-}} \n\
\\tkzTabVar{{-/$-\\infty$ ,+/${latex(hien_phan_so(y_0))}$,-/${latex(hien_phan_so(c))}$,+/${latex(hien_phan_so(y_0))}$,-/$-\\infty$}}\n\
                \\end{{tikzpicture}}\n"
   
    return code

# Code latex bảng biến thiên bậc bốn trùng phương
def codelatex_bbt_bac4(a,b,c):
    x=sp.symbols("x")
    f=a*x**4+b*x**2+c
    if a>0 and b>=0:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}} \n\
                \\usetikzlibrary{{calc,intersections,patterns}}\n\
                \\begin{{document}} \n\
                \\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1, $y'$ /1 ,$y$ /2}}\n\
                {{$-\\infty$, $0$, $+\\infty$}} \n\
                \\tkzTabLine{{,-,0,+}} \n\
                \\tkzTabVar{{+/$+\\infty$ , -/ ${latex(hien_phan_so(c))}$ /, +/$+\\infty$ /}} \n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
    if a<0 and b<=0:
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}}\n\
                \\begin{{document}}\n \
                \\begin{{tikzpicture}} \n\
                \\tkzTabInit[nocadre=false, lgt=1, espcl=1.5] \n\
                {{$x$ /1,$y'$ /1, $y$ /2}}\n\
                {{$-\\infty$,$0$,$+\\infty$}}\n\
                \\tkzTabLine{{,+,0,-}} \n\
                \\tkzTabVar{{-/$-\\infty$ , +/ ${latex(hien_phan_so(c))}$ /, -/$-\\infty$ /}}\n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
    if a>0 and b<0:
        x_0=(sqrt(-b*2*a))/(2*a)        
        y_0=f.subs(x,x_0)
        if check_so_nguyen(x_0):
            x_1=latex(hien_phan_so(rutgon_can(sqrt(-b*2*a)/(2*a))))
        else:
            x_1=f"{latex((sqrt(-b*2*a))/(2*a))}"
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}}\n\
                \\begin{{document}}\n \
                \\begin{{tikzpicture}} \n\
                \\tkzTabInit[lgt=0.6,espcl=1.5]\n\
                {{$x$  /0.7, $y'$ /0.7,$y$  /1.3}}\n\
                {{$-\\infty$ , $-{x_1}$,$0$,${x_1}$, $+\\infty$}}\n\
                \\tkzTabLine{{,-,0,+,0,-,0,+}} \n\
                \\tkzTabVar{{+/$+\\infty$ ,-/${latex(hien_phan_so(y_0))}$,+/${latex(hien_phan_so(c))}$,-/${latex(hien_phan_so(y_0))}$,+/$+\\infty$}}\n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
    if a<0 and b>0:
        x_0=(sqrt(-b*2*a))/(-2*a) 
        y_0=f.subs(x,x_0)
        if check_so_nguyen(x_0):
            x_1=latex(hien_phan_so(rutgon_can((sqrt(-b*2*a))/(-2*a))))
        else:
            x_1=f"{latex(rutgon_can(sqrt(-b*2*a)/(-2*a)))}"
        code =f"\\documentclass[border=2pt]{{standalone}}\n\
                \\usepackage{{tkz-tab,tikz}}\n\
                \\begin{{document}}\n \
                \\begin{{tikzpicture}} \n\
\\tkzTabInit[lgt=0.6,espcl=1.5]\n\
{{$x$  /0.7, $y'$ /0.7,$y$  /1.3}}\n\
{{$-\\infty$ , $-{x_1}$,$0$,${x_1}$, $+\\infty$}}\n\
\\tkzTabLine{{,+,0,-,0,+,0,-}} \n\
\\tkzTabVar{{-/$-\\infty$ ,+/${latex(hien_phan_so(y_0))}$,-/${latex(hien_phan_so(c))}$,+/${latex(hien_phan_so(y_0))}$,-/$-\\infty$}}\n\
                \\end{{tikzpicture}}\n\
                \\end{{document}}\n"
   
    return code

# Code latex bảng biến thiên tìm tiệm cận
def codelatex_bbt_tim_tiemcan(x_1,x_2,x_3,y_1,y_2,y_3,y_4,y_5,y_6):
    #y_1,y_3 có thể là -\\infty hoặc số
    #y_4,y_6 có thể là +\\infty hoặc số
    code =f"\\begin{{tikzpicture}}[>=stealth,scale=0.8]\n\
\\tkzTabInit[lgt=1.3,espcl=3] \n\
{{$x$/1.2,$f’(x)$/1.2,$f(x)$/2.5}}\n\
{{$-\\infty$,${x_1}$,${x_2}$,${x_3}$,$+\\infty$}}\n\
\\tkzTabLine{{ ,+,z,-,d,-,z,+, }}\n\
\\tkzTabVar{{-/${y_1}$,+/${y_2}$,-D+/${y_3}$/${y_4}$,-/${y_5}$,+/${y_6}$}}\n\
\\end{{tikzpicture}}\n"    
    return code
def code_bbt_tim_tiemcan(x_1,x_2,x_3,y_1,y_2,y_3,y_4,y_5,y_6):
    #y_1,y_3 có thể là -\\infty hoặc số
    #y_4,y_6 có thể là +\\infty hoặc số
    code =f"\\begin{{tikzpicture}}[>=stealth,scale=0.8]\n\
\\tkzTabInit[lgt=1.3,espcl=3] \n\
{{$x$/1.2,$f’(x)$/1.2,$f(x)$/2.5}}\n\
{{$-\\infty$,${x_1}$,${x_2}$,${x_3}$,$+\\infty$}}\n\
\\tkzTabLine{{ ,+,z,-,d,-,z,+, }}\n\
\\tkzTabVar{{-/${y_1}$,+/${y_2}$,-D+/${y_3}$/${y_4}$,-/${y_5}$,+/${y_6}$}}\n\
\\end{{tikzpicture}}\n"    
    return code 

#Code latex vẽ hình học
def codelatex_hinh_hop(a,b,c,d,e,f,g,h):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.8] \n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (3,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (4,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({e}) at (0,2)   node at ({e}) [left] {{${e}$}};\n\
\\coordinate ({f}) at (-1,1) node at ({f}) [left] {{${f}$}};\n\
\\coordinate ({g}) at (3,1)  node at ({g}) [right] {{${g}$}};\n\
\\coordinate ({h}) at (4,2)   node at ({h}) [right] {{${h}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({e})--({a});\n\
\\draw ({e})--({f})--({g})--({h})--({e}) ({f})--({b}) ({g})--({c}) ({h})--({d});\n\
\\draw ({b})--({c})--({d});\n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code


#Code latex vẽ hình học
def ve_hinh_lapphuong(a,b,c,d,e,f,g,h):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (2,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (3,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({e}) at (0,2)   node at ({e}) [left] {{${e}$}};\n\
\\coordinate ({f}) at (-1,1) node at ({f}) [left] {{${f}$}};\n\
\\coordinate ({g}) at (2,1)  node at ({g}) [right] {{${g}$}};\n\
\\coordinate ({h}) at (3,2)   node at ({h}) [right] {{${h}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({e})--({a});\n\
\\draw ({e})--({f})--({g})--({h})--({e}) ({f})--({b}) ({g})--({c}) ({h})--({d});\n\
\\draw ({b})--({c})--({d});\n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

def code_hinh_lapphuong(a,b,c,d,e,f,g,h):
    code=f"\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (2,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (3,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({e}) at (0,2)   node at ({e}) [left] {{${e}$}};\n\
\\coordinate ({f}) at (-1,1) node at ({f}) [left] {{${f}$}};\n\
\\coordinate ({g}) at (2,1)  node at ({g}) [right] {{${g}$}};\n\
\\coordinate ({h}) at (3,2)   node at ({h}) [right] {{${h}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({e})--({a});\n\
\\draw ({e})--({f})--({g})--({h})--({e}) ({f})--({b}) ({g})--({c}) ({h})--({d});\n\
\\draw ({b})--({c})--({d});\n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n"
    return code

def code_hinh_hop(a,b,c,d,e,f,g,h):
    code=f"\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (3,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (4,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({e}) at (0,2)   node at ({e}) [left] {{${e}$}};\n\
\\coordinate ({f}) at (-1,1) node at ({f}) [left] {{${f}$}};\n\
\\coordinate ({g}) at (3,1)  node at ({g}) [right] {{${g}$}};\n\
\\coordinate ({h}) at (4,2)   node at ({h}) [right] {{${h}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({e})--({a});\n\
\\draw ({e})--({f})--({g})--({h})--({e}) ({f})--({b}) ({g})--({c}) ({h})--({d});\n\
\\draw ({b})--({c})--({d});\n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n"
    return code

def codelatex_hinh_langtruxien_tamgiac(a,b,c,a1,b1,c1):
    code=f"\\documentclass[border=2pt]{{standalone}} \n\
\\usepackage{{tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}} \n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}} \n\
    \\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}}; \n\
    \\coordinate ({b}) at (2,-1) node at ({b}) [below] {{${b}$}}; \n\
    \\coordinate ({c}) at (4,0)  node at ({c}) [below] {{${c}$}}; \n\
    \\coordinate ({a1}) at (1,4)   node at ({a1}) [above] {{${a1}$}}; \n\
    \\coordinate ({b1}) at (3,3)   node at ({b1}) [above] {{${b1}$}}; \n\
    \\coordinate ({c1}) at (5,4)   node at ({c1}) [above] {{${c1}$}};\n\
    \\draw [dashed] ({a})--({c}); \n\
    \\draw ({a})--({b})--({b1})--({a1})--({a}); \n\
    \\draw ({b1})--({c1})--({c})--({b}) ({a1})--({c1})--({c}); \n\
\\end{{scriptsize}} \n\
\\end{{tikzpicture}} \n\
\\end{{document}}"
    return code

def code_hinh_langtruxien_tamgiac(a,b,c,a1,b1,c1):
    code=f"\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}} \n\
    \\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}}; \n\
    \\coordinate ({b}) at (2,-1) node at ({b}) [below] {{${b}$}}; \n\
    \\coordinate ({c}) at (4,0)  node at ({c}) [below] {{${c}$}}; \n\
    \\coordinate ({a1}) at (1,4)   node at ({a1}) [above] {{${a1}$}}; \n\
    \\coordinate ({b1}) at (3,3)   node at ({b1}) [above] {{${b1}$}}; \n\
    \\coordinate ({c1}) at (5,4)   node at ({c1}) [above] {{${c1}$}};\n\
    \\draw [dashed] ({a})--({c}); \n\
    \\draw ({a})--({b})--({b1})--({a1})--({a}); \n\
    \\draw ({b1})--({c1})--({c})--({b}) ({a1})--({c1})--({c}); \n\
\\end{{scriptsize}} \n\
\\end{{tikzpicture}} \n"
    return code

def code_hinh_langtrudung_tamgiac(a,b,c,a1,b1,c1):
    code=f"\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}} \n\
    \\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}}; \n\
    \\coordinate ({b}) at (2,-1) node at ({b}) [below] {{${b}$}}; \n\
    \\coordinate ({c}) at (4,0)  node at ({c}) [below] {{${c}$}}; \n\
    \\coordinate ({a1}) at (0,4)   node at ({a1}) [above] {{${a1}$}}; \n\
    \\coordinate ({b1}) at (2,3)   node at ({b1}) [above] {{${b1}$}}; \n\
    \\coordinate ({c1}) at (4,4)   node at ({c1}) [above] {{${c1}$}};\n\
    \\draw [dashed] ({a})--({c}); \n\
    \\draw ({a})--({b})--({b1})--({a1})--({a}); \n\
    \\draw ({b1})--({c1})--({c})--({b}) ({a1})--({c1})--({c}); \n\
\\end{{scriptsize}} \n\
\\end{{tikzpicture}} \n"
    return code

def ve_hinh_langtrudung_tamgiac(a,b,c,a1,b1,c1):
    code=f"\\documentclass[border=2pt]{{standalone}} \n\
\\usepackage{{tikz}} \n\
\\usetikzlibrary{{calc,intersections,patterns}} \n\
\\begin{{document}} \n\
\\begin{{tikzpicture}}[scale=0.7] \n\
\\begin{{scriptsize}} \n\
    \\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}}; \n\
    \\coordinate ({b}) at (2,-1) node at ({b}) [below] {{${b}$}}; \n\
    \\coordinate ({c}) at (4,0)  node at ({c}) [below] {{${c}$}}; \n\
    \\coordinate ({a1}) at (0,4)   node at ({a1}) [above] {{${a1}$}}; \n\
    \\coordinate ({b1}) at (2,3)   node at ({b1}) [above] {{${b1}$}}; \n\
    \\coordinate ({c1}) at (4,4)   node at ({c1}) [above] {{${c1}$}};\n\
    \\draw [dashed] ({a})--({c}); \n\
    \\draw ({a})--({b})--({b1})--({a1})--({a}); \n\
    \\draw ({b1})--({c1})--({c})--({b}) ({a1})--({c1})--({c}); \n\
\\end{{scriptsize}} \n\
\\end{{tikzpicture}} \n\
\\end{{document}}"
    return code

def codelatex_hinhchop_tamgiac_canhvg(s,a,b,c):
    code=f"\\begin{{tikzpicture}}[scale=0.7]\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (1,-2) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (4,0)   node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({s}) at (0,4)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({a})--({c}) ; \n\
\\draw ({a})--({b}) ({b})--({c}) ({s})--({a}) ({s})--({b}) ({s})--({c}); \n\
\\end{{tikzpicture}}\n"
    return code
# Vẽ hình phẳng
def code_latex_hinhbinhhanh(a,b,c,d):
    code=rf"""\begin{{tikzpicture}}[line join=round, line cap=round,thick]
\coordinate ({a}) at (1,3);
\coordinate ({b}) at (6,3);
\coordinate ({d}) at (0,0);
\coordinate ({c}) at ($({b})+({d})-({a})$);
\draw({a})--({b})--({c})--({d})--cycle;
\foreach \i/\g in {{{a}/90,{b}/90,{c}/-90,{d}/-90}}{{\draw[fill=white](\i) circle (1.5pt) ($(\i)+(\g:3mm)$) node[scale=1]{{$\i$}};}}
\end{{tikzpicture}}"""
    return code

#Vẽ các hình không gian
def ve_hinhchop_tamgiac_canhvg(s,a,b,c):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7]\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (1,-2) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (4,0)   node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({s}) at (0,4)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({a})--({c}) ; \n\
\\draw ({a})--({b}) ({b})--({c}) ({s})--({a}) ({s})--({b}) ({s})--({c});\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

def ve_hinhchop_tamgiac_matvg(s,a,b,c):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7]\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (2,-2) node at ({b}) [below] {{${b}$}};\n\
\\coordinate ({c}) at (4,0)   node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({s}) at (1,3)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({a})--({c}) ; \n\
\\draw ({a})--({b}) ({b})--({c}) ({s})--({a}) ({s})--({b}) ({s})--({c});\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

def ve_hinhchop_tamgiac_deu(S,A,B,C):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,angles}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,thick]\n\
\\coordinate ({A}) at (-2,0);\n\
\\coordinate ({B}) at (-1,-3);\n\
\\coordinate ({C}) at (3,0);\n\
\\coordinate (M) at ($({A})!0.5!({B})$);\n\
\\coordinate (O) at ($({C})!2/3!(M)$);\n\
\\coordinate ({S}) at ($(O)+(0,5)$);\n\
\\pic[draw,thin,angle radius=1.5mm] {{right angle = {A}--O--{S}}} pic[draw,thin,angle radius=1.5mm] {{right angle = {S}--O--{C}}};\n\
\\draw({S})--({A}) ({S})--({B}) ({S})--({C}) ({B})--({C}) ({A})--({B});\n\
\\draw[dashed,thin]({A})--({C}) ({A})--(O)--({B}) ({C})--(O)--({S});\n\
\\foreach \\i/\\g in {{{S}/90,{A}/180,{B}/-90,{C}/0,O/-70}}{{\\draw[fill=white](\\i) circle (1.5pt) ($(\\i)+(\\g:3mm)$) node[scale=1]{{$\\i$}};}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}"
    return code

def ve_hinh_tu_dien(s,a,b,c):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7]\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (1,-2) node at ({b}) [below] {{${b}$}};\n\
\\coordinate ({c}) at (4,0)   node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({s}) at (1,3)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({a})--({c}) ; \n\
\\draw ({a})--({b}) ({b})--({c}) ({s})--({a}) ({s})--({b}) ({s})--({c});\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

def ve_hinhchop_tugiac(S,A,B,C,D):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,thick]\n\
\\coordinate ({S}) at (1,4);\n\
\\coordinate ({A}) at (0,0);\n\
\\coordinate ({B}) at (1,-2);\n\
\\coordinate ({C}) at (3,-3);\n\
\\coordinate ({D}) at (5,0);\n\
\\draw({S})--({A}) ({S})--({B}) ({S})--({C}) ({S})--({D}) ({B})--({C}) ({A})--({B})--({C})--({D});\n\
\\draw[dashed,thin]({A})--({D});\n\
\\foreach \\i/\\g in {{{S}/90,{A}/180,{B}/-90,{C}/-90,{D}/0}}{{\\draw[fill=white](\\i) circle (1.5pt) ($(\\i)+(\\g:3mm)$) node[scale=1]{{$\\i$}};}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}"
    return code

def ve_hinhchop_hbh_canhvg(S,A,B,C,D):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,angles}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,thick]\n\
\\coordinate ({A}) at (0,0);\n\
\\coordinate ({B}) at (-2,-3);\n\
\\coordinate ({D}) at (7,0);\n\
\\coordinate ({C}) at ($({B})+({D})-({A})$);\n\
\\coordinate ({S}) at ($({A})+(0,5)$);\n\
\\draw({S})--({B}) ({S})--({C}) ({S})--({D}) ({B})--({C})--({D});\n\
\\draw[dashed,thin]({S})--({A}) ({A})--({B}) ({A})--({D});\n\
\\pic[draw,thin,angle radius=3mm] {{right angle = {S}--{A}--{D}}} pic[draw,thin,angle radius=3mm] {{right angle = {S}--{A}--{B}}};\n\
\\foreach \\i/\\g in {{{S}/90,{A}/-90,{B}/-90,{C}/-90,{D}/0}}{{\\draw[fill=white](\\i) circle (1.5pt) ($(\\i)+(\\g:3mm)$) node[scale=1]{{$\\i$}};}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}} "
    return code

def ve_hinhchop_hbh_matvg(S,A,B,C,D):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,angles}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,thick]\n\
\\coordinate ({A}) at (0,0);\n\
\\coordinate ({B}) at (-2,-3);\n\
\\coordinate ({D}) at (7,0);\n\
\\coordinate ({C}) at ($({B})+({D})-({A})$);\n\
\\coordinate ({S}) at (-1,4);\n\
\\draw({S})--({B}) ({S})--({C}) ({S})--({D}) ({B})--({C})--({D});\n\
\\draw[dashed,thin]({S})--({A}) ({A})--({B}) ({A})--({D});\n\
\\foreach \\i/\\g in {{{S}/90,{A}/-90,{B}/-90,{C}/-90,{D}/0}}{{\\draw[fill=white](\\i) circle (1.5pt) ($(\\i)+(\\g:3mm)$) node[scale=1]{{$\\i$}};}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}} "
    return code

def ve_hinhchop_tugiacdeu(S,A,B,C,D):
    code=f" \\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,angles}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[line join=round, line cap=round,thick]\n\
\\coordinate ({A}) at (0,0);\n\
\\coordinate ({B}) at (2,-2);\n\
\\coordinate ({D}) at (5,0);\n\
\\coordinate ({C}) at ($({B})+({D})-({A})$);\n\
\\coordinate (O) at ($({A})!0.5!({C})$);\n\
\\coordinate ({S}) at ($(O)+(0,7)$);\n\
\\draw({S})--({A}) ({S})--({B}) ({S})--({C}) ({A})--({B}) ({B})--({C});\n\
\\draw[dashed,thin]({A})--({C}) ({A})--({D}) ({C})--({D}) ({S})--({D}) ({S})--(O) ({B})--({D});\n\
\\pic[draw,thin,angle radius=2mm] {{right angle = {S}--O--{D}}};\n\
\\pic[draw,thin,angle radius=2mm] {{right angle = {S}--O--{A}}};\n\
\\foreach \\i/\\g in {{{S}/90,{A}/180,{B}/-90,{C}/-90,{D}/0,O/-90}}{{\\draw[fill=white](\\i) circle (1.5pt) ($(\\i)+(\\g:3mm)$) node[scale=1]{{$\\i$}};}}\n\
\\end{{tikzpicture}}\n\
\\end{{document}}" 
    return code

  
def codelatex_hinhchop_hbh_canhvg(s,a,b,c,d):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7]\n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (3,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (4,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({s}) at (0,4)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({s})--({a}); \n\
\\draw ({b})--({c})--({d}) ({s})--({b}) ({s})--({c}) ({s})--({d}); \n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

def code_hinhchop_hbh_canhvg(s,a,b,c,d):
    code=f"\\begin{{tikzpicture}}[scale=0.7]\n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (-1,-1) node at ({b}) [left] {{${b}$}};\n\
\\coordinate ({c}) at (3,-1)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (4,0)   node at ({d}) [right] {{${d}$}};\n\
\\coordinate ({s}) at (0,4)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({b})--({a})--({d}) ({s})--({a}); \n\
\\draw ({b})--({c})--({d}) ({s})--({b}) ({s})--({c}) ({s})--({d}); \n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n"
    return code

def codelatex_hinhchop_hinhthang(s,a,b,c,d):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
\\usepackage{{tikz}}\n\
\\usetikzlibrary{{calc,intersections,patterns}}\n\
\\begin{{document}}\n\
\\begin{{tikzpicture}}[scale=0.7]\n\
\\begin{{scriptsize}}\n\
\\coordinate ({a}) at (0,0)   node at ({a}) [left] {{${a}$}};\n\
\\coordinate ({b}) at (5,0) node at ({b}) [right] {{${b}$}};\n\
\\coordinate ({c}) at (3,-3)  node at ({c}) [right] {{${c}$}};\n\
\\coordinate ({d}) at (1,-3)   node at ({d}) [left] {{${d}$}};\n\
\\coordinate ({s}) at (1.5,2)   node at ({s}) [above] {{${s}$}};\n\
\\draw [dashed] ({b})--({a}); \n\
\\draw ({s})--({a}) ({b})--({c})--({d}) ({s})--({b}) ({s})--({c}) ({s})--({d})--({a}); \n\
\\end{{scriptsize}}\n\
\\end{{tikzpicture}}\n\
 \\end{{document}}"
    return code

#Xử lý mẫu số liệu ghép nhóm
def codelatex_bang_ghep_nhom(ten_nhom,list_khoang_gia_tri,ten_tan_so,list_tan_so):
    code=f"\\documentclass[border=2pt]{{standalone}}\n\
    \\usepackage[utf8]{{inputenc}}\n\
    \\usepackage[vietnamese]{{babel}}\n\
    \\begin{{document}}\n\
    \\begin{{tabular}}{{|c|c|c|c|c|c|c|}}\n\
        \\hline\n\
        {ten_nhom}   & {list_khoang_gia_tri}\\\\  \n\
        \\hline \n\
        {ten_tan_so} & {list_tan_so} \\\\ \n\
        \\hline \n\
    \\end{{tabular}}\n\
    \\end{{document}}"
    return code

def tao_ten_mau_ghep_nhom():
    list_nhom=[["Khoảng tuổi","Số người"], ["Cân nặng","Số người"], ["Điểm số","Số học sinh"], ["Lương","Số nhân viên"]]
    so_nhom=len(list_nhom)
    #Chọn ngẫu nhiên một nhóm
    i =random.randint(0,so_nhom-1)
    nhom=list_nhom[i]
    ten_nhom=nhom[0]
    ten_tan_so=nhom[1]
    #
    if ten_nhom=="Khoảng tuổi":
        u1 = random.randint(15,25)
        d = random.randint(5,10)
        tan_so_min =1
        tan_so_max =35

    if ten_nhom=="Cân nặng":
        u1 = random.randint(35,45)
        d = random.randint(5,10)
        tan_so_min =1
        tan_so_max =35

    if ten_nhom=="Điểm số":
        u1 = random.choice([0,1,2])
        d = random.choice([2, 2.5])
        tan_so_min =1
        tan_so_max =10

    if ten_nhom=="Lương":
        u1 = random.randint(5,10)
        d = random.randint(3,6)
        tan_so_min =1
        tan_so_max =15
    return ten_nhom,ten_tan_so,u1,d,tan_so_min,tan_so_max

def tao_mau_ghep_nhom(so_nhom,so_bat_dau,khoang_cach,tan_so_min,tan_so_max):
    #Tạo số nhóm, số bắt đầu, khoảng cách
    u1= so_bat_dau
    d = khoang_cach

    #Tạo list các cặp giá trị và tần số
    gia_tri= [u1+i*d for i in range(so_nhom)]

    dau_noi=" & "
    khoang_gt=dau_noi.join(f"[{gia_tri[i-1]} ; {gia_tri[i]})" for i in range(1,so_nhom))
    khoang_gt=khoang_gt.replace(".",",")
    khoang_gt=khoang_gt.replace(",0","")

    tan_so=[random.randint(tan_so_min,tan_so_max) for i in range(so_nhom-1)]
    list_tan_so = dau_noi.join(f"{str(tan_so[i])}" for i in range(so_nhom-1))
    return gia_tri,khoang_gt,list_tan_so, tan_so

#Thay thế dấu ngoặc sin(x) thành sinx
def thay_the_ngoac_sincos(st):
    str_thaythe=st.replace(f"\\cos{{\\left(x \\right)}}",f"\\cos x")
    str_thaythe=str_thaythe.replace(f"\\sin 1x",f"\\sin x")
    str_thaythe=str_thaythe.replace(f"\\cos 1x",f"\\cosx")
    str_thaythe=str_thaythe.replace(f"\\sin{{\\left(x \\right)}}",f"\\sin x")
    str_thaythe=str_thaythe.replace(f"\\tan{{\\left(x \\right)}}",f"\\tan x")
    str_thaythe=str_thaythe.replace(f"\\cot{{\\left(x \\right)}}",f"\\cot x")
    str_thaythe=str_thaythe.replace(f"\\log{{\\left(x \\right)}}",f"\\ln |x|")   
    for i in range (1,15):
        str_thaythe=str_thaythe.replace(f"\\sin^{{{i}}}{{\\left(x \\right)}}",f"\\sin ^{{{i}}}x")
        str_thaythe=str_thaythe.replace(f"\\cos^{{{i}}}{{\\left(x \\right)}}",f"\\cos ^{{{i}}}x")
        str_thaythe=str_thaythe.replace(f"\\tan^{{{i}}}{{\\left(x \\right)}}",f"\\tan ^{{{i}}}x")
        str_thaythe=str_thaythe.replace(f"\\cot^{{{i}}}{{\\left(x \\right)}}",f"\\cot ^{{{i}}}x")
        str_thaythe=str_thaythe.replace(f"\\sin{{\\left({i} x \\right)}}",f"\\sin {i}x")
        str_thaythe=str_thaythe.replace(f"\\cos{{\\left({i} x \\right)}}",f"\\cos {i}x")
        str_thaythe=str_thaythe.replace(f"\\tan{{\\left({i} x \\right)}}",f"\\tan {i}x")
        str_thaythe=str_thaythe.replace(f"\\cot{{\\left({i} x \\right)}}",f"\\cot {i}x")
        str_thaythe=str_thaythe.replace(f"\\log{{\\left({i} x \\right)}}",f"\\ln {i}|x|")
        str_thaythe=str_thaythe.replace(f"\\ln x^{{{i}}}",f"\\ln^{{{i}}}|x|")
        for j in range (1,15):
            str_thaythe=str_thaythe.replace(f"\\sin^{{{i}}}{{\\left({j} x \\right)}}",f"\\sin ^{{{i}}}{j}x")
            str_thaythe=str_thaythe.replace(f"\\cos^{{{i}}}{{\\left({j} x \\right)}}",f"\\cos ^{{{i}}}{j}x")
            str_thaythe=str_thaythe.replace(f"\\tan^{{{i}}}{{\\left({j} x \\right)}}",f"\\tan ^{{{i}}}{j}x")
            str_thaythe=str_thaythe.replace(f"\\cot^{{{i}}}{{\\left({j} x \\right)}}",f"\\cot ^{{{i}}}{j}x")

    str_thaythe=str_thaythe.replace("(1x)","x").replace("(-1x)","-x")
    str_thaythe=str_thaythe.replace("-+","-").replace("--","+").replace("+-","-").replace("++","+").replace("+ +","+").replace("+ -","-").replace("- -","+").replace("- +","-")

    return str_thaythe

#Thay thế dấu ngoặc của log
def thay_the_ngoac_log(st):
    str_thaythe=st.replace("\\log","\\ln")
    str_thaythe=str_thaythe.replace(f"{{\\left(",f"|")
    str_thaythe=str_thaythe.replace(f"\\right)}}",f"|")       
    return str_thaythe

def tu_phan_vi(sample_data):
    # Tính tứ phân vị cho số lượng chẵn    
    sorted_data  = np.sort(sample_data)

    #Lấy vị trí chính giữa
    t=len(sample_data)

    if t%2==0:
        t1=int(len(sorted_data)/2)
        sorted_data_right=sorted_data[t1+1:]
    else:
        t1=int((len(sorted_data)-1)/2)
        sorted_data_right=sorted_data[t1+1:]
    sorted_data_left=sorted_data[:t1]

    # Tìm tứ phân vị thứ nhất, thứ hai và thứ ba
    Q1 = np.percentile(sorted_data_left, 50)
    Q2 = np.percentile(sorted_data, 50)
    Q3 = np.percentile(sorted_data_right, 50)
    return Q1,Q2,Q3

#Nhóm code liên quan hình học
def trung_diem(x_1,y_1,x_2,y_2):
    x_m=(x_1+x_2)/2
    y_m=(y_1+y_2)/2
    return x_m,y_m

#print(get_image_dimensions("H:\\One Drive\\OneDrive\\NGAN HANG DE\\NGAN HANG SUU TAM\\PYTHON\\HINH VE\\2024-07-25_20-35-27.png"))

