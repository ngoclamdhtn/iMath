from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt, QTimer, QEventLoop, QUrl
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTreeWidget, QTreeWidgetItem, QPushButton, QMenu, QVBoxLayout, QWidget, QMessageBox, QInputDialog,
        QFrame, QDialog, QFileDialog, QProgressBar, QToolTip, QScrollBar, QDesktopWidget,  QCheckBox, QTableWidget, QTableWidgetItem)
from PyQt5.QtGui import QPainter, QPalette, QColor, QPixmap, QImage, QIcon, QDesktopServices
import math
import random
from sympy import *
from tabulate import tabulate
import sympy as sp
import my_module, license
import D10_C1,D10_C2,D10_C3,D10_C4,D10_C5,D10_C6,D10_C7,D10_C10,D10_C8,D10_C9, D11_C1, D11_C2, D11_C3, D11_C4, D11_C5, D11_C6, D11_C7, D11_C8,D11_C9, D12_C1,D12_C2, D12_C4, D12_C3, D12_C5, D12_C7
import pyperclip
import os, shutil,re, sys, subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import requests, webbrowser
from bs4 import BeautifulSoup
from openpyxl import Workbook, load_workbook
import hashlib, winreg, psutil, subprocess, base64, glob
from cryptography.fernet import Fernet
from datetime import datetime

class ShowMessageBox(QMessageBox):
    def __init__(self,icon, title, text):
        super().__init__()
        self.setIcon(QMessageBox.Information)        
        self.setText(text)
        self.setWindowTitle(title)
        folder_path = my_module.get_folder_icon()    
        folder_path= f"{folder_path}/icon.png"  
        self.setWindowIcon(QIcon(folder_path))

class Form_tieude(QWidget):
    def __init__(self, parent):
        super().__init__()
        self.parent = parent
        self.initUI()

    def initUI(self):
        font = QtGui.QFont()
        font.setPointSize(11)
        font.setFamily("Arial")
        layout = QVBoxLayout()
        self.setLayout(layout)
        self.setWindowTitle('Thiết lập tiêu đề chung')

        self.textEdit =QtWidgets.QTextEdit()
        layout.addWidget(self.textEdit)        

        self.btn_ok = QtWidgets.QPushButton()
        #self.btn_ok.setGeometry(QtCore.QRect(520, 370, 150, 30))
        self.btn_ok.setFont(font)
        self.btn_ok.setObjectName("btn_ok")
        self.btn_ok.setText("Apply")
        self.btn_ok.clicked.connect(self.closeEvent)
        layout.addWidget(self.btn_ok)

    def closeEvent(self, event):
        self.parent.update_label(self.textEdit.toPlainText())
        self.close()

class Ui_MainWindow(object):
        def setupUi(self, MainWindow):
       
                MainWindow.setObjectName("MainWindow")
                MainWindow.resize(1920, 1440) 
                MainWindow.setWindowTitle("iMath")
                MainWindow.showMaximized()
                MainWindow.setStyleSheet("QMainWindow:title { color: red; background-color: #2B579A; }")
                folder_path = my_module.get_folder_icon()    
                folder_path= f"{folder_path}/icon.png" 
                MainWindow.setWindowIcon(QIcon(folder_path))               
        

                self.centralwidget = QtWidgets.QWidget(parent=MainWindow)
                self.centralwidget.setObjectName("centralwidget")
                self.tab_main = QtWidgets.QTabWidget(parent=self.centralwidget)
                #self.tab_main.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                self.tab_main.setGeometry(QtCore.QRect(0, 0, 1920, 1440))
                self.tab_main.setAccessibleName("")
                self.tab_main.setObjectName("tab_main")

                MainWindow.setCentralWidget(self.centralwidget)      
                self.tab_main.setCurrentIndex(1)
                QtCore.QMetaObject.connectSlotsByName(MainWindow)

                
                self.tab_taode = QWidget()
                self.tab_thongtin_dethi = QWidget()
                self.tab_dothi = QWidget()
                self.tab_bbt = QWidget()
                self.tab_vehinh = QWidget()
                # self.tab_run_latex = QWidget()
                self.tab_bang_so_lieu = QWidget()
                self.tab_ban_quyen = QWidget()
                self.tab_huongdan = QWidget()

                self.tab_main.addTab(self.tab_taode, "Thiết lập ma trận")
                self.tab_taode.setObjectName("tab_taode")
                #self.tab_taode.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                self.tab_taode.setGeometry(QtCore.QRect(0, 0, 1920, 1440))

                self.tab_main.addTab(self.tab_thongtin_dethi, "Tạo đề")
                self.tab_thongtin_dethi.setObjectName("tab_thongtin_dethi")
                #self.tab_thongtin_dethi.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                self.tab_taode.setGeometry(QtCore.QRect(0, 0, 1920, 1440))

                self.tab_main.addTab(self.tab_dothi, "Đồ thị")
                self.tab_dothi.setObjectName("tab_dothi")
                #self.tab_dothi.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                self.tab_dothi.setGeometry(QtCore.QRect(0, 0, 1920, 1440))

                self.tab_main.addTab(self.tab_bbt, "Bảng biến thiên")
                self.tab_bbt.setObjectName("tab_bbt")
                #self.tab_bbt.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))  
                self.tab_bbt.setGeometry(QtCore.QRect(0, 0, 1920, 1440))

                self.tab_main.addTab(self.tab_vehinh, "Vẽ hình")
                self.tab_vehinh.setObjectName("tab_vehinh")                
                self.tab_vehinh.setGeometry(QtCore.QRect(0, 0, 1920, 1440)) 

                self.tab_main.addTab(self.tab_bang_so_lieu, "Bảng số liệu")
                self.tab_bang_so_lieu.setObjectName("tab_bang_so_lieu")
                #self.tab_bang_so_lieu.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height())) 
                self.tab_bang_so_lieu.setGeometry(QtCore.QRect(0, 0, 1920, 1440)) 

                # self.tab_main.addTab(self.tab_run_latex, "Biên dịch Latex")
                # self.tab_run_latex.setObjectName("tab_run_latex")
                # #self.tab_run_latex.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                # self.tab_bang_so_lieu.setGeometry(QtCore.QRect(0, 0, 1920, 1440))        
           

                self.tab_main.addTab(self.tab_ban_quyen, "Bản quyền-Cập nhật")
                self.tab_ban_quyen.setObjectName("tab_ban_quyen")
                #self.tab_ban_quyen.setGeometry(QtCore.QRect(0, 0, screen_rect.width(), screen_rect.height()))
                self.tab_bang_so_lieu.setGeometry(QtCore.QRect(0, 0, 1920, 1440))


                self.tab_main.addTab(self.tab_huongdan, "Hướng dẫn")
                self.tab_huongdan.setObjectName("tab_huongdan")                
                self.tab_huongdan.setGeometry(QtCore.QRect(0, 0, 1920, 1440))

          

        #Thiết lập font cho MainWindow
                font_8 = QtGui.QFont()
                font_8.setPointSize(8)
                font_8.setFamily("Arial")

                font_9 = QtGui.QFont()
                font_9.setPointSize(9)
                font_9.setFamily("Arial")

                font_10 = QtGui.QFont()
                font_10.setPointSize(10)
                font_10.setFamily("Arial")

                font = QtGui.QFont()
                font.setPointSize(11)
                font.setFamily("Arial")

                font_12 = QtGui.QFont()
                font_12.setPointSize(12)
                font_12.setFamily("Arial")

                font_tieude=QtGui.QFont()
                font_tieude.setPointSize(11)
                font_tieude.setFamily("Arial")
                font_tieude.setBold(True)
        
        
        #Tab thông tin đề thi
                le_trai=20
                le_top=20               

        #Tab Tác giả - ver up
                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)               
                self.label.setGeometry(QtCore.QRect(600, 100, 250, 20))                
                self.label.setFont(font_12)        
                self.label.setText(f"iMath\u00A92023 ver 17.11.2024")

                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label.setGeometry(QtCore.QRect(400, 120, 800, 20))
                self.label.setFont(font)        
                self.label.setText("Tác giả: Trần Ngọc Lam - GV Toán, trường THPT Trần Đại Nghĩa, tỉnh Đắk Lắk.")

                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label.setGeometry(QtCore.QRect(400, 140, 700, 20))
                self.label.setFont(font)        
                self.label.setText("Email: ngoclamdhtn@gmail.com")

                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label.setGeometry(QtCore.QRect(400, 160, 700, 20))
                self.label.setFont(font)        
                self.label.setText("Zalo: 0974.940.049")

                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label.setGeometry(QtCore.QRect(350, 200, 700, 40))
                self.label.setFont(font)        
                self.label.setText("Mã đăng kí:")

                text_ma_may=license.encrypt_serial_number()
                self.tab_ban_quyen_ma_may= QtWidgets.QTextEdit(parent=self.tab_ban_quyen)
                self.tab_ban_quyen_ma_may.setGeometry(QtCore.QRect(500, 200, 600, 40))
                self.tab_ban_quyen_ma_may.setObjectName("tab_ban_quyen_ma_may")
                self.tab_ban_quyen_ma_may.setFont(font_12)
                self.tab_ban_quyen_ma_may.setText(text_ma_may)
                if license.kiemtra_banquyen():
                        self.tab_ban_quyen_ma_may.setText("")

                
                self.label= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label.setGeometry(QtCore.QRect(350, 260, 700, 20))
                self.label.setFont(font)        
                self.label.setText("Mã bản quyền:")


                self.tab_ban_quyen_ma_dangki= QtWidgets.QTextEdit(parent=self.tab_ban_quyen)
                self.tab_ban_quyen_ma_dangki.setGeometry(QtCore.QRect(500, 260, 600, 60))
                self.tab_ban_quyen_ma_dangki.setObjectName("tab_ban_quyen_ma_dangki")
                self.tab_ban_quyen_ma_dangki.setFont(font_12)
                self.tab_ban_quyen_ma_dangki.setText("")               

                

                #Thời gian còn lại
                thoi_gian = f"Thời gian sử dụng còn lại là {license.so_ngay_banquyen()} ngày."
                self.label_thoigian= QtWidgets.QLabel(parent=self.tab_ban_quyen)
                self.label_thoigian.setGeometry(QtCore.QRect(600, 330, 600, 20))
                self.label_thoigian.setFont(font)        
                self.label_thoigian.setText(thoi_gian)

                self.tab_ban_quyen_dangki = QtWidgets.QPushButton(parent=self.tab_ban_quyen)
                self.tab_ban_quyen_dangki.setGeometry(QtCore.QRect(520, 370, 150, 30))
                self.tab_ban_quyen_dangki.setFont(font)
                self.tab_ban_quyen_dangki.setObjectName("tab_ban_quyen_copy_ma_may")
                self.tab_ban_quyen_dangki.setText("Đăng ký bản quyền")
                self.tab_ban_quyen_dangki.clicked.connect(self.dang_ki_ban_quyen)

                self.tab_ban_quyen_copy_ma_may = QtWidgets.QPushButton(parent=self.tab_ban_quyen)
                self.tab_ban_quyen_copy_ma_may.setGeometry(QtCore.QRect(800, 370, 150, 30))
                self.tab_ban_quyen_copy_ma_may.setFont(font)
                self.tab_ban_quyen_copy_ma_may.setObjectName("tab_ban_quyen_copy_ma_may")
                self.tab_ban_quyen_copy_ma_may.setText("Copy mã máy")
                self.tab_ban_quyen_copy_ma_may.clicked.connect(self.copy_ma_may)

                #Button check update
                self.btn_check_update = QtWidgets.QPushButton(parent=self.tab_ban_quyen)
                self.btn_check_update.setGeometry(QtCore.QRect(520, 430, 150, 30))
                self.btn_check_update.setFont(font)
                self.btn_check_update.setObjectName("btn_check_update")
                self.btn_check_update.setText("Check Update")
                self.btn_check_update.clicked.connect(self.btn_check_update_click)

                #Button Lịch sử cập nhậtp
                self.btn_version_update = QtWidgets.QPushButton(parent=self.tab_ban_quyen)
                self.btn_version_update.setGeometry(QtCore.QRect(800, 430, 150, 30))
                self.btn_version_update.setFont(font)
                self.btn_version_update.setObjectName("btn_version_update")
                self.btn_version_update.setText("Lịch sử cập nhật")
                self.btn_version_update.clicked.connect(self.btn_version_update_click) 
              

        #Giao diện tab Bảng số liệu:
                texbox_left = 160
                textbox_top =20
                d=30
                #Hàng 1
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top, 200, 20))
                self.label.setFont(font)        
                self.label.setText("Tên nhóm giá trị")

                self.tab_bang_so_lieu_tennhom= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_tennhom.setGeometry(QtCore.QRect(texbox_left, textbox_top, 100, 20))
                self.tab_bang_so_lieu_tennhom.setObjectName("tab_bang_so_lieu_tennhom")
                self.tab_bang_so_lieu_tennhom.setFont(font_12)
                self.tab_bang_so_lieu_tennhom.setText("Cân nặng")

                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(290, textbox_top, 100, 20))
                self.label.setFont(font)        
                self.label.setText("Số nhóm")

                self.tab_bang_so_lieu_sonhom= QtWidgets.QSpinBox(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_sonhom.setGeometry(QtCore.QRect(360, textbox_top, 40, 20))
                self.tab_bang_so_lieu_sonhom.setFont(font_12)
                self.tab_bang_so_lieu_sonhom.setObjectName("tab_bang_so_lieu_sonhom")
                self.tab_bang_so_lieu_sonhom.setValue(5)

                #Hàng 2
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+d, 200, 20))
                self.label.setFont(font)        
                self.label.setText("Tên tần số")

                self.tab_bang_so_lieu_tentanso= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_tentanso.setGeometry(QtCore.QRect(texbox_left, textbox_top+d, 100, 20))
                self.tab_bang_so_lieu_tentanso.setObjectName("tab_bang_so_lieu_tentanso")
                self.tab_bang_so_lieu_tentanso.setFont(font_12)
                self.tab_bang_so_lieu_tentanso.setText("Số người")

                #Hàng 3
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+2*d, 200, 20))
                self.label.setFont(font)        
                self.label.setText("Giá trị bắt đầu")

                self.tab_bang_so_lieu_giatribatdau= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_giatribatdau.setGeometry(QtCore.QRect(texbox_left, textbox_top+2*d, 50, 20))
                self.tab_bang_so_lieu_giatribatdau.setFont(font_12)
                self.tab_bang_so_lieu_giatribatdau.setObjectName("tab_bang_so_lieu_giatribatdau")
                self.tab_bang_so_lieu_giatribatdau.setText("45")

                #Hàng 4
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+3*d, 300, 20))
                self.label.setFont(font)        
                self.label.setText("Khoảng cách 2 giá trị")

                self.tab_bang_so_lieu_khoangcach= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_khoangcach.setGeometry(QtCore.QRect(texbox_left, textbox_top+3*d, 50, 20))
                self.tab_bang_so_lieu_khoangcach.setFont(font_12)
                self.tab_bang_so_lieu_khoangcach.setObjectName("tab_bang_so_lieu_khoangcach")
                self.tab_bang_so_lieu_khoangcach.setText("5")

                #Hàng 5
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+4*d, 300, 20))
                self.label.setFont(font)        
                self.label.setText("Tần số (cách nhau bằng dấu \",\")")

                self.tab_bang_so_lieu_tanso= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_tanso.setGeometry(QtCore.QRect(texbox_left+80, textbox_top+4*d, 200, 20))
                self.tab_bang_so_lieu_tanso.setObjectName("tab_bang_so_lieu_tanso")
                self.tab_bang_so_lieu_tanso.setFont(font_12)
                self.tab_bang_so_lieu_tanso.setText("15,12,5,7,16")

                #Hàng 6
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+5*d, 300, 20))
                self.label.setFont(font)        
                self.label.setText("Tần số ngẫu nhiên min")

                self.tab_bang_so_lieu_tansomin= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_tansomin.setGeometry(QtCore.QRect(texbox_left+20, textbox_top+5*d, 50, 20))
                self.tab_bang_so_lieu_tansomin.setObjectName("tab_bang_so_lieu_tansomin")
                self.tab_bang_so_lieu_tansomin.setFont(font_12)
                self.tab_bang_so_lieu_tansomin.setText("10")

                #Hàng 7
                self.label= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label.setGeometry(QtCore.QRect(10, textbox_top+6*d, 300, 20))
                self.label.setFont(font)        
                self.label.setText("Tần số ngẫun nhiên max")

                self.tab_bang_so_lieu_tansomax= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_tansomax.setGeometry(QtCore.QRect(texbox_left+20, textbox_top+6*d, 50, 20))
                self.tab_bang_so_lieu_tansomax.setFont(font_12)
                self.tab_bang_so_lieu_tansomax.setObjectName("tab_bang_so_lieu_tansomax")
                self.tab_bang_so_lieu_tansomax.setText("20")

                frame = QtWidgets.QFrame(parent=self.tab_bang_so_lieu)
                # Thiết lập màu nền của QFrame
                frame.setStyleSheet("background-color: #fdfdfd;")
                # Đặt kích thước và vị trí của QFrame
                frame.setGeometry(500, 15, 600, 100)

                self.label_tab_bang_so_lieu_getbang= QtWidgets.QLabel(parent=self.tab_bang_so_lieu)
                self.label_tab_bang_so_lieu_getbang.setGeometry(QtCore.QRect(500, 15, 650, 100))
                self.label_tab_bang_so_lieu_getbang.setFont(font)
                self.label_tab_bang_so_lieu_getbang.setObjectName("label_tab_bang_so_lieu_getbang")

                self.btn_tao_bang_so_lieu = QtWidgets.QPushButton(parent=self.tab_bang_so_lieu)
                self.btn_tao_bang_so_lieu.setGeometry(QtCore.QRect(10, textbox_top+8*d, 150, 30))
                self.btn_tao_bang_so_lieu.setFont(font)
                self.btn_tao_bang_so_lieu.setObjectName("btn_tao_bang_so_lieu")
                self.btn_tao_bang_so_lieu.setText("Tạo bảng ghép nhóm")
                self.btn_tao_bang_so_lieu.clicked.connect(self.tao_bang_so_lieu)

                self.btn_tao_bang_so_lieu_runcode = QtWidgets.QPushButton(parent=self.tab_bang_so_lieu)
                self.btn_tao_bang_so_lieu_runcode.setGeometry(QtCore.QRect(texbox_left+50, textbox_top+8*d, 150, 30))
                self.btn_tao_bang_so_lieu_runcode.setFont(font)
                self.btn_tao_bang_so_lieu_runcode.setObjectName("btn_tao_bang_so_lieu_runcode")
                self.btn_tao_bang_so_lieu_runcode.setText("Biên dịch code")
                self.btn_tao_bang_so_lieu_runcode.clicked.connect(self.tao_bang_so_lieu_runcode)

                self.btn_copy_bang_so_lieu = QtWidgets.QPushButton(parent=self.tab_bang_so_lieu)
                self.btn_copy_bang_so_lieu.setGeometry(QtCore.QRect(10, textbox_top+9*d+5, 150, 30))
                self.btn_copy_bang_so_lieu.setFont(font)
                self.btn_copy_bang_so_lieu.setObjectName("btn_copy_bang_so_lieu")
                self.btn_copy_bang_so_lieu.setText("Copy bảng số liệu")
                self.btn_copy_bang_so_lieu.clicked.connect(self.copy_bang_so_lieu)

                self.tab_bang_so_lieu_code= QtWidgets.QTextEdit(parent=self.tab_bang_so_lieu)
                self.tab_bang_so_lieu_code.setGeometry(QtCore.QRect(500, 150, 600, 200))
                self.tab_bang_so_lieu_code.setObjectName("tab_bang_so_lieu_code")  

         #Tab Tạo đề  
                le_trai=720
                le_top=50    
                          
                #Tạo nút Tiêu đề nhóm
                # self.combo_title_nhom = QtWidgets.QComboBox(parent=self.tab_taode)
                # self.combo_title_nhom.setGeometry(QtCore.QRect(le_trai, le_top, 150, 30))
                # self.combo_title_nhom.setFont(font_10)
                # self.combo_title_nhom.setObjectName("combo_title_nhom")
                # self.combo_title_nhom.addItem("Tiêu đề nhóm")
                # self.combo_title_nhom.addItem("Nhóm trắc nghiệm")
                # self.combo_title_nhom.addItem("Nhóm đúng sai")
                # self.combo_title_nhom.addItem("Nhóm trả lời ngắn")
                # self.combo_title_nhom.currentIndexChanged.connect(self.tao_title_nhom)                  

                #Table dạng toán
                self.tableWidget =  QtWidgets.QTableWidget(parent=self.tab_taode)
                self.tableWidget.setGeometry(QtCore.QRect(le_trai, le_top, 550, 375))
                self.tableWidget.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOn)                
                self.tableWidget.setRowCount(0)
                self.tableWidget.setColumnCount(4)
                self.tableWidget.setColumnWidth(0, 360)
                self.tableWidget.setColumnWidth(1, 50)
                self.tableWidget.setColumnWidth(2, 50)
                self.tableWidget.setColumnWidth(3, 50)
                self.tableWidget.itemChanged.connect(self.thongke)

                # Set column headers
                headers = ["Dạng toán", "Loại \n câu", "Mức \n độ", "Số \n câu lấy"]
                for i, header in enumerate(headers):                    
                    self.tableWidget.setHorizontalHeaderItem(i, QTableWidgetItem(header))                  
                

                current_directory = os.path.dirname(os.path.abspath(__file__))
                icon_folder_path = os.path.join(current_directory, 'ICON')

                #Tạo nút Move Up
                self.btn_move_row_up = QtWidgets.QPushButton(parent=self.tab_taode)
                self.btn_move_row_up.setGeometry(QtCore.QRect(le_trai+575, le_top, 30, 30))
                self.btn_move_row_up.setFont(font_10)
                self.btn_move_row_up.setObjectName("btn_move_up")
                self.btn_move_row_up.setText("")
                self.btn_move_row_up.clicked.connect(self.move_row_up)
                self.btn_move_row_up.setStyleSheet("color: white;background-color: #4D82B8;")
                self.btn_move_row_up.setIcon(QIcon(f"{icon_folder_path}\\up-arrow.png"))

                #Tạo nút Move Down
                self.btn_move_row_down = QtWidgets.QPushButton(parent=self.tab_taode)
                self.btn_move_row_down.setGeometry(QtCore.QRect(le_trai+575, le_top+30, 30, 30))
                self.btn_move_row_down.setFont(font_10)
                self.btn_move_row_down.setObjectName("btn_move_down")
                self.btn_move_row_down.setText("")
                self.btn_move_row_down.clicked.connect(self.move_row_down)
                self.btn_move_row_down.setStyleSheet("color: white;background-color: #4D82B8;")
                self.btn_move_row_down.setIcon(QIcon(f"{icon_folder_path}\\down-arrow.png"))

                #Tạo nút Chuyển tự luận
                self.btn_chuyen_tuluan = QtWidgets.QPushButton(parent=self.tab_taode)
                self.btn_chuyen_tuluan.setGeometry(QtCore.QRect(le_trai+555, le_top+70, 100, 30))
                self.btn_chuyen_tuluan.setFont(font_10)
                self.btn_chuyen_tuluan.setObjectName("btn_chuyen_tuluan")
                self.btn_chuyen_tuluan.setText("TN => TL")
                self.btn_chuyen_tuluan.clicked.connect(self.chuyen_tuluan)
                self.btn_chuyen_tuluan.setStyleSheet("color: white;background-color: #4D82B8;")

                #Tạo nút Xóa dòng
                self.btn_xoa_dong = QtWidgets.QPushButton(parent=self.tab_taode)
                self.btn_xoa_dong.setGeometry(QtCore.QRect(le_trai+555, le_top+110, 100, 30))
                self.btn_xoa_dong.setFont(font_10)
                self.btn_xoa_dong.setObjectName("btn_xoa_dong")
                self.btn_xoa_dong.setText("Xóa dòng")
                self.btn_xoa_dong.clicked.connect(self.btn_xoa_dong_click)
                self.btn_xoa_dong.setStyleSheet("color: white;background-color: #4D82B8;")                

                #Nút xóa ma trận
                self.btn_xoa_matran = QtWidgets.QPushButton(parent=self.tab_taode)        
                self.btn_xoa_matran.setGeometry(QtCore.QRect(le_trai+555, le_top+150, 100, 30))
                self.btn_xoa_matran.setFont(font_10)
                self.btn_xoa_matran.setObjectName("btn_xoa_matran")
                self.btn_xoa_matran.setText("Xóa toàn bộ")
                self.btn_xoa_matran.clicked.connect(self.clear_dangtoan)
                self.btn_xoa_matran.setStyleSheet("color: white;background-color: #4385F6;")

                #Nút mở ma trận
                self.btn_load_matran = QtWidgets.QPushButton(parent=self.tab_taode)        
                self.btn_load_matran.setGeometry(QtCore.QRect(le_trai+555, le_top+300, 100, 30))
                self.btn_load_matran.setFont(font_10)
                self.btn_load_matran.setObjectName("btn_load_matran")
                self.btn_load_matran.setText("Mở ma trận")
                self.btn_load_matran.clicked.connect(self.load_matran)
                self.btn_load_matran.setStyleSheet("color: white;background-color: #4385F6;")

                #Nút lưu ma trận
                self.btn_luu_matran = QtWidgets.QPushButton(parent=self.tab_taode)        
                self.btn_luu_matran.setGeometry(QtCore.QRect(le_trai+555, le_top+340, 100, 30))
                self.btn_luu_matran.setFont(font_10)
                self.btn_luu_matran.setObjectName("btn_luu_matran")
                self.btn_luu_matran.setText("Lưu ma trận")
                self.btn_luu_matran.clicked.connect(self.luu_matran)
                self.btn_luu_matran.setStyleSheet("color: white;background-color: #4385F6;")

                

                

                le_top=420
                
                #Table thống kê
                self.table_thongke =  QtWidgets.QTableWidget(parent=self.tab_taode)
                self.table_thongke.setGeometry(QtCore.QRect(le_trai, le_top, 640, 230))
                self.table_thongke.setRowCount(6)
                self.table_thongke.setColumnCount(12)
                self.table_thongke.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                self.table_thongke.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
                for i in range(12):
                    self.table_thongke.setColumnWidth(i, 54)
                
                self.table_thongke.horizontalHeader().hide()
                self.table_thongke.verticalHeader().hide()                
                labels = ["Nhận biết", "Thông hiểu", "Vận dụng thấp", "Vận dụng cao"]
                col = 0
                for label in labels:
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    if label=="Nhận biết":
                        item.setBackground(QColor(173, 216, 230))
                    if label=="Thông hiểu":
                        item.setBackground(QColor(255, 255, 153))
                    if label=="Vận dụng thấp":
                        item.setBackground(QColor(152, 251, 152))
                    if label=="Vận dụng cao":
                        item.setBackground(QColor(255, 192, 203))                 
                    self.table_thongke.setItem(0, col, item)
                    self.table_thongke.setSpan(0, col, 1, 3)

                    item = QTableWidgetItem(f"{0} câu")
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(3, col, item)
                    self.table_thongke.setSpan(3, col, 1, 3)
                    col += 3

                for col in [0,3,6,9]:
                    label="TN"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(1, col, item)

                    label="0"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(2, col, item)
                for col in [1,4,7,10]:
                    label="Đ-S"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(1, col, item)

                    label="0"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(2, col, item)
                for col in [2,5,8,11]:
                    label="TL"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(1, col, item)

                    label="0"
                    item = QTableWidgetItem(label)
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_thongke.setItem(2, col, item)

                item = QTableWidgetItem(f"Trắc nghiệm: {0} câu")
                item.setTextAlignment(Qt.AlignCenter)
                self.table_thongke.setItem(4, 0, item)
                self.table_thongke.setSpan(4, 0, 1, 4)

                item = QTableWidgetItem(f"Đúng-Sai: {0} câu")
                item.setTextAlignment(Qt.AlignCenter)
                self.table_thongke.setItem(4, 4, item)
                self.table_thongke.setSpan(4, 4, 1, 4)

                item = QTableWidgetItem(f"Trả lời ngắn: {0} câu")
                item.setTextAlignment(Qt.AlignCenter)
                self.table_thongke.setItem(4, 8, item)
                self.table_thongke.setSpan(4, 8, 1, 4)

                item = QTableWidgetItem(f"Tổng cộng: {0} câu")
                item.setTextAlignment(Qt.AlignCenter)
                item.setBackground(QtGui.QColor(251, 243, 221))
                self.table_thongke.setItem(5, 0, item)
                self.table_thongke.setSpan(5, 0, 1, 12)


                top_btn = 100
                le_trai=20
                le_top=50         
                
                #label dạng câu hỏi cần tạo đề
                self.label_dangcauhoi = QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label_dangcauhoi.setGeometry(QtCore.QRect(le_trai, le_top-30, 200, 20))
                self.label_dangcauhoi.setFont(font_tieude)
                self.label_dangcauhoi.setStyleSheet("color: #697DBA;")
                self.label_dangcauhoi.setObjectName("label_socau")
                self.label_dangcauhoi.setText("1. Thông tin đề thi")

                #Tên sở
                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)               
                self.label.setGeometry(QtCore.QRect(le_trai, le_top, 250, 30))
                self.label.setFont(font_10)        
                self.label.setText(f"Tên Sở GD: ")

                self.ten_sogd= QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.ten_sogd.setGeometry(QtCore.QRect(le_trai+200, le_top, 300, 30))
                self.ten_sogd.setObjectName("ten_sogd")
                self.ten_sogd.setFont(font_10)                
                self.ten_sogd.setText(license.read_reg_thongtin("So_gd"))
                
                #Tên trường THPT
                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)               
                self.label.setGeometry(QtCore.QRect(le_trai, le_top+40, 250, 30))
                self.label.setFont(font_10)      
                self.label.setText(f"Tên trường THPT: ")

                self.ten_truong= QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.ten_truong.setGeometry(QtCore.QRect(le_trai+200, le_top+40, 300, 30))
                self.ten_truong.setObjectName("ten_truong")
                self.ten_truong.setFont(font_10)
                self.ten_truong.setText(license.read_reg_thongtin("Truong_THPT"))               
     

                #Tên đề thi
                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)               
                self.label.setGeometry(QtCore.QRect(le_trai+600, le_top, 250, 30))
                self.label.setFont(font_10)        
                self.label.setText(f"Tên kỳ thi")


                self.ten_kythi= QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.ten_kythi.setGeometry(QtCore.QRect(le_trai+700, le_top, 300, 30))
                self.ten_kythi.setObjectName("ten_kythi")
                self.ten_kythi.setFont(font_10)
                self.ten_kythi.setText("ĐỀ ÔN TẬP")
                

                #Tên môn thi
                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)               
                self.label.setGeometry(QtCore.QRect(le_trai+600, le_top+40, 250, 30))
                self.label.setFont(font_10)        
                self.label.setText(f"Môn thi")

                self.ten_monthi= QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.ten_monthi.setGeometry(QtCore.QRect(le_trai+700, le_top+40, 300, 30))
                self.ten_monthi.setObjectName("ten_monthi")
                self.ten_monthi.setFont(font_10)
                self.ten_monthi.setText(license.read_reg_thongtin("ten_monthi"))
                

                #Thời gian
                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)               
                self.label.setGeometry(QtCore.QRect(le_trai, le_top+80, 250, 30))
                self.label.setFont(font_10)        
                self.label.setText(f"Thời gian")

                self.ten_thoigian= QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.ten_thoigian.setGeometry(QtCore.QRect(le_trai+200, le_top+80, 60, 30))
                self.ten_thoigian.setObjectName("ten_thoigian")
                self.ten_thoigian.setFont(font_10)

                

                le_top=270

                #label dạng câu hỏi cần tạo đề
                self.label_noidung_cauhoi = QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label_noidung_cauhoi.setGeometry(QtCore.QRect(le_trai, le_top-100, 200, 20))
                self.label_noidung_cauhoi.setFont(font_tieude)
                self.label_noidung_cauhoi.setStyleSheet("color: #697DBA;")
                self.label_noidung_cauhoi.setObjectName("label_socau")
                self.label_noidung_cauhoi.setText("2. Tùy chọn Đề bài")

                #Label số đề 
                self.label_soluongde = QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label_soluongde.setGeometry(QtCore.QRect(le_trai+20, le_top-75, 180, 20))
                self.label_soluongde.setFont(font_10)
                self.label_soluongde.setObjectName("label_soluongde")
                self.label_soluongde.setText("Số lượng đề")
                       
                #Spin số đề
                self.spin_soluong_de = QtWidgets.QSpinBox(parent=self.tab_thongtin_dethi)
                self.spin_soluong_de.setGeometry(QtCore.QRect(le_trai+125, le_top-75, 45, 20))
                self.spin_soluong_de.setFont(font_10)
                self.spin_soluong_de.setObjectName("spin_soluong_de")
                self.spin_soluong_de.setValue(1)

                #Tạo nút Tiêu đề chung
                self.btn_title_chung = QtWidgets.QPushButton(parent=self.tab_taode)
                self.btn_title_chung.setGeometry(QtCore.QRect(le_trai+380, le_top-70, 110, 20))
                self.btn_title_chung.setFont(font_10)
                self.btn_title_chung.setObjectName("btn_title_chung")
                self.btn_title_chung.setText("Thông tin đề thi")
                self.btn_title_chung.clicked.connect(self.thietlap_tieude)
                self.btn_title_chung.setStyleSheet("color: white;background-color: #4D82B8;")
                # Lấy kích thước gợi ý dựa trên nội dung của nút
                # size =  self.btn_title_chung.sizeHint()                
                # self.btn_title_chung.setFixedSize(size)              

                

                #Label chứa mã đề tự nhập
                self.label_nhapmade= QtWidgets.QTextBrowser(parent=self.tab_thongtin_dethi)
                self.label_nhapmade.setGeometry(QtCore.QRect(le_trai+400, le_top-310, 400, 20))
                self.label_nhapmade.setFont(font_10)
                self.label_nhapmade.setObjectName("label_nhapmade")
                self.label_nhapmade.setText("")
                self.label_nhapmade.setVisible(False)

                #Nút tạo đề word
                self.combo_taode = QtWidgets.QComboBox(parent=self.tab_thongtin_dethi)
                self.combo_taode.setGeometry(QtCore.QRect(le_trai+20, le_top-50, 160, 30))
                self.combo_taode.setFont(font_10)
                self.combo_taode.setObjectName("combo_taode")
                self.combo_taode.addItem("Tạo đề Word - Equation")
                self.combo_taode.addItem("Tạo đề Word - MathType")
                self.combo_taode.addItem("Tạo đề Latex - PDF")
                self.combo_taode.addItem("Tạo code Latex")

                #Nút tạo mã đề ngẫu nhiên
                self.checkbox_made_random = QtWidgets.QCheckBox(parent=self.tab_thongtin_dethi)        
                self.checkbox_made_random.setGeometry(QtCore.QRect(le_trai+200, le_top-50, 170, 30))
                self.checkbox_made_random.setFont(font_10)
                self.checkbox_made_random.setObjectName("checkbox_made_random")
                self.checkbox_made_random.setText("Mã đề ngẫu nhiên") 

                #Tạo nút nhập mã đề
                self.btn_nhapmade = QtWidgets.QPushButton(parent=self.tab_thongtin_dethi)
                self.btn_nhapmade.setGeometry(QtCore.QRect(le_trai+200, le_top-20, 120, 30))
                self.btn_nhapmade.setFont(font_10)
                self.btn_nhapmade.setObjectName("btn_nhapmade")
                self.btn_nhapmade.setText("Tự nhập mã đề")
                self.btn_nhapmade.clicked.connect(self.nhapmade)
                self.btn_nhapmade.setStyleSheet("color: white;background-color: #4D82B8;")


                #Nút tự trộn dạng toán
                self.checkbox_shuffle_dangtoan = QtWidgets.QCheckBox(parent=self.tab_thongtin_dethi)        
                self.checkbox_shuffle_dangtoan.setGeometry(QtCore.QRect(le_trai+400, le_top-50, 170, 30))
                self.checkbox_shuffle_dangtoan.setFont(font_10)
                self.checkbox_shuffle_dangtoan.setObjectName("checkbox_shuffle_dangtoan")
                self.checkbox_shuffle_dangtoan.setText("Trộn câu hỏi")

                #Khung text chứa câu hỏi xuất ra
                self.text_taode = QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.text_taode.setGeometry(QtCore.QRect(le_trai, le_top+20, 600, 260))
                self.text_taode.setObjectName("text_taode")

                self.text_taode_HDG = QtWidgets.QTextEdit(parent=self.tab_thongtin_dethi)
                self.text_taode_HDG.setGeometry(QtCore.QRect(le_trai+650, le_top+20, 600, 260))
                self.text_taode_HDG.setFont(font)
                self.text_taode_HDG.setObjectName("text_taode_HDG")
                self.text_taode_HDG.setVisible(False)

                #Label hướng dẫn
                letop_hd=le_top+10

                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label.setGeometry(QtCore.QRect(le_trai+700, letop_hd, 600, 30))
                self.label.setFont(font_tieude)
                self.label.setStyleSheet("color: #697DBA;")
                self.label.setObjectName("label_socau")   
                self.label.setText("*Quy trình chung để tạo đề:")

                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label.setGeometry(QtCore.QRect(le_trai+700, letop_hd+40, 700, 30))
                self.label.setFont(font)        
                self.label.setText("1. Chọn dạng toán và nhập số câu cho mỗi dạng ở tab Thiết lập ma trận.")

                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label.setGeometry(QtCore.QRect(le_trai+700, letop_hd+80, 600, 30))
                self.label.setFont(font)        
                self.label.setText("2. Nhập thông tin đề ở tab Tạo đề và bấm nút Tạo đề.")

                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label.setGeometry(QtCore.QRect(le_trai+700, letop_hd+120, 700, 30))
                self.label.setFont(font)        
                self.label.setText("3. Chọn thư mục để xuất đề. Nếu muốn iMath tự đặt tên thư mục thì chọn Cancel.")

                self.label= QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label.setGeometry(QtCore.QRect(le_trai+700, letop_hd+160, 700, 30))
                self.label.setFont(font)        
                self.label.setText("4. Mở file Word. Chọn Word to PowerPoint. Chọn Xử lí câu hỏi iMath.")

                #Nút tạo đề
                
                self.btn_taode = QtWidgets.QPushButton(parent=self.tab_thongtin_dethi)        
                self.btn_taode.setGeometry(QtCore.QRect(le_trai, le_top+300, 100, 30))
                self.btn_taode.setFont(font)
                self.btn_taode.setObjectName("btn_taode")
                self.btn_taode.setText("Tạo đề")
                self.btn_taode.clicked.connect(self.tao_de)
                self.btn_taode.setStyleSheet("color: white;background-color: #4385F6;")


                #Thanh Progress bar                
                self.progress_bar = QProgressBar(parent=self.tab_thongtin_dethi)
                self.progress_bar.setGeometry(le_trai+200, le_top+330, 305, 15) 

                #Tạo label đang xử lí
                self.label_dangxuli = QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label_dangxuli.setGeometry(QtCore.QRect(le_trai+200, le_top+350, 400, 20))
                self.label_dangxuli.setFont(font)
                self.label_dangxuli.setObjectName("label_dangxuli")

                #Tạo label đang xử lí
                self.label_socau_daxuli = QtWidgets.QLabel(parent=self.tab_thongtin_dethi)
                self.label_socau_daxuli.setGeometry(QtCore.QRect(le_trai+200, le_top+380, 450, 20))
                self.label_socau_daxuli.setFont(font)
                self.label_socau_daxuli.setObjectName("label_socau_daxuli")

       #Tab Đồ thị
                self.label_chonhamso = QtWidgets.QLabel(parent=self.tab_dothi)
                self.label_chonhamso.setGeometry(QtCore.QRect(30, 15, 150, 30))
                self.label_chonhamso.setFont(font)
                self.label_chonhamso.setObjectName("label_chonhamso")
                self.label_chonhamso.setText("Chọn hàm số")

                self.combo_dothihamso = QtWidgets.QComboBox(parent=self.tab_dothi)
                self.combo_dothihamso.setGeometry(QtCore.QRect(140, 15, 250, 30))
                self.combo_dothihamso.setFont(font)
                self.combo_dothihamso.setObjectName("combo_dothihamso")
                self.combo_dothihamso.addItem("Hàm số bậc nhất")
                self.combo_dothihamso.addItem("Hàm số bậc 2")
                self.combo_dothihamso.addItem("Hàm số bậc ba")
                self.combo_dothihamso.addItem("Hàm số bậc bốn trùng phương")
                self.combo_dothihamso.addItem("Hàm số bậc 1/bậc 1")
                self.combo_dothihamso.addItem("Hàm số bậc 2/bậc 1")


                self.tab_dothi_textcode = QtWidgets.QTextEdit(parent=self.tab_dothi)
                self.tab_dothi_textcode.setGeometry(QtCore.QRect(350, 60, 495, 455))
                self.tab_dothi_textcode.setFont(font)
                self.tab_dothi_textcode.setObjectName("tab_dothi_textcode")
               
                
                #text chứa hàm số xuất ra
                self.text_hamso = QtWidgets.QTextBrowser(parent=self.tab_dothi)
                self.text_hamso.setGeometry(QtCore.QRect(410, 15, 320, 30))
                self.text_hamso.setFont(font)
                self.text_hamso.setObjectName("text_hamso")

                # Tạo một đối tượng QFrame
                frame = QtWidgets.QFrame(parent=self.tab_dothi)
                # Thiết lập màu nền của QFrame
                frame.setStyleSheet("background-color: #fdfdfd;")
                # Đặt kích thước và vị trí của QFrame
                frame.setGeometry(850, 15, 500, 500)

                #Text chứa đồ thị xuất ra
                self.label_showpic= QtWidgets.QLabel(parent=self.tab_dothi)
                self.label_showpic.setGeometry(QtCore.QRect(850, 15, 500, 500))    

                self.btn_vedothi = QtWidgets.QPushButton(parent=self.tab_dothi)
                self.btn_vedothi.setGeometry(QtCore.QRect(30, 250, 100, 35))
                self.btn_vedothi.setFont(font)
                self.btn_vedothi.setObjectName("btn_vedothi")
                self.btn_vedothi.setText("Vẽ đồ thị")
                self.btn_vedothi.clicked.connect(self.btn_vedothi_click)

                self.btn_tab_dothi_velaidothi = QtWidgets.QPushButton(parent=self.tab_dothi)
                self.btn_tab_dothi_velaidothi.setGeometry(QtCore.QRect(30, 300, 100, 35))
                self.btn_tab_dothi_velaidothi.setFont(font)
                self.btn_tab_dothi_velaidothi.setObjectName("btn_vedothi")
                self.btn_tab_dothi_velaidothi.setText("Biên dịch lại")
                self.btn_tab_dothi_velaidothi.clicked.connect(self.tab_dothi_velaidothi)

                self.btn_copydothi = QtWidgets.QPushButton(parent=self.tab_dothi)
                self.btn_copydothi.setGeometry(QtCore.QRect(150, 250, 100, 35))
                self.btn_copydothi.setFont(font)
                self.btn_copydothi.setObjectName("btn_copydothi")
                self.btn_copydothi.setText("Copy đồ thị")
                self.btn_copydothi.clicked.connect(self.btn_copydothi_click)

        #Tab Vẽ hình
                self.label_chonhinh = QtWidgets.QLabel(parent=self.tab_vehinh)
                self.label_chonhinh.setGeometry(QtCore.QRect(30, 15, 150, 30))
                self.label_chonhinh.setFont(font)
                self.label_chonhamso.setObjectName("label_chonhamso")
                self.label_chonhinh.setText("Chọn hình")

                self.combo_loaihinh= QtWidgets.QComboBox(parent=self.tab_vehinh)
                self.combo_loaihinh.setGeometry(QtCore.QRect(140, 15, 250, 30))
                self.combo_loaihinh.setFont(font)
                self.combo_loaihinh.setObjectName("combo_loaihinh")
                self.combo_loaihinh.addItem("Hình chóp tam giác cạnh vuông góc")
                self.combo_loaihinh.addItem("Hình chóp tam giác mặt vuông góc")
                self.combo_loaihinh.addItem("Hình chóp tam giác đều")
                self.combo_loaihinh.addItem("Hình tứ diện")
                self.combo_loaihinh.addItem("Hình chóp tứ giác")
                self.combo_loaihinh.addItem("Hình chóp đáy h.b.h cạnh vuông góc")
                self.combo_loaihinh.addItem("Hình chóp đáy h.b.h mặt vuông góc")
                self.combo_loaihinh.addItem("Hình chóp tứ giác đều")
                # self.combo_loaihinh.addItem("Hình chóp đáy hình thang")
                # self.combo_loaihinh.addItem("Hình chóp đáy hình thang vuông")                
                self.combo_loaihinh.addItem("Hình lăng trụ xiên tam giác")
                self.combo_loaihinh.addItem("Hình lăng trụ đứng tam giác")
                self.combo_loaihinh.addItem("Hình hộp")
                self.combo_loaihinh.addItem("Hình lập phương")
               


                self.tab_vehinh_textcode = QtWidgets.QTextEdit(parent=self.tab_vehinh)
                self.tab_vehinh_textcode.setGeometry(QtCore.QRect(350, 60, 495, 455))
                self.tab_vehinh_textcode.setFont(font)
                self.tab_vehinh_textcode.setObjectName("tab_vehinh_textcode")
               
                
                #text chứa hàm số xuất ra
                self.text_hinhve = QtWidgets.QTextBrowser(parent=self.tab_vehinh)
                self.text_hinhve.setGeometry(QtCore.QRect(410, 15, 320, 30))
                self.text_hinhve.setFont(font)
                self.text_hinhve.setObjectName("text_hinhve")

                # Tạo một đối tượng QFrame
                frame = QtWidgets.QFrame(parent=self.tab_vehinh)
                # Thiết lập màu nền của QFrame
                frame.setStyleSheet("background-color: #fdfdfd;")
                # Đặt kích thước và vị trí của QFrame
                frame.setGeometry(850, 15, 500, 500)

                #Text chứa đồ thị xuất ra
                self.label_showpic_vehinh= QtWidgets.QLabel(parent=self.tab_vehinh)
                self.label_showpic_vehinh.setGeometry(QtCore.QRect(850, 15, 500, 500))    

                self.btn_vehinh = QtWidgets.QPushButton(parent=self.tab_vehinh)
                self.btn_vehinh.setGeometry(QtCore.QRect(30, 250, 100, 35))
                self.btn_vehinh.setFont(font)
                self.btn_vehinh.setObjectName("btn_vehinh")
                self.btn_vehinh.setText("Vẽ hình")
                self.btn_vehinh.clicked.connect(self.btn_vehinh_click)

                self.btn_tab_vehinh_velaihinh = QtWidgets.QPushButton(parent=self.tab_vehinh)
                self.btn_tab_vehinh_velaihinh.setGeometry(QtCore.QRect(30, 300, 100, 35))
                self.btn_tab_vehinh_velaihinh.setFont(font)
                self.btn_tab_vehinh_velaihinh.setObjectName("btn_vehinh")
                self.btn_tab_vehinh_velaihinh.setText("Biên dịch lại")
                self.btn_tab_vehinh_velaihinh.clicked.connect(self.tab_vehinh_velaihinh)

                self.btn_copyhinh = QtWidgets.QPushButton(parent=self.tab_vehinh)
                self.btn_copyhinh.setGeometry(QtCore.QRect(150, 250, 100, 35))
                self.btn_copyhinh.setFont(font)
                self.btn_copyhinh.setObjectName("btn_copyhinh")
                self.btn_copyhinh.setText("Copy hình")
                self.btn_copyhinh.clicked.connect(self.btn_copyhinh_click)


        #Tab bbt 
                self.label_chonhamso = QtWidgets.QLabel(parent=self.tab_bbt)
                self.label_chonhamso.setGeometry(QtCore.QRect(30, 15, 150, 30))
                self.label_chonhamso.setFont(font)
                self.label_chonhamso.setObjectName("label_chonhamso")
                self.label_chonhamso.setText("Chọn hàm số")

                self.combo_dothihamso_BBT = QtWidgets.QComboBox(parent=self.tab_bbt)
                self.combo_dothihamso_BBT.setGeometry(QtCore.QRect(140, 15, 250, 30))
                self.combo_dothihamso_BBT.setFont(font)
                self.combo_dothihamso_BBT.setObjectName("combo_dothihamso_BBT")  
                self.combo_dothihamso_BBT.addItem("Hàm số bậc 2")
                self.combo_dothihamso_BBT.addItem("Hàm số bậc ba")
                self.combo_dothihamso_BBT.addItem("Hàm số bậc bốn trùng phương")
                self.combo_dothihamso_BBT.addItem("Hàm số bậc 1/bậc 1")
                self.combo_dothihamso_BBT.addItem("Hàm số bậc 2/bậc 1")
        

                self.tab_bbt_textcode = QtWidgets.QTextEdit(parent=self.tab_bbt)
                self.tab_bbt_textcode.setGeometry(QtCore.QRect(350, 60, 495, 455))
                self.tab_bbt_textcode.setFont(font)
                self.tab_bbt_textcode.setObjectName("tab_bbt_textcode")      
                   

                #text chứa hàm số xuất ra
                self.text_hamso_BBT = QtWidgets.QTextBrowser(parent=self.tab_bbt)
                self.text_hamso_BBT.setGeometry(QtCore.QRect(410, 15, 320, 30))
                self.text_hamso_BBT.setFont(font)
                self.text_hamso_BBT.setObjectName("text_hamso_BBT")

                # Tạo một đối tượng QFrame
                frame = QtWidgets.QFrame(parent=self.tab_bbt)
                frame.setStyleSheet("background-color: #fdfdfd;")    
                frame.setGeometry(850, 15, 500, 500)

                #Text chứa đồ thị xuất ra
                self.label_showpic_BBT= QtWidgets.QLabel(parent=self.tab_bbt)
                self.label_showpic_BBT.setGeometry(QtCore.QRect(850, 15, 700, 500))    

                self.btn_veBBT = QtWidgets.QPushButton(parent=self.tab_bbt)
                self.btn_veBBT.setGeometry(QtCore.QRect(30, 250, 100, 35))
                self.btn_veBBT.setFont(font)
                self.btn_veBBT.setObjectName("btn_veBBT")
                self.btn_veBBT.setText("Vẽ BBT")
                self.btn_veBBT.clicked.connect(self.btn_veBBT_click)

                self.btn_tab_bbt_velaiBBT= QtWidgets.QPushButton(parent=self.tab_bbt)
                self.btn_tab_bbt_velaiBBT.setGeometry(QtCore.QRect(30, 300, 100, 35))
                self.btn_tab_bbt_velaiBBT.setFont(font)
                self.btn_tab_bbt_velaiBBT.setObjectName("btn_velaiBBT")
                self.btn_tab_bbt_velaiBBT.setText("Biên dịch lại")
                #self.btn_tab_bbt_velaidothi.clicked.connect(self.tab_bbt_velaidothi)

                self.btn_copyBBT = QtWidgets.QPushButton(parent=self.tab_bbt)
                self.btn_copyBBT.setGeometry(QtCore.QRect(150, 250, 100, 35))
                self.btn_copyBBT.setFont(font)
                self.btn_copyBBT.setObjectName("btn_copyBBT")
                self.btn_copyBBT.setText("Copy BBT")
                self.btn_copyBBT.clicked.connect(self.btn_copyBBT_click)

        #Tab Hướng dẫn
                self.btn_hd_taode_video= QtWidgets.QPushButton(parent=self.tab_huongdan)
                self.btn_hd_taode_video.setGeometry(QtCore.QRect(30, 50, 250, 35))
                self.btn_hd_taode_video.setFont(font_10)
                self.btn_hd_taode_video.setObjectName("btn_copyBBT")
                self.btn_hd_taode_video.setText("Video hướng dẫn tạo đề")
                self.btn_hd_taode_video.clicked.connect(self.hd_taode_video)

                self.btn_hd_taode_pdf= QtWidgets.QPushButton(parent=self.tab_huongdan)
                self.btn_hd_taode_pdf.setGeometry(QtCore.QRect(30, 100, 250, 35))
                self.btn_hd_taode_pdf.setFont(font_10)
                self.btn_hd_taode_pdf.setObjectName("btn_copyBBT")
                self.btn_hd_taode_pdf.setText("Hướng dẫn tạo đề - PDF")
                self.btn_hd_taode_pdf.clicked.connect(self.hd_taode_pdf)

                self.btn_show_tool_Word2PPT= QtWidgets.QPushButton(parent=self.tab_huongdan)
                self.btn_show_tool_Word2PPT.setGeometry(QtCore.QRect(30, 150, 250, 35))
                self.btn_show_tool_Word2PPT.setFont(font_10)
                self.btn_show_tool_Word2PPT.setObjectName("btn_copyBBT")
                self.btn_show_tool_Word2PPT.setText("Hiện lại tool Word to PPT")
                self.btn_show_tool_Word2PPT.clicked.connect(self.show_tool_Word2PPT)

                letop_hd=200

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd, 600, 30))
                self.label.setFont(font_tieude)
                self.label.setStyleSheet("color: #697DBA;")
                self.label.setObjectName("label_socau")  
                self.label.setText("*Quy trình chung để tạo đề:")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+40, 600, 30))
                self.label.setFont(font)        
                self.label.setText("1. Chọn dạng toán và nhập số câu cho mỗi dạng ở Tab Thiết lập ma trận.")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+80, 600, 30))
                self.label.setFont(font)        
                self.label.setText("2. Thiết lập tùy chọn ở tab Tạo đề và bấm nút Tạo đề.")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+120, 600, 30))
                self.label.setFont(font)        
                self.label.setText("3. Mở file Word. Chọn Word to PowerPoint. Chọn Xử lí câu hỏi iMath.")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+160, 600, 30))
                self.label.setFont(font)        
                self.label.setText("*Chú ý:")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+200, 1250, 30))
                self.label.setFont(font)        
                self.label.setText("1) Lần đầu tạo đề nếu thấy hộp cảnh báo Alway show this dialog và Install thì: tắt Alway show this dialog rồi bấm Install và chờ đến khi xuất được đề.")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+240, 1000, 30))
                self.label.setFont(font)        
                self.label.setText("2) Lần đầu tạo đề chứa hình vẽ máy sẽ treo vì tải cập nhật gói vẽ hình cần chờ đến khi xuất được đề.")

                self.label= QtWidgets.QLabel(parent=self.tab_huongdan)
                self.label.setGeometry(QtCore.QRect(30, letop_hd+280, 1000, 30))
                self.label.setFont(font)        
                self.label.setText("3) Tạo đề cho học sinh làm bài tập nên chọn Tạo đề Word - Equation để xuất đề nhanh nhất.")




        #Tab latextoimage

                # self.text_run_latex = QtWidgets.QTextEdit(parent=self.tab_run_latex)
                # self.text_run_latex.setGeometry(QtCore.QRect(10, 40, 500, 500))
                # self.text_run_latex.setFont(font)
                # self.text_run_latex.setObjectName("text_run_latex")
                # code=my_module.moi_truong_latex("")
                # self.text_run_latex.setText(code)

                # # Tạo QFrame
                # frame = QtWidgets.QFrame(parent=self.tab_run_latex)
                # # Thiết lập màu nền của QFrame
                # frame.setStyleSheet("background-color: #fdfdfd;")
                # frame.setStyleSheet("border: 1px solid grey; padding: 10px;")
                # # Đặt kích thước và vị trí của QFrame
                # frame.setGeometry(650, 40, 700, 500)

                # #Label chứa đồ thị xuất ra
                # self.label_latextoimage= QtWidgets.QLabel(parent=self.tab_run_latex)
                # self.label_latextoimage.setGeometry(QtCore.QRect(660, 40, 700, 500))

                # top_biendichlatex=40
                # left_biendichlatex=515
                # #Button chạy code biên dịch latex 
                # self.btn_moitruong_latex = QtWidgets.QPushButton(parent=self.tab_run_latex)
                # self.btn_moitruong_latex.setGeometry(QtCore.QRect(left_biendichlatex, top_biendichlatex, 130, 35))
                # self.btn_moitruong_latex.setFont(font)
                # self.btn_moitruong_latex.setObjectName("btn_moitruong_latex")
                # self.btn_moitruong_latex.setText("Môi trường Latex")
                # self.btn_moitruong_latex.clicked.connect(self.btn_moitruong_latex_click)

                # #Button chạy code biên dịch latex 
                # self.btn_run_latex = QtWidgets.QPushButton(parent=self.tab_run_latex)
                # self.btn_run_latex.setGeometry(QtCore.QRect(left_biendichlatex, top_biendichlatex+40, 130, 35))
                # self.btn_run_latex.setFont(font)
                # self.btn_run_latex.setObjectName("btn_run_latex")
                # self.btn_run_latex.setText("Biên dịch Latex")
                # self.btn_run_latex.clicked.connect(self.btn_run_latex_click)

        # Tab Tạo đề
        #label chọn dạng toán
                self.label_dangcauhoi = QtWidgets.QLabel(parent=self.tab_taode)
                self.label_dangcauhoi.setGeometry(QtCore.QRect(0, 5, 200, 30))
                self.label_dangcauhoi.setFont(font_tieude)
                self.label_dangcauhoi.setStyleSheet("color: #697DBA;")
                self.label_dangcauhoi.setObjectName("label_socau")
                self.label_dangcauhoi.setText("1. Chọn dạng toán")

        # Checkbox chọn ngẫu nhiên từ thư mục con
                self.checkbox_tree_random = QtWidgets.QLabel(parent=self.tab_taode)        
                self.checkbox_tree_random.setGeometry(QtCore.QRect(25, 25, 205, 30))
                self.checkbox_tree_random.setFont(font_10)
                self.checkbox_tree_random.setObjectName("checkbox_tree_random")
                self.checkbox_tree_random.setText("Số dạng ngẫu nhiên từ mục con")
        
                self.soluong_dangtoan = QtWidgets.QTextEdit(parent=self.tab_taode)
                self.soluong_dangtoan.setGeometry(QtCore.QRect(235, 25, 30, 30))
                self.soluong_dangtoan.setFont(font)
                self.soluong_dangtoan.setObjectName("soluong_dangtoan")

        #label thông tin dạng toán
                self.label_dangcauhoi = QtWidgets.QLabel(parent=self.tab_taode)
                self.label_dangcauhoi.setGeometry(QtCore.QRect(720, 5, 200, 30))
                self.label_dangcauhoi.setFont(font_tieude)
                self.label_dangcauhoi.setStyleSheet("color: #697DBA;")
                self.label_dangcauhoi.setObjectName("label_socau")
                self.label_dangcauhoi.setText("2. Thông tin dạng toán")

        # Cây thư mục
                self.treeWidget = QtWidgets.QTreeWidget(parent=self.tab_taode)
                self.treeWidget.setGeometry(QtCore.QRect(0, 50, 720, 600))
                self.treeWidget.setObjectName("treeWidget")
                
               
                self.treeWidget.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOn)
                self.treeWidget.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOn)        
                 

                #self.treeWidget.setHeaderLabels(["Chọn dạng toán"])
                self.treeWidget.setHeaderHidden(True)
                L10 = QTreeWidgetItem(self.treeWidget, ["Lớp 10"])
                L10.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1 = QTreeWidgetItem(L10, ["Chương 1 -  Mệnh đề và tập hợp"])
                L10_C1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B1 = QTreeWidgetItem(L10_C1, ["Bài 1 - Mệnh đề"])
                L10_C1_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B1_1 = QTreeWidgetItem(L10_C1_B1, ["Trắc nghiệm - Trả lời ngắn"])
                L10_C1_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_1, ["[D10_C1_B1_04]-M2. Tìm mệnh đề đúng (cảm thán, số là số gì, số chia hết cho, nghiệm của phương trình)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_1, ["[D10_C1_B1_05]-M2. Tìm mệnh đề sai (tỉnh thuộc miền, số là số gì, số chia hết cho, nghiệm của phương trình)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_1, ["[D10_C1_B1_07]-M2. Tìm mềnh đề phủ định của mệnh đề không chứa biến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_1, ["[D10_C1_B1_06]-M2. Tìm mềnh đề phủ định của mệnh đề chứa biến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_1, ["[D10_C1_B1_08]-M3. Tìm mệnh đề kéo theo đúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B1_2 = QTreeWidgetItem(L10_C1_B1, ["Đúng-Sai"])
                L10_C1_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_2, ["[D10_C1_B1_02]-TF-M2. Xét Đ-S về số chính phương, số hữu tỉ, số chẵn, lẻ, số chia hết.."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_2, ["[D10_C1_B1_01]-TF-M3. Xét tính Đ-S của mệnh đề, mệnh đề phủ định, mệnh đề chứa biến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B1_2, ["[D10_C1_B1_03]-TF-M3. Cho P(n)=an^2+bn+c. Xét Đ-S: P(x_0) chẵn (lẻ),P(x_1) chia hết cho, P(ta)-P(a), ..."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B2 = QTreeWidgetItem(L10_C1, ["Bài 2 - Tập hợp"])
                L10_C1_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B2_1 = QTreeWidgetItem(L10_C1_B2, ["Trắc nghiệm - Trả lời ngắn"])
                L10_C1_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_11]-M2. Liệt kê phần tử thỏa mãn điều kiện tùy ý"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_01]-M2. Liệt kê các phần tử của tập hợp số nguyên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)        

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_02]-M2. Liệt kê các phần tử là nghiệm của phương trình bậc 2"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_03]-M2. Liệt kê các phần tử là ước của n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_04]-M2. Liệt kê các phần tử là bội của n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_09]-M2. Liệt kê các phần tử là số nguyên tố"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_10]-M2. Liệt kê các phần tử là số chẵn(lẻ)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_1, ["[D10_C1_B2_05]-M2. Chỉ ra tính chất đặc trưng của tập hợp"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B2_2 = QTreeWidgetItem(L10_C1_B2, ["Đúng-Sai"])
                L10_C1_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_2, ["[D10_C1_B2_06]-TF-M2. Cho tập hợp dạng đặc trưng. Xét Đ-S: liệt kê, số phần tử, số tập con, xác định tập con"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_2, ["[D10_C1_B2_08]-TF-M2. Cho tập hợp dạng liệt kê. Xét Đ-S: liệt kê, số phần tử, số tập con, xác định tập con"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B2_2, ["[D10_C1_B2_07]-TF-M3. Xét đúng sai về số phần tử của: a<x<b, ax^2+bx+c=0, phương trình tích, m<ax+b<n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #BÀI 3         

                L10_C1_B3 = QTreeWidgetItem(L10_C1, ["Bài 3 - Các phép toán trên tập hợp"])
                L10_C1_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B3_1 = QTreeWidgetItem(L10_C1_B3, ["Trắc nghiệm - Trả lời ngắn"])
                L10_C1_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_01]-M2. Tìm giao hai tập hợp dạng liệt kê"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_02]-M2. Tìm hợp hai tập hợp dạng liệt kê"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_03]-M2. Tìm hiệu hai tập hợp dạng liệt kê"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_04]-M2. Tìm phần bù hai tập hợp dạng liệt kê"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_16]-M2. Tìm một phép toán của hai tập hợp dạng liệt kê"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_05]-M2. Tìm giao hai tập hợp dạng tính chất đặc trưng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_06]-M2. Tìm hợp hai tập hợp dạng tính chất đặc trưng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_07]-M2. Tìm hiệu hai tập hợp dạng tính chất đặc trưng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_08]-M1. Tìm số tập hợp con của một tập hợp"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_1, ["[D10_C1_B3_09]-M2. Tìm số tập hợp con gồm k phần tử của một tập hợp"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B3_2 = QTreeWidgetItem(L10_C1_B3, ["Đúng-Sai"])
                L10_C1_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_2, ["[D10_C1_B3_10]-TF-M2. Cho 2 tập hợp dạng liệt kê. Xét Đ-S: Giao, Hợp, Hiệu, Kiểm tra tập con"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_2, ["[D10_C1_B3_11]-TF-M2. Cho 2 tập hợp dạng đặc trưng. Xét Đ-S: Giao, Hợp, Hiệu, Kiểm tra tập con"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_2, ["[D10_C1_B3_17]-TF-M3. Cho A-đặc trưng, B-liệt kê. Xét Đ-S: Liệt kê A, Số tập hợp con, phép toán, A con X con B"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_2, ["[D10_C1_B3_15]-TF-M3.  Bài thực tế: Cho n(A), n(B), n(A giao B), n_khôngAB. Xét Đ-S: Chỉ A, Chỉ B, A hoặc B, Tổng số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B3_3 = QTreeWidgetItem(L10_C1_B3, ["Trả lời ngắn"])
                L10_C1_B3_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B3_3.setCheckState(0, Qt.CheckState.PartiallyChecked)  

                item = QTreeWidgetItem(L10_C1_B3_3, ["[D10_C1_B3_12]-TL-M3.  Bài thực tế: Cho n(A), n(B). Tính n(A giao B)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_3, ["[D10_C1_B3_13]-TL-M3.  Bài thực tế: Cho n(A), n(B), n(A giao B). Tính tổng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_3, ["[D10_C1_B3_14]-TL-M3.  Bài thực tế: Cho n(A), n(B), n(A giao B), n_khong(AB). Tính tổng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_3, ["[D10_C1_B3_18]-TL-M4.  Bài toán thực tế liên quan 3 tập hợp (biết ít nhất 1 môn)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B3_3, ["[D10_C1_B3_19]-TL-M4.  Bài toán thực tế liên quan 3 tập hợp (biết đúng 1 môn)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                


                

                L10_C1_B4 = QTreeWidgetItem(L10_C1, ["Bài 4 - Các tập hợp con đặc biệt của R"])
                L10_C1_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B4_1 = QTreeWidgetItem(L10_C1_B4, ["Trắc nghiệm"])
                L10_C1_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_01]-M2. Cho kí hiệu khoảng,đoạn. Mô tả tính chất đặc trưng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_02]-M2. Cho tính chất đặc trưng tìm kí hiệu khoảng đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_14]-M3.  Tìm kí hiệu khoảng, đoạn ứng với tập hợp chứa điều kiện."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_03]-M2. Tìm giao các khoảng, đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_04]-M2. Tìm hợp các khoảng, đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_05]-M2. Tìm hiệu các khoảng, đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_06]-M2. Tìm phần bù của khoảng, đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_1, ["[D10_C1_B4_16]-M2. Tìm giao (hoặc hợp, hiệu, phần bù) các khoảng, đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B4_2 = QTreeWidgetItem(L10_C1_B4, ["Đúng-Sai"])
                L10_C1_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_2, ["[D10_C1_B4_07]-TF-M2. Cho A, B là các khoảng đoạn. Xét Đ-S: A giao B, A hợp B, A\\B, B\\A, C_R(A), C_R(B)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_2, ["[D10_C1_B4_13]-TF-M2. Xét Đ-S: Kí hiệu khoảng, đoạn, nữa khoảng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_2, ["[D10_C1_B4_15]-TF-M2. Xét Đ-S: Kí hiệu, phép toán, giao với N-Z, tìm m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C1_B4_3 = QTreeWidgetItem(L10_C1_B4, ["Trả lời ngắn"])
                L10_C1_B4_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C1_B4_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_18]-TL-M2. Tìm số nguyên thuộc một phép toán khoảng đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_08]-TL-M3. Cho A, B. Cho A, B. Tìm m để A giao B có đúng 1 phần tử"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_09]-TL-M3. Cho A, B. Tìm m để A giao B khác rỗng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_10]-TL-M3. Cho A, B. Tìm m để A giao B bằng rỗng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_11]-TL-M3. Cho A, B. Tìm m để (A U B)=A."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_12]-TL-M3. Cho A, B. Tìm m để A con B."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C1_B4_3, ["[D10_C1_B4_17]-TL-M3. Tìm m thỏa điều kiện về phép toán khoảng đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #CHƯƠNG 2 - BẤT PHƯƠNG TRÌNH VÀ HỆ BẤT PHƯƠNG TRÌNH BẬC NHẤT 2 ẨN
                L10_C2 = QTreeWidgetItem(L10, ["Chương 2 -  Bất phương trình và hệ bất phương trình bậc nhất 2 ẩn"])
                L10_C2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                #BÀI 1 - BẤT PHƯƠNG TRÌNH BẬC NHẤT 2 ẨN
                L10_C2_B1 = QTreeWidgetItem(L10_C2, ["Bài 1 - Bất phương trình bậc nhất 2 ẩn"])
                L10_C2_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C2_B1_1 = QTreeWidgetItem(L10_C2_B1, ["Trắc nghiệm"])
                L10_C2_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_01]-M2. Tìm cặp số là nghiệm của BPT ax+by+c>0 (>=0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_02]-M2. Tìm cặp số là nghiệm của BPT ax+by+c<0 (<=0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_03]-M1. Nhận dạng bất phương trình bậc nhất 2 ẩn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_04]-M2. Miền nghiệm của ax+by+c>0 (<0) là nữa mặt phẳng chứa điểm nào"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_05]-M2. Miền nghiệm của ax+by+c >(<) ex+fy+g là nữa mặt phẳng chứa điểm nào"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_06]-M2. Cặp số nào là nghiệm thuộc nữa mặt phẳng có hình vẽ minh họa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_07]-M2. Tìm miền nghiệm của bất phương trình ax+by+c>0 (<0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_1, ["[D10_C2_B1_09]-M2. Lập BPT bậc nhất 2 ẩn từ bài toán thực tế."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C2_B1_2 = QTreeWidgetItem(L10_C2_B1, ["Đúng-Sai"])
                L10_C2_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B1_2, ["[D10_C2_B1_08]-TF-M2. Cho BPT ax+by+c>0 (<0). Xét Đ-S: nghiệm, không là nghiệm, nghiệm chứa bờ, miền nghiệm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                #BÀI 2 - HỆ BẤT PHƯƠNG TRÌNH BẬC NHẤT 2 ẨN

                L10_C2_B2 = QTreeWidgetItem(L10_C2, ["Bài 2 - Hệ bất phương trình bậc nhất 2 ẩn"])
                L10_C2_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C2_B2_1 = QTreeWidgetItem(L10_C2_B2, ["Trắc nghiệm"])
                L10_C2_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B2_1, ["[D10_C2_B2_01]-M2. Cho Hệ 2 BPT. Tìm cặp số là nghiệm của hệ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B2_1, ["[D10_C2_B2_02]-M2. Cho Hệ 2 BPT. Tìm cặp số không là nghiệm của hệ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # item = QTreeWidgetItem(L10_C2_B2_1, ["[D10_C2_B2_03]-M3. Cho Hệ 3 BPT. Tìm cặp số là nghiệm của hệ."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C2_B2_2 = QTreeWidgetItem(L10_C2_B2, ["Đúng-Sai"])
                L10_C2_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C2_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C2_B2_2, ["[D10_C2_B2_04]-TF-M2. Cho hệ 2 BPT. Xét đúng sai về nghiệm của hệ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                #CHƯƠNG 3 - HÀM SỐ VÀ ĐỒ THỊ       

                L10_C3 = QTreeWidgetItem(L10, ["Chương 3 -  Hàm số và đồ thị"])
                L10_C3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled) 
                L10_C3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C3_B1 = QTreeWidgetItem(L10_C3, ["Bài 1 - Hàm số và đồ thị"])
                L10_C3_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C3_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_01]-M1. Tính giá trị của hàm số tại một điểm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_02]-M2. Tính giá trị của hàm số cho bởi 2 biểu thức tại một điểm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giá trị của hàm số cho bởi 2 biểu thức tại một điểm')

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_03]-M1. Tìm tập xác định y =(ax+b)/(cx+d)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm tập xác định y =(ax+b)/(cx+d).')

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_04]-M1. Tìm tập xác định y = căn(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm tập xác định y = căn(ax+b).')


                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_05]-M2. Tìm tập xác định y = căn(A) + B/C."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm tập xác định y = căn(A) + B/C.') 

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_06]-M2. Tìm tập giá trị dựa vào hình vẽ đồ thị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm tập giá trị dựa vào hình vẽ đồ thị.")       

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_07]-M2. Tìm tập xác định y =căn(ax+b) + căn(cx+d)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm tập xác định y =căn(ax+b) + căn(cx+d).")

                item = QTreeWidgetItem(L10_C3_B1, ["[D10_C3_B1_08]-M2. Tìm tập xác định hàm số y =A/(ax^2+bx+c)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm tập xác định hàm số y =A/(ax^2+bx+c)")

                L10_C3_B2 = QTreeWidgetItem(L10_C3, ["Bài 2 - Hàm số bậc 2"])
                L10_C3_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C3_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C3_B2_1 = QTreeWidgetItem(L10_C3_B2, ["10.2.1 Trắc Nghiệm"])
                L10_C3_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C3_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_01]-M1. Tìm tọa độ đỉnh của hàm số bậc 2"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm tọa độ đỉnh của hàm số bậc 2")

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_02]-M1.Tìm trục đối xứng của đồ thị hàm số bậc 2"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm trục đối xứng của đồ thị hàm số bậc 2")

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_03]-M2. Cho hàm số bậc 2. Tìm khoảng biến thiên."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Cho hàm số bậc 2. Tìm khoảng biến thiên.")

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_04]-M2. Cho BBT hàm số bậc 2. Tìm khoảng biến thiên."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Cho BBT hàm số bậc 2. Tìm khoảng biến thiên.")        

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_05]-M2. Cho BBT. Tìm hàm số bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Cho BBT. Tìm hàm số bậc 2.") 

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_06]-M2. Cho đồ thị. Tìm hàm số bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Cho đồ thị. Tìm hàm số bậc 2.") 

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_07]-M2. Tìm hàm số bậc 2 có đồ thị đi qua một điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm hàm số bậc 2 có đồ thị đi qua một điểm.") 

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_08]-M2. Tìm hàm số bậc 2 có đồ thị đi qua hai điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm hàm số bậc 2 có đồ thị đi qua hai điểm.") 

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_09]-M3. Tìm hàm số bậc 2 có đồ thị đi qua điểm và trục đối xứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, "Tìm hàm số bậc 2 có đồ thị đi qua điểm và trục đối xứng.")

                item = QTreeWidgetItem(L10_C3_B2_1, ["[D10_C3_B2_10]-M3. Cho hàm số bậc 2. Xét dấu các hệ số a,b,c."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Câu hỏi đúng sai hàm số bậc 2

                L10_C3_B2_2 = QTreeWidgetItem(L10_C3_B2, ["10.2.2 Đúng-Sai"])
                L10_C3_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C3_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B2_2, ["[D10_C3_B2_11]-TF-M2. Cho hàm số bậc 2. Tạo Đúng-Sai hỗn hợp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B2_2, ["[D10_C3_B2_12]-TF-M2. Cho hàm số bậc 2. Tạo Đúng-Sai về đồ thị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C3_B2_2, ["[D10_C3_B2_13]-TF-M2. Cho BBT hàm số bậc 2. Tạo Đúng-Sai."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Câu hỏi trả lời ngắn

                L10_C3_B2_3 = QTreeWidgetItem(L10_C3_B2, ["10.2.3 Câu hỏi trả lời ngắn"])
                L10_C3_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C3_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # LỚP 10 - CHƯƠNG 4 - HỆ THỨC LƯỢNG TRONG TAM GIÁC
                L10_C4 = QTreeWidgetItem(L10, ["Chương 4 - Hệ thức lượng trong tam giác"])
                L10_C4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #BÀI 1- Giá trị lượng giác của một góc từ 0 đến 180
                L10_C4_B1 = QTreeWidgetItem(L10_C4, ["Bài 1 - Giá trị lượng giác của một góc từ 0 đến 180"])
                L10_C4_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B1, ["[D10_C4_B1_01]-M1. Tìm dấu của một giá trị lượng giác từ 0 đến 180 độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm dấu của một giá trị lượng giác từ 0 đến 180 độ')

                item = QTreeWidgetItem(L10_C4_B1, ["[D10_C4_B1_02]-M1. Cho giá trị lượng giác, tìm góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho giá trị lượng giác, tìm góc')

                #BÀI 2 - ĐỊNH LÍ COSIN VÀ ĐỊNH LÍ SIN
                L10_C4_B2 = QTreeWidgetItem(L10_C4, ["Bài 2 - Định lí cosin và định lí sin"])
                L10_C4_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C4_B2_1 = QTreeWidgetItem(L10_C4_B2, ["Trắc nghiệm"])
                L10_C4_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_02]-M2. Cho 2 cạnh và góc xen giữa. Tính độ dài cạnh thứ 3."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_03]-M2. Tính số đo một góc biết độ dài 3 cạnh."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)          

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_04]-M2. Tam giác có một cạnh và góc đối diện. Tính bán kính đường tròn ngoại tiếp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)         

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_01]-M1. Cho a,b,C. Tính diện tích."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)      

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_05]-M2. Cho a,b,c. Tính diện tích tam giác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_1, ["[D10_C4_B2_12]-M2. Cho A,B và cạnh a. Tính cạnh b hoặc c"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L10_C4_B2_2 = QTreeWidgetItem(L10_C4_B2, ["Đúng-Sai"])
                L10_C4_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_06]-TF-M2. Cho a,b và góc C. Xét Đ-S: cạnh c, cos A, góc B, diện tích."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_07]-TF-M2. Cho AC,BC và góc C. Xét Đ-S: cạnh AB, cos A, góc B, diện tích."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_08]-TF-M2. Cho a,b c. Xét Đ-S: p, S, r, R"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_09]-TF-M2. Cho a,b,c. Xét Đ-S: cosA, B, S, sinC, R, r"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_10]-TF-M2. Cho AB,BC,AC. Xét Đ-S: cosA, B, S, sinC, R, r"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_2, ["[D10_C4_B2_11]-TF-M2. Cho A,B,a. Xét Đ-S: góc C, cạnh b, cạnh c, diện tích"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C4_B2_3 = QTreeWidgetItem(L10_C4_B2, ["Trả lời ngắn"])
                L10_C4_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C4_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C4_B2_3, ["[D10_C4_B2_13]-TL-M2.  Cho b,c và góc C. Tính cạnh a"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 10 - Chương 5
                L10_C5 = QTreeWidgetItem(L10, ["Chương 5 - Véctơ"])
                L10_C5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Lớp 11 - Chương1 - Bài 1 - Góc lượng giác
                L10_C5_B1 = QTreeWidgetItem(L10_C5, ["Bài 1 - Véctơ"])
                L10_C5_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B1_1 = QTreeWidgetItem(L10_C5_B1, ["Trắc nghiệm"])
                L10_C5_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_01]-M2. Cho tam giác đều. Tìm khẳng định sai về hướng, phương, độ dài"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_02]-M2. Cho hình dạng h.b.h. Tìm khẳng định sai về vectơ bằng nhau, độ dài bằng nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_03]-M2. Cho hình dạng h.b.h. Tìm khẳng định đúng về vectơ hướng, phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_04]-M2. Cho trung điểm của đoạn. Tìm khẳng định đúng về hướng, phương, véctơ bằng nhau, độ dài bằng nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_05]-M2. Cho 3 điểm. Tìm khẳng định về hướng và phương"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_06]-M1. Cho hình vuông (chữ nhật). Tính độ dài vectơ cạnh."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B1_1, ["[D10_C5_B1_07]-M2. Cho hình vuông (chữ nhật). Tính độ dài vectơ đường chéo."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B2 = QTreeWidgetItem(L10_C5, ["Bài 2 - Tổng hiệu các véctơ"])
                L10_C5_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B2_1 = QTreeWidgetItem(L10_C5_B2, ["Trắc nghiệm"])
                L10_C5_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_08]-M2. Nhận dạng quy tắc hình bình hành."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_07]-M2. Cho các điểm. Tính tổng các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_03]-M2. Tìm đẳng thức đúng liên quan đến 3 điểm (phép cộng trừ)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_01]-M2. Cho trung điểm của đoạn. Tìm khẳng định tổng hiệu các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_02]-M2. Cho các điểm. Tính tổng hiệu các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_04]-M2. Cho tứ giác. Tìm quy tắc đúng về phép cộng trừ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_05]-M2. Cho hình bình hành. Tìm khẳng định sai về phép toán vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_1, ["[D10_C5_B2_06]-M2. Cho hình bình hành. Tính tổng hiệu các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B2_2 = QTreeWidgetItem(L10_C5_B2, ["Đúng-Sai"])
                L10_C5_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_2, ["[D10_C5_B2_09]-TF-M3. Cho h.b.h. Xét Đ-S: quy tắc h.b.h, tổng-hiệu vectơ, độ dài."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_2, ["[D10_C5_B2_11]-TF-M3. Cho h.c.n. Xét Đ-S:hướng-phương-bằng, quy tắc h.b.h, tổng-hiệu, độ dài các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B2_2 = QTreeWidgetItem(L10_C5_B2, ["Trả lời ngắn"])
                L10_C5_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_2, ["[D10_C5_B2_10]-TL-M2. Cho hai lực hợp góc. Tính độ lớn của tổng lực.."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_2, ["[D10_C5_B2_12]-TL-M3. Cho tam giác vuông. Tính độ dài của vectơ có điểm thỏa mãn đẳng thức vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B2_2, ["[D10_C5_B2_13]-TL-M3. Cho h.c.n. Tính độ dài tổng-hiệu các vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B3 = QTreeWidgetItem(L10_C5, ["Bài 3 - Tích véctơ với một số"])
                L10_C5_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B3_1 = QTreeWidgetItem(L10_C5_B3, ["Trắc nghiệm"])
                L10_C5_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B3_1, ["[D10_C5_B3_01]-M2. Cho tứ giác. Tìm khẳng định đúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B3_1, ["[D10_C5_B3_02]-M2. Cho tam giác có trọng tâm. Tìm khẳng định đúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B3_1, ["[D10_C5_B3_03]-M2. Cho MA=kMB. Tìm đẳng thức đúng về vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B3_1, ["[D10_C5_B3_04]-M2. Cho hình bình hành. Tìm đẳng thức đúng về vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B3_2 = QTreeWidgetItem(L10_C5_B3, ["Đúng-Sai"])
                L10_C5_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B3_2, ["[D10_C5_B3_05]-TF-M2. Cho hình bình hành. Xét đúng sai các đẳng thức vectơ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B4 = QTreeWidgetItem(L10_C5, ["Bài 4 - Tích vô hướng của hai vectơ"])
                L10_C5_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B4_1 = QTreeWidgetItem(L10_C5_B4, ["Trắc nghiệm"])
                L10_C5_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_1, ["[D10_C5_B4_01]-M1.  Cho hai véctơ có độ dài và góc. Tính tích vô hướng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_1, ["[D10_C5_B4_02]-M2.  Cho hai véctơ có độ dài và tích vô hướng. Tìm góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_1, ["[D10_C5_B4_03]-M3.  Cho hai véctơ có độ dài và góc. Tính |vta + vtb|."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_1, ["[D10_C5_B4_04]-M3.  Cho hai véctơ có độ dài và góc. Tính |vta - vtb|."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C5_B4_2 = QTreeWidgetItem(L10_C5_B4, ["Đúng-Sai"])
                L10_C5_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C5_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_2, ["[D10_C5_B4_05]-TF-M2. Cho hình vuông. Xét Đ-S: các phép toán về vectơ, tích vô hướng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C5_B4_2, ["[D10_C5_B4_06]-TF-M2. Cho tam giác. Xét Đ-S: Tích vô hướng, phép toán vectơ, độ dài."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                

        #Lớp 10 - Chương 6
                L10_C6 = QTreeWidgetItem(L10, ["Chương 6 - Thống kê"])
                L10_C6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C6.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 
                L10_C6_B1 = QTreeWidgetItem(L10_C6, ["Bài 1 - Số gần đúng và sai số"])
                L10_C6_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C6_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B1, ["[D10_C6_B1_01]-M2. Viết số quy tròn của số nguyên có độ chính xác cho trước."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 2

                L10_C6_B2 = QTreeWidgetItem(L10_C6, ["Bài 2 - Mô tả và biểu diễn dữ liệu trên các bảng và biểu đồ"])
                L10_C6_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C6_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 3
                L10_C6_B3 = QTreeWidgetItem(L10_C6, ["Bài 3 - Các số đặc trưng đo xu thế trung tâm"])
                L10_C6_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C6_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_01]-M2. Cho dãy số liệu. Tính số trung bình."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_02]-M2. Cho dãy số liệu. Tính số trung vị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_03]-M2. Cho dãy số liệu. Tính mốt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_04]-M2. Cho dãy số liệu. Tính tứ phân vị thứ nhất"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_05]-M2. Cho dãy số liệu. Tính tứ phân vị thứ hai"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B3, ["[D10_C6_B3_06]-M2. Cho dãy số liệu. Tính tứ phân vị thứ ba"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 4
                L10_C6_B4 = QTreeWidgetItem(L10_C6, ["Bài 4 - Các số đặc trưng đo xu thế mức độ phân tán"])
                L10_C6_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C6_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B4, ["[D10_C6_B4_01]-M2. Cho dãy số liệu. Tính độ lệch chuẩn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C6_B4, ["[D10_C6_B4_02]-M2. Cho dãy số liệu. Tính phương sai"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 10 - Chương 7
                L10_C7 = QTreeWidgetItem(L10, ["Chương 7 - Bất phương trình bậc 2 một ẩn"])
                L10_C7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 
                L10_C7_B1 = QTreeWidgetItem(L10_C7, ["Bài 1 - Dấu của tam thức bậc 2"])
                L10_C7_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C7_B1_1 = QTreeWidgetItem(L10_C7_B1, ["7.1. Trắc Nghiệm"])
                L10_C7_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_05]-M1. Xác định biểu thức nào là tam thức bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_04]-M1. Tìm m để biểu thức là bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_01]-M1. Cho biểu thức bậc 2 vô nghiệm. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_02]-M1. Cho biểu thức bậc 2 có nghiệm kép. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_03]-M2. Cho biểu thức bậc 2 có 2 nghiệm. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_06]-M1. Cho đồ thị bậc 2 luôn âm hoặc luôn dương.Tìm khẳng định đúng về dấu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_07]-M1. Cho đồ thị bậc 2 có nghiệm kép.Tìm khẳng định đúng về dấu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_08]-M1. Cho đồ thị bậc 2 có 2 nghiệm.Tìm khẳng định đúng về dấu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_09]-M1. Cho BXD bậc 2 luôn âm hoặc luôn dương. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_10]-M1. Cho BXD bậc 2 có 1 nghiệm. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_11]-M1. Cho BXD bậc 2 có 2 nghiệm. Tìm khẳng định đúng về dấu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_14]-M2. Cho BXD bậc 2 có 2 nghiệm. Tìm biểu thức bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_1, ["[D10_C7_B1_15]-M2. Cho BXD bậc 2 có 1 nghiệm. Tìm biểu thức bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C7_B1_2 = QTreeWidgetItem(L10_C7_B1, ["7.2. Đúng-Sai"])
                L10_C7_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_2, ["[D10_C7_B1_12]-TF-M2. Cho biểu thức bậc 2 có 2 nghiệm. Tạo câu đúng-sai."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B1_2, ["[D10_C7_B1_13]-TF-M2. Cho bảng xét dấu hai nghiệm. Tạo câu đúng-sai."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C7_B2 = QTreeWidgetItem(L10_C7, ["Bài 2 - Bất phương trình bậc 2 một ẩn"])
                L10_C7_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_01]-M2. Giải BPT bậc 2, tam thức vô nghiệm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_02]-M2. Giải BPT bậc 2, tam thức nghiệm kép."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_03]-M2. Giải BPT bậc 2, tam thức 2 nghiệm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_04]-M2. Giải BPT có 2 vế là bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_05]-M3. Tìm m để ax^2 + bx + c >0 (<0) với mọi x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_06]-M3. Tìm m để ax^2 + bx + c >=0 (<=0) với mọi x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_07]-M3. Tìm m để ax^2 + bx + c >=0 (<=0) vô nghiệm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_08]-M3. Tìm m để ax^2 + bx + c >0 (<0) vô nghiệm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B2, ["[D10_C7_B2_09]-M2. Tìm tập xác định của y=căn(ax^2+bx+c)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C7_B3 = QTreeWidgetItem(L10_C7, ["Bài 3 - Phương trình quy về bậc 2"])
                L10_C7_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C7_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B3, ["[D10_C7_B3_01]-M2. Giải PT căn(ax^2 + bx^2 + c)=căn(dx^2 + ex + f)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C7_B3, ["[D10_C7_B3_02]-M2. Giải PT căn(ax^2 + bx^2 + c)=dx+e."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 10 - Chương 8
                L10_C8 = QTreeWidgetItem(L10, ["Chương 8 - Đại số tổ hợp"])
                L10_C8.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 
                L10_C8_B1 = QTreeWidgetItem(L10_C8, ["Bài 1 - Quy tắc cộng và quy tắc nhân"])
                L10_C8_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B1_1 = QTreeWidgetItem(L10_C8_B1, ["8.1.1 Quy tắc cộng"])
                L10_C8_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_1, ["[D10_C8_B1_01]-M1. Cho 2 nhóm đồ vật. Tìm số cách chọn 1 vật."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_1, ["[D10_C8_B1_02]-M1. Cho số lượng học sinh. Chọn 1 bạn để giữ chức vụ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_1, ["[D10_C8_B1_03]-M1. Chọn một thức uống từ các loại nước."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_1, ["[D10_C8_B1_04]-M1. Chọn một địa điểm đi du lịch."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B1_2 = QTreeWidgetItem(L10_C8_B1, ["8.1.2 Quy tắc nhân"])
                L10_C8_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_05]-M1. Cho 2 nhóm đồ vật. Tìm số cách chọn 2 vật."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_06]-M1. Cho số lượng học sinh. Chọn 2 bạn để giữ chức vụ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_07]-M1. Tìm số đường đi từ A đến B và từ B đến C."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_08]-M2. Từ n chữ số lập số có k chữ số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_09]-M2. Từ n chữ số lập số có k chữ số khác nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_10]-M2. Từ n chữ số có chứa số 0 lập số có k chữ số khác nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B1_2, ["[D10_C8_B1_11]-M3. Chọn 2 quả cầu khác màu và có tổng là chẵn hoặc lẻ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                #Bài 2
                L10_C8_B2 = QTreeWidgetItem(L10_C8, ["Bài 2 - Hoán vị, chỉnh hợp, tổ hợp"])
                L10_C8_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #8.2.1
                L10_C8_B2_1 = QTreeWidgetItem(L10_C8_B2, ["8.2.1. Hoán vị"])
                L10_C8_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_1, ["[D10_C8_B2_01]-M1. Xếp n bạn vào một hàng ngang(dọc)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_1, ["[D10_C8_B2_07]-M2. Xếp k nhóm đồ vật vào một hàng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_1, ["[D10_C8_B2_12]-M2. Xếp k người vào một hàng có 2 bạn đứng ở 2 đầu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_1, ["[D10_C8_B2_13]-M2. Xếp k người vào một hàng có 2 bạn đứng cạnh nhau(không cạnh nhau)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #8.2.1
                L10_C8_B2_2 = QTreeWidgetItem(L10_C8_B2, ["8.2.2. Chỉnh hợp"])
                L10_C8_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_2, ["[D10_C8_B2_02]-M1. Chọn k vật từ n đồ vật khác nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_2, ["[D10_C8_B2_03]-M1. Chọn k người từ n người và xếp chức vụ trong lớp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_2, ["[D10_C8_B2_04]-M2. Chọn k người từ n người phụ trách n công việc khác nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_2, ["[D10_C8_B2_05]-M2. Lập số có k chữ số khác nhau từ n chữ số (không chứa số 0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_2, ["[D10_C8_B2_06]-M2. Lập số có k chữ số khác nhau từ n chữ số (chứa số 0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #8.3.1
                L10_C8_B2_3 = QTreeWidgetItem(L10_C8_B2, ["8.2.3. Tổ hợp"])
                L10_C8_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B2_3_1 = QTreeWidgetItem(L10_C8_B2_3, ["1. Trắc Nghiệm"])
                L10_C8_B2_3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2_3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_09]-M2. Chọn k đối tượng tự n đối tượng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_10]-M2. Chọn k đối tượng tự từ 2 nhóm đối tượng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_11]-M2. Chọn k đối tượng tự từ 2 nhóm đối tượng thỏa mãn điều kiện."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_14]-M2. Số đoạn thẳng tạo bởi n điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_15]-M2. Số tam giác tạo bởi n điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_16]-M1. Số tập hợp con chứa k phần tử của n phần tử."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_17]-M3. Số tập hợp con có số phần tử < hơn k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_18]-M2. Số giao điểm của n đường thẳng phân biệt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_19]-M3. Số hình chữ nhật tạo từ n đường // và m đường vg."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_20]-M2. Số giao điểm của n đường tròn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_1, ["[D10_C8_B2_21]-M3. Số giao điểm của m đường thẳng và n đường tròn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B2_3_2 = QTreeWidgetItem(L10_C8_B2_3, ["2. Đúng-Sai"])
                L10_C8_B2_3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B2_3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_08]-TF-M2. Bài toán chọn 2 nhóm đối tượng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_22]-TF-M2. Bài toán chọn 2 cuốn sách."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_23]-TF-M2. Bài toán chọn 2 cuốn truyện."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_24]-TF-M2. Bài toán chọn 2 bức tranh."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_25]-TF-M2. Bài toán chọn 2 viên bi."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B2_3_2, ["[D10_C8_B2_26]-TF-M2. Bài toán chọn 2 người."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 3
                L10_C8_B3 = QTreeWidgetItem(L10_C8, ["Bài 3 - Nhị thức Niutơn"])
                L10_C8_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B3_1 = QTreeWidgetItem(L10_C8_B3, ["8.3.1. Đúng-Sai"])
                L10_C8_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C8_B3_2 = QTreeWidgetItem(L10_C8_B3, ["8.3.2. Trắc Nghiệm"])
                L10_C8_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C8_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_01]-M2. Viết khai triển của (x+a)^n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_01]-M2. Viết khai triển của (a-x)^n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_03]-M2. Viết khai triển của (ax+b)^n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_06]-M2. Tìm hệ số của x^k của (ax+b)^n, n=4,5."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_07]-M2. Tìm số hạng chứa x^k của (ax+b)^n, n=4,5."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_04]-M2. Tìm hệ số của x^k của (ax+b)^n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_05]-M2. Tìm số hạng chứa x^k của (ax+b)^n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C8_B3_2, ["[D10_C8_B3_08]-M2. Tìm số hạng thứ k của (ax+b)^n khai triển giảm dần."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
               

        #Lớp 10 - Chương 9 - Xác suất
                L10_C9 = QTreeWidgetItem(L10, ["Chương 9 - Xác Suất"])
                L10_C9.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #Bài 1 
                L10_C9_B1 = QTreeWidgetItem(L10_C9, ["Bài 1 - Không gian mẫu và biến cố"])
                L10_C9_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                L10_C9_B1_1 = QTreeWidgetItem(L10_C9_B1, ["Trắc Nghiệm"])
                L10_C9_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_01]-M1. Chọn 1 vật từ 2 nhóm đồ vật. Tính số phần tử không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_02]-M1. Chọn  2 vật từ 2 nhóm đồ vật. Tínhsố phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_03]-M2. Chọn k quả cầu khác màu. Tính số phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_04]-M2. Gieo con xúc xắc k lần. Tính số phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_05]-M2. Gieo đồng xu k lần. Tính số phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_06]-M2. Chọn k vật từ n vật. Tính số phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_1, ["[D10_C9_B1_07]-M2. Chọn k đối tượng từ 2 nhóm đối tượng. Tính số phần tử của không gian mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C9_B1_2 = QTreeWidgetItem(L10_C9_B1, ["Đúng-Sai"])
                L10_C9_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_2, ["[D10_C9_B1_08]-TF-M2. Chọn k vật từ 2 nhóm. Xét Đ-S: không gian mẫu, k vật được chọn, cả 2 loại được chọn, ít nhất 1 được chọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_2, ["[D10_C9_B1_09]-TF-M2. Gieo xúc sắc 2 lần. Xét Đ-S: không gian mẫu, biến cố liên quan tổng, tích số chấm 2 lần gieo."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_2, ["[D10_C9_B1_10]-TF-M3. Chọn k vật từ 3 nhóm. Xét Đ-S: không gian mẫu, biến cố ít nhất một, có đúng m được chọn, chỉ một loại được chọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B1_2, ["[D10_C9_B1_11]-TF-M3. Có n tấm thẻ, lấy ngẫu nhiên k thẻ. Đúng-Sai: không gian mẫu, các thẻ đều chẵn (lẻ), ít nhất một thẻ chia hết cho m, m thẻ chẵn (lẻ) được chọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C9_B2 = QTreeWidgetItem(L10_C9, ["Bài 2 - Xác suất của biến cố"])
                L10_C9_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C9_B2_1 = QTreeWidgetItem(L10_C9_B2, ["Trắc Nghiệm"])
                L10_C9_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_01]-M2. Chọn n vật từ 2 nhóm. Xác suất có đúng k vật được chọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_02]-M3. Chọn n vật từ 2 nhóm. Xác suất cả 2 nhóm được chọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_05]-M2. Cho một hộp chứa n tấm thẻ đánh số. Xác suất thẻ được chọn ghi số thuộc [a;b]."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_06]-M2. Gieo xúc xắc 2 lần. Xác suất để  i + j>k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_07]-M2. Gieo xúc xắc 2 lần. Xác suất  để  i + j<k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_08]-M2. Gieo xúc xắc 2 lần. Xác suất  để  i + j=k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_03]-M2. Chọn một quả cầu từ n quả cầu được đánh số. Xác suất quả là số chẵn(lẻ)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_04]-M3. Chọn một quả cầu từ n quả cầu được đánh số. Xác suất quả là số chẵn (lẻ) và chia hết cho k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_1, ["[D10_C9_B2_09]-M2. Gieo xúc xắc 2 lần. Xác suất: i + j = k và i.j <(>) m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C9_B2_2 = QTreeWidgetItem(L10_C9_B2, ["Đúng-Sai"])
                L10_C9_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C9_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_2, ["[D10_C9_B1_12]-TF-M3. Chọn 2 nhóm đồ vật. Đúng-sai: không gian mẫu, xác suất"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_2, ["[D10_C9_B1_13]-TF-M3. Gieo xúc sắc 2 lần. Đúng-Sai: không gian mẫu, xác suất."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C9_B2_2, ["[D10_C9_B1_14]-TF-M3. Chọn k vật từ 3 nhóm. Xét Đ-S: không gian mẫu, xác suất."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 10 - Chương 10 - Phương pháp tọa độ trong mặt phẳng 
                L10_C10 = QTreeWidgetItem(L10, ["Chương 10 - Phương pháp tọa độ trong mặt phẳng"])
                L10_C10.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B0 = QTreeWidgetItem(L10_C10, ["Bài 1 - Tọa độ vectơ"])
                L10_C10_B0.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B0.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B0_1 = QTreeWidgetItem(L10_C10_B0, ["10.1.1. Tọa độ vectơ"])
                L10_C10_B0_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B0_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_1, ["[D10_CX_B0_01]-M1. Cho 2 điểm. Tìm tọa độ vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_1, ["[D10_CX_B0_02]-M1. Cho 2 vectơ. Tìm tọa độ vectơ tổng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_1, ["[D10_CX_B0_03]-M1. Cho 2 vectơ. Tìm tọa độ vectơ hiệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L10_C10_B0_1, ["[D10_CX_B0_05]-M1. Cho 1 vectơ. Tính độ dài."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_1, ["[D10_CX_B0_11]-M2. Cho hai véctơ. Tìm m để 2 vectơ cùng phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                L10_C10_B0_2 = QTreeWidgetItem(L10_C10_B0, ["10.1.2. Tọa độ điểm"])
                L10_C10_B0_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B0_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_06]-M1. Cho 2 điểm. Tính độ dài đoạn thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_07]-M1. Cho hai điểm. Tìm tọa độ trung điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_08]-M1. Cho tam giác. Tìm tọa độ trọng tâm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_09]-M2. Cho hai điểm A,B. Tìm C để AC nhận B làm trung điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_10]-M2. Cho ba điểm A,B,G. Tìm tọa độ C để ABC nhận G làm trọng tâm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_2, ["[D10_CX_B0_13]-M3. Tìm A thuộc Ox cách B một khoảng cho trước."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B0_3 = QTreeWidgetItem(L10_C10_B0, ["10.1.2. Biểu thức tọa độ tích vô hướng"])
                L10_C10_B0_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B0_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_3, ["[D10_CX_B0_04]-M1. Cho 2 vectơ. Cho hai véctơ tính tích vô hướng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B0_3, ["[D10_CX_B0_12]-M2. Cho hai véctơ. Tìm m để 2 vectơ vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 
                L10_C10_B1 = QTreeWidgetItem(L10_C10, ["Bài 2 - Phương trình đường thẳng"])
                L10_C10_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B1_1 = QTreeWidgetItem(L10_C10_B1, ["10.2.1. Véctơ chỉ phương, véctơ pháp tuyến, điểm thuộc đường thẳng"])
                L10_C10_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_01]-M1. Cho VTPT tìm VTCP và ngược lại."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_02]-M2. Cho phương trình tổng quát tìm VTPT(VTCP)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_03]-M2. Cho phương trình tham số tìm VTPT(VTCP)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_17]-M2. Tìm điểm thuộc đường thẳng biết PTTS."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_18]-M2. Tìm điểm thuộc đường thẳng biết PTTQ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_1, ["[D10_CX_B1_31]-M3. Tìm điểm A thuộc đường thẳng cách điểm B một khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #10.1.2 Phương trình tham số

                L10_C10_B1_2 = QTreeWidgetItem(L10_C10_B1, ["10.2.2. Phương trình tham số"])
                L10_C10_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_2, ["[D10_CX_B1_12]-M2. Lập PTTS của d qua điểm và có VT chỉ phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_2, ["[D10_CX_B1_13]-M2. Lập PTTS của d qua điểm và có VT pháp tuyến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_2, ["[D10_CX_B1_14]-M2. Lập PTTS của d qua 2 điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_2, ["[D10_CX_B1_15]-M2. Lập PTTS của d từ PTTQ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L10_C10_B1_3 = QTreeWidgetItem(L10_C10_B1, ["10.2.3. Phương trình tổng quát"])
                L10_C10_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_04]-M1. Lập PTTQ của d qua điểm và có VT pháp tuyến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_05]-M2. Lập PTTQ của d qua điểm và có VT chỉ phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_16]-M2. Lập PTTQ của d từ PTTS."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_06]-M2. Lập PTTQ của d qua điểm và song song với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_07]-M2. Lập PTTQ của d qua điểm và vuông góc với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_08]-M2. Lập PTTQ của d qua 2 điểm A và B."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_09]-M2. Lập PTTQ của đường cao trong tam giác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_10]-M2. Lập PTTQ của đường trung trực trong tam giác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_3, ["[D10_CX_B1_11]-M2. Lập PTTQ của đường trung tuyến trong tam giác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                L10_C10_B1_4 = QTreeWidgetItem(L10_C10_B1, ["10.2.4. Vị trí tương đối"])
                L10_C10_B1_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_4, ["[D10_CX_B1_19]-M2. Cho PTTQ của 2 đường thẳng, xét vị trí tương đối."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_4, ["[D10_CX_B1_25]-M3. Cho PTTS của 2 đường thẳng, xét vị trí tương đối."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                
                L10_C10_B1_5 = QTreeWidgetItem(L10_C10_B1, ["10.2.5. Góc giữa 2 đường thẳng"])
                L10_C10_B1_5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_5, ["[D10_CX_B1_20]-M2. Cho PTTQ của 2 đường thẳng, tính cosin của góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_5, ["[D10_CX_B1_23]-M2. Cho PTTS của 2 đường thẳng, tính cosin của góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_5, ["[D10_CX_B1_24]-M2. Cho PTTS và PTTQ của 2 đường thẳng, tính cosin của góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L10_C10_B1_6 = QTreeWidgetItem(L10_C10_B1, ["10.2.6. Khoảng cách từ một điểm đến đường thẳng"])
                L10_C10_B1_6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_6.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_6, ["[D10_CX_B1_21]-M1. Cho điểm và PTTQ của đường thẳng, tính khoảng cách."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_6, ["[D10_CX_B1_22]-M2. Cho điểm và PTTS của đường thẳng, tính khoảng cách."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B1_7 = QTreeWidgetItem(L10_C10_B1, ["10.2.7. Đúng-Sai"])
                L10_C10_B1_7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B1_7.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_7, ["[D10_CX_B1_26]-TF-M2. Tạo câu Đ-S: Cho PTTQ, xét đúng-sai về VTCP, VTPT, điểm thuộc đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_7, ["[D10_CX_B1_27]-TF-M2. Tạo câu Đ-S: Cho PTTS, xét đúng-sai về VTCP, VTPT, điểm thuộc đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_7, ["[D10_CX_B1_28]-TF-M2. Tạo câu Đ-S: Cho PTTS, xét đúng-sai về VTCP, VTPT, PTTS, PTTQ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_7, ["[D10_CX_B1_29]-TF-M2. Tạo câu Đ-S: Cho PTTQ, xét đúng-sai về vị trí tương đối của 3 đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B1_7, ["[D10_CX_B1_30]-TF-M2. Tạo câu đúng-sai: Cho 1 điểm và PTTQ , xét đúng-sai về vị trí, khoảng cách."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 3
                L10_C10_B3 = QTreeWidgetItem(L10_C10, ["Bài 3 - Phương trình đường tròn"])
                L10_C10_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                L10_C10_B3_1 = QTreeWidgetItem(L10_C10_B3, ["1. Trắc Nghiệm"])
                L10_C10_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B3_1_1 = QTreeWidgetItem(L10_C10_B3_1, ["1.1. Viết phương trình đường tròn"])
                L10_C10_B3_1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3_1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_01]-M1. Viết phương trình đường tròn có tâm và bán kính."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_02]-M1. Đọc tọa độ tâm từ PT đường tròn thu gọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_03]-M1. Đọc bán kính từ PT đường tròn thu gọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_04]-M2. Đọc tọa độ tâm từ PT đường tròn khai triển."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_05]-M2. Đọc bán kính từ PT đường tròn khai triển."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_06]-M2. Viết phương trình đường tròn có tâm và đường kính."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_07]-M3. Viết phương trình đường tròn có đường kính AB."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_08]-M3. Viết phương trình đường tròn đi qua 3 điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_14]-M2. Viết phương trình đường tròn có tâm và đi qua điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_15]-M2. Viết phương trình đường tròn có tâm và tiếp xúc đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_1, ["[D10_CX_B3_16]-M2. Viết phương trình đường tròn có tâm và tiếp xúc trục Ox(Oy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                L10_C10_B3_1_2 = QTreeWidgetItem(L10_C10_B3_1, ["1.2. Tiếp tuyến của đường tròn"])
                L10_C10_B3_1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3_1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_2, ["[D10_CX_B3_09]-M2. Tiếp tuyến tại điểm (x_0;y_0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_2, ["[D10_CX_B3_10]-M3. Tiếp tuyến song song với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_2, ["[D10_CX_B3_11]-M3. Viết tiếp tuyến của đường tròn vuông góc với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B3_1_3 = QTreeWidgetItem(L10_C10_B3_1, ["1.3. Các bài toán khác"])
                L10_C10_B3_1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3_1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_1_3, ["[D10_CX_B3_17]-M2. Xét vị trí tương đối giữa điểm và đường tròn có phương trình thu gọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B3_2 = QTreeWidgetItem(L10_C10_B3, ["2. Đúng-Sai"])
                L10_C10_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_2, ["[D10_CX_B3_12]-TF-M2. Cho phương trình đường tròn dạng thu gọn. Tạo Đúng-Sai."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B3_2, ["[D10_CX_B3_13]-TF-M2. Cho phương trình đường tròn dạng khai triển. Tạo Đúng-Sai."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                #Bài 3 
                L10_C10_B4 = QTreeWidgetItem(L10_C10, ["Bài 4 - Phương trình 3 đường conic"])
                L10_C10_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_1 = QTreeWidgetItem(L10_C10_B4, ["4.1. Phương trình đường elip"])
                L10_C10_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_1_1 = QTreeWidgetItem(L10_C10_B4_1, ["4.1.1. Đọc các thông tin của đường elip"])
                L10_C10_B4_1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_1, ["[D10_CX_B4_01]-M2. Cho phương trình elip. Đọc độ dài trục lớn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_1, ["[D10_CX_B4_02]-M2. Cho phương trình elip. Đọc độ dài trục nhỏ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_1, ["[D10_CX_B4_03]-M2. Cho phương trình elip. Đọc độ dài tiêu cự."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_1, ["[D10_CX_B4_04]-M2. Cho phương trình elip. Đọc tọa độ tiêu điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_1, ["[D10_CX_B4_05]-M2. Cho phương trình elip. Đọc tọa độ đỉnh."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_1_2 = QTreeWidgetItem(L10_C10_B4_1, ["4.1.1. Viết phương trình đường elip"])
                L10_C10_B4_1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_2, ["[D10_CX_B4_06]-M2. Cho trục lớn, trục nhỏ. Viết phương trình Elip."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_2, ["[D10_CX_B4_07]-M2. Cho trục lớn, tiêu cự. Viết phương trình Elip."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_2, ["[D10_CX_B4_08]-M2. Cho trục nhỏ, tiêu cự. Viết phương trình Elip."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_2, ["[D10_CX_B4_09]-M2. Cho đỉnh thuộc 2 trục. Viết phương trình Elip."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_1_2, ["[D10_CX_B4_10]-M2. Cho đỉnh thuộc trục lớn và tiêu điểm. Viết phương trình Elip."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)





                L10_C10_B4_2 = QTreeWidgetItem(L10_C10_B4, ["4.2 - Phương trình đường hypebol"])
                L10_C10_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_2_1 = QTreeWidgetItem(L10_C10_B4_2, ["4.2.1. Đọc các thông tin của đường hypebol"])
                L10_C10_B4_2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_1, ["[D10_CX_B4_11]-M2. Cho phương trình Hypebol. Tìm độ dài trục thực."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_1, ["[D10_CX_B4_12]-M2. Cho phương trình Hypebol. Tìm độ dài trục ảo."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_1, ["[D10_CX_B4_13]-M2. Cho phương trình Hypebol. Tìm độ dài tiêu cự."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_1, ["[D10_CX_B4_14]-M2. Cho phương trình Hypebol. Đọc tọa độ tiêu điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_1, ["[D10_CX_B4_15]-M2. Cho phương trình Hypebol. Đọc tọa độ đỉnh."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_2_2 = QTreeWidgetItem(L10_C10_B4_2, ["4.2.2. Viết phương trình đường hypebol"])
                L10_C10_B4_2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_2, ["[D10_CX_B4_16]-M2. Cho trục thực, trục ảo. Viết phương trình Hypebol."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_2, ["[D10_CX_B4_17]-M2. Cho trục thực, tiêu cự. Viết phương trình Hypebol."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_2, ["[D10_CX_B4_18]-M2. Cho trục ảo, tiêu cự. Viết phương trình Hypebol."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_2_2, ["[D10_CX_B4_19]-M2. Cho đỉnh thuộc trục thực, tiêu điểm. Viết phương trình Hypebol."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L10_C10_B4_3 = QTreeWidgetItem(L10_C10_B4, ["4.3 - Phương trình đường parabol"])
                L10_C10_B4_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L10_C10_B4_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_20]-M2. Cho parabol. Tìm tham số tiêu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_21]-M2. Cho parabol. Tìm đường chuẩn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_22]-M2. Cho parabol. Tìm tiêu điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_23]-M2. Viết phương trình parabol có tham số tiêu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_24]-M2. Viết phương trình parabol có tiêu điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L10_C10_B4_3, ["[D10_CX_B4_25]-M2. Viết phương trình parabol có đường chuẩn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                L10_C10_B4_4 = QTreeWidgetItem(L10_C10_B4, ["4.4 - Câu hỏi đúng sai"])
                L10_C10_B4_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                

        #Tạo cây lớp 11
                L11 = QTreeWidgetItem(self.treeWidget, ["Lớp 11"])
                L11.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11.setCheckState(0, Qt.CheckState.PartiallyChecked)
        #Lớp 11 - Chương 1
                L11_C1 = QTreeWidgetItem(L11, ["Chương 1 - Phương trình lượng giác"])
                L11_C1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Lớp 11 - Chương1 - Bài 1 - Góc lượng giác
                L11_C1_B1 = QTreeWidgetItem(L11_C1, ["Bài 1-Góc lượng giác"])
                L11_C1_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B1_1 = QTreeWidgetItem(L11_C1_B1, ["Trắc nghiệm-Trả lời ngắn"])
                L11_C1_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_02]-M1. Đổi số đo từ độ sang radian"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_03]-M1. Đổi số đo từ radian sang độ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_04]-M2. Tìm góc có điểm biểu diễn trùng nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_05]-M2. Tìm góc lượng giác có điểm biểu diễn cho trước"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_06]-M2. Cho bán kính và góc radian. Tính độ dài của cung."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_1, ["[D11_C1_B1_07]-M2. Cho bán kính và góc độ. Tính độ dài của cung."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B1_2 = QTreeWidgetItem(L11_C1_B1, ["Đúng-Sai"])
                L11_C1_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_2, ["[D11_C1_B1_08]-TF-M2. Cho góc radian. Xét Đ-S: Đổi sang độ, Điểm biểu diễn thuộc phần tư, Góc cùng điểm biểu diễn, Đếm số điểm biểu diễn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_2, ["[D11_C1_B1_09]-TF-M2. Cho góc độ. Xét Đ-S: Đổi sang radian, Điểm biểu diễn thuộc phần tư, Góc cùng điểm biểu diễn, Đếm số điểm biểu diễn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B1_3 = QTreeWidgetItem(L11_C1_B1, ["Trả lời ngắn"])
                L11_C1_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_3, ["[D11_C1_B1_10]-TL-M2. Cho số vòng quay bánh xe sau t1 giây. Tính góc radian sau khi quay trong t2 giây"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_3, ["[D11_C1_B1_11]-TL-M3. Cho số vòng quay bánh xe sau t1 giây. Tính quãng đường đi được sau t2 giây"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B1_3, ["[D11_C1_B1_12]-TL-M3. Cho đường kính của bánh trước và bánh sau, vận tốc n vòng/phút. Tính quãng đường xe đi."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                
                
                #Lớp 11 - Chương1 - Bài 2 - Giá trị lượng giác của một góc lượng giác
                L11_C1_B2 = QTreeWidgetItem(L11_C1, ["Bài 2-Giá trị lượng giác của một góc lượng giác"])
                L11_C1_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B2_1 = QTreeWidgetItem(L11_C1_B2, ["Trắc nghiệm-Trả lời ngắn"])
                L11_C1_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_1, ["[D11_C1_B2_01]-M1. Tính giá trị đặc biệt của một góc lượng giác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B2_2 = QTreeWidgetItem(L11_C1_B2, ["Đúng-Sai"])
                L11_C1_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_2, ["[D11_C1_B2_06]-TF-M2. Cho góc lượng giác thuộc cung phần tư. Xét Đ-S dấu của sin, cos, tan, cot."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_2, ["[D11_C1_B2_07]-TF-M3. Cho tanx. Xét Đ-S: cotx, sin^2 x, cos^2x, (asinx+bcosx)/(csinx+dcosx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_2, ["[D11_C1_B2_08]-TF-M3. Cho sinx (a<x<b). Xét Đ-S: dấu của cosx, cosx, sin(x+kpi/2), P=f(tanx)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_2, ["[D11_C1_B2_09]-TF-M3. Cho cosx (a<x<b). Xét Đ-S: dấu của sinx, sinx, sin(x+kpi/2), P=f(tanx)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B2_3 = QTreeWidgetItem(L11_C1_B2, ["Trả lời ngắn"])
                L11_C1_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_3, ["[D11_C1_B2_02]-TL-M2. Cho sinx (hoặc cosx), x thuộc (a;b). Tìm cosx (hoặc sinx)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_3, ["[D11_C1_B2_03]-TL-M3. Cho sinx (hoặc cosx), x thuộc (a;b). Tìm tanx (hoặc cotx)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_3, ["[D11_C1_B2_04]-TL-M3. Cho tanx (hoặc cotx). Tìm P=(asinx+bcosx)/(csinx+dcosx)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B2_3, ["[D11_C1_B2_05]-TL-M3. Cho tanx (hoặc cotx). Tìm P=(asin^2 x+bsinxcosx)/(csin^2 x+dcos^2 x)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Lớp 11 - Chương1 - Bài 3 - Các công thức lượng giác
                L11_C1_B3 = QTreeWidgetItem(L11_C1, ["Bài 3-Các công thức lượng giác"])
                L11_C1_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B3_1 = QTreeWidgetItem(L11_C1_B3, ["Trắc nghiệm"])
                L11_C1_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_03]-M2. Tìm khẳng định đúng về hai góc đối nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_04]-M2. Tìm khẳng định đúng về hai góc bù nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_05]-M2. Tìm khẳng định đúng về hai góc phụ nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_06]-M2. Tìm khẳng định đúng về hai góc hơn kém pi"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_07]-M2. Tìm khẳng định đúng về hai góc liên quan tùy ý"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_08]-M2. Tìm khẳng định đúng về công thức nhân đôi"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_09]-M2. Tìm khẳng định đúng về công thức cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_10]-M2. Tìm khẳng định đúng về công thức tích thành tổng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_11]-M2. Tìm khẳng định đúng về công thức tổng thành tích"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
             
                item= QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_01]-M1. Cho sina, cosa. Tính sin2a."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
             
                item = QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_02]-M2. Cho sina hoặc cosa. Tính cos2a."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B3_1, ["[D11_C1_B3_14]-M3. Cho sinx. Tính sin(x+a) hoặc cos(x+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B3_3 = QTreeWidgetItem(L11_C1_B3, ["Đúng-Sai"])
                L11_C1_B3_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B3_3.setCheckState(0, Qt.CheckState.PartiallyChecked)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B3_3, ["[D11_C1_B3_12]-TF-M2. Cho tanx. Xét Đ-S: cotx, cos2x, sin2x, tan2x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B3_3, ["[D11_C1_B3_13]-TF-M2. Cho sinx. Xét Đ-S: cosx, sin2x, sin(x+a), cos(x+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Lớp 11 - Chương1 - Bài 4 - Hàm số lượng giác và đồ thị
                L11_C1_B4 = QTreeWidgetItem(L11_C1, ["Bài 4-Hàm số lượng giác và đồ thị"])
                L11_C1_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B4_1 = QTreeWidgetItem(L11_C1_B4, ["Trắc nghiệm"])
                L11_C1_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_01]-M2. Tìm tập xác định hàm số y=a/sinx hoặc y=a/cosx"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_02]-M2. Tìm tập xác định hàm số y=a/sin(mx) hoặc y=a/cos(mx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_03]-M2. Tìm tập xác định hàm số y=tan(ax+bpi)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_04]-M2. Tìm GTLN,GTNN của hàm số sin, cos"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_08]-M2. Tìm tập giá trị của HSLG"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_09]-M2. Tìm chu kì tuần hoàn của HSLG"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B4_1 = QTreeWidgetItem(L11_C1_B4, ["Đúng-Sai"])
                L11_C1_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_05]-TF-M2. Cho y=acosu+b hoặc y=asinu+b. Xét Đ-S: TXĐ, GTLN, GTNN, TGT"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_06]-TF-M2. Cho y=acos^2u+b hoặc y=asinu+b. Xét Đ-S: TXĐ, GTLN, GTNN, TGT"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B4_1, ["[D11_C1_B4_07]-TF-M2. Cho y=acosmx+b hoặc y=asinmx+b. Xét Đ-S: TXĐ, Chẵn-Lẻ, TGT, Cắt trục Oy"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
              

                #Lớp 11 - Chương1 - Bài 5 - Hàm số lượng giác và đồ thị
                L11_C1_B5 = QTreeWidgetItem(L11_C1, ["Bài 5-Phương trình lượng giác cơ bản"])
                L11_C1_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C1_B5_1 = QTreeWidgetItem(L11_C1_B5, ["Trắc nghiệm"])
                L11_C1_B5_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B5_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_05]-M1. Giải phương trình cos(x) = b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)           

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_01]-M2. Giải phương trình cos(ax) = b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_06]-M1. Giải phương trình sin(x) = b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)        

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_02]-M2. Giải phương trình sin(ax) = b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_03]-M2. Giải phương trình tan(ax) = m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_07]-M2. Giải phương trình cot(ax) = m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_08]-M3. Giải phương trình cos(ax+m)=cos(bx+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_09]-M3. Giải phương trình cos(ax+m)=sin(bx+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_10]-M3. Giải phương trình sin(ax+m)=sin(bx+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_11]-M2. Giải phương trình tan(ax+m)=tan(bx+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_12]-M2. Giải phương trình cot(ax+m)=cot(bx+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
    
                item = QTreeWidgetItem(L11_C1_B5_1, ["[D11_C1_B5_04]-M3. Tìm m để phương trình sin, cos có nghiệm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                L11_C1_B5_3 = QTreeWidgetItem(L11_C1_B5, ["Trả lời ngắn"])
                L11_C1_B5_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C1_B5_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_3, ["[D11_C1_B5_13]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của sinu=m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_3, ["[D11_C1_B5_14]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của cosu=m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C1_B5_3, ["[D11_C1_B5_15]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của tanu=m"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



        #Lớp 11 - Chương 2 - Dãy số. Cấp số cộng. Cấp số nhân
                L11_C2 = QTreeWidgetItem(L11, ["Chương 2 - Dãy số. Cấp số cộng. Cấp số nhân"])
                L11_C2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B1 = QTreeWidgetItem(L11_C2, ["Bài 1-Dãy số"])
                L11_C2_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B1_1 = QTreeWidgetItem(L11_C2_B1, ["Trắc nghiệm"])
                L11_C2_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B1_1, ["[D11_C2_B1_02]-M1. Cho dãy số. Tìm 3 số đầu tiên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B1_1, ["[D11_C2_B1_01]-M2. Cho dãy số có SHTQ. Hỏi số là số hạng thứ mấy"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B1_1, ["[D11_C2_B1_03]-M2. Cho dãy số dạng truy hồi. Tìm SHTQ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B1_1, ["[D11_C2_B1_04]-M2. Cho dãy số có SHQT xét tính tăng giảm, bị chặn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                


                L11_C2_B2 = QTreeWidgetItem(L11_C2, ["Bài 2-Cấp số cộng"])
                L11_C2_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B2_1 = QTreeWidgetItem(L11_C2_B2, ["Trắc nghiệm"])
                L11_C2_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_01]-M1. Cấp số cộng có u_1 và d. Tìm u_k"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_02]-M2. Cấp số cộng có u_1 và d. Tìm S_n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
             

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_03]-M2. Cấp số cộng có u_m và u_k. Tìm u_1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)               

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_04]-M2. Cấp số cộng có u_m và u_k. Tìm d."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_06]-M2. Cấp số cộng có SHTQ. Tìm d."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_07]-M2. Cho các SHTQ. Tìm dãy số là cấp số cộng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_1, ["[D11_C2_B2_08]-M2. Nhận dạng dãy các số hữu hạn là cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B2_2 = QTreeWidgetItem(L11_C2_B2, ["Đúng-Sai"])
                L11_C2_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_2, ["[D11_C2_B2_05]-TF-M2. CSC có u1,d. Đ-S: u_k, số hạng thứ mấy, S_k, u_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_2, ["[D11_C2_B2_09]-TF-M2. CSC có u_m,u_n. Xét Đ-S: u_1,d, số hạng thứ mấy, S_k, u_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_2, ["[D11_C2_B2_10]-TF-M2. CSC có u_1,S_n. Xét Đ-S: u_1,d, u_k, S_k, u_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_2, ["[D11_C2_B2_11]-TF-M2. Xét đúng-sai từ bài toán thực tế liên quan cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B2_3 = QTreeWidgetItem(L11_C2_B2, ["Trả lời ngắn"])
                L11_C2_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_12]-TL-M2. CSC có u_k,u_m. Tính u_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_13]-TL-M3. CSC có u_k,u_m. Tính S_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_14]-TL-M3. CSC có u_m+u_n=S. Tính u_p+u_q"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_15]-TL-M2. Tính giá của một cái máy sau n năm sử dụng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_16]-TL-M3. Bài toán trồng cây theo cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_18]-TL-M2. Tính số vật khi xếp chồng và giảm dần theo cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_17]-TL-M3. Tính số lượng theo cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_19]-TL-M3. Tính số hàng khi xếp đồ vật và giảm dần theo cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B2_3, ["[D11_C2_B2_20]-TL-M3. Tính số hàng khi xếp đồ vật và tăng dần theo cấp số cộng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                

                L11_C2_B3 = QTreeWidgetItem(L11_C2, ["Bài 3-Cấp số nhân"])
                L11_C2_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B3_1 = QTreeWidgetItem(L11_C2_B3, ["Trắc nghiệm"])
                L11_C2_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_01]-M1. Cho u_1 và q. Tìm u_k"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            


                item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_02]-M1. Cho u_n và u_(n+1). Tìm q"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_03]-M2. Cho cấp số nhân có u1, q. Tính tổng vài số hạng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_04]-M2. Cho cấp số nhân có u1, q. Tính tổng S_n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_08]-M2. Nhận dạng dãy số hữu hạn là 1 cấp số nhân."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                # item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_05]-M2. Tính số virut tạo ra sau một số khoảng thời gian."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                # item = QTreeWidgetItem(L11_C2_B3_1, ["[D11_C2_B3_06]-M2. Tính số virut tạo ra sau một số khoảng thời gian."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C2_B3_2 = QTreeWidgetItem(L11_C2_B3, ["Đúng-Sai"])
                L11_C2_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C2_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C2_B3_2, ["[D11_C2_B3_07]-TF-M2. Cho CSN có un,u_m. Xét Đ-S: q, u_k, S_k, u_n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

        #Lớp 11 - Chương 3 - Giới hạn
                L11_C3 = QTreeWidgetItem(L11, ["Chương 3 - Giới hạn. Hàm số liên tục"])
                L11_C3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 - Giới hạn dãy số
                L11_C3_B1 = QTreeWidgetItem(L11_C3, ["Bài 1 - Giới hạn dãy số"])
                L11_C3_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C3_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_07]-M1. Cho giới hạn của (u_n) và (v_n). Tính giới hạn tổng hiệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho giới hạn của (u_n) và (v_n). Tính giới hạn tổng hiệu')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_01]-M2. Giới hạn dãy số bậc tử = bậc mẫu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Giới hạn dãy số bậc tử = bậc mẫu')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_02]-M2. Giới hạn dãy số bậc tử < bậc mẫu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Giới hạn dãy số bậc tử < bậc mẫu')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_03]-M2. Tính giới hạn phân thức can(an^2+bn+c)/(dn+e)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn phân thức can(an^2+bn+c)/(dn+e)')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_04]-M2. Tính giới hạn phân thức chứa lũy thừa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn phân thức chứa lũy thừa')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_05]-M2. Tính tổng cấp số nhân lùi vô hạn có u_1=1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính tổng cấp số nhân lùi vô hạn có u_1=1')

                item = QTreeWidgetItem(L11_C3_B1, ["[D11_C3_B1_06]-M2. Tính tổng cấp số nhân lùi vô hạn có u_1 tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính tổng cấp số nhân lùi vô hạn có u_1 tùy ý ')

                #Bài 2 - Giới hạn hàm số
                L11_C3_B2 = QTreeWidgetItem(L11_C3, ["Bài 2 - Giới hạn hàm số"])
                L11_C3_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C3_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_09]-M1. Cho limf(x) và limg(x). Tính lim f(x)g(x) hoặc lim f(x)/g(x)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho limf(x) và limg(x). Tính lim f(x)g(x) hoặc lim f(x)/g(x)')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_01]-M1. Tính giới hạn tại điểm: thay số trực tiếp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn tại điểm: thay số trực tiếp')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_02]-M2. Tính giới hạn 0/0: Bậc 2/Bậc 1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn 0/0: Bậc 2/Bậc 1')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_11]-M2. Tính giới hạn 0/0: Bậc 1/Bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn 0/0: Bậc 1/Bậc 2')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_03]-M2. Tính giới hạn 0/0: Bậc 2/Bậc 2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn 0/0: Bậc 2/Bậc 2')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_04]-M3. Tính giới hạn 0/0: (Căn(A)-B)/(mx+n)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn 0/0: (Căn(A)-B)/(mx+n)')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_05]-M2. Tính giới hạn x-->oo: bậc tử = bậc mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn x-->oo: bậc tử = bậc mẫu')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_06]-M2. Tính giới hạn x-->oo: bậc tử < bậc mẫu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn x-->oo: bậc tử < bậc mẫu')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_07]-M2. Tính giới hạn x-->oo: Căn (A)/B."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn x-->oo: Căn (A)/B')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_08]-M2. Tính giới hạn x-->oo: A/Căn (B)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn x-->oo: A/Căn (B)')

                item = QTreeWidgetItem(L11_C3_B2, ["[D11_C3_B2_10]-M2. Tính giới hạn x-->oo: Đa thức."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tính giới hạn x-->oo: Đa thức')

                #Bài 3 - Hàm số liên tục
                L11_C3_B3 = QTreeWidgetItem(L11_C3, ["Bài 3 - Hàm số liên tục"])
                L11_C3_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C3_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C3_B3, ["[D11_C3_B3_01]-M2. Cho f(x)=căn(ax+b). Xét tính liên tục tại điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho f(x)=căn(ax+b). Xét tính liên tục tại điểm')

                item = QTreeWidgetItem(L11_C3_B3, ["[D11_C3_B3_02]-M2. Cho f(x)=(ax+b)/(cx+d). Tìm khoảng liên tục."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho f(x)=(ax+b)/(cx+d). Tìm khoảng liên tục')

                item = QTreeWidgetItem(L11_C3_B3, ["[D11_C3_B3_03]-M2. Cho f(x) có >=,<. Xét tính liên tục tại điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho f(x) có >=,<. Xét tính liên tục tại điểm')

                item = QTreeWidgetItem(L11_C3_B3, ["[D11_C3_B3_05]-M2. Cho f(x) có phân thức bậc 2. Tìm m để f(x) liên tục tại x_0."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho f(x) có phân thức bậc 2. Tìm m để f(x) liên tục tại x_0')

                item = QTreeWidgetItem(L11_C3_B3, ["[D11_C3_B3_04]-M2. Cho f(x) có phân thức bậc 3. Tìm m để f(x) liên tục tại x_0."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho f(x) có phân thức bậc 3. Tìm m để f(x) liên tục tại x_0')

        #Lớp 11 - Chương 4 - Quan hệ song song trong không gian
                L11_C4 = QTreeWidgetItem(L11, ["Chương 4 - Quan hệ song song trong không gian"])
                L11_C4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 1 - Điểm, đường thẳng và mặt phẳng
                L11_C4_B1 = QTreeWidgetItem(L11_C4, ["Bài 1 - Điểm, đường thẳng và mặt phẳng."])
                L11_C4_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C4_B1_1 = QTreeWidgetItem(L11_C4_B1, ["Trắc nghiệm"])
                L11_C4_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_01]-M2. Cho tứ diện. Xét quan hệ điểm-đường thẳng-mặt phẳng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_02]-M2. H.chóp đáy h.b.h. Xét quan hệ điểm-đường thẳng-mặt phẳng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_03]-M2. H.chóp đáy tứ giác. Xét quan hệ điểm-đường thẳng-mặt phẳng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_04]-M1. Cho tứ diện hoặc hình chóp tứ giác. Tìm giao tuyến giữa các mặt bên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_05]-M1. Cho tứ diện. Tìm giao tuyến của các mặt phẳng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B1_1, ["[D11_C4_B1_06]-M2. H.chóp đáy h.b.h. Tìm giao điểm đường thẳng - mặt phẳng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 2 - Hai đường thẳng song song.
                L11_C4_B2 = QTreeWidgetItem(L11_C4, ["Bài 2 - Hai đường thẳng song song."])
                L11_C4_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B2, ["[D11_C4_B2_01]-M2. Cho hình chóp-hbh. Xét hai đường thẳng song song - đường trung bình."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
          


                #Bài 3 - Đường thẳng và mặt phẳng song song.
                L11_C4_B3 = QTreeWidgetItem(L11_C4, ["Bài 3 - Đường thẳng và mặt phẳng song song."])
                L11_C4_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B3, ["[D11_C4_B3_01]-M2. Cho hình chóp đáy h.b.h. Xét sự song song của một đường thẳng với các mp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho hình chóp đáy h.b.h. Xét sự song song của một đường thẳng với các mp.')

                item = QTreeWidgetItem(L11_C4_B3, ["[D11_C4_B3_02]-M2. Cho hình chóp đáy h.thang. Xét sự song song của một đường thẳng với các mp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho hình chóp đáy h.thang. Xét sự song song của một đường thẳng với các mp.')

                #Bài 4 - Hai mặt phẳng song song.
                L11_C4_B4 = QTreeWidgetItem(L11_C4, ["Bài 4 - Hai mặt phẳng song song."])
                L11_C4_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_03]-M2. Cho lăng trụ tam giác. Xét sự song song của 2 mp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho lăng trụ tam giác. Xét sự song song của 2 mp.')

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_09]-M2. Cho lăng trụ tam giác. Xét sự song song của 2 mp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho lăng trụ tam giác. Xét sự song song của 2 mp.')

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_05]-M2. Tìm khẳng định đúng về hình lăng trụ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Tìm khẳng định đúng về hình lăng trụ.')

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_06]-M1. Tìm cặp điểm đối diện của hình hộp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Tìm cặp điểm đối diện của hình hộp.')

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_07]-M2. Cho hình hộp. Xét sự song song của một đường với các đường."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho hình hộp. Xét sự song song của một đường với các đường.')

                item = QTreeWidgetItem(L11_C4_B4, ["[D11_C4_B4_08]-M2. Cho hình hộp. Xét sự song song của hai đường tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                ##item.setToolTip(0, 'Cho hình hộp. Xét sự song song của hai đường tùy ý.')

                #Bài 5 - Phép chiếu song song.
                L11_C4_B5 = QTreeWidgetItem(L11_C4, ["Bài 5 - Phép chiếu song song."])
                L11_C4_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C4_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 11 - Chương 5 -  MẪU SỐ LIỆU GHÉP NHÓM
                L11_C5 = QTreeWidgetItem(L11, ["Chương 5 - Mẫu số liệu ghép nhóm"])
                L11_C5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C5_B1 = QTreeWidgetItem(L11_C5, ["Bài 1 - Số trung bình và Mốt của mẫu số liệu ghép nhóm"])
                L11_C5_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C5_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B1, ["[D11_C5_B1_01]-M1. Cho bảng số liệu ghép nhóm khách hàng. Tìm giá trị đại diện."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho bảng số liệu ghép nhóm khách hàng. Tìm giá trị đại diện.')

                item = QTreeWidgetItem(L11_C5_B1, ["[D11_C5_B1_02]-M1. Cho bảng số liệu ghép nhóm. Tìm giá trị đại diện."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Cho bảng số liệu ghép nhóm ngẫu nhiên. Tìm giá trị đại diện.')

                item = QTreeWidgetItem(L11_C5_B1, ["[D11_C5_B1_03]-M2. Cho bảng số liệu ghép nhóm. Tính số trung bình."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C5_B2 = QTreeWidgetItem(L11_C5, ["Bài 2 - Trung vị và Tứ phân vị của mẫu số liệu ghép nhóm"])
                L11_C5_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C5_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B2, ["[D11_C5_B2_01]-M3. Cho bảng số liệu ghép nhóm. Tính số trung vị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B2, ["[D11_C5_B2_02]-M2. Cho bảng số liệu ghép nhóm. Tính mốt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B2, ["[D11_C5_B2_03]-M2. Cho bảng số liệu ghép nhóm. Tính Q1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B2, ["[D11_C5_B2_04]-M2. Cho bảng số liệu ghép nhóm. Tính Q2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C5_B2, ["[D11_C5_B2_05]-M2. Cho bảng số liệu ghép nhóm. Tính Q3."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                



        #Lớp 11 - Chương 6 -  Hàm số mũ - Hàm số logarit      
                L11_C6 = QTreeWidgetItem(L11, ["Chương 6 - Hàm số mũ. Hàm số Lôgarit"])
                L11_C6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #Bài 1 - Phép tính lũy thừa

                L11_C6_B1 = QTreeWidgetItem(L11_C6, ["Bài 1 - Phép tính lũy thừa"])
                L11_C6_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B1_1 = QTreeWidgetItem(L11_C6_B1, ["6.1.1. Đúng-Sai"])
                L11_C6_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_1, ["[D11_C6_B1_07]-TF-M1. Tính chất của lũy thừa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_1, ["[D11_C6_B1_08]-TF-M1. Tính chất của căn bậc n."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B1_2 = QTreeWidgetItem(L11_C6_B1, ["6.1.2. Trắc Nghiệm"])
                L11_C6_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_01]-M1. Rút gọn a^m.a^n với m,n là phân số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_02]-M2. Rút gọn (a^m.a^n)/a^p với m,n,p là phân số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_03]-M1. Biểu diễn căn bậc thành lũy thừa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_04]-M2. Biểu diễn tích chứa 2 căn thành lũy thừa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_05]-M3. Biểu diễn tích chứa 3 căn thành lũy thừa."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B1_2, ["[D11_C6_B1_06]-M3. Tìm k để tích 3 căn = x^m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                L11_C6_B2 = QTreeWidgetItem(L11_C6, ["Bài 2 - Phép tính Lôgarit"])
                L11_C6_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B2_1 = QTreeWidgetItem(L11_C6_B2, ["6.2.1. Đúng-Sai"])
                L11_C6_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B2_2 = QTreeWidgetItem(L11_C6_B2, ["6.2.2. Câu hỏi lý thuyết công thức"])
                L11_C6_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_2, ["[D11_C6_B2_01]-M1. Tìm khẳng định đúng về log_a (a^m)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_2, ["[D11_C6_B2_02]-M1. Tìm khẳng định đúng về log_a (1/a^m)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_2, ["[D11_C6_B2_03]-M2. Tìm khẳng định đúng về log_can[n](a) (1/a^m)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_2, ["[D11_C6_B2_06]-M2. Tìm khẳng định đúng của log_a(a^m.b^n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B2_3 = QTreeWidgetItem(L11_C6_B2, ["6.2.3. Tính giá trị biểu thức chứa logarit"])
                L11_C6_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

                item = QTreeWidgetItem(L11_C6_B2_3, ["[D11_C6_B2_04]-M1. Tính giá trị biểu thức chứa logarit bằng máy tính"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L11_C6_B2_4 = QTreeWidgetItem(L11_C6_B2, ["6.2.4 Biểu diễn logarit theo các loagrit khác"])
                L11_C6_B2_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B2_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_4, ["[D11_C6_B2_05]-M2. Biễu diễn một logarit theo một logarit khác"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B2_4, ["[D11_C6_B2_07]-M3. Biểu diễn một logarit theo 2 logarit khác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 3
                L11_C6_B3 = QTreeWidgetItem(L11_C6, ["Bài 3 - Hàm số mũ - Hàm số Lôgarit"])
                L11_C6_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B3_1 = QTreeWidgetItem(L11_C6_B3, ["3.1.1. Đúng-Sai"])
                L11_C6_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_12]-TF-M2. Cho y=a^x, a>1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_13]-TF-M2. Cho y=a^x, 0<a<1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_14]-TF-M2. Cho y=log_a (x), a>1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_15]-TF-M2. Cho y=log_a (x), 0<a<1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_25]-TF-M2. Cho y=a^x, a>1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_22]-TF-M2. Cho đồ thị y=a^x, 0<a<1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_23]-TF-M2. Cho đồ thị y=log_a x, a>1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_1, ["[D11_C6_B3_24]-TF-M2. Cho đồ thị y=log_a x, 0<a<1. Tạo câu Đ-S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B3_2 = QTreeWidgetItem(L11_C6_B3, ["3.1.2. Trắc Nghiệm"])
                L11_C6_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B3_2_1 = QTreeWidgetItem(L11_C6_B3_2, ["3.1.2.1. Tập xác định của hàm số"])
                L11_C6_B3_2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3_2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_01]-M1. TXĐ hàm số y=(ax+b)^n với n là số nguyên âm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)               

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_02]-M1. TXĐ hàm số y=(ax+b)^n với n không nguyên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
               

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_03]-M2. TXĐ hàm số y=(ax^2+bx+c)^n với n là số nguyên âm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_04]-M2. TXĐ hàm số y=(ax^2+bx+c)^n với n không nguyên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_05]-M1. TXĐ hàm số y=log(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'TXĐ hàm số y=log(ax+b)')

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_06]-M2. TXĐ hàm số y=log(ax^2+bx+c)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_10]-M3. Tìm m để log (ax^2 +bx+c) có tập xác định là R."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_1, ["[D11_C6_B3_11]-M3. Tìm m để a^(m/căn(ax^2 +bx+c)) có tập xác định là R."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B3_2_2 = QTreeWidgetItem(L11_C6_B3_2, ["3.1.2.2. Đồ thị của hàm số"])
                L11_C6_B3_2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3_2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_07]-M2. Đồ thị hàm số y=a^x, a tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_16]-M2. Đồ thị hàm số y=a^x, a>1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_17]-M2. Đồ thị hàm số y=a^x, 0<a<1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_08]-M1. Đồ thị hàm số y=log_a(x), a tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_18]-M1. Đồ thị hàm số y=log_a(x), a>1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_19]-M1. Đồ thị hàm số y=log_a(x), 0<a<1."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_2, ["[D11_C6_B3_09]-M1. Đồ thị hàm số y=căn(a)^x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B3_2_3 = QTreeWidgetItem(L11_C6_B3_2, ["3.1.2.3. Các bài toán khác"])
                L11_C6_B3_2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B3_2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_3, ["[D11_C6_B3_20]-M2. Xét tính đơn điệu của hàm số y=a^x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B3_2_3, ["[D11_C6_B3_21]-M2. Xét tính đơn điệu của hàm số y=log_a x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                
                #Bài 4           
                
                L11_C6_B4 = QTreeWidgetItem(L11_C6, ["Bài 4 - Phương trình mũ - Phương trình Lôgarit"])
                L11_C6_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B4_1 = QTreeWidgetItem(L11_C6_B4, ["6.4.1. Đúng-Sai"])
                L11_C6_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_1, ["[D11_C6_B4_09]-TF-M2. Tạo câu Đ-S: Phương trình a^x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_1, ["[D11_C6_B4_10]-TF-M2. Tạo câu Đ-S: Phương trình log_a (x)=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B4_2 = QTreeWidgetItem(L11_C6_B4, ["6.4.2. Trắc Nghiệm"])
                L11_C6_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_01]-M1. Phương trình dạng a^x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_02]-M2. Phương trình dạng a^(x-m)=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Phương trình dạng a^(x-m)=b"')

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_03]-M2. Phương trình dạng a^(mx+n)=b^(px+q)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Phương trình dạng a^(mx+n)=b^(px+q)')

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_04]-M2. Giải phương trình bậc 2 đối với a^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_11]-M3. Giải PT: m.a^x + n.b^x = p.a^x +q.b^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_12]-M3. Giải PT: a^x + a^(x+m) +a^(x-n) = b^x + b^(x+p) + b^(x-q)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_05]-M1. Giải phương trình log_a(x)=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_06]-M2. Giải phương trình log_a(mx+n)=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_07]-M3. Giải phương trình log_m(ax+b)-log_m(cx+d)=e"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B4_2, ["[D11_C6_B4_08]-M3. Giải phương trình log_m(ax+b)+log_m(cx+d)=e"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                L11_C6_B5 = QTreeWidgetItem(L11_C6, ["Bài 5 - Bất phương trình mũ - Bất phương trình Lôgarit"])
                L11_C6_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B5_1 = QTreeWidgetItem(L11_C6_B5, ["6.5.1. Đúng-Sai"])
                L11_C6_B5_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B5_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_1, ["[D11_C6_B5_10]-TF-M2. Tạo câu Đ-S: BPT chứa a^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C6_B5_2 = QTreeWidgetItem(L11_C6_B5, ["6.5.2. Trắc Nghiệm"])
                L11_C6_B5_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B5_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_04]-M1. Giải BPT a^x >(>=) b, (a>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_05]-M1. Giải BPT a^x <(<=) b, (a>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_06]-M2. Giải BPT a^x <(<=) b, (0<a<1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_07]-M2. Giải BPT a^x >(>=) b, (0<a<1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_08]-M3. Giải BPT m^(ax+b) >(>=) m^(cx+d), (m>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_09]-M3. Giải BPT m^(ax+b) >(>=) m^(cx+d), (0<m<1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_01]-M1. Giải BPT log_a(x) > b, (a>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
            

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_02]-M2. Giải BPT log_a(x+m) > b, (a>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
            

                item = QTreeWidgetItem(L11_C6_B5_2, ["[D11_C6_B5_03]-M2. Giải BPT log_m(ax+b) > log_m(cx+d), (a>1)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                

                #BÀI 6 - BÀI TOÁN LÃI SUẤT- TĂNG TRƯỞNG 

                L11_C6_B6 = QTreeWidgetItem(L11_C6, ["Bài 6 - Bài toán lãi suất, tăng trưởng"])
                L11_C6_B6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C6_B6.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_01]-M1. Cho số tiền và lãi suất theo năm. Tính tổng tiền sau n năm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
      

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_02]-M1. Cho số tiền và lãi suất theo tháng. Tính tổng tiền sau n tháng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
            

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_03]-M2. Cho số tiền và lãi suất/năm. Tính số năm để thu được khoản tiền nào đó."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
             

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_04]-M2. Cho số tiền và kì hạn theo quý, lãi suất theo năm. Tính tổng tiền thu được sau n năm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_05]-M2. Cho số dân và tốc độ tăng trưởng theo năm. Tính số dân sau n năm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C6_B6, ["[D11_C6_B6_06]-M2. Cho mức tiền lương và tỉ lệ tăng lương. Tính mức lương nhận được sau n năm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                 

        #Lớp 11 - Chương 7 -  Đạo hàm    
                L11_C7 = QTreeWidgetItem(L11, ["Chương 7 - Đạo hàm"])
                L11_C7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B1 = QTreeWidgetItem(L11_C7, ["Bài 1 - Đạo hàm"])
                L11_C7_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B1_1 = QTreeWidgetItem(L11_C7_B1, ["7.1.1. Tính đạo hàm bằng định nghĩa"])
                L11_C7_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B1_1, ["[D11_C7_B1_01]-M1. Biểu diễn đạo hàm tại điểm x_0 của f(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B1_1, ["[D11_C7_B1_02]-M1. Cho hàm số đa thức. Biễu diễn công thức đạo hàm tại x_0"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2 = QTreeWidgetItem(L11_C7, ["Bài 2 - Các quy tắc đạo hàm"])
                L11_C7_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_1 = QTreeWidgetItem(L11_C7_B2, ["2.1. Đạo hàm chứa đa thức, 1/x, căn(x)"])
                L11_C7_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_1, ["[D11_C7_B2_01]-M2. Đạo hàm đa thức"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_1, ["[D11_C7_B2_02]-M2. Đạo hàm  đa thức + a/x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_1, ["[D11_C7_B2_03]-M2. Đạo hàm đa thức + acăn(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_1, ["[D11_C7_B2_04]-M2. Đạo hàm hàm số đa thức + acăn(x) + b/x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_2 = QTreeWidgetItem(L11_C7_B2, ["2.2. Đạo hàm hàm số lượng giác"])
                L11_C7_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_05]-M2. Đạo hàm y=asin(bx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_06]-M2. Đạo hàm y=acos(bx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_07]-M2. Đạo hàm y=asinx+bcosx"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_14]-M2. Đạo hàm y=atanx + bcotx"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_15]-M2. Đạo hàm y=atan(bx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)               

                item = QTreeWidgetItem(L11_C7_B2_2, ["[D11_C7_B2_16]-M2. Đạo hàm y=acot(bx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_3 = QTreeWidgetItem(L11_C7_B2, ["2.3. Đạo hàm hàm số mũ, lôgarit"])
                L11_C7_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_08]-M1. Đạo hàm y=a^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_09]-M2. Đạo hàm y=a^(đa thức)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_12]-M2. Đạo hàm y= đa thức + e^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_13]-M2. Đạo hàm y= e^(đa thức)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_21]-M1. Đạo hàm y= log_a x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_22]-M2. Đạo hàm y= log_a (đa thức)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_3, ["[D11_C7_B2_23]-M2. Đạo hàm y= ln (đa thức)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                L11_C7_B2_4 = QTreeWidgetItem(L11_C7_B2, ["2.4. Đạo hàm của tích, thương"])
                L11_C7_B2_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_17]-M2. Đạo hàm y=(ax+b).căn(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_18]-M2. Đạo hàm y=(ax+b).sin(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_19]-M2. Đạo hàm y=(ax+b).cos(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_24]-M2. Đạo hàm y=(đa thức).a^x"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_10]-M2. Đạo hàm y=m/(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_35]-M2. Đạo hàm y=(ax+b)/(cx+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_11]-M2. Đạo hàm y=(mx^2+nx+p)/(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_20]-M2. Đạo hàm y=(ax+b)/căn(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_4, ["[D11_C7_B2_25]-M2. Đạo hàm y=(ax+b)/cos(x) hoặc y=(ax+b)/sin(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_5 = QTreeWidgetItem(L11_C7_B2, ["2.5. Đạo hàm của hàm hợp"])
                L11_C7_B2_5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_5, ["[D11_C7_B2_30]-M2. Đạo hàm y=(đa thức)^n"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_5, ["[D11_C7_B2_26]-M2. Đạo hàm y=căn(đa thức)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_5, ["[D11_C7_B2_27]-M2. Đạo hàm y=căn(sin(ax+b))"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_5, ["[D11_C7_B2_28]-M2. Đạo hàm y=căn(cos(ax+b))"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_5, ["[D11_C7_B2_29]-M2. Đạo hàm y=k/u"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_6 = QTreeWidgetItem(L11_C7_B2, ["2.6. Đạo hàm cấp hai"])
                L11_C7_B2_6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_6.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_6, ["[D11_C7_B2_31]-M2. Đạo hàm cấp 2: đa thức"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_6, ["[D11_C7_B2_37]-M2. Cho hàm số đa thức.Tính đạo hàm cấp hai tại điểm x_0"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_6, ["[D11_C7_B2_32]-M2. Đạo hàm cấp 2: y=sin(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_6, ["[D11_C7_B2_33]-M2. Đạo hàm cấp 2: y=cos(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_6, ["[D11_C7_B2_34]-M3. Đạo hàm cấp 2: y=(ax+b)/(cx+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                
                L11_C7_B2_7 = QTreeWidgetItem(L11_C7_B2, ["2.7. Bài toán vận tốc, gia tốc"])
                L11_C7_B2_7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_7.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_7, ["[D11_C7_B2_38]-M2. Cho hàm số quãng đường. Tính vận tốc tại thời điểm t_0."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_7, ["[D11_C7_B2_39]-M2. Cho hàm số quãng đường. Tính gia tốc tốc tại thời điểm t_0."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B2_8 = QTreeWidgetItem(L11_C7_B2, ["2.8. Phương trình, Bất phương trình chứa f'(x)"])
                L11_C7_B2_8.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B2_8.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_8, ["[D11_C7_B2_40]-M2. Cho f(x) là hàm số bậc 3. Giải PT f'(x)=a."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_8, ["[D11_C7_B2_36]-M2. Cho hàm số đa thức. Giải phương trình chứa đạo hàm cấp hai"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_8, ["[D11_C7_B2_41]-M2. Cho f(x) là hàm số bậc 3. Giải BPT f'(x)+a>0 (<0,>=0,<=0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_8, ["[D11_C7_B2_42]-M3. f(x)=a.căn(x)+ bx +c. Giải BPT f'(x) > 0 (<0,>=0,<=0)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B2_8, ["[D11_C7_B2_43]-M4. Cho f(x) là hàm bậc 3. Tìm m để f'(x)>0 (<0) với mọi x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                

                L11_C7_B3 = QTreeWidgetItem(L11_C7, ["Bài 3 - Phương trình tiếp tuyến"])
                L11_C7_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B3, ["[D11_C7_B3_04]-M1. Cho hàm số. Tính hệ số góc của tiếp tuyến tại điểm x_0"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B3, ["[D11_C7_B3_01]-M2. Cho hàm số đa thức. Viết PTTT tại điểm M(x0;y0)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B3, ["[D11_C7_B3_02]-M2. Cho hàm số đa thức. Viết PTTT tại điểm có hoành độ x_0"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B3, ["[D11_C7_B3_03]-M2. Cho hàm số đa thức. Viết PTTT tại điểm có tung độ y_0"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C7_B4 = QTreeWidgetItem(L11_C7, ["Đúng-Sai"])
                L11_C7_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C7_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B4, ["[D11_C7_B4_01]-TF-M2. Tạo câu đúng-sai: Đạo hàm cấp 1,2,tiếp tuyến của hàm số đa thức"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B4, ["[D11_C7_B4_02]-TF-M2. Tạo câu đúng-sai: Đạo hàm cấp 1,2,tiếp tuyến của y=(ax+b)/(cx+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B4, ["[D11_C7_B4_03]-TF-M2. Tạo câu đúng-sai: Đạo hàm của y=msin(ax)+cos(x+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B4, ["[D11_C7_B4_04]-TF-M2. Tạo câu đúng-sai: Đạo hàm của hàm số lượng giác"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C7_B4, ["[D11_C7_B4_05]-TF-M2. Tạo câu đúng-sai: Bài toán chuyển động theo hàm bậc 2"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                



        #Lớp 11 - Chương 8 -  Quan hệ vuông góc trong không gian      
                L11_C8 = QTreeWidgetItem(L11, ["Chương 8 - Quan hệ vuông góc trong không gian"])
                L11_C8.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B1 = QTreeWidgetItem(L11_C8, ["Bài 1 - Hai đường thẳng vuông góc"])
                L11_C8_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B1_1 = QTreeWidgetItem(L11_C8_B1, ["8.1.1 - Đúng-Sai"])
                L11_C8_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B1_1, ["[D11_C8_B1_04]-TF-M2. Cho hình lập phương. Tạo câu đúng-sai: góc giữa 2 đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B1_2 = QTreeWidgetItem(L11_C8_B1, ["8.1.2 - Trắc Nghiệm"])
                L11_C8_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)


                item = QTreeWidgetItem(L11_C8_B1_2, ["[D11_C8_B1_01]-M1. Cho hình lập phương. Tính góc giữa 2 đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B1_2, ["[D11_C8_B1_02]-M2. S.ABCD: cạnh v.g đáy h.c.n. Tính góc giữa 2 đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B1_2, ["[D11_C8_B1_03]-M2. S.ABCD: cạnh v.g đáy h.vuông. Tìm 2 đường vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 2

                L11_C8_B2 = QTreeWidgetItem(L11_C8, ["Bài 2 - Đường thẳng vuông góc với mặt phẳng"])
                L11_C8_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B2_1 = QTreeWidgetItem(L11_C8_B2, ["8.2.1. Đúng-Sai"])
                L11_C8_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_1, ["[D11_C8_B2_08]-TF-M2. S.ABCD: cạnh v.g đáy h.vuông. Tạo câu đúng-sai: đường vuông góc đường."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                item = QTreeWidgetItem(L11_C8_B2_1, ["[D11_C8_B2_10]-TF-M2. S.ABCD: cạnh v.g đáy h.chữ nhật. Tạo câu đúng-sai: đường vuông góc đường."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_1, ["[D11_C8_B2_09]-TF-M2. S.ABCD: cạnh v.g đáy h.vuông. Tạo câu đúng-sai: đường vuông góc mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_1, ["[D11_C8_B2_11]-TF-M2. S.ABCD: cạnh v.g đáy h.chữ nhật. Tạo câu đúng-sai: đường vuông góc mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_1, ["[D11_C8_B2_12]-TF-M2. S.ABCD: cạnh v.g đáy h.vuông. Tạo câu đúng-sai:lập luận đường vuông góc mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B2_2 = QTreeWidgetItem(L11_C8_B2, ["8.2.2. Xác định đường thẳng vuông góc với mặt phẳng"])
                L11_C8_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_01]-M2. S.ABCD: cạnh v.g đáy h.vuông. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_02]-M2. S.ABCD: cạnh v.g đáy h.chữ nhật. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_03]-M2. S.ABCD: cạnh v.g đáy h.thoi. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_04]-M2. S.ABC: cạnh v.g đáy tam giác vuông. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_05]-M2. S.ABC: cạnh v.g đáy tam giác đều. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_2, ["[D11_C8_B2_06]-M3. S.ABC: ABC đều, có hình chiếu lên cạnh bên. Tìm đường vg mặt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B2_3 = QTreeWidgetItem(L11_C8_B2, ["8.2.3. Xác định đường thẳng vuông góc với đường thẳng"])
                L11_C8_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_07]-M3. S.ABC: ABC đều, có hình chiếu lên cạnh bên. Tìm đường vg đường."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B2_3 = QTreeWidgetItem(L11_C8_B2, ["8.2.4. Hình chiếu của điểm trên mặt phẳng"])
                L11_C8_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_13]-M2. S.ABCD: ABCD h.vuông. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_14]-M2. S.ABCD: ABCD h.chữ nhật. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_15]-M2. S.ABC: ABC tam giác vuông. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_16]-M2. S.ABC: ABC tam giác đều. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_17]-M3. S.ABCD: ABCD h.vuông. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_18]-M3. S.ABCD: ABCD h.chữ nhật. Tìm hình chiếu của điểm trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B2_3 = QTreeWidgetItem(L11_C8_B2, ["8.2.4. Hình chiếu của đường thẳng trên mặt phẳng"])
                L11_C8_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_19]-M1. S.ABCD: ABCD h.vuông. Tìm hình chiếu của đường thẳng trên mặt đáy."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B2_3, ["[D11_C8_B2_20]-M2. S.ABCD: ABCD h.vuông. Tìm hình chiếu của đường thẳng trên mặt đứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B3 = QTreeWidgetItem(L11_C8, ["Bài 3 - Hai mặt phẳng vuông góc"])
                L11_C8_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B3_1 = QTreeWidgetItem(L11_C8_B3, ["8.3.1. Đúng-Sai"])
                L11_C8_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B3_2 = QTreeWidgetItem(L11_C8_B3, ["8.3.2. Trắc Nghiệm"])
                L11_C8_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B3_2_1 = QTreeWidgetItem(L11_C8_B3_2, ["8.3.2.1. Góc giữa 2 mặt phẳng"])
                L11_C8_B3_2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B3_2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_1, ["[D11_C8_B3_06]-M2. S.ABCD: ABCD h.vuông. Tìm vị trí góc (mặt nghiêng, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_1, ["[D11_C8_B3_07]-M2. S.ABCD: ABCD h.chữ nhật. Tìm vị trí góc (mặt nghiêng, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_1, ["[D11_C8_B3_08]-M2. S.ABC: ABC đều. Tính góc (mặt nghiêng, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_1, ["[D11_C8_B3_09]-M2. S.ABC: ABC vuông. Tính góc (mặt nghiêng, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L11_C8_B3_2_2 = QTreeWidgetItem(L11_C8_B3_2, ["8.3.2.2. Hai mặt phẳng vuông góc"])
                L11_C8_B3_2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B3_2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_2, ["[D11_C8_B3_01]-M2. S.ABCD: ABCD h.vuông. Tìm 2 MP vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_2, ["[D11_C8_B3_03]-M2. S.ABCD: ABCD h.chữ nhật. Tìm 2 MP vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_2, ["[D11_C8_B3_02]-M3. S.ABCD: ABCD h.vuông. Tìm 2 MP vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L11_C8_B3_2_2, ["[D11_C8_B3_04]-M3. S.ABCD: ABCD h.chữ nhật. Tìm 2 MP vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B3_2_2, ["[D11_C8_B3_05]-M3. S.ABC: ABC vuông tại A. Tìm 2 MP vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B4 = QTreeWidgetItem(L11_C8, ["Bài 4 - Khoảng cách trong không gian"])
                L11_C8_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B4_1 = QTreeWidgetItem(L11_C8_B4, ["8.4.1 Khoảng cách từ một điểm đến đường thẳng"])
                L11_C8_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B4_2 = QTreeWidgetItem(L11_C8_B4, ["8.4.2 Khoảng cách từ một điểm đến mặt phẳng"])
                L11_C8_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_2, ["[D11_C8_B4_01]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c từ điểm thuộc đáy đến mặt đứng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_2, ["[D11_C8_B4_02]-M3. S.ABCD: ABCD h.chữ nhật. Tính k.c từ điểm thuộc cạnh bên đến mặt đứng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_2, ["[D11_C8_B4_03]-M3. S.ABCD: ABCD h.chữ nhật. Tính k.c từ chân đường cao đến mặt nghiêng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_2, ["[D11_C8_B4_04]-M2. Lăng trụ ABC.A'B'C': ABC vuông. Tính k.c từ điểm đến mặt đứng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_2, ["[D11_C8_B4_09]-M3. Lăng trụ ABC.A'B'C'. Tính k.c từ điểm đến mặt nghiêng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                L11_C8_B4_3 = QTreeWidgetItem(L11_C8_B4, ["8.4.3 Khoảng cách giữa 2 đường thẳng chéo nhau"])
                L11_C8_B4_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B4_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_3, ["[D11_C8_B4_05]-M2. Hình lập phương. Tính k.c 2 đường chéo nhau(kẻ đường vuông góc chung)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_3, ["[D11_C8_B4_06]-M2. Hình lập phương. Tính k.c 2 đường chéo nhau(tạo mặt song song đường)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_3, ["[D11_C8_B4_07]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c 2 đường chéo nhau(kẻ đường vuông góc chung)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B4_3, ["[D11_C8_B4_08]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c 2 đường chéo nhau(tạo mặt song song đường)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # L11_C8_B4_4 = QTreeWidgetItem(L11_C8_B4, ["8.4.4 Khoảng cách giữa đường thẳng và mặt phẳng"])
                # L11_C8_B4_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L11_C8_B4_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # L11_C8_B4_5 = QTreeWidgetItem(L11_C8_B4, ["8.4.5 Khoảng cách giữa 2 mặt phẳng"])
                # L11_C8_B4_5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L11_C8_B4_5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B5 = QTreeWidgetItem(L11_C8, ["Bài 5 - Góc giữa đường thẳng và mặt phẳng. Góc nhị diện"])
                L11_C8_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B5_1 = QTreeWidgetItem(L11_C8_B5, ["8.5.1 - Góc giữa đường thẳng và mặt phẳng."])
                L11_C8_B5_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B5_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_01]-M1. S.ABCD: ABCD h.vuông. Tìm vị trí góc(đường, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_02]-M2. S.ABCD: ABCD h.vuông. Tìm vị trí góc(đường, mặt đứng)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_03]-M2. S.ABCD: ABCD h.chữ nhật. Xác định vị trí góc(đường, mặt đứng)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_04]-M1. S.ABC: ABC tam giác vuông. Xác định vị trí góc(đường, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_05]-M2. S.ABC: ABC tam giác vuông. Xác định vị trí góc(đường, mặt đứng)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_08]-M2. S.ABCD: đáy hình vuông. Tính số đo góc(đường, mặt đáy)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_1, ["[D11_C8_B5_09]-M3. S.ABCD: đáy hình vuông. Tính số đo góc(đường, mặt bên)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L11_C8_B5_2 = QTreeWidgetItem(L11_C8_B5, ["8.5.2 - Góc nhị diện."])
                L11_C8_B5_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B5_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_2, ["[D11_C8_B5_06]-M2. S.ABCD, đáy h.c.n. Tính số đo góc phẳng nhị diện bởi mặt nghiêng và đáy."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B5_2, ["[D11_C8_B5_07]-M3. S.ABCD, đáy h.c.n. Tính số đo góc phẳng nhị diện bởi mặt nghiêng và đáy."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
        


                L11_C8_B6 = QTreeWidgetItem(L11_C8, ["Bài 6 - Thể tích"])
                L11_C8_B6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B6.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B6_1 = QTreeWidgetItem(L11_C8_B6, ["6.1 Thể tích của khối chóp"])
                L11_C8_B6_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B6_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_1, ["[D11_C8_B6_01]-M1. H.chóp có S_đáy và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_1, ["[D11_C8_B6_02]-M2. H.chóp có đáy tam giác đều và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_1, ["[D11_C8_B6_03]-M2. H.chóp tam giác vuông và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_1, ["[D11_C8_B6_04]-M2. H.chóp đáy tam giác đều và góc(cạnh bên, đáy). Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_1, ["[D11_C8_B6_09]-M3. H.chóp đáy h.chữ nhật có góc giữa 2 đường thẳng. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)  

                L11_C8_B6_2 = QTreeWidgetItem(L11_C8_B6, ["6.2 Thể tích của khối lăng trụ, khối hộp"])
                L11_C8_B6_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B6_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_05]-M1. H.lăng trụ có S_đáy và h Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_06]-M2. H.lăng trụ có đáy tam giác đều và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_07]-M1. H.lăng trụ có đáy tam giác vuông và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_08]-M1. H.lăng trụ có đáy h.vuông, h.chữ nhật và h. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_11]-M1. Cho hình lập phương có độ dài cạnh. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_2, ["[D11_C8_B6_12]-M1. Cho hình hộp chữ nhật có độ dài cạnh. Tính V."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B6_3 = QTreeWidgetItem(L11_C8_B6, ["6.3 Nhận dạng tính chất khối chóp, khối lăng trụ, khối hộp"])
                L11_C8_B6_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B6_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B6_3, ["[D11_C8_B6_10]-M1. Nhận dạng tính chất của lăng trụ đều."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C8_B7 = QTreeWidgetItem(L11_C8, ["Đúng-Sai tổng hợp"])
                L11_C8_B7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C8_B7.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B7, ["[D11_C8_B7_01]-TF-M2. S.ABC: ABC đều. Đúng-Sai: góc(đường,mặt), góc(mặt,mặt), mặt vg mặt, thể tích"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C8_B7, ["[D11_C8_B7_02]-TF-M3. S.ABC: ABC đều. Đúng-Sai: góc(đường,mặt), góc(mặt,mặt), thể tích, khoảng cách"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 11 - Chương 9 - Xác suất
                L11_C9 = QTreeWidgetItem(L11, ["Chương 9 - Xác suất"])
                L11_C9.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C9_B1 = QTreeWidgetItem(L11_C9, ["Bài 1 - Biến cố giao và quy tắc nhân xác suất"])
                L11_C9_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C9_B1_1 = QTreeWidgetItem(L11_C9_B1, ["9.1.1 Biến cố giao"])
                L11_C9_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_01]-M2. Hộp chứa n quả cầu. Biến cố giao: vừa là số chẵn (lẻ), vừa chia hết cho k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_02]-M2. Hộp chứa n tấm thẻ. Biến cố giao: thẻ vừa thuộc đoạn [a;b], vừa chia hết cho k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_03]-M2. Hộp chứa n tấm thẻ. Biến cố giao: thẻ vừa chia hết cho m, vừa chia hết cho k."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_04]-M2. Gieo xúc xắc 2 lần. Biến cố giao: k < i + j và i+j <m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_05]-M2. Gieo xúc xắc 2 lần. Biến cố giao: mặt k chấm xuất hiện và i+j =,>,< m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_06]-M2. Gieo xúc xắc 2 lần. Biến cố giao: i + j = k và i.j <(>) m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_1, ["[D11_C9_B1_09]-M2. Hộp chứa các viên bi có 2 màu. Phát biểu biến cố giao."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C9_B1_2 = QTreeWidgetItem(L11_C9_B1, ["9.1.2 Quy tắc nhân xác suất"])
                L11_C9_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_07]-M1. Cho A, B độc lập và P(A), P(B). Tính P(AB)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_08]-M2. Cho A, B độc lập và P(A), P(B). Tính x.s của biến cố giao chứa biến cố đối."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_10]-M2. Cho x.s bị bệnh của 2 bệnh nhân. Tính x.s cả hai không bị biến chứng nặng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_15]-M2. Cho x.s bị bệnh của 2 bệnh nhân. Tính x.s cả hai bị biến chứng nặng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_11]-M2. Cho x.s  bị bệnh của 2 bệnh nhân. Tính x.s chỉ một trong hai không bị biến chứng nặng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_14]-M2. Cho x.s bắn trúng bia. Tính x.s cả hai lần bắn trúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_12]-M2. Cho x.s bắn trúng bia. Tính x.s cả hai lần bắn trượt."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B1_2, ["[D11_C9_B1_13]-M2. Cho x.s bắn trúng bia. Tính x.s chỉ một lần bắn trúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                #BÀI 2: BIẾN CỐ HỢP-QUY TẮC CỘNG XÁC SUẤT

                L11_C9_B2 = QTreeWidgetItem(L11_C9, ["Bài 2 - Quy tắc cộng xác suất"])
                L11_C9_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                L11_C9_B2_1 = QTreeWidgetItem(L11_C9_B2, ["9.2.1 Biến cố hợp"])
                L11_C9_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_1, ["[D11_C9_B2_01]-M2. Hộp chứa các viên bi có 2 màu. Phát biểu biến cố hợp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_1, ["[D11_C9_B2_02]-M2. Hộp chứa các viên bi, lấy ra rồi trả lại. Phát biểu biến cố hợp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_1, ["[D11_C9_B2_03]-M2. Gieo 1 xúc xắc + 1 đồng xu. Tính số phần tử biến cố hợp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L11_C9_B2_2 = QTreeWidgetItem(L11_C9_B2, ["9.2.2 Quy tắc cộng xác suất"])
                L11_C9_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L11_C9_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_04]-M1. Cho A, B xung khắc và P(A), P(B). Tính P(AUB)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_05]-M2. Gieo 1 xúc xắc và 1 đồng xu. Tính P(biến cố hợp)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_06]-M2. Cho 2 nhóm đồ vật. Tính x.s để số vật được chọn thuộc cùng 1 nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_07]-M2. Cho 2 nhóm học sinh. Tính xác suất để học sinh được chọn cùng giới tính."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_08]-M2. Cho các viên bi. Tính xác suất để các bi được chọn cùng màu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L11_C9_B2_2, ["[D11_C9_B2_09]-M2. Cho các cuốn truyện. Tính xác suất để các cuốn được chọn cùng thể loại."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


        #Lớp 12 - Chương 1 - Khảo sát hàm số 
                L12 = QTreeWidgetItem(self.treeWidget, ["Lớp 12"])
                L12.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12.setCheckState(0, Qt.CheckState.PartiallyChecked)
                

                L12_C1 = QTreeWidgetItem(L12, ["Chương 1 - Khảo sát hàm số"])
                L12_C1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B1 = QTreeWidgetItem(L12_C1, ["Bài 1 - Tính đơn điệu và cực trị của hàm số"])
                L12_C1_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B1_1 = QTreeWidgetItem(L12_C1_B1, ["Trắc nghiệm - Trả lời ngắn"])
                L12_C1_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_08]-M2. Cho bảng xét dấu f'(x). Tìm khoảng đơn điệu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_01]-M1. Cho bảng biến thiên. Tìm khoảng đơn điệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_28]-M2. Cho dấu f'(x) trên khoảng (a,b) và (c,d). Tìm khẳng định đúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_02]-M2. Cho hàm số bậc 3. Tìm khoảng đơn điệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_04]-M2. Cho hàm số y=(ax+b)/(cx+d). Tìm khoảng đơn điệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_05]-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Tìm khoảng đơn điệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_21]-M2. Tìm khoảng đơn điệu của các hàm số: y=cănu, y=e^u, y=ax/(x^2+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_06]-M2. Cho đồ thị bậc 3. Tìm khoảng đơn điệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_03]-M2. Cho đồ thị bậc 4. Tìm khoảng đơn điệu."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_07]-M2. Cho hàm số f'(x). Tìm khoảng đơn điệu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_10]-M2. Cho đồ thị f'(x). Tìm khoảng đơn điệu"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_09]-M2. Tìm m để hàm bậc 3 đơn điệu trên R"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_15]-M1. Cho bảng biến thiên bậc 3. Tìm điểm cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_16]-M1. Cho bảng biến thiên bậc 3. Tìm giá trị cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_34]-M2. Cho f'(x) tìm số điểm cực trị của f(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_35]-M2. Cho đồ thị f'(x). Tìm điểm cực trị"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_17]-M2. Cho hàm bậc 3. Tìm điểm cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_18]-M2. Cho hàm bậc 3. Tìm giá trị cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_19]-M2. Cho y=(ax^2+bx+c)/(dx+e). Tìm điểm cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_1, ["[D12_C1_B1_20]-M2. Cho y=(ax^2+bx+c)/(dx+e). Tìm giá trị cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B1_2 = QTreeWidgetItem(L12_C1_B1, ["Đúng-Sai"])
                L12_C1_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_14]-TF-M2. Cho BBT. Xét Đ-S: đơn điệu, so sánh giá trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_13]-TF-M2. Cho BBT. Xét Đ-S: đơn điệu, so sánh giá trị, điểm cực trị, giá trị cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_11]-TF-M2. Cho BXD f'(x). Xét Đ-S: đơn điệu, so sánh giá trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_12]-TF-M2. Cho BXD f'(x). Xét Đ-S: đơn điệu, so sánh giá trị, điểm cực trị, giá trị cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_26]-TF-M2. Cho hàm số bậc 3. Xét Đ-S: TXĐ, y', đơn điệu, y_CD, y_CT, đường thẳng qua cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_27]-TF-M2. Cho y=(ax^2+bx+c)/(dx+e). Xét Đ-S: TXD, y', đơn điệu, cực trị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_30]-TF-M3. Cho f(x)=(ax+b)/(x+m). Xét Đ-S: đạo hàm, cực trị, đồng biến, nghịch biến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_2, ["[D12_C1_B1_31]-TF-M3. Cho BXD f'(x). Xét Đ-S: đơn điệu và cực trị của f(x), đơn điệu và cực trị của f(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B1_3 = QTreeWidgetItem(L12_C1_B1, ["Trả lời ngắn"])
                L12_C1_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_33]-TL-M3. Tìm m để y=(ax+b)/(x+m) đơn điệu trên các khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_29]-TL-M3. Tìm m để y=(ax+b)/(x+m) đơn điệu trên K."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_32]-TL-M3. Tìm m để hàm số đồng biến, nghịch biến trên khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_22]-TL-M2. Cho hàm bậc 3 có điểm cực trị x_1,x_2. Tính P=ax_1 + bx_2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_23]-TL-M2. Cho hàm bậc 3 có giá trị cực trị y_1,y_2. Tính P=ay_1 + by_2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_24]-TL-M2. Cho y=(ax^2+bx+c)/(dx+e) có điểm cực trị x_1,x_2. Tính P=ax_1 + bx_2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B1_3, ["[D12_C1_B1_25]-TL-M2. Cho y=(ax^2+bx+c)/(dx+e) có giá trị cực trị y_1,y_2. Tính P=ay_1 + by_2."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                


                # BÀI 3: GIÁ TRỊ LỚN NHẤT - GIÁ TRỊ NHỎ NHẤT  
                L12_C1_B2 = QTreeWidgetItem(L12_C1, ["Bài 2 - Giá trị lớn nhất và Giá trị nhỏ nhất"])
                L12_C1_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B2_1 = QTreeWidgetItem(L12_C1_B2, ["Trắc nghiệm-Trả lời ngắn"])
                L12_C1_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_24]-M2. Cho đồ thị f(x). Tìm GTLN-GTNN."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_01]-M2. Tìm GTLN-GTNN của hàm bậc 3 trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_02]-M2. Tìm GTLN-GTNN của hàm y=(ax+b)/(cx+d) trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_10]-M3. Tìm GTLN-GTNN của hàm y=e^x.(ax^2+bx+c) trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_11]-M3. Tìm GTLN-GTNN của hàm y=e^(ax^2+bx+c) trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_12]-M3. Tìm GTLN-GTNN của hàm y=ln(ax^2+bx+c) trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_03]-M2. Tìm GTLN-GTNN của hàm bậc 3 trên khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_04]-M2. Tìm GTLN-GTNN của y=(ax+b)/(cx+d) trên khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_05]-M2. Tìm GTLN-GTNN của y=ax+b/x trên khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_06]-M2. Tìm GTLN-GTNN của y=(ax^2+bx+c)/(dx+e) trên khoảng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_13]-M3. Tìm GTLN-GTNN của y=căn(x-a)+căn(b-x)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_1, ["[D12_C1_B2_14]-M3. Tìm GTLN-GTNN của y=b.x.căn(a^2-x^2)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B2_2 = QTreeWidgetItem(L12_C1_B2, ["Đúng-Sai"])
                L12_C1_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_2, ["[D12_C1_B2_07]-TF-M2. Cho hàm bậc 3. Xét Đ-S: y', y'=0, f(x_0), Min-Max trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_2, ["[D12_C1_B2_08]-TF-M2. y=(ax+b)/(cx+d). Xét Đ-S: y', y'=0, f(x_0), Min-Max trên đoạn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_2, ["[D12_C1_B2_09]-TF-M2. Cho hàm số bậc 3. Xét Đ-S: y', y'=0, f(x_0), Min-Max trên khoảng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B2_3 = QTreeWidgetItem(L12_C1_B2, ["Trả lời ngắn"])
                L12_C1_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_15]-TL-M2. Cho hàm số. Tìm GTLN-GTNN trên đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_16]-TL-M3. Cho hàm số. Tìm GTLN-GTNN trên đoạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_17]-TL-M3. Cho hàm số. Tìm GTLN-GTNN trên khoảng, nửa khoảng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_18]-TL-M3. Tìm m để y=(ax+b)/(cx+d) đạt min y=T hoặc (max y = T."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_20]-TL-M3. Tìm m để y=(ax+b)/(cx+d) có min y + max y=T."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_19]-TL-M3. Tìm m để y=|f(x)|+am đạt GTLN(GTNN) bằng y_0."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_21]-TL-M3. Tìm vận tốc lớn nhất, nhỏ nhất của chuyển động"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_22]-TL-M3. Cắt hình chữ nhật từ nửa đường tròn. Tìm diện tích lớn nhất của hình chữ nhật."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B2_3, ["[D12_C1_B2_23]-TL-M3. Tìm chi phí thấp nhất để xây cái bể hình hộp chữ nhật."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                # BÀI 3: ĐƯỜNG TIỆM CẬN                    

                L12_C1_B3 = QTreeWidgetItem(L12_C1, ["Bài 3 - Đường tiệm cận"])
                L12_C1_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B3_1 = QTreeWidgetItem(L12_C1_B3, ["Trắc nghiệm-Trả lời ngắn"])
                L12_C1_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_01]-M2. Tìm đường tiệm cận của đồ thị y=(ax+b)/(cx+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_03]-M1. Tìm đường tiệm cận đứng của đồ thị y=(ax^2+bx+c)/(dx+e)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_04]-M2. Tìm đường tiệm cận xiên của đồ thị y=(ax^2+bx+c)/(dx+e)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_05]-M2. Tìm số đường tiệm cận đứng của đồ thị hàm số khác"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_06]-M2. Tìm số đường tiệm cận của đồ thị hàm số khác"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_07]-M2. Tìm đường tiệm cận từ giới hạn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_02]-M2. Cho BBT tìm số đường tiệm cận."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_14]-M2. Cho đường tiệm cận xiên. Tìm khẳng định đúng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_15]-M2. Cho giới hạn. Tìm tiệm cận xiên."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_1, ["[D12_C1_B3_16]-M2. Đọc đường tiệm cận từ đồ thị hàm số y=(ax+b)/(cx+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B3_2 = QTreeWidgetItem(L12_C1_B3, ["Đúng-Sai"])
                L12_C1_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_08]-TF-M2. Cho y=(ax+b)/(cx+d). Xét Đ-S: lim, tiệm cận đứng, tiệm cận ngang"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_09]-TF-M2. Cho y=(ax^2+bx+c)/(dx+e). Xét Đ-S: lim, tiệm cận đứng, tiệm cận xiên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_17]-TF-M2. Cho y=(ax^2+bx+c)/(dx+e). Xét Đ-S: y',y'=0, tiệm cận đứng, tiệm cận xiên"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_11]-TF-M3. Cho BBT của y=f(x). Xét Đ-S: Tìm TCĐ, Tìm TCN, Số TCD+TCN, TCN(TCD) của y= m/(f(x)+n)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_12]-TF-M3. Cho hàm số y=(ax+b)/(cx+d). Xét Đ-S: TCĐ, TCN, giao điểm của TC, HCN tạo bởi 2 tiệm cận"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_2, ["[D12_C1_B3_13]-TF-M3. Cho limf. Xét Đ-S: TCĐ, TCN, Số TCD và TCN, TNC của y=(af+b)/(cf+d)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B3_3 = QTreeWidgetItem(L12_C1_B3, ["Trả lời ngắn"])
                L12_C1_B3_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B3_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_3, ["[D12_C1_B3_10]-TL-M3. Tìm m để y=(ax+b)/(cx+d) có tiệm cận"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B3_3, ["[D12_C1_B3_18]-TL-M2. Cho y=(ax^2+bx+c)/(dx+e) có TCX là y=ax+b. Tính ma+nb."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)





                # BÀI 4: ĐỒ THỊ HÀM SỐ

                L12_C1_B4 = QTreeWidgetItem(L12_C1, ["Bài 4 - Đồ thị hàm số"])
                L12_C1_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B4_1 = QTreeWidgetItem(L12_C1_B4, ["Trắc nghiệm-Trả lời ngắn"])
                L12_C1_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_01]-M2. Cho đồ thị. Tìm hàm số bậc 3."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_02]-M2. Cho đồ thị. Tìm hàm số bậc 4."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_03]-M2. Cho đồ thị. Tìm hàm số y=(ax+b)/(cx+d)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_04]-M2. Cho bảng biến thiên. Tìm hàm số y=(ax+b)/(cx+d)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_09]-M2. Cho bảng biến thiên. Tìm hàm số bậc 3."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_05]-M2. Cho đồ thị. Tìm hàm số y=(ax^2+bx+c)/(dx+e)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_13]-M2. Tìm số giao điểm của hàm số đa thức với trục Ox."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_10]-M2. Cho hàm số y=(ax+b)/(cx+d) và y=mx+n. Tìm số giao điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_14]-M2. Cho y=(ax+b)/(cx+d). Tìm tọa độ tâm đối xứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_15]-M2. Cho y=(ax^2+bx+c)/(dx+e). Tìm tọa độ tâm đối xứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_16]-M2. Cho y=ax^3+bx^2+cx+d. Tìm tọa độ tâm đối xứng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_11]-M3. Cho đồ thị. Tìm số nghiệm của phương trình f(x)=m."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_1, ["[D12_C1_B4_12]-M2. Tìm số giao điểm của đồ thị với trục Ox, Oy."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B4_2 = QTreeWidgetItem(L12_C1_B4, ["Đúng-Sai"])
                L12_C1_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_2, ["[D12_C1_B4_06]-TF-M2. Cho hàm bậc 3. Xét Đ-S: y', dấu y', BBT, Đồ thị"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_2, ["[D12_C1_B4_07]-TF-M2. Cho hàm số y=(ax+b)/(cx+d). Xét Đ-S: y', dấu y', BBT, Đồ thị"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_2, ["[D12_C1_B4_08]-TF-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Xét Đ-S: y', dấu y', BBT, Đồ thị"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B4_3 = QTreeWidgetItem(L12_C1_B4, ["Trả lời ngắn"])
                L12_C1_B4_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B4_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_3, ["[D12_C1_B4_17]-TL-M2. Cho đồ thị hàm bậc 3. Tính tổng các hệ số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B4_3, ["[D12_C1_B4_18]-TL-M2. Cho đồ thị hàm y=(ax+b)/(cx+d). Tính tổng các hệ số"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                # BÀI 5: TOÁN THỰC TẾ

                L12_C1_B5 = QTreeWidgetItem(L12_C1, ["Bài 5 - Bài toán thực tế"])
                L12_C1_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C1_B5_3 = QTreeWidgetItem(L12_C1_B5, ["Trả lời ngắn"])
                L12_C1_B5_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C1_B5_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B5_3, ["[D12_C1_B5_04]-TL-M2. Bài toán tìm tốc độ trung bình để chi phí xăng nhỏ nhất."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B5_3, ["[D12_C1_B5_02]-TL-M3. Cho hàm số chi phí. Tìm chi phí trung bình thấp nhất."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B5_3, ["[D12_C1_B5_01]-TL-M4. Bài toán thực tế tìm chi phí thấp nhất"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C1_B5_3, ["[D12_C1_B5_03]-TL-M3. Bài toán tìm số giếng dầu có thể khai thác thêm để sản lượng là lớn nhất"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                


                

        #Lớp 12 - Chương 2 - Tọa độ của vectơ trong không gian            

                L12_C2 = QTreeWidgetItem(L12, ["Chương 2 - Tọa độ của vectơ trong không gian"])
                L12_C2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #Bài 1
                L12_C2_B1 = QTreeWidgetItem(L12_C2, ["Bài 1 - Vectơ và các phép toán vectơ trong không gian"])
                L12_C2_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B1_1 = QTreeWidgetItem(L12_C2_B1, ["Trắc nghiệm"])
                L12_C2_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_01]-M2. Cho hình hộp. Tìm vectơ bằng vectơ cho trước"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_02]-M2. Tìm đẳng thức đúng liên quan đến 3 điểm (phép cộng trừ)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_03]-M1. Nhận dạng quy tắc hình hộp"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_04]-M1. Cho tứ diện. Tìm khẳng định đúng về quy tắc cộng trừ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_05]-M2. Cho tứ diện. Tìm khẳng định về phép toán vectơ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_06]-M3. Cho hình lập phương. Tính độ dài vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_11]-M2. Cho hình chóp có đáy h.b.h. Tìm khẳng định sai về phép toán vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_18]-M2. Cho hình lập phương. Tìm góc giữa hai vectơ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_19]-M2. Cho tứ diện đều. Tính tích vô hướng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_12]-M2. Cho hai vectơ có độ dài và góc. Tính tích vô hướng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_13]-M2. Cho hai vectơ có độ dài và tích vô hướng. Tính góc giữa 2 vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_1, ["[D12_C2_B1_14]-M3. Cho hai vectơ a,b có độ dài và tích vô hướng. Tính |ma+nb|."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                L12_C2_B1_2 = QTreeWidgetItem(L12_C2_B1, ["Đúng-Sai"])
                L12_C2_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_2, ["[D12_C2_B1_07]-TF-M2. Cho hình hộp chữ nhật. Xét Đ-S: phép toán về vectơ, quy tắc hình bình hành, hình hộp, độ dài vectơ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_2, ["[D12_C2_B1_08]-TF-M2. Cho tứ diện đều. Xét Đ-S: phép toán vectơ, góc, tích vô hướng"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_2, ["[D12_C2_B1_10]-TF-M2. Cho hình chóp. Xét Đ-S: Tích vô hướng, phép toán vectơ, độ dài."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_2, ["[D12_C2_B1_17]-TF-M3. Cho hình lập phương. Xét Đ-S: hướng, phương, góc, tích vô hướng, độ dài vectơ"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B1_3 = QTreeWidgetItem(L12_C2_B1, ["Trả lời ngắn"])
                L12_C2_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_3, ["[D12_C2_B1_09]-TL-M3. Cho 3 lực đôi một vuông góc. Tính tổng hợp lực."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_3, ["[D12_C2_B1_15]-TL-M3. Cho hai vectơ a,b và độ dài |ma+nb|. Tính cos(a,b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B1_3, ["[D12_C2_B1_16]-TL-M3. Cho hai vectơ a,b có |a|, |b|, ab. Đặt x=ma+nb, y=pa+qb. Tính cos(x,y)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)



                #Bài 2
                L12_C2_B2 = QTreeWidgetItem(L12_C2, ["Bài 2 - Tọa độ của vectơ"])
                L12_C2_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B2_1 = QTreeWidgetItem(L12_C2_B2, ["Trắc nghiệm"])
                L12_C2_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_01]-M1. Đọc tọa độ vectơ theo vectơ đơn vị i, j, k"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_02]-M1. Cho vectơ OA= vecto(u). Tìm tọa độ điểm A."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_03]-M1. Cho hai điểm. Tìm tọa độ vectơ tạo bởi 2 điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_04]-M2. Cho vectơ u và điểm A. Tìm tọa độ điểm B sao cho vto(AB)=vto(u)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_05]-M1. Cho điểm A. Tìm hình chiếu của điểm A trên các mặt phẳng tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_1, ["[D12_C2_B2_06]-M2. Cho các điểm A,B,C. Tìm D sao cho ABCD là hình bình hành."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B2_2 = QTreeWidgetItem(L12_C2_B2, ["Đúng-Sai"])
                L12_C2_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_2, ["[D12_C2_B2_08]-TF-M2. H.hộp chữ nhật tọa độ hóa gốc A. Xét Đ-S: tọa độ vectơ, tọa độ điểm, tọa độ tâm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_2, ["[D12_C2_B2_09]-TF-M2. H.hộp lập phương tọa độ hóa gốc O. Xét Đ-S: tọa độ vectơ, tọa độ điểm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B2_3 = QTreeWidgetItem(L12_C2_B2, ["Trả lời ngắn"])
                L12_C2_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_3, ["[D12_C2_B2_07]-TL-M2. Cho h.hộp chữ nhật được tọa độ hóa. Tìm tọa độ điểm hoặc vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B2_3, ["[D12_C2_B2_10]-TL-M2. Bài toán thực tế: Tìm tọa độ máy bay"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 3
                L12_C2_B3 = QTreeWidgetItem(L12_C2, ["Bài 3 - Biểu thức tọa độ của phép toán vectơ"])
                L12_C2_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B3_1 = QTreeWidgetItem(L12_C2_B3, ["Trắc nghiệm"])
                L12_C2_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B3_1_1 = QTreeWidgetItem(L12_C2_B3_1, ["Tọa độ vectơ"])
                L12_C2_B3_1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3_1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_19]-M1. Cho A,B. Tìm tọa độ vectơ AB."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)       

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_01]-M1. Cho hai vectơ. Tìm tọa độ vectơ tổng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_02]-M1. Cho hai vectơ. Tìm tọa độ vectơ hiệu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_07]-M2. Cho hai vectơ a và b. Tìm tọa độ ma+nb."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_05]-M1. Cho vectơ. Tính độ dài."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_20]-M2. Cho A,B. Tính độ dài đoạn thẳng AB."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_03]-M1. Cho hai vectơ. Tính tích vô hướng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_06]-M2. Cho hai vectơ. Tính tích có hướng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_1, ["[D12_C2_B3_04]-M3. Tìm m để hai vectơ vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B3_1_2 = QTreeWidgetItem(L12_C2_B3_1, ["Tọa độ điểm"])
                L12_C2_B3_1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3_1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_24]-M1. Tìm hình chiếu lên trục tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_25]-M1. Tìm hình chiếu lên mặt phẳng tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_26]-M2. Tìm điểm đối xứng qua tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_27]-M2. Tìm điểm đối xứng qua mặt phẳng tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_15]-M1. Cho hai điểm A,B. Tìm tọa độ trung điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_12]-M2. Cho hai điểm A,B. Tìm tọa độ C để AC nhận B làm trung điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_16]-M1. Cho tam giác. Tìm tọa độ trọng tâm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_17]-M2. Cho A, vto_AB, vto_AC. Tìm tọa độ trọng tâm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_1_2, ["[D12_C2_B3_18]-M3. Cho A,B,G. Tìm C để ABC nhận G làm trọng tâm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)          

                
                L12_C2_B3_2 = QTreeWidgetItem(L12_C2_B3, ["Đúng-Sai"])
                L12_C2_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_2, ["[D12_C2_B3_08]-TF-M2. Cho h.b.h có A,B,C. Xét Đ-S: tọa độ vectơ, tọa độ điểm, tọa độ tâm ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_2, ["[D12_C2_B3_10]-TF-M2. Cho hai vectơ. Xét Đ-S: tổng, hiệu, độ dài, cos"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_2, ["[D12_C2_B3_28]-TF-M2. Cho tam giác. Xét Đ-S: Độ dài, tọa độ vectơ, R, S."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C2_B3_3 = QTreeWidgetItem(L12_C2_B3, ["Trả lời ngắn"])
                L12_C2_B3_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C2_B3_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_09]-TL-M2. Cho h.hộp chữ nhật được tọa độ hóa. Tìm tọa độ điểm hoặc vectơ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_11]-TL-M2. Tìm vectơ x để m.vec(a)+nvec(x)=p.vec(b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_13]-TL-M2. Cho hai điểm A,B. Tìm tọa độ C để AC nhận B làm trung điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_14]-TL-M2. Cho hai véctơ. Tìm m để 2 vectơ vuông góc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_21]-TL-M3. Cho A,B,C(a;y_C;b),G(x_G;c;z_G). Tính a+b+c."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_22]-TL-M3. Tìm m để ba điểm lập thành tam giác vuông"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C2_B3_3, ["[D12_C2_B3_23]-TL-M3. Bài toán liên quan trung điểm"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Lớp 12 - Chương 3 - Các số đặc trưng đo mức độ phân tán cho mẫu số liệu ghép nhóm            

                L12_C3 = QTreeWidgetItem(L12, ["Chương 3 - Các số đặc trưng đo mức độ phân tán cho mẫu số liệu ghép nhóm"])
                L12_C3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #Bài 1
                L12_C3_B1 = QTreeWidgetItem(L12_C3, ["Bài 1 - Khoảng biến thiên và khoảng tứ phân vị của mẫu số liệu ghép nhóm"])
                L12_C3_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B1_1 = QTreeWidgetItem(L12_C3_B1, ["Trắc nghiệm"])
                L12_C3_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B1_1, ["[D12_C3_B1_01]-M1. Tính khoảng biến thiên của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B1_1, ["[D12_C3_B1_02]-M3. Tính khoảng tứ phân vị của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B1_2 = QTreeWidgetItem(L12_C3_B1, ["Đúng-Sai"])
                L12_C3_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B1_2, ["[D12_C3_B1_04]-TF-M3. Xét Đ-S: Khoảng biến thiên, Q1, Q3, khoảng tứ phân vị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B1_3 = QTreeWidgetItem(L12_C3_B1, ["Trả lời ngắn"])
                L12_C3_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B1_3, ["[D12_C3_B1_03]-TL-M3. Tính khoảng tứ phân vị của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                #Bài 2
                L12_C3_B2 = QTreeWidgetItem(L12_C3, ["Bài 2 - Phương sai và độ lệch chuẩn của mẫu số liệu ghép nhóm"])
                L12_C3_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B2_1 = QTreeWidgetItem(L12_C3_B2, ["Trắc nghiệm"])
                L12_C3_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B2_1, ["[D12_C3_B2_01]-M2. Tính phương sai của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B2_1, ["[D12_C3_B2_02]-M2. Tính độ lệch chuẩn của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B2_2 = QTreeWidgetItem(L12_C3_B2, ["Đúng-Sai"])
                L12_C3_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B2_2, ["[D12_C3_B2_05]-TF-M3. Xét Đ-S: số trung bình, phương sai, độ lệch chuẩn, khoảng tứ phân vị."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C3_B2_3 = QTreeWidgetItem(L12_C3_B2, ["Trả lời ngắn"])
                L12_C3_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C3_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B2_3, ["[D12_C3_B2_03]-TL-M2. Tính phương sai của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C3_B2_3, ["[D12_C3_B2_04]-TL-M2. Tính độ lệch chuẩn của mẫu số liệu ghép nhóm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Chương 4 - Nguyên hàm tích phân
                L12_C4 = QTreeWidgetItem(L12, ["Chương 4 - Nguyên Hàm - Tích phân"])
                L12_C4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B1 = QTreeWidgetItem(L12_C4, ["Bài 1 - Nguyên hàm cơ bản"])
                L12_C4_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_01]-M1. Tìm nguyên hàm đa thức."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm của hàm số đa thức.')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_06]-M1. Tìm nguyên hàm đa thức + a/x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)             
              

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_02]-M2. Tìm nguyên hàm của hàm số đa thức thỏa mãn F(a)=b."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm của hàm số đa thức thỏa mãn F(a)=b.')        

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_04]-M1. Tìm nguyên hàm asinu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm chứa sinu, cosu.')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_05]-M1. Tìm nguyên hàm acosu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm chứa sinu, cosu.')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_03]-M2. Tìm nguyên hàm asinu + bcosu."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm chứa sinu, cosu.')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_07]-M1. Tìm nguyên hàm m/(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm m/(ax+b).')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_08]. Tìm nguyên hàm (mx+n)/(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm (mx+n)/(ax+b).')

                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_09]-M2. Tìm nguyên hàm (mx^2+nx+p)/(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                
                item = QTreeWidgetItem(L12_C4_B1, ["[D12_C4_B1_10]-M2. Tìm nguyên hàm e^(ax+b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B2 = QTreeWidgetItem(L12_C4, ["Bài 2 - Nguyên hàm đổi biến"])
                L12_C4_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B2, ["[D12_C4_B2_01]-M2. Nguyên hàm đổi biến chứa căn(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Nguyên hàm đổi biến chứa căn(ax+b)')

                item = QTreeWidgetItem(L12_C4_B2, ["[D12_C4_B2_02]-M2. Nguyên hàm đổi biến chứa căn(ax^2+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Nguyên hàm đổi biến chứa căn(ax^2+b)')

                item = QTreeWidgetItem(L12_C4_B2, ["[D12_C4_B2_04]-M2. Nguyên hàm đổi biến (acosx+b)sinx."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Nguyên hàm đổi biến (acosx+b)sinx.')

                item = QTreeWidgetItem(L12_C4_B2, ["[D12_C4_B2_03]-M2. Nguyên hàm đổi biến sinx/căn(acosx+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Nguyên hàm đổi biến sinx/căn(acosx+b)')      

                item = QTreeWidgetItem(L12_C4_B2, ["[D12_C4_B2_05]-M2. Nguyên hàm đổi biến (alnx + b)/x."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Nguyên hàm đổi biến (alnx + b)/x.')

                L12_C4_B3 = QTreeWidgetItem(L12_C4, ["Bài 3 - Nguyên hàm từng phần"])
                L12_C4_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B3, ["[D12_C4_B3_01]-M2. Tìm nguyên hàm P(x).sin hoặc P(x).cos"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm P(x).sin hoặc P(x).cos')
                

                item = QTreeWidgetItem(L12_C4_B3, ["[D12_C4_B3_02]-M2. Tìm nguyên hàm P(x).e^(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                #item.setToolTip(0, 'Tìm nguyên hàm P(x).e^(ax+b)')

                item = QTreeWidgetItem(L12_C4_B3, ["[D12_C4_B3_03]-M2. Tìm nguyên hàm P(x).ln(ax)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B4 = QTreeWidgetItem(L12_C4, ["Bài 4 - Tích phân"])
                L12_C4_B4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B4_1 = QTreeWidgetItem(L12_C4_B4, ["4.1 Tích phân lý thuyết"])
                L12_C4_B4_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B4_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_01]-M1. Cho F(a) và F(b). Tính tích phân từ a đến b."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_05]-M2. Cho tích phân f và F(a). Tính F(b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_04]-M2. Cho tích phân f. Tính tích phân (m.f+n)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_02]-M1. Cho 2 tích phân của f và g. Tính tích phân (m.f+n.g)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_06]-M1. Cho f(a) và f(b). Tính tích phân f'(x)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_07]-M2. Cho tích phân f'(x) và f(a). Tính f(b)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_03]-M1. Tính tổng tích phân trên 2 đoạn nối tiếp."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_08]-M2. Cho tphan_a^c(f) và tphan_a^b(f). Tính tphan_b^c(f), a<b<c."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_1, ["[D12_C4_B4_09]-M2. Cho tphan_a^d(f) và tphan_b^c(f). Tính tphan_a^b(f) + tphan_c^d(f), a<b<c<d."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B4_2 = QTreeWidgetItem(L12_C4_B4, ["4.1 Tích phân đa thức, phân thức"])
                L12_C4_B4_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B4_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_2, ["[D12_C4_B4_12]-M1. Tính tích phân của đa thức"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_2, ["[D12_C4_B4_10]-M2. Tính tích phân của m/(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_2, ["[D12_C4_B4_11]-M2. Tính tích phân của (mx+n)/(ax+b)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B4_3 = QTreeWidgetItem(L12_C4_B4, ["4.3 Tích phân hàm số lượng giác"])
                L12_C4_B4_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B4_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B4_3, ["[D12_C4_B4_13]-M2. Tính tích phân của asin(bx)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)


                L12_C4_B5 = QTreeWidgetItem(L12_C4, ["Bài 5 - Ứng dụng của tích phân"])
                L12_C4_B5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B5_1 = QTreeWidgetItem(L12_C4_B5, ["4.5.1. Diện tích hình phẳng"])
                L12_C4_B5_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B5_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_1, ["[D12_C4_B5_03]-M1. Tìm công thức tính diện tích từ hình vẽ có 1 đồ thị"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_1, ["[D12_C4_B5_01]-M1. Diện tích hình phẳng: y=f(x),Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_1, ["[D12_C4_B5_02]-M2. Diện tích hình phẳng: y=f(x),y=g(x),x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_1, ["[D12_C4_B5_04]-M2. Diện tích hình phẳng: y=f(x),y=g(x)"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B5_2 = QTreeWidgetItem(L12_C4_B5, ["4.5.2. Thể tích khối tròn xoay"])
                L12_C4_B5_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B5_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_05]-M2. V_khối tròn xoay: y=f(x),Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_06]-M2. V_khối tròn xoay: y=ax+b,Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_07]-M2. V_khối tròn xoay: y=ax^2+bx+c,Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_08]-M2. V_khối tròn xoay: y=căn(ax+b),Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)            
        

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_09]-M3. V_khối tròn xoay: y=ax.e^(bx),Ox,x=a,x=b"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_10]-M2. V_khối tròn xoay: y=ax^2+bx+c và Ox"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_14]-M2. Thể tích vật thể có thiết diện là hình chữ nhật"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_15]-M2. Thể tích vật thể có thiết diện là hình vuông"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_2, ["[D12_C4_B5_16]-M2. Thể tích vật thể có thiết diện là hình tam giác đều"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C4_B5_3 = QTreeWidgetItem(L12_C4_B5, ["4.5.3. Bài toán chuyển động"])
                L12_C4_B5_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C4_B5_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_3, ["[D12_C4_B5_11]-M2. Cho hàm số vận tốc. Tính quãng đường đi được từ t1 đến t2"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_3, ["[D12_C4_B5_12]-M2. Xe chạy chậm dần đều. Tính quãng đường đi được đến khi dừng hẳn"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C4_B5_3, ["[D12_C4_B5_13]-M3. Xe tăng tốc với gia tốc. Tính quãng đường đi được trong khoảng thời gian"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)




        #Chương 5 - Phương trình mặt phẳng, đường thẳng, mặt cầu trong không gian
                L12_C5 = QTreeWidgetItem(L12, ["Chương 5 - Phương trình mặt phẳng, đường thẳng, mặt cầu trong không gian"])
                L12_C5.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B1 = QTreeWidgetItem(L12_C5, ["Bài 1 - Phương trình mặt phẳng"])
                L12_C5_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B1_1 = QTreeWidgetItem(L12_C5_B1, ["5.1.1. Đọc thông tin VTPT, điểm thuộc mặt phẳng"])
                L12_C5_B1_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B1_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_1, ["[D12_C5_B1_14]-M1. Cho PTMP, đọc véctơ pháp tuyến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_1, ["[D12_C5_B1_15]-M2. Cho PTMP, tìm điểm thuộc mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_1, ["[D12_C5_B1_16]-M2. Cho PTMP, tìm điểm không thuộc mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B1_2 = QTreeWidgetItem(L12_C5_B1, ["5.1.2. Viết phương trình mặt phẳng"])
                L12_C5_B1_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B1_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_01]-M2. Viết PTMP qua điểm và có véctơ pháp tuyến."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_02]-M2. Viết PTMP qua điểm và vuông góc với đường thẳng qua 2 điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_03]-M2. Viết PTMP qua điểm và vuông góc trục tọa độ."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_04]-M2. Viết PTMP qua điểm và song song với mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_05]-M3.  Viết PTMP trung trực của đoạn AB."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_06]-M3. Viết PTMP đi qua 3 điểm A,B,C."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_2, ["[D12_C5_B1_17]-M1.  Viết PTMP đoạn chắn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B1_3 = QTreeWidgetItem(L12_C5_B1, ["5.1.3. Vị trí tương đối giữa các mặt phẳng"])
                L12_C5_B1_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B1_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_3, ["[D12_C5_B1_08]-M2. Xét vị trí giữa 2 mặt phẳng tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_3, ["[D12_C5_B1_09]-M2. Xét vị trí giữa 2 mặt phẳng. KQ: trùng nhau"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_3, ["[D12_C5_B1_10]-M2. Xét vị trí giữa 2 mặt phẳng. KQ: song song"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_3, ["[D12_C5_B1_11]-M2. Xét vị trí giữa 2 mặt phẳng. KQ: vuông góc"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_3, ["[D12_C5_B1_12]-M2. Xét vị trí giữa 2 mặt phẳng. KQ: cắt và không vuông góc"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B1_4 = QTreeWidgetItem(L12_C5_B1, ["5.1.4. Khoảng cách từ điểm đến mặt phẳng"])
                L12_C5_B1_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B1_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_4, ["[D12_C5_B1_07]-M1. Tính khoảng cách từ điểm đến mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B1_4, ["[D12_C5_B1_13]-M2. Tính khoảng cách giữa 2 mặt phẳng song song."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

               

                L12_C5_B2 = QTreeWidgetItem(L12_C5, ["Bài 2 - Phương trình đường thẳng"])
                L12_C5_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_1 = QTreeWidgetItem(L12_C5_B2, ["2.1. Đọc thông tin VTCP, điểm thuộc đường thẳng"])
                L12_C5_B2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_1, ["[D12_C5_B2_13]-M2. Đọc VTCP từ phương trình tham số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_1, ["[D12_C5_B2_15]-M2. Đọc VTCP từ phương trình tham số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_1, ["[D12_C5_B2_14]-M2. Đọc tọa độ điểm từ phương trình tham số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_1, ["[D12_C5_B2_16]-M2. Đọc tọa độ điểm từ phương trình chính tắc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_2 = QTreeWidgetItem(L12_C5_B2, ["2.2. Viết phương trình đường thẳng"])
                L12_C5_B2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_2_1 = QTreeWidgetItem(L12_C5_B2_2, ["2.2.1. Viết phương trình chính tắc"])
                L12_C5_B2_2_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_2_1.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_01]-M1. Viết PT chính tắc: d qua điểm và có VTCP."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_02]-M2. Viết PT chính tắc: d qua điểm A và nhận véctơ BC làm VTCP."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_03]-M2. Viết PT chính tắc: d qua điểm và song song với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_04]-M2. Viết PT chính tắc: d qua điểm và vuông góc với mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_05]-M2. Viết PT chính tắc: d qua 2 điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_12]-M2. Viết PT chính tắ từ phương trình tham số."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_06]-M3. Viết PT chính tắc: d qua vuông góc d' và song song với (P)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_1, ["[D12_C5_B2_07]-M3. Viết PT chính tắc đường thẳng d và song song với (P) và (Q)."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_2_2 = QTreeWidgetItem(L12_C5_B2_2, ["2.2.2. Viết phương trình tham số"])
                L12_C5_B2_2_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_2_2.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_2, ["[D12_C5_B2_08]-M1. Viết PTTS đường thẳng d qua điểm và có véctơ chỉ phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_2, ["[D12_C5_B2_09]-M2. Viết PTTS đường thẳng d qua điểm A và nhận vectơ BC làm véctơ chỉ phương."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_2_2, ["[D12_C5_B2_10]-M2. Viết PTTS đường thẳng qua điểm và song song với đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                item = QTreeWidgetItem(L12_C5_B2_2_2, ["[D12_C5_B2_11]-M2. Viết PTTS đường thẳng từ phương trình chính tắc."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_3 = QTreeWidgetItem(L12_C5_B2, ["2.3. Vị trí tương đối giữa các đường thẳng, đường thẳng và mặt phẳng"])
                L12_C5_B2_3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_3, ["[D12_C5_B2_19]-M2. Xét vị trí 2 đường thẳng chính tắc. Kết quả là song song."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_3, ["[D12_C5_B2_20]-M2. Xét vị trí 2 đường thẳng chính tắc. Kết quả là trùng nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_3, ["[D12_C5_B2_21]-M2. Xét vị trí 2 đường thẳng chính tắc. Kết quả là cắt nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_3, ["[D12_C5_B2_22]-M2. Xét vị trí 2 đường thẳng chính tắc. Kết quả là chéo nhau."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_3, ["[D12_C5_B2_23]-M2. Xét vị trí 2 đường thẳng chính tắc. Kết quả tùy ý."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B2_4 = QTreeWidgetItem(L12_C5_B2, ["2.4. Các bài toán khác về đường thẳng"])
                L12_C5_B2_4.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B2_4.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_4, ["[D12_C5_B2_17]-M3. Tìm giao điểm của đường thẳng và mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_4, ["[D12_C5_B2_18]-M3. Tìm hình chiếu của điểm trên đường thẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_4, ["[D12_C5_B2_24]-M4. Tìm đường thẳng đi qua điểm cắt và vuông góc với đường thẳng khác."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_4, ["[D12_C5_B2_25]-M4. Tìm đường thẳng đi qua điểm, cắt đường thẳng khác và song song mặt phẳng."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B2_4, ["[D12_C5_B2_26]-M4. Viết PTMP qua A,B cách M một khoảng lớn nhất"])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B3 = QTreeWidgetItem(L12_C5, ["Bài 3 - Phương trình mặt cầu"])
                L12_C5_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                L12_C5_B3.setCheckState(0, Qt.CheckState.PartiallyChecked)

                L12_C5_B3_1 = QTreeWidgetItem(L12_C5_B3, ["5.3.1. Đúng-Sai"])
                L12_C5_B3_1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

                L12_C5_B3_2 = QTreeWidgetItem(L12_C5_B3, ["5.3.2. Trắc Nghiệm"])
                L12_C5_B3_2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_02]-M1. Đọc tọa độ tâm từ PTMC thu gọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_03]-M1. Đọc bán kính từ PTMC thu gọn."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_04]-M2. Đọc tọa độ tâm từ PTMC khai triển."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_05]-M2. Đọc bán kính từ PTMC khai triển."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_01]-M1. Viết phương trình mặt cầu có tâm và bán kính."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_06]-M2. Viết phương trình mặt cầu có tâm và đường kính."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_07]-M3. Viết phương trình mặt cầu có đường kính AB."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

                item = QTreeWidgetItem(L12_C5_B3_2, ["[D12_C5_B3_08]-M2. Viết phương trình mặt cầu có tâm và đi qua điểm."])
                item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                item.setCheckState(0, Qt.CheckState.PartiallyChecked)

        #Chương 6 - Một số yếu tố xác suất
                L12_C6 = QTreeWidgetItem(L12, ["Chương 6 - Một số yếu tố xác suất"])
                L12_C6.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

                L12_C6_B1 = QTreeWidgetItem(L12_C6, ["Bài 1 - Xác suất có điều kiện"])
                L12_C6_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

                L12_C6_B2 = QTreeWidgetItem(L12_C6, ["Bài 2 - Công thức xác suất toàn phần, xác suất Bayes"])
                L12_C6_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)

        #Chương 7 - Mặt nón - Mặt trụ - Mặt cầu
                # L12_C7 = QTreeWidgetItem(L12, ["Chương 7 - Mặt nón - Mặt trụ - Mặt cầu"])
                # L12_C7.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L12_C7.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                # L12_C7_B1 = QTreeWidgetItem(L12_C7, ["Bài 1 - Mặt Nón"])
                # L12_C7_B1.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L12_C7_B1.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                # item = QTreeWidgetItem(L12_C7_B1, ["[D12_C7_B1_01]-M1. Tính diện tích xung quanh hình nón biết r và l"])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked) 
                # #item.setToolTip(0, 'Tính diện tích xung quanh hình nón biết r và l')
            
                # item = QTreeWidgetItem(L12_C7_B1, ["[D12_C7_B1_02]-M1. Tính thể tích khối nón biết r và h"])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Tính thể tích khối nón biết r và h')

                # item = QTreeWidgetItem(L12_C7_B1, ["[D12_C7_B1_03]-M2. Cho hình nón có chu vi và đường sinh. Tính chiều cao h."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho hình nón có chu vi và đường sinh. Tính chiều cao h')

                # item = QTreeWidgetItem(L12_C7_B1, ["[D12_C7_B1_04]-M2. Cho hình nón có diện tích đáy và đường sinh. Tính chiều cao h."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)  
                # #item.setToolTip(0, 'Cho hình nón có diện tích đáy và đường sinh. Tính chiều cao h') 


                # L12_C7_B2 = QTreeWidgetItem(L12_C7, ["Bài 2 - Mặt Trụ"])
                # L12_C7_B2.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L12_C7_B2.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                # item = QTreeWidgetItem(L12_C7_B2, ["[D12_C7_B2_01]-M1. Cho hình trụ có bán kính r và đường sinh l. Tính S_xq."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho hình trụ có bán kính r và đường sinh l. Tính S_xq')

                # item = QTreeWidgetItem(L12_C7_B2, ["[D12_C7_B2_02]-M1. Cho hình trụ có bán kính r và đường sinh l. Tính V."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho hình trụ có bán kính r và đường sinh l. Tính V.')

                # item = QTreeWidgetItem(L12_C7_B2, ["[D12_C7_B2_03]-M3. Cho hình trụ có thiết diện qua trục là hình vuông. Tính S_tp."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho hình trụ có thiết diện qua trục là hình vuông. Tính S_tp')

                # item = QTreeWidgetItem(L12_C7_B2, ["[D12_C7_B2_04]-M3. Cho hình trụ có thiết diện qua trục là hình chữ nhật. Tính S_tp."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho hình trụ có thiết diện qua trục là hình chữ nhật. Tính S_tp')


                # L12_C7_B3 = QTreeWidgetItem(L12_C7, ["Bài 3 - Mặt Cầu"])
                # L12_C7_B3.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # L12_C7_B3.setCheckState(0, Qt.CheckState.PartiallyChecked) 

                # item = QTreeWidgetItem(L12_C7_B3, ["[D12_C7_B3_01]-M1. Cho mặt cầu có bán kính r. Tính S_mc."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)
                # #item.setToolTip(0, 'Cho mặt cầu có bán kính r. Tính S_mc')

                # item = QTreeWidgetItem(L12_C7_B3, ["[D12_C7_B3_02]-M2. Cho mặt cầu có thể tích. Tính bán kính R."])
                # item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
                # item.setCheckState(0, Qt.CheckState.PartiallyChecked)                

                #Check_Update
                #self.check_update()

                #Thêm dạng toán sau khi chọn vào cây thư mục                          
                self.treeWidget.itemChanged.connect(self.handleItemChanged)
                self.treeWidget.itemChanged.connect(self.get_chi_muc)

        
#begin
        def tao_de(self):
                try:
                        if license.kiemtra_banquyen():                
                            self.text_taode.clear()
                            self.save_thongtin_dethi()                            
                            current_directory = os.path.dirname(os.path.abspath(__file__))
                            hinh_folder_path = os.path.join(current_directory, 'HINH VE')                            
                            os.chdir(hinh_folder_path)                            
                            files_to_delete = os.listdir()
                            for file_to_delete in files_to_delete:
                                file_path = os.path.join(hinh_folder_path, file_to_delete)
                                os.remove(file_path)                     

                            max_value = 100
                            self.progress_bar.setMaximum(max_value)    
                            self.timer = QTimer(parent=self.tab_taode)
                            self.timer.timeout.connect(self.updateProgressBar)
                            self.timer.start(100)

                            #Đặt tên thư mục chứa file                            
                            if self.combo_taode.currentText() in ["Tạo đề Word - Equation", "Tạo đề Word - MathType", "Tạo đề Latex - PDF"]:
                                #desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
                                current_directory = os.path.dirname(os.path.abspath(__file__))
                                doc_folder_path = os.path.join(current_directory, 'DOC')
                                folder = QFileDialog.getExistingDirectory(None, "Chọn thư mục xuất đề", doc_folder_path)
                                if folder:
                                        name_thu_muc=f"{folder}"
                                else:
                                        current_directory = os.path.dirname(os.path.abspath(__file__))
                                        doc_folder_path = os.path.join(current_directory, 'DOC')
                                        name_thu_muc=f"De_{datetime.now().strftime("%d-%m__%H-%M-%S")}"
                                        new_folder_path = os.path.join(doc_folder_path, name_thu_muc)
                                        if not os.path.exists(new_folder_path):
                                                os.makedirs(new_folder_path)
                                        name_thu_muc=new_folder_path
                                print(name_thu_muc)         
                            
                            #Tạo list mã đề ngẫu nhiên
                            unique_digits_set = set()
                            while len(unique_digits_set) < self.spin_soluong_de.value():
                                digit = random.randint(11, 99)  
                                unique_digits_set.add(digit)
                            t_random=list(unique_digits_set)                            
                            #begin
                            code_bang_dap_an=""
                            list_ma_de=[]
                            sum_toanbo=self.thongke()

                            #Tạo list dạng toán                            
                            list_tonghop=""
                            list_tonghop_HDG=""
                            list_dapan_word=[]

                            for j in range(self.spin_soluong_de.value()):
                                self.text_taode.setText("")
                                self.text_taode_HDG.setText("")

                                list_tracnghiem, list_tracnghiem_HDG=[],[]
                                list_dungsai, list_dungsai_HDG=[],[]
                                list_tuluan, list_tuluan_HDG=[],[]
                                list_dapan_TN, list_dapan_TF, list_dapan_TL=[], [], []
                                

                                list_noi_dung, list_noi_dung_HDG="",""                      
                                ghep_tracnghiem, ghep_tracnghiem_HDG="",""
                                ghep_dungsai, ghep_dungsai_HDG="", ""
                                ghep_tuluan, ghep_tuluan_HDG="",""
                                ghep_dapan_TN, ghep_dapan_TF, ghep_dapan_TL= "","",""                                       
                                    

                                #Thiết lập mã đề tự nhập
                                text_nhapmade=self.label_nhapmade.toPlainText()
                                if len(text_nhapmade)>0:                                    
                                    text_nhapmade=text_nhapmade.split(',')
                                    ma_de=text_nhapmade[j]
                                    name_de=f"{ma_de}"
                                else:
                                #Thiết lập mã đề ngẫu nhiên hoặc thứ tự
                                    t=j+1
                                    if j>9: t=int(j/10)                              

                                    if self.checkbox_made_random.isChecked():                                
                                        
                                        ma_de= t*100 + t_random[j]
                                        if ma_de>1000:
                                            ma_de= t*10 + t_random[j]
                                        name_de=f"{ma_de}"
                                    else:
                                        ma_de = j+1                                        
                                        if ma_de>9: 
                                            name_de=f"0{ma_de}"
                                        else:
                                            name_de=f"00{ma_de}"
                                    
                                list_ma_de.append(name_de)

                                #Tao_qrcode
                                st_qrcode=""
                                                        
                                socau_daxuli=0
                                self.label_dangxuli.setText("Chương trình đang xử lý. Vui lòng đợi...")                                        

                                if self.combo_taode.currentText() == "Tạo đề Latex - PDF":                                                  
           
                                    list_noi_dung+=f"\\tieude{{{name_de}}}\n"\
                                    f"\\chantrang{{\\pageref{{LastPage}}}}{{{name_de}}}\n"\
                                    f"\\setcounter{{page}}{{1}}\n"                                


                                if self.combo_taode.currentText() == "Tạo đề Word - Equation":

                                    list_noi_dung+=f"\\begin{{tabular}}{{cc}}\n"\
                                    f"{{\\bf {self.ten_sogd.toPlainText()}}} & {{\\bf {self.ten_kythi.toPlainText()}}}\\\\ \n"\
                                    f"{{\\bf {self.ten_truong.toPlainText()}}} & {{\\bf Môn: {self.ten_monthi.toPlainText()}}}\\\\  \n"\
                                    f"& {{\\bf Thời gian: {self.ten_thoigian.toPlainText()} phút}}\\\\ \n"\
                                    f"& {{\\bf Mã đề: {name_de}}}\\\\ \n"\
                                    f"\\end{{tabular}}\n\n"\
                                    f"{{Họ tên HS: ..................................................Số báo danh:..................................................}}\n\n\n"


                                    list_noi_dung_HDG+=f"\\begin{{tabular}}{{cc}}\n"\
                                    f"{{\\bf {self.ten_sogd.toPlainText()}}} & {{\\bf {self.ten_kythi.toPlainText()}}}\\\\ \n"\
                                    f"{{\\bf {self.ten_truong.toPlainText()}}} & {{\\bf Môn: {self.ten_monthi.toPlainText()}}}\\\\  \n"\
                                    f"& {{\\bf Thời gian: {self.ten_thoigian.toPlainText()} phút}}\\\\ \n"\
                                    f"& {{\\bf Mã đề: {name_de}}}\\\\ \n"\
                                    f"\\end{{tabular}}\n\n"\
                                    f"{{\\bf Họ tên HS: ..................................................Số báo danh:..................................................}}\n\n\n"
                                   

                                row_count = self.tableWidget.rowCount()
    
                                for i in range(row_count):
                                    dong_dangtoan= self.tableWidget.item(i, 0)                                    

                                    if dong_dangtoan is not None:
                                        if dong_dangtoan.text()!="": 
                                            dong_dangtoan=dong_dangtoan.text()
                                        else:
                                            dong_dangtoan=""
                                    else:
                                        dong_dangtoan=""

                                    if dong_dangtoan[0]!="[":
                                        #self.text_taode.append(f"{dong_dangtoan}\n\n")                                        
                                        loai_cau=""
                                        muc_do=""
                                        so_cau=""

                                    if len(dong_dangtoan)>=14:
                                        if dong_dangtoan[14]=="-":                                
                                            dang_toan=dong_dangtoan[0:14]
                                        else:
                                            dang_toan=""
                                        if dong_dangtoan[14]=="-":   
                                            loai_cau= self.tableWidget.item(i, 1)                                            
                                            muc_do = self.tableWidget.item(i, 2)  
                                            so_cau = self.tableWidget.item(i, 3)

                                            if loai_cau is not None:
                                                if loai_cau.text() != "":
                                                    loai_cau = loai_cau.text()
                                                else:
                                                    loai_cau = ""
                                            else:
                                                loai_cau = ""

                                            if muc_do is not None:
                                                if muc_do.text()!="": 
                                                    muc_do=muc_do.text()
                                                else:
                                                    muc_do=""
                                            else:
                                                muc_do=""

                                            if so_cau is not None:
                                                if so_cau.text()!="":
                                                    so_cau=int(so_cau.text())
                                                else:
                                                    so_cau=""
                                            else:
                                                    so_cau=""                                                                                                       

                                            for socau in range(so_cau):
                                            ######### Toán 12_ Chương 1 _ KHẢO SÁT HÀM SỐ  ########
                                         #Bài 1. Sự đồng biến và nghịch biến của hàm số 
                                                #[D12_C1_B1_01]. Cho bảng biến thiên. Tìm khoảng đồng biến, nghịch biến
                                                if dang_toan == "[D12_C1_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_01()                               
                                                    

                                                #[D12_C1_B1_02]. Cho hàm số bậc 3. Tìm khoảng đồng biến, nghịch biến
                                                if dang_toan == "[D12_C1_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_02()                                     
                                                    

                                                #[D12_C1_B1_03]. Cho đồ thị bậc 4. Tìm khoảng đồng biến, nghịch biến
                                                if dang_toan == "[D12_C1_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_03()

                                                #[D12_C1_B1_04]. Cho hàm số y =(ax+b)/(cx+d). Tìm khoảng đồng biến, nghịch biến
                                                if dang_toan == "[D12_C1_B1_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_04()

                                                #[D12_C1_B1_04]. Cho hàm số y=(ax^2+bx+c)/(dx+e). Tìm khoảng đơn điệu.
                                                if dang_toan == "[D12_C1_B1_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_05()

                                                #[D12_C1_B1_06]-M2. Cho đồ thị bậc 3. Tìm khoảng đơn điệu.
                                                if dang_toan == "[D12_C1_B1_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_06()

                                                #[D12_C1_B1_07]-M2. Cho hàm số f'(x). Tìm khoảng đơn điệu.
                                                if dang_toan == "[D12_C1_B1_07]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_07()

                                                #[D12_C1_B1_08]-M2. Cho bảng xét dấu f'(x). Tìm khoảng đơn điệu.
                                                if dang_toan == "[D12_C1_B1_08]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_08()

                                                #[D12_C1_B1_09]-M2. Tìm m để hàm số bậc 3 đơn điệu trên R.
                                                if dang_toan == "[D12_C1_B1_09]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_09()

                                                #[D12_C1_B1_10]-M3. Cho đồ thị f'(x). Tìm khoảng đơn điệu f(x)
                                                if dang_toan == "[D12_C1_B1_10]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_10()

                                                #[D12_C1_B1_11]-TF-M2. Cho bảng xét dấu f'(x). Xét đúng-sai:đồng biến, nghịch biến, so sánh giá trị.     
                                                if dang_toan == "[D12_C1_B1_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_11()

                                                #[D12_C1_B1_12]-TF-M2. Cho bảng xét dấu f'(x). Xét đúng-sai:đơn điệu, so sánh, cực trị.    
                                                if dang_toan == "[D12_C1_B1_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_12()

                                                #[D12_C1_B1_13]-TF-M2. Cho BBT. Xét đúng-sai: đơn điệu, so sánh giá trị, cực trị.     
                                                if dang_toan == "[D12_C1_B1_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_13()

                                                #[D12_C1_B1_14]-TF-M2. Cho BBT. Xét đúng-sai: đơn điệu, so sánh giá trị.     
                                                if dang_toan == "[D12_C1_B1_14]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_14()           

                                        #Cực trị của hàm số 
                                                #[D12_C1_B1_15]. Cho bảng biến thiên bậc 3. Tìm điểm cực trị."
                                                if dang_toan == "[D12_C1_B1_15]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_15()                                     

                                                #[D12_C1_B1_16]. Cho bảng biến thiên bậc 3. Tìm giá trị cực trị."
                                                if dang_toan == "[D12_C1_B1_16]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_16()

                                                #[D12_C1_B1_17]-M2. Cho hàm bậc 3. Tìm điểm cực trị."
                                                if dang_toan == "[D12_C1_B1_17]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_17()

                                                #[D12_C1_B1_18]-M2. Cho hàm số bậc 3. Tìm giá trị cực trị.
                                                if dang_toan == "[D12_C1_B1_18]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_18()

                                                #[D12_C1_B1_19]-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Tìm điểm cực trị. 
                                                if dang_toan == "[D12_C1_B1_19]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_19()

                                                #[D12_C1_B1_20]-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Tìm giá trị cực trị. 
                                                if dang_toan == "[D12_C1_B1_20]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_20()

                                                #[D12_C1_B1_21]-M2. Tìm khoảng đơn điệu của các hàm số khác: y=căn(u), y=e^u, y=ax/(x^2+b)
                                                if dang_toan == "[D12_C1_B1_21]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_21()

                                                #[D12_C1_B1_22]-M2. Cho hàm số bậc 3 có hai điểm cực trị x_1,x_2. Tính P=ax_1 + bx_2. 
                                                if dang_toan == "[D12_C1_B1_22]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_22()

                                                #[D12_C1_B1_23]-M2. Cho hàm số bậc 3 có giá trị cực trị y_1,y_2. Tính P=ay_1 + by_2. 
                                                if dang_toan == "[D12_C1_B1_23]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_23()

                                                #[D12_C1_B1_24]-M2. Cho y=(ax^2+bx+c)/(dx+e) có điểm cực trị x_1,x_2. Tính P=ax_1 + bx_2.
                                                if dang_toan == "[D12_C1_B1_24]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_24() 

                                                #[D12_C1_B1_25]-M2. Cho y=(ax^2+bx+c)/(dx+e) có giá trị cực trị y_1,y_2. Tính P=ay_1 + by_2.
                                                if dang_toan == "[D12_C1_B1_25]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_25()

                                                #[D12_C1_B1_26]-TF-M2. Cho hàm số bậc 3. Xét Đ-S: TXD, y', đơn điệu, cực trị
                                                if dang_toan == "[D12_C1_B1_26]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_26() 

                                                #[D12_C1_B1_27]-TF-M2. Cho y=(ax^2+bx+c)/(dx+e). Xét Đ-S: TXD, y', đơn điệu, cực trị
                                                if dang_toan == "[D12_C1_B1_27]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_27()

                                                #[D12_C1_B1_28]-M2. Cho dấu f'(x) trên khoảng (a,b) và (c,d). Tìm khẳng định đúng?
                                                if dang_toan == "[D12_C1_B1_28]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_28()
                                                
                                                #[D12_C1_B1_29]-TL-M3. Cho f(x)=(ax+b)/(x+m). Tìm số giá trị nguyên của m để hàm số đồng biến,nghịch biến trên K.
                                                if dang_toan == "[D12_C1_B1_29]": 
                                                     debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_29()

                                                #[D12_C1_B1_30]-TF-M3. Cho f(x)=(ax+b)/(x+m). Xét Đ-S: đạo hàm, cực trị, đồng biến, nghịch biến
                                                if dang_toan == "[D12_C1_B1_30]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_30()

                                                #[D12_C1_B1_31]-TF-M3. Cho BXD f'(x). Xét đúng-sai:đơn điệu và cực trị của f(x), đơn điệu và cực trị của f(ax+b).
                                                if dang_toan == "[D12_C1_B1_31]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B1_31()

                                                #[D12_C1_B1_32]-TL-M3. Tìm m để hàm số đồng biến, nghịch biến trên khoảng.
                                                if dang_toan == "[D12_C1_B1_32]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_32()

                                                #[D12_C1_B1_33]-TL-M3. Tìm m để y=(ax+b)/(x+m) đơn điệu trên các khoảng.
                                                if dang_toan == "[D12_C1_B1_33]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B1_33()

                                                #[D12_C1_B1_34]-M2. Cho hàm số f'(x). Tìm số cực trị của y=f(x)
                                                if dang_toan == "[D12_C1_B1_34]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_34()

                                                #[D12_C1_B1_35]-M2. Cho đồ thị f'(x). Tìm điểm cực trị
                                                if dang_toan == "[D12_C1_B1_35]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B1_35()                                        

                                        #Bài 2. GIÁ TRỊ LỚN NHẤT  - GIÁ TRỊ NHỎ NHẤT
                                                #[D12_C1_B2_01]-M2. Tìm GTLN-GTNN của hàm bậc 3 trên đoạn
                                                if dang_toan == "[D12_C1_B2_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_01()                                    

                                                #[D12_C1_B2_02]-M2. Tìm GTLN-GTNN của hàm y=(ax+b)/(cx+d)
                                                if dang_toan == "[D12_C1_B2_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_02()

                                                #[D12_C1_B2_03]-M2. Tìm GTLN-GTNN của hàm bậc 3 trên khoảng
                                                if dang_toan == "[D12_C1_B2_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_03()

                                                #[D12_C1_B2_04]-M2. Tìm GTLN-GTNN của hàm y=(ax+b)/(cx+d) trên khoảng
                                                if dang_toan == "[D12_C1_B2_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_04()

                                                #[D12_C1_B2_05]-M2. Tìm GTLN-GTNN của hàm y=ax+b/x trên khoảng
                                                if dang_toan == "[D12_C1_B2_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_05()

                                                #[D12_C1_B2_06]-M2. Tìm GTLN-GTNN của hàm y=(ax^2+bx+c)/(dx+e) trên khoảng
                                                if dang_toan == "[D12_C1_B2_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_06()

                                                #[D12_C1_B2_07]-TF-M2. Cho hàm số bậc 3. Xét Đ-S: y', y'=0, f(x_0), Min-Max trên đoạn  
                                                if dang_toan == "[D12_C1_B2_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B2_07()

                                                #[D12_C1_B2_08]-TF-M2. Cho y=(ax+b)/(cx+d). Xét Đ-S: y', y'=0, f(x_0), Min-Max trên đoạn  
                                                if dang_toan == "[D12_C1_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B2_08()

                                                #[D12_C1_B2_09]-TF-M2. Cho hàm số bậc 3. Xét Đ-S: y', y'=0, f(x_0), Min-Max trên khoảng  
                                                if dang_toan == "[D12_C1_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B2_09()

                                                #[D12_C1_B2_10]-M3. Tìm GTLN-GTNN của y=e^x.(ax^2+bx+c) trên đoạn
                                                if dang_toan == "[D12_C1_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_10()

                                                #[D12_C1_B2_11]-M3. Tìm GTLN-GTNN của y=e^(ax^2+bx+c) trên đoạn
                                                if dang_toan == "[D12_C1_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_11()

                                                #[D12_C1_B2_12]-M3. Tìm GTLN-GTNN của y=ln(ax^2+bx+c) trên đoạn
                                                if dang_toan == "[D12_C1_B2_12]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_12()

                                                #[D12_C1_B2_13]-M3. Tìm GTLN-GTNN của y=căn(x-a) + căn(b-x)
                                                if dang_toan == "[D12_C1_B2_13]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_13()

                                                #[D12_C1_B2_14]-M3. Tìm GTLN-GTNN của y=b.x.căn(a^2-x^2)
                                                if dang_toan == "[D12_C1_B2_14]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_14()

                                                #[D12_C1_B2_15]-TL-M2. Cho hàm số. Tìm GTLN-GTNN trên đoạn
                                                if dang_toan == "[D12_C1_B2_15]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_15()

                                                #[D12_C1_B2_16]-TL-M3. Cho hàm số. Tìm GTLN-GTNN trên đoạn
                                                if dang_toan == "[D12_C1_B2_16]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_16()

                                                #[D12_C1_B2_17]-TL-M3. Cho hàm số. Tìm GTLN-GTNN trên khoảng, nửa khoảng
                                                if dang_toan == "[D12_C1_B2_17]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_17()

                                                #[D12_C1_B2_18]-TL-M3. Tìm m để y=(ax+b)/(cx+d) đạt GTLN(GTNN) bằng y_0.
                                                if dang_toan == "[D12_C1_B2_18]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_18()

                                                #[D12_C1_B2_19]-TL-M3. Tìm m để $y=|f(x)|+am$ đạt GTLN(GTNN) bằng y_0.
                                                if dang_toan == "[D12_C1_B2_19]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_19()

                                                #[D12_C1_B2_20]-TL-M3. Tìm m để y=(ax+b)/(cx+d) có min y + max y=T.
                                                if dang_toan == "[D12_C1_B2_20]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_20()

                                                #[D12_C1_B2_21]-TL-M3. Tìm vận tốc lớn nhất, nhỏ nhất của chuyển động
                                                if dang_toan == "[D12_C1_B2_21]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_21()

                                                #[D12_C1_B2_22]-TL-M3. Cắt hình chữ nhật từ nửa đường tròn. Tìm diện tích lớn nhất của hình chữ nhật.
                                                if dang_toan == "[D12_C1_B2_22]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_22()

                                                #[D12_C1_B2_23]-TL-M3. Cắt hình chữ nhật từ nửa đường tròn. Tìm diện tích lớn nhất của hình chữ nhật.
                                                if dang_toan == "[D12_C1_B2_23]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B2_23()

                                                #[D12_C1_B2_24]-M3. Cho đồ thị f(x).Tìm GTLN-GTNN
                                                if dang_toan == "[D12_C1_B2_24]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B2_24()
                                                    

                                        #Bài 3. ĐƯỜNG TIỆM CẬN
                                                #[D12_C1_B3_01]. Đọc đường tiệm cận của đồ thị hàm số y=(ax+b)/(cx+d)
                                                if dang_toan == "[D12_C1_B3_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_01()                                     

                                                #[D12_C1_B3_02]. Cho bảng biến thiên. Đọc số đường tiệm cận
                                                if dang_toan == "[D12_C1_B3_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_02()

                                                #[D12_C1_B3_03]-M1. Đọc đường tiệm cận đứng của đồ thị hàm số y=(ax^2+bx+c)/(dx+e)
                                                if dang_toan == "[D12_C1_B3_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_03()

                                                #[D12_C1_B3_04]-M2. Đọc đường tiệm cận xiên của đồ thị hàm số y=(ax^2+bx+c)/(dx+e)
                                                if dang_toan == "[D12_C1_B3_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_04()                                  
                                                
                                                #[D12_C1_B3_05]-M2. Tìm số đường tiệm cận đứng của các hàm số khác
                                                if dang_toan == "[D12_C1_B3_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_05()

                                                #[D12_C1_B3_06]-M2. Tìm số đường tiệm cận của các hàm số khác
                                                if dang_toan == "[D12_C1_B3_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_06()

                                                #[D12_C1_B3_07]-M2. Tìm đường tiệm cận từ giới hạn
                                                if dang_toan == "[D12_C1_B3_07]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_07()

                                                #[D12_C1_B3_08]-TF-M2. Cho hàm số y=(ax+b)/(cx+d). Xét Đ-S: lim, tiệm cận đứng, tiệm cận ngang    
                                                if dang_toan == "[D12_C1_B3_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_08()

                                                #[D12_C1_B3_09]-TF-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Xét Đ-S: lim, tiệm cận đứng, tiệm cận ngang    
                                                if dang_toan == "[D12_C1_B3_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_09()

                                                #[D12_C1_B3_10]-M3. Tìm m để y=(ax+b)/(cx+d) có tiệm cận 
                                                if dang_toan == "[D12_C1_B3_10]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B3_10()

                                                #[D12_C1_B3_11]-TF-M2.  Cho BBT của y=f(x). Xét Đ-S: Tìm TCĐ, Tìm TCN, Số TCD+TCN, Số TCN(TCD) của y= m/(f(x)+n)   
                                                if dang_toan == "[D12_C1_B3_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_11()

                                                #[D12_C1_B3_12]-TF-M2. Cho hàm số y=(ax+b)/(cx+d). Xét Đ-S: TCĐ, TCN, giao điểm của TC, HCN tạo bởi 2 tiệm cận   
                                                if dang_toan == "[D12_C1_B3_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_12()

                                                #[D12_C1_B3_13]-TF-M2. Cho limf. Xét Đ-S: TCĐ, TCN, Số TCD và TCN, TNC của y=(af+b)/(cf+d)  
                                                if dang_toan == "[D12_C1_B3_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_13()

                                                #[D12_C1_B3_14]-M2. Cho đường tiệm cận xiên. Tìm khẳng định đúng.
                                                if dang_toan == "[D12_C1_B3_14]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_14()

                                                #[D12_C1_B3_15]-M2. Cho giới hạn. Tìm tiệm cận xiên.
                                                if dang_toan == "[D12_C1_B3_15]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_15()

                                                #[D12_C1_B3_16]-M2. Đọc đường tiệm cận từ đồ thị hàm số y=(ax+b)/(cx+d)
                                                if dang_toan == "[D12_C1_B3_16]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B3_16()

                                                #[D12_C1_B3_17]-TF-M2. Cho y=(ax^2+bx+c)/(dx+e). Xét Đ-S: y',y'=0, tiệm cận đứng, tiệm cận xiên 
                                                if dang_toan == "[D12_C1_B3_17]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B3_17()

                                                #[D12_C1_B3_18]-TL-M2. Cho y=(ax^2+bx+c)/(dx+e) có TCX là y=ax+b. Tính ma+nb.
                                                if dang_toan == "[D12_C1_B3_18]": 
                                                   debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B3_18()





                                        #Bài 4. ĐỒ THỊ HÀM SỐ 
                                                #[D12_C1_B4_01]-M2. Cho đồ thị. Tìm hàm số bậc 3.
                                                if dang_toan == "[D12_C1_B4_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_01()                                     

                                                #[D12_C1_B4_02]-M2. Cho đồ thị. Tìm hàm số bậc 4.
                                                if dang_toan == "[D12_C1_B4_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_02()                                     

                                                #[D12_C1_B4_03]-M2. Cho đồ thị. Tìm hàm số y=(ax+b)/(cx+d).
                                                if dang_toan == "[D12_C1_B4_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_03()                                     

                                                #[D12_C1_B4_04]-M2. Cho bảng biến thiên. Tìm hàm số y=(ax+b)/(cx+d).
                                                if dang_toan == "[D12_C1_B4_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_04()

                                                #[D12_C1_B4_05]-M2. Cho đồ thị. Tìm hàm số y=(ax^2+bx+c)/(dx+e) tương ứng.
                                                if dang_toan == "[D12_C1_B4_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_05()

                                                #[D12_C1_B4_06]-TF-M2. Cho hàm bậc 3. Xét Đ-S: y', dấu y', BBT, Đồ thị
                                                if dang_toan == "[D12_C1_B4_06]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B4_06()

                                                #[D12_C1_B4_07]-TF-M2. Cho hàm số y=(ax+b)/(cx+d). Xét Đ-S: y', đơn điệu, BBT, đồ thị
                                                if dang_toan == "[D12_C1_B4_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B4_07()

                                                #[D12_C1_B4_08]-TF-M2. Cho hàm số y=(ax^2+bx+c)/(dx+e). Xét Đ-S: y', đơn điệu, BBT, đồ thị
                                                if dang_toan == "[D12_C1_B4_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C1.prt_34_L12_C1_B4_08()

                                                #[D12_C1_B4_09]-M2. Cho bảng biến thiên, tìm hàm số bậc 3.
                                                if dang_toan == "[D12_C1_B4_09]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_09()

                                                #[D12_C1_B4_10]. Cho hai hàm số. Tìm số giao điểm của 2 đồ thị.
                                                if dang_toan == "[D12_C1_B4_10]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_10()                                     
  
                                                #[D12_C1_B4_11]. Cho đồ thị. Tìm số nghiệm của phương trình f(x)=m
                                                if dang_toan == "[D12_C1_B4_11]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_11()                                     

                                                #[D12_C1_B4_12]. Tìm số giao điểm của đồ thị với trục Ox, Oy
                                                if dang_toan == "[D12_C1_B4_12]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_12()

                                                #[D12_C1_B4_13]. Cho hàm số đa thức. Tìm số giao điểm với trục Ox, Oy
                                                if dang_toan == "[D12_C1_B4_13]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_13()

                                                #[D12_C1_B4_14]. Cho y=(ax+b)/(cx+d). Tìm tọa độ tâm đối xứng.
                                                if dang_toan == "[D12_C1_B4_14]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_14()

                                                #[D12_C1_B4_15]. Cho y=(ax^2+bx+c)/(dx+e). Tìm tọa độ tâm đối xứng.
                                                if dang_toan == "[D12_C1_B4_15]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_15()

                                                #[D12_C1_B4_16]-M2. Cho y=ax^3+bx^2+cx+d. Tìm tọa độ tâm đối xứng.
                                                if dang_toan == "[D12_C1_B4_16]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C1.prt_34_L12_C1_B4_16()

                                                #[D12_C1_B4_17]-TL-M3. Cho đồ thị hàm bậc 3. Tính tổng các hệ số
                                                if dang_toan == "[D12_C1_B4_17]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B4_17()

                                                #[D12_C1_B4_18]-TL-M3. Cho đồ thị hàm y=(ax+b)/(cx+d). Tính tổng các hệ số
                                                if dang_toan == "[D12_C1_B4_18]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B4_18()

                                                #Toán 12 - Chương 1 - Bài 5------------------------------------------------->
                                                #[D12_C1_B5_01]-TL-M3. Bài toán thực tế tìm chi phí thấp nhất
                                                if dang_toan == "[D12_C1_B5_01]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B5_01()

                                                #[D12_C1_B5_02]-TL-M3. Cho hàm số chi phí. Tìm chi phí trung bình thấp nhất.
                                                if dang_toan == "[D12_C1_B5_02]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B5_02()

                                                #[D12_C1_B5_03]-TL-M3. Bài toán tìm số giếng dầu có thể khai thác thêm để sản lượng là lớn nhất
                                                if dang_toan == "[D12_C1_B5_03]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B5_03()

                                                #[D12_C1_B5_04]-TL-M3. Bài toán tìm tốc độ trung bình để chi phí xăng nhỏ nhất.
                                                if dang_toan == "[D12_C1_B5_04]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C1.prt_34_L12_C1_B5_04()

                                                    

                                    ######### Toán 12_ Chương 2 - TỌA ĐỘ VECTƠ TRONG KHÔNG GIAN ########
                                    #------------------------Toán 12 - Chương 2 - Bài 1------------------------------------------------------------->
                                                #[D12_C2_B1_01]-M2. Cho hình hộp. Tìm vectơ bằng vectơ cho trước
                                                if dang_toan == "[D12_C2_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_01()


                                                #[D12_C2_B1_02]-M2. Tìm đẳng thức đúng liên quan đến 3 điểm (phép cộng trừ)
                                                if dang_toan == "[D12_C2_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_02()

                                                #[D12_C2_B1_03]-M1. Nhận dạng quy tắc hình hộp
                                                if dang_toan == "[D12_C2_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_03()

                                                #[D12_C2_B1_04]-M1. Cho tứ diện. Tìm khẳng định đúng về quy tắc cộng trừ.
                                                if dang_toan == "[D12_C2_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_04()

                                                #[D12_C2_B1_05]-M1. Cho tứ diện. Tìm khẳng định về phép toán vectơ 
                                                if dang_toan == "[D12_C2_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_05()

                                                #[D12_C2_B1_06]-M1. Cho hình lập phương. Tính độ dài vectơ.
                                                if dang_toan == "[D12_C2_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_06()

                                                #[D12_C2_B1_07]-TF-M2. Cho hình hộp. Xét Đ-S: phép toán về vectơ, quy tắc hình bình hành, hình hộp, độ dài vectơ 
                                                if dang_toan == "[D12_C2_B1_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_07()

                                                #[D12_C2_B1_08]-TF-M2. Cho tứ diện đều. Xét Đ-S: các phép toán về vectơ, tích vô hướng
                                                if dang_toan == "[D12_C2_B1_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_08()

                                                #[D12_C2_B1_09]-TL-M3. Cho 3 lực đôi một vuông góc. Tính tổng hợp lực.
                                                if dang_toan == "[D12_C2_B1_09]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_09()

                                                #[D12_C2_B1_10]-TF-M2. Cho tứ diện. Xét Đ-S: Tích vô hướng, phép toán vectơ,  độ dài.
                                                if dang_toan == "[D12_C2_B1_10]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_10()

                                                #[D12_C2_B1_11]-M2. Cho hình chóp có đáy h.b.h. Tìm khẳng định sai về phép toán vectơ 
                                                if dang_toan == "[D12_C2_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_11()

                                                #[D12_C2_B1_12]-M2. Cho hai vectơ có độ dài và góc. Tính tích vô hướng.
                                                if dang_toan == "[D12_C2_B1_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_12()

                                                #[D12_C2_B1_13]-M2. Cho hai vectơ có độ dài và tích vô hướng. Tính góc giữa 2 vectơ.
                                                if dang_toan == "[D12_C2_B1_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_13()

                                                #[D12_C2_B1_14]-M2. Cho hai vectơ a,b có độ dài và tích vô hướng. Tính |ma+nb|.
                                                if dang_toan == "[D12_C2_B1_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_14()

                                                #[D12_C2_B1_15]-M2. Cho hai vectơ a,b và độ dài |ma+nb|. Tính cos(a,b)
                                                if dang_toan == "[D12_C2_B1_15]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_15()

                                                #[D12_C2_B1_16]-M2. Cho hai vectơ a,b có |a|, |b|, ab. Đặt x=ma+nb, y=pa+qb. Tính cos(x,y)
                                                if dang_toan == "[D12_C2_B1_16]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_16()

                                                #[D12_C2_B1_17]-TF-M2. Cho hình lập phương. Xét Đ-S: hướng, phương, góc, tích vô hướng, độ dài vectơ 
                                                if dang_toan == "[D12_C2_B1_17]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_17()

                                                #[D12_C2_B1_18]-M2. Cho hình lập phương. Tìm góc giữa hai vectơ
                                                if dang_toan == "[D12_C2_B1_18]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_18()

                                                #[D12_C2_B1_19]-M2. Cho hình tứ diện đều. Tính tích vô hướng
                                                if dang_toan == "[D12_C2_B1_19]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B1_19()

                                    #------------------------Toán 12 - Chương 2 - Bài 2------------------------------------------------------------->
                                                #[D12_C2_B2_01]. Đọc tọa độ vectơ theo vectơ đơn vị i, j, k
                                                if dang_toan == "[D12_C2_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_01()                                     
                                                    

                                                #[D12_C2_B2_02]. Cho vectơ u=OA. Tìm tọa độ điểm A.
                                                if dang_toan == "[D12_C2_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_02()                                     
                                                    

                                                #[D12_C2_B2_03]. Cho hai điểm. Tìm tọa độ vectơ tạo bởi 2 điểm.
                                                if dang_toan == "[D12_C2_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_03()                                     
                                                    

                                                #[D12_C2_B2_04]. Cho vectơ u và điểm A. Tìm tọa độ điểm B sao cho vto(AB)=vto(u)
                                                if dang_toan == "[D12_C2_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_04()                                     
                                                    

                                                #[D12_C2_B2_05]. Cho điểm A. Tìm hình chiếu của điểm A trên các mặt phẳng tọa độ.
                                                if dang_toan == "[D12_C2_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_05()                                     
                                                    

                                                #[D12_C2_B2_06]. Cho các điểm A,B,C. Tìm D sao cho ABCD là hình bình hành.
                                                if dang_toan == "[D12_C2_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_06()

                                                #[D12_C2_B2_07]-TL-M3. Cho h.hộp chữ nhật được tọa độ hóa. Tìm tọa độ điểm hoặc vectơ.
                                                if dang_toan == "[D12_C2_B2_07]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_07()

                                                #[D12_C2_B2_08]-TF-M2. Cho h.b.h có A,B,C. Xét Đ-S: tọa độ vectơ, tọa độ điểm, tọa độ tâm 
                                                if dang_toan == "[D12_C2_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_08()

                                                #[D12_C2_B2_09]-TF-M2. Cho hộp chữ nhật tọa độ hóa gốc O. Xét Đ-S: tọa độ vectơ, tọa độ điểm, tọa độ tâm 
                                                if dang_toan == "[D12_C2_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_09()

                                                #[D12_C2_B2_10]-TL-M3. Bài toán thực tế: Tìm tọa độ máy bay
                                                if dang_toan == "[D12_C2_B2_10]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B2_10()

                                                                 
                                                    

                                            #Bài 3 -  Các phép toán vectơ

                                                #[D12_C2_B3_01]. Cho hai vectơ. Tìm tọa độ vectơ tổng.
                                                if dang_toan == "[D12_C2_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_01()                                     
                                                    

                                                #[D12_C2_B3_02]. Cho hai vectơ. Tìm tọa độ vectơ tổng.
                                                if dang_toan == "[D12_C2_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_02()                                     
                                                    

                                            #[D12_C2_B3_03]. Cho hai vectơ. Tính tích vô hướng.
                                                if dang_toan == "[D12_C2_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_03()                                     

                                            #[D12_C2_B3_04]. Tìm m để 2 vectơ vuông góc.
                                                if dang_toan == "[D12_C2_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_04()                                     
                                                    

                                            #[D12_C2_B3_05]. Cho vectơ. Tính độ dài.
                                                if dang_toan == "[D12_C2_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_05()                                     

                                            #[D12_C2_B3_06]. Cho hai véctơ tính tích có hướng.
                                                if dang_toan == "[D12_C2_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_06()

                                            #[D12_C2_B3_07]. Cho hai véctơ a và b. Tính ma+nb
                                                if dang_toan == "[D12_C2_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_07()

                                            #[D12_C2_B3_08]-TF-M2. Cho h.b.h có A,B,C. Xét Đ-S: tọa độ vectơ, tọa độ điểm, tọa độ tâm 
                                                if dang_toan == "[D12_C2_B3_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_08()

                                            #[D12_C2_B3_09]-TL-M3. Cho h.hộp chữ nhật được tọa độ hóa. Tìm tọa độ điểm hoặc vectơ.
                                                if dang_toan == "[D12_C2_B3_09]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_09()

                                            #[D12_C2_B3_10]-TF-M2. Cho hai vectơ. Xét Đ-S: tổng, hiệu, độ dài, cos 
                                                if dang_toan == "[D12_C2_B3_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_10()

                                            #[D12_C2_B3_11]-TL-M3. Tìm vectơ x để m.vec(a)+nvec(x)=p.vec(b)
                                                if dang_toan == "[D12_C2_B3_11]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_11()

                                            #[D12_C2_B3_12]. Cho hai điểm A,B. Tìm tọa độ C để AC nhận B làm trung điểm.
                                                if dang_toan == "[D12_C2_B3_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_12()

                                            #[D12_C2_B3_13]-TL-M2. Cho hai điểm A,B. Tìm tọa độ C để AC nhận B làm trung điểm.
                                                if dang_toan == "[D12_C2_B3_13]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_13()

                                            #[D12_C2_B3_14]-TL-M2. Cho hai véctơ. Tìm m để 2 vectơ vuông góc.
                                                if dang_toan == "[D12_C2_B3_14]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_14()

                                            #[D12_C2_B3_15]-M1. Cho hai điểm A,B. Tìm tọa độ trung điểm.
                                                if dang_toan == "[D12_C2_B3_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_15()

                                            #[D12_C2_B3_16]-M1. Cho tam giác ABC. Tìm tọa độ trọng tâm.
                                                if dang_toan == "[D12_C2_B3_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_16()

                                            #[D12_C2_B3_17]-M1. Cho A, vto_AB, vto_AC. Tìm tọa độ trọng tâm.
                                                if dang_toan == "[D12_C2_B3_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_17()

                                            #[D12_C2_B3_18]-M1. Cho A,B,G. Tìm C để ABC nhận G làm trọng tâm
                                                if dang_toan == "[D12_C2_B3_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_18()

                                            #[D12_C2_B3_19]-M1. Cho A,B. Tìm tọa độ vectơ AB
                                                if dang_toan == "[D12_C2_B3_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_19()

                                            #[D12_C2_B3_20]-M2. Cho A,B. Tính độ dài đoạn thẳng AB
                                                if dang_toan == "[D12_C2_B3_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_20()
                                            
                                            #[D12_C2_B3_21]-TL-M3. Cho A,B,C(a;y_C;b),G(x_G;c;z_G). Tính a+b+c
                                                if dang_toan == "[D12_C2_B3_21]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_21()

                                            #[D12_C2_B3_22]-TL-M3. Tìm m để ba điểm lập thành tam giác vuông
                                                if dang_toan == "[D12_C2_B3_22]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_22()

                                            #[D12_C2_B3_23]-TL-M3. Bài toán liên quan trung điểm (VDT)
                                                if dang_toan == "[D12_C2_B3_23]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_23()

                                            #[D12_C2_B3_24]-M2. Tìm hình chiếu của điểm lên trục tọa độ
                                                if dang_toan == "[D12_C2_B3_24]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_24()

                                            #[D12_C2_B3_25]-M2. Tìm hình chiếu của điểm lên mặt phẳng tọa độ
                                                if dang_toan == "[D12_C2_B3_25]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_25()

                                            #[D12_C2_B3_26]-M2. Tìm điểm đối xứng của điểm qua trục tọa độ
                                                if dang_toan == "[D12_C2_B3_26]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_26()

                                            #[D12_C2_B3_27]-M2. Tìm điểm đối xứng của điểm qua mặt phẳng tọa độ
                                                if dang_toan == "[D12_C2_B3_27]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_27()

                                            #[D12_C2_B3_28]-TF-M3. Cho tam giác. Xét Đ-S: Độ dài, tọa độ vectơ, R, S.
                                                if dang_toan == "[D12_C2_B3_28]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D12_C2.mnj_34_jkl_L12_C2_B3_28()



                                        #------------Toán 12 - Chương 3 - Mẫu số liệu ghép nhóm ----------------------------->
                                        #Bài 1: Khoảng biến thiên, khoảng tứ phân vị
                                           #[D12_C3_B1_01]-M1. Tìm khoảng biến thiên của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B1_01]":                                        
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C3.ytrzz_L12_C3_B1_01()

                                           #[D12_C3_B1_02]-M3. Tìm khoảng tứ phân vị của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B1_02]":                                        
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C3.ytrzz_L12_C3_B1_02()

                                           #[D12_C3_B1_03]-M3. Tìm khoảng tứ phân vị của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B1_03]": 
                                                   debai_word,loigiai_word,latex_tuluan,dap_an=D12_C3.ytrzz_L12_C3_B1_03()

                                           #[D12_C3_B1_04]-TF-M3. Xét Đ-S: Khoảng biến thiên, Q1, Q3, khoảng tứ phân vị.
                                                if dang_toan == "[D12_C3_B1_04]": 
                                                   debai_word,debai_latex,loigiai_word,dap_an=D12_C3.ytrzz_L12_C3_B1_04()

                                        #Bài 2: Phương sai, độ lệch chuẩn
                                           #[D12_C3_B2_01]-M2. Tìm phương sai của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B2_01]":                                        
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C3.ytrzz_L12_C3_B2_01()

                                           #[D12_C3_B2_02]-M2. Tìm độ lệch chuẩn của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B2_02]":                                        
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C3.ytrzz_L12_C3_B2_02()
                                           
                                           #[D12_C3_B2_03]-TL-M2. Tính phương sai của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B2_03]": 
                                                   debai_word,loigiai_word,latex_tuluan,dap_an=D12_C3.ytrzz_L12_C3_B2_03()

                                           #[D12_C3_B2_04]-TL-M2. Tính độ lệch chuẩn của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D12_C3_B2_04]": 
                                                   debai_word,loigiai_word,latex_tuluan,dap_an=D12_C3.ytrzz_L12_C3_B2_04()

                                           #[D12_C3_B2_05]-TF-M2. Xét Đ-S: số trung bình, phương sai, độ lệch chuẩn, khoảng tứ phân vị
                                                if dang_toan == "[D12_C3_B2_05]": 
                                                   debai_word,debai_latex,loigiai_word,dap_an=D12_C3.ytrzz_L12_C3_B2_05()

                                                                       
                                                    

                                    ######### Toán 12_ Chương 4 - NGUYÊN HÀM - TÍCH PHÂN ########
                                            #Bài 1. Nguyên hàm cơ bản

                                                #[D12_C4_B1_01]. Tìm nguyên hàm của hàm số đa thức
                                                if dang_toan == "[D12_C4_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_01()                                     
                                                    

                                                #[D12_C4_B1_02]. Tìm nguyên hàm của hàm số đa thức có F(a)=b
                                                if dang_toan == "[D12_C4_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_02()                                     
                                                    

                                                #[D12_C4_B1_03]. Tìm nguyên hàm chứa sinu, cosu
                                                if dang_toan == "[D12_C4_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_03()                                     
                                                    

                                                #[D12_C4_B1_04]. Tìm nguyên hàm asinu
                                                if dang_toan == "[D12_C4_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_04()                                     
                                                    

                                                #[D12_C4_B1_05]. Tìm nguyên hàm acosu
                                                if dang_toan == "[D12_C4_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_05()                                     
                                                    

                                                #[D12_C4_B1_06]. Tìm nguyên hàm đa thức +a/x
                                                if dang_toan == "[D12_C4_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_06()                                     
                                                    

                                                #[D12_C4_B1_07]. Tìm nguyên hàm đa thức m/(ax+b)
                                                if dang_toan == "[D12_C4_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_07()                                     
                                                    

                                                #[D12_C4_B1_08]. Tìm nguyên hàm đa thức (mx+n)/(ax+b)
                                                if dang_toan == "[D12_C4_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_08()                                     
                                                    

                                                #[D12_C4_B1_09]. Tìm nguyên hàm đa thức (mx^2+nx+p)/(ax+b)
                                                if dang_toan == "[D12_C4_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_09()                                     
                                                    

                                                #[D12_C4_B1_10]. Tìm nguyên hàm của e^(ax+b)
                                                if dang_toan == "[D12_C4_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B1_10()                                     
                                                    

                                            #Bài 2. Nguyên hàm đổi biến
                                                #[D12_C4_B2_01]. Nguyên hàm đổi biến chứa căn(ax^+b)
                                                if dang_toan == "[D12_C4_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B2_01()                                     
                                                    

                                                #[D12_C4_B2_02]. Nguyên hàm đổi biến chứa căn(ax^2+b)
                                                if dang_toan == "[D12_C4_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B2_02()                                     
                                                    

                                                #[D12_C4_B2_03]. Nguyên hàm đổi biến sinx/căn(acosx+b)dx
                                                if dang_toan == "[D12_C4_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B2_03()                                     
                                                    

                                                #[D12_C4_B2_04]. Nguyên hàm đổi biến (acosx+b)sinx
                                                if dang_toan == "[D12_C4_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B2_04()                                     
                                                    

                                                #[D12_C4_B2_05]. Nguyên hàm đổi biến (alnx + b)/x
                                                if dang_toan == "[D12_C4_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B2_05()                                     
                                                    

                                            #Bài 3. Nguyên hàm từng phần
                                                #[D12_C4_B3_01]. Nguyên hàm P(x).sin hoặc P(x).cos
                                                if dang_toan == "[D12_C4_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B3_01()                                     
                                                                   

                                                #[D12_C4_B3_02]. Nguyên hàm P(x).e^(ax+b)
                                                if dang_toan == "[D12_C4_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B3_02()                                     
                                                    

                                                #[D12_C4_B3_03]. Nguyên hàm P(x).ln(mx)
                                                if dang_toan == "[D12_C4_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B3_03()                                     
                                                    

                                            #Bài 4. Tích phân
                                            #[D12_C4_B4_01]-M1. Cho F(a) và F(b). Tính tích phân từ a đến b.
                                                if dang_toan == "[D12_C4_B4_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_01()                                     
                                                    

                                            #[D12_C4_B4_02]-M1. Cho 2 tích phân theo f và g. Tính tích phân (m.f+n.g).
                                                if dang_toan == "[D12_C4_B4_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_02()                                     
                                                    

                                            #[D12_C4_B4_03]-M1. Tính tổng tích phân trên 2 đoạn nối tiếp.
                                                if dang_toan == "[D12_C4_B4_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_03()                                     
                                                    

                                            #[D12_C4_B4_04]-M1. Cho theo f. Tính tích phân (m.f+n).
                                                if dang_toan == "[D12_C4_B4_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_04()                                     
                                                    

                                            #[D12_C4_B4_05]-M2. Cho tích phân f và F(a). Tính F(b).
                                                if dang_toan == "[D12_C4_B4_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_05()                                     
                                                    

                                            #[D12_C4_B4_06]-M1. Cho f(a) và f(b). Tính tích phân f'(x).
                                                if dang_toan == "[D12_C4_B4_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_06()                                     
                                                    

                                            #[D12_C4_B4_07]-M1. Cho f(a) và f(b). Tính tích phân f'(x).
                                                if dang_toan == "[D12_C4_B4_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_07()                                     
                                                    

                                            #[D12_C4_B4_08]-M1. Cho tphan_a^c(f) và tphan_a^b(f). Tính tphan_b^c(f), a<b<c.
                                                if dang_toan == "[D12_C4_B4_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_08()                                     
                                                    

                                            #[D12_C4_B4_09]-M2. Cho tphan_a^d(f) và tphan_b^c(f). Tính tphan_a^b(f) + tphan_c^d(f), a<b<c<d.
                                                if dang_toan == "[D12_C4_B4_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_09()                                     
                                                    

                                            #Tích phân đa thức, phân thức

                                            #[D12_C4_B4_10]-M2. Tính tích phân của m/(ax+b).
                                                if dang_toan == "[D12_C4_B4_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_10()                                     
                                                    

                                            #[D12_C4_B4_11]-M2. Tính tích phân của (mx+n)/(ax+b).
                                                if dang_toan == "[D12_C4_B4_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_11()                                     
                                                    

                                            #[D12_C4_B4_12]-M2. Tính tích phân của đa thức.
                                                if dang_toan == "[D12_C4_B4_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_12()                                     
                                                    

                                            #[D12_C4_B4_13]-M2. Tính tích phân của asinbx.
                                                if dang_toan == "[D12_C4_B4_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B4_13()                                     
                                                    

                                            #Bài 5. Ứng dụng tích phân
                                            #[D12_C4_B5_03]. Tìm công thức tính diện tích từ hình vẽ có 1 đồ thị.
                                                if dang_toan == "[D12_C4_B5_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_03()                                     
                                                    

                                     
                                            #[D12_C4_B5_01]. Diện tích hình phẳng: y=f(x),Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_01()                                     
                                                    

                                            #[D12_C4_B5_02]. Diện tích hình phẳng: y=f(x),y=g(x),x=a,x=b
                                                if dang_toan == "[D12_C4_B5_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_02()                                     
                                                    

                                            #[D12_C4_B5_04]. Diện tích hình phẳng: y=f(x),y=g(x)
                                                if dang_toan == "[D12_C4_B5_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_04()                                     
                                                    

                                            #[D12_C4_B5_05]. Thể tích quay quanh Ox bởi hình: y=f(x),Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_05()                                     
                                                     

                                            #[D12_C4_B5_06]. Thể tích quay quanh Ox bởi hình: y=ax+b,Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_06()                                     
                                                    

                                            #[D12_C4_B5_07]. Thể tích quay quanh Ox bởi hình: y=ax^2+bx+c,Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_07()                                     
                                                    

                                            #[D12_C4_B5_08]. Thể tích quay quanh Ox bởi hình: y=căn(ax+b),Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_08()                                     
                                                    

                                            #[D12_C4_B5_09]. Thể tích quay quanh Ox bởi hình: y=căn(ax+b),Ox,x=a,x=b
                                                if dang_toan == "[D12_C4_B5_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_09()                                     
                                                    

                                            #[D12_C4_B5_10]. Thể tích quay quanh Ox bởi hình: y=ax^2+bx+c, Ox
                                                if dang_toan == "[D12_C4_B5_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_10()                                     
                                                    

                                            #[D12_C4_B5_11]. Cho hàm số vận tốc. Tính quãng đường đi được từ t1 đến t2.
                                                if dang_toan == "[D12_C4_B5_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_11()                                     
                                                    

                                            #[D12_C4_B5_12]. Cho hàm số vận tốc. Tính quãng đường đi được từ t1 đến t2.
                                                if dang_toan == "[D12_C4_B5_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_12()                                     
                                                    

                                            #[D12_C4_B5_13]. Xe tăng tốc với gia tốc. Tính quãng đường đi được trong khoảng thời gian
                                                if dang_toan == "[D12_C4_B5_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_13()                                     
                                                    

                                            #[D12_C4_B5_14]. Thể tích vật thể có thiết diện là hình chữ nhật.
                                                if dang_toan == "[D12_C4_B5_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_14()                                     
                                                    

                                            #[D12_C4_B5_15]. Thể tích vật thể có thiết diện là hình vuông.
                                                if dang_toan == "[D12_C4_B5_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_15()                                     
                                                    

                                            #[D12_C4_B5_16]. Thể tích vật thể có thiết diện là hình tam giác đều.
                                                if dang_toan == "[D12_C4_B5_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C4.zz8zz_L12_C4_B5_16()                                     
                                                    

                                                ##################################################################
                                    #Toán 12 - Chương 5 - Phương trình mặt phẳng, đường thẳng, mặt cầu
                                    #BÀI 1 - PHƯƠNG TRÌNH MẶT PHẲNG
                                            #[D12_C5_B1_01]-M1. Viết PTMP qua điểm và có véctơ pháp tuyến
                                                if dang_toan == "[D12_C5_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_01()                                     
                                                    

                                            #[D12_C5_B1_02]-M2. Viết PTMP qua điểm A và nhận vectơ BC làm véctơ pháp tuyến.
                                                if dang_toan == "[D12_C5_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_02()                                     
                                                    

                                            #[D12_C5_B1_03]-M2. Viết PTMP qua điểm A và vuông góc trục tọa độ.
                                                if dang_toan == "[D12_C5_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_03()                                     
                                                    

                                            #[D12_C5_B1_04]-M2. Viết PTMP qua điểm và song song với mặt phẳng.
                                                if dang_toan == "[D12_C5_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_04()                                     
                                                    

                                            #[D12_C5_B1_05]-M2. Viết PTMP trung trực của đoạn AB.
                                                if dang_toan == "[D12_C5_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_05()                                     
                                                    

                                            #[D12_C5_B1_06]-M3. Viết PTMP đi qua 3 điểm A,B,C.
                                                if dang_toan == "[D12_C5_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_06()                                     
                                                    

                                            #[D12_C5_B1_07]-M1. Tính khoảng cách từ điểm đến mặt phẳng
                                                if dang_toan == "[D12_C5_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_07()                                     
                                                    

                                            #[D12_C5_B1_08]-M2. Xét vị trí tương đối giữa 2 mặt phẳng
                                                if dang_toan == "[D12_C5_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_08()                                     
                                                    

                                            #[D12_C5_B1_09]-M2. Xét vị trí tương đối giữa 2 mặt phẳng. KQ: trùng nhau
                                                if dang_toan == "[D12_C5_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_09()                                     
                                                    

                                            #[D12_C5_B1_10]-M2. Xét vị trí tương đối giữa 2 mặt phẳng. KQ: song song
                                                if dang_toan == "[D12_C5_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_10()                                     
                                                    

                                            #[D12_C5_B1_11]-M2. Xét vị trí tương đối giữa 2 mặt phẳng. KQ: vuông góc
                                                if dang_toan == "[D12_C5_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_11()                                     
                                                    

                                            #[D12_C5_B1_12]-M2. Xét vị trí tương đối giữa 2 mặt phẳng. KQ: cắt và không vuông góc
                                                if dang_toan == "[D12_C5_B1_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_12()                                     
                                                    

                                            #[D12_C5_B1_13]-M2. Tính khoảng cách giữa 2 mặt phẳng song song
                                                if dang_toan == "[D12_C5_B1_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_13()                                     
                                                    

                                            #[D12_C5_B1_14]-M2. Cho PTMP, đọc véctơ pháp tuyến
                                                if dang_toan == "[D12_C5_B1_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_14()                                     
                                                    

                                            #[D12_C5_B1_15]-M2. Cho PTMP, tìm điểm thuộc mặt phẳng 
                                                if dang_toan == "[D12_C5_B1_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_15()                                     
                                                    

                                            #[D12_C5_B1_16]-M2. Cho PTMP, tìm điểm không thuộc mặt phẳng 
                                                if dang_toan == "[D12_C5_B1_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_16()                                     
                                                    

                                            #[D12_C5_B1_17]-M2.  Viết phương trình mặt phẳng đoạn chắn
                                                if dang_toan == "[D12_C5_B1_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B1_17()                                     
                                                    


                                            #BÀI 2- PHƯƠNG TRÌNH ĐƯỜNG THẲNG
                                            #[D12_C5_B2_01]-M1. Viết PTDT qua điểm và có véctơ chỉ phương 
                                                if dang_toan == "[D12_C5_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_01()                                     
                                                    

                                            #[D12_C5_B2_02]-M1. Viết PTĐT qua điểm A và nhận vectơ BC làm véctơ chỉ phương.
                                                if dang_toan == "[D12_C5_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_02()                                     
                                                    

                                            #[D12_C5_B2_03]-M2. Viết PT chính tắc của đường thẳng qua điểm A và song song với đường thẳng.
                                                if dang_toan == "[D12_C5_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_03()                                     
                                                    

                                            #[D12_C5_B2_04]-M2. Viết PT chính tắc của đường thẳng qua điểm A và vuông góc với mặt phẳng
                                                if dang_toan == "[D12_C5_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_04()                                     
                                                    

                                            #[D12_C5_B2_05]-M2. Viết PT chính tắc của đường thẳng qua 2 điểm 
                                                if dang_toan == "[D12_C5_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_05()                                     
                                                    

                                            #[D12_C5_B2_06]-M2. Viết PT chính tắc đường thẳng d song song d' và vuông góc với (P).
                                                if dang_toan == "[D12_C5_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_06()                                     
                                                    

                                            #[D12_C5_B2_07]-M2. Viết PT chính tắc đường thẳng d và vuông góc với (P) và (Q).
                                                if dang_toan == "[D12_C5_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_07()                                     
                                                    

                                            #[D12_C5_B2_08]-M2. Viết PT tham số đường thẳng qua điểm và có véctơ chỉ phương.
                                                if dang_toan == "[D12_C5_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_08()                                     
                                                    

                                            #[D12_C5_B2_09]-M2. Viết PTTS đường thẳng qua điểm A và nhận vectơ BC làm véctơ chỉ phương.
                                                if dang_toan == "[D12_C5_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_09()                                     
                                                    

                                            #[D12_C5_B2_10]-M2. Viết PTTS đường thẳng qua điểm và song song với đường thẳng
                                                if dang_toan == "[D12_C5_B2_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_10()                                     
                                                    

                                            #[D12_C5_B2_11]-M2. Cho phương trình chính tắc viết phương trình tham số
                                                if dang_toan == "[D12_C5_B2_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_11()                                     
                                                    

                                            #[D12_C5_B2_12]-M2. Cho phương trình tham số viết phương trình chính tắc
                                                if dang_toan == "[D12_C5_B2_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_12()                                     
                                                    

                                            #[D12_C5_B2_13]-M2. Đọc VTCP từ phương trình tham số
                                                if dang_toan == "[D12_C5_B2_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_13()                                     
                                                    

                                            #[D12_C5_B2_14]-M2. Đọc tọa độ điểm từ phương trình tham số
                                                if dang_toan == "[D12_C5_B2_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_14()                                     
                                                    

                                            #[D12_C5_B2_15]-M2. Đọc véctơ chỉ phương từ phương trình chính tắc
                                                if dang_toan == "[D12_C5_B2_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_15()                                     
                                                    

                                            #[D12_C5_B2_16]-M2. Đọc tọa độ điểm từ phương trình chính tắc
                                                if dang_toan == "[D12_C5_B2_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_16()                                     
                                                    

                                            #[D12_C5_B2_17]-M3. Tìm giao điểm của đường thẳng và mặt phẳng
                                                if dang_toan == "[D12_C5_B2_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_17()                                     
                                                    

                                            #[D12_C5_B2_18]-M3. Tìm hình chiếu của điểm trên đường thẳng
                                                if dang_toan == "[D12_C5_B2_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_18()                                     
                                                    

                                            #[D12_C5_B2_19]-M2. Xét vị trí 2 đường thẳng. Kết quả là song song
                                                if dang_toan == "[D12_C5_B2_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_19()                                     
                                                    

                                            #[D12_C5_B2_20]-M2. Xét vị trí 2 đường thẳng. Kết quả là trùng nhau
                                                if dang_toan == "[D12_C5_B2_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_20()                                     
                                                    

                                            #[D12_C5_B2_21]-M2. Xét vị trí 2 đường thẳng. Kết quả là cắt nhau
                                                if dang_toan == "[D12_C5_B2_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_21()                                     
                                                    

                                            #[D12_C5_B2_22]-M2. Xét vị trí 2 đường thẳng. Kết quả là chéo nhau
                                                if dang_toan == "[D12_C5_B2_22]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_22()                                     
                                                    

                                            #[D12_C5_B2_23]-M2. Xét vị trí 2 đường thẳng tùy ý
                                                if dang_toan == "[D12_C5_B2_23]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_23()

                                            #[D12_C5_B2_24]-M4. Tìm đường thẳng đi qua điểm cắt và vuông góc với đường thẳng khác
                                                if dang_toan == "[D12_C5_B2_24]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_24()

                                            #[D12_C5_B2_25]-M4. Tìm đường thẳng đi qua điểm, cắt đường thẳng khác và song song mặt phẳng
                                                if dang_toan == "[D12_C5_B2_25]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_25()

                                            #[D12_C5_B2_26]-M4. Viết PTMP qua A,B cách M một khoảng lớn nhất
                                                if dang_toan == "[D12_C5_B2_26]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B2_26()                                    
                                                    

                                            #BÀI 3 - PHƯƠNG TRÌNH MẶT CẦU
                                            #[D12_C5_B3_01]-M1. Đọc tọa độ tâm từ phương trình mặt cầu
                                                if dang_toan == "[D12_C5_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_01() 

                                            #[D12_C5_B3_02]-M1. Đọc tọa độ tâm từ phương trình mặt cầu
                                                if dang_toan == "[D12_C5_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_02()                                     
                                                    

                                            #[D12_C5_B3_03]-M1. Đọc bán kính từ phương trình mặt cầu thu gọn
                                                if dang_toan == "[D12_C5_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_03()                                     
                                                    

                                            #[D12_C5_B3_04]-M1. Đọc tọa độ tâm từ phương trình mặt cầu khai triển
                                                if dang_toan == "[D12_C5_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_04()                                     
                                                    

                                            #[D12_C5_B3_05]-M1. Đọc bán kính từ phương trình mặt cầu khai triển
                                                if dang_toan == "[D12_C5_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_05()                                     
                                                    

                                            #[D12_C5_B3_06]-M2. Viết phương trình mặt cầu có tâm và đường kính
                                                if dang_toan == "[D12_C5_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_06()                                     
                                                    

                                            #[D12_C5_B3_07]-M3. Viết phương trình mặt cầu có đường kính AB
                                                if dang_toan == "[D12_C5_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_07()                                     
                                                    

                                            #[D12_C5_B3_08]-M2. Viết phương trình mặt cầu có tâm và đi qua điểm
                                                if dang_toan == "[D12_C5_B3_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C5.zz8zz_L12_C5_B3_08()                                     
                                                    
                                       
                                    ######### Toán 12_ Chương 6 - MẶT TRÒN XOAY ########
                                            #Bài 1. MẶT NÓN
                                                #[D12_C7_B1_01]. Cho hình nón có bán kính r và đường sinh l. Tính S_xq.
                                                if dang_toan == "[D12_C7_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B1_01()                                     
                                                    
                                                #[D12_C7_B1_02]. Cho hình nón có bán kính r và chiều cao h. Tính V.

                                                if dang_toan == "[D12_C7_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B1_02()                                     
                                                    

                                                #[D12_C7_B1_03].Cho hình nón có chu vi và đường sinh. Tính chiều cao h.
                                                if dang_toan == "[D12_C7_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B1_03()                                     
                                                    

                                                #[D12_C7_B1_04].Cho hình nón có diện tích đáy và đường sinh. Tính chiều cao h.
                                                if dang_toan == "[D12_C7_B1_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B1_03()                                     
                                                    

                                        #Bài 2. MẶT TRỤ

                                                 #[D12_C7_B2_01]. Cho hình trụ có bán kính r và đường sinh l. Tính S_xq. 
                                                if dang_toan == "[D12_C7_B2_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B2_01()                                     
                                                    

                                                #[D12_C7_B2_02]. Cho hình trụ có bán kính r và đường sinh l. Tính S_xq. 
                                                if dang_toan == "[D12_C7_B2_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B2_02()                                     
                                                    

                                                #[D12_C7_B2_03]. Cho hình trụ có thiết diện qua trục là hình vuông. Tính S_tp.
                                                if dang_toan == "[D12_C7_B2_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B2_03()                                     
                                                    

                                                #[D12_C7_B2_04]. Cho hình trụ có thiết diện qua trục là hình chữ nhật. Tính S_tp. 
                                                if dang_toan == "[D12_C7_B2_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B2_04()                                     
                                                    

                                        #Bài 3. MẶT CẦU

                                                #[D12_C7_B3_01]. Cho mặt cầu có bán kính r. Tính S_mc. 
                                                if dang_toan == "[D12_C7_B3_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B3_01()                                     
                                                    

                                                #[D12_C7_B3_02]. Cho mặt cầu có thể tích. Tính bán kính R.                 
                                                if dang_toan == "[D12_C7_B3_02]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D12_C7.zz8zz_L12_C7_B3_02()                                     
                                                    

                                        #[PT_DE2024_01]. Cho bảng biến thiên. Tìm giá trị cực trị. 
                                                if dang_toan == "[PT_DE2024_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_01()                                     
                                                    

                                        #[PT_DE2024_02]. Tìm nguyên hàm của hàm số đa thức
                                                if dang_toan == "[PT_DE2024_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_02()                                     
                                                    

                                        #[PT_DE2024_03]. Giải phương trình log_a(mx^2+nx+p)=b
                                                if dang_toan == "[PT_DE2024_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_03()                                     
                                                    

                                        #[PT_DE2024_04]. Cho hai điểm. Tìm tọa độ vectơ tạo bởi 2 điểm.
                                                if dang_toan == "[PT_DE2024_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_04()                                     
                                                    

                                        #[PT_DE2024_05]. Đọc đường tiệm cận từ đồ thị hàm số y=(ax+b)/(cx+d)
                                                if dang_toan == "[PT_DE2024_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_05()                                     
                                                    

                                        #[PT_DE2024_06]. Cho bảng biến thiên, tìm hàm số.
                                                if dang_toan == "[PT_DE2024_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_06()                                     
                                                    

                                        #[PT_DE2024_07]. Tìm tập xác định của y=[f(x)]^n
                                                if dang_toan == "[PT_DE2024_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_07()                                     
                                                    

                                        #[PT_DE2024_08]. Đọc véctơ chỉ phương từ phương trình chính tắc
                                                if dang_toan == "[PT_DE2024_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_08()                                     
                                                    

                                        #[PT_DE2024_09]. Đọc số phức từ điểm biểu diễn
                                                if dang_toan == "[PT_DE2024_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_09()                                     
                                                    

                                        #[PT_DE2024_10]. Viết phương trình mặt cầu có tâm và bán kính
                                                if dang_toan == "[PT_DE2024_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_10()                                     
                                                    

                                        #[PT_DE2024_11]. Tìm khẳng định đúng về tính chất logarit
                                                if dang_toan == "[PT_DE2024_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_11()                                     
                                                    

                                        #[PT_DE2024_12]. Cho đồ thị. Tìm khoảng đồng biến, nghịch biến. 
                                                if dang_toan == "[PT_DE2024_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_12()                                     
                                                    

                                        #[PT_DE2024_13]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_13()                                     
                                                    

                                        #[PT_DE2024_14]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_14()                                     
                                                    

                                        #[PT_DE2024_15]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_15()                                     
                                                    

                                        #[PT_DE2024_16]. Đọc véctơ pháp tuyến của mặt phẳng tọa độ
                                                if dang_toan == "[PT_DE2024_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_16()                                     
                                                    

                                        #[PT_DE2024_17]. Cho f'(x) đếm số điểm cực trị 
                                                if dang_toan == "[PT_DE2024_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_17()                                     
                                                    

                                        #[PT_DE2024_18]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_18()                                     
                                                    

                                        #[PT_DE2024_19]. Cho tích phân từ a đến b, tính tích phân từ b đến a
                                                if dang_toan == "[PT_DE2024_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_19()                                     
                                                    

                                        #[PT_DE2024_20]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_20()                                     
                                                    

                                        #[PT_DE2024_21]. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[PT_DE2024_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_21()                                     
                                                    

                                        #[PT_DE2024_22]. Tìm công thức đúng về l,S_xq,S_tp của hình nón, hình trụ
                                                if dang_toan == "[PT_DE2024_22]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_22()                                     
                                                    

                                        #[PT_DE2024_23]. Xếp k bạn vào một hàng
                                                if dang_toan == "[PT_DE2024_23]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_23()                                     
                                                    

                                        #[PT_DE2024_24]. Xếp k bạn vào một hàng
                                                if dang_toan == "[PT_DE2024_24]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_24()                                     
                                                    

                                        #[PT_DE2024_25]. Xếp k bạn vào một hàng
                                                if dang_toan == "[PT_DE2024_25]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_25()                                     
                                                    

                                        #[PT_DE2024_26]. Tìm công thức đúng về l,r,h của hình trụ, hình nón theo S_xq
                                                if dang_toan == "[PT_DE2024_26]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_26()                                     
                                                    

                                        #[PT_DE2024_27]. Xếp k bạn vào một hàng
                                                if dang_toan == "[PT_DE2024_27]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_27()                                     
                                                    

                                        #[PT_DE2024_28]. Tìm phần thực, phần ảo, liên hợp, môđun số phức.
                                                if dang_toan == "[PT_DE2024_28]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_28()                                     
                                                    

                                        #[PT_DE2024_29]. Tìm phần thực, phần ảo, liên hợp của tích 2 số phức.
                                                if dang_toan == "[PT_DE2024_29]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_29()                                     
                                                    

                                        #[PT_DE2024_30]. Cho hình lập phương. Xác định góc giữa hai đường thẳng.
                                                if dang_toan == "[PT_DE2024_30]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_30()                                     
                                                    

                                        #[PT_DE2024_31]. S.ABCD: ABCD h.chữ nhật. Tính k.c từ chân đường cao đến mặt nghiêng.
                                                if dang_toan == "[PT_DE2024_31]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_31()                                     
                                                    

                                        #[PT_DE2024_32]. Cho f'(x). Tìm khoảng đồng biến, nghịch biến.
                                                if dang_toan == "[PT_DE2024_32]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_32()                                     
                                                    

                                        #[PT_DE2024_33]. Cho 2 nhóm đồ vật. Tính xác suất để số vật được chọn thuộc cùng 1 nhóm.
                                                if dang_toan == "[PT_DE2024_33]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_33()                                     
                                                    

                                        #[PT_DE2024_34]. Cho tích phân theo f. Tính tích phân (m.f+n).
                                                if dang_toan == "[PT_DE2024_34]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_34()                                     
                                                    

                                        #[PT_DE2024_35]. Tìm GTLN-GTNN của hàm số trên đoạn.
                                                if dang_toan == "[PT_DE2024_35]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_35()                                     
                                                    

                                        #[PT_DE2024_36]. Tìm khẳng định đúng về biến đổi logarit
                                                if dang_toan == "[PT_DE2024_36]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_36()                                     
                                                    

                                        #[PT_DE2024_37]. Viết phương trình mặt cầu có tâm và bán kính.
                                                if dang_toan == "[PT_DE2024_37]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_37()                                     
                                                    

                                        #[PT_DE2024_38]. Viết phương trình đường thẳng qua điểm và song song với đường thẳng.
                                                if dang_toan == "[PT_DE2024_38]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_38()                                     
                                                    

                                        #[PT_DE2024_39]. Cho log_m a= log_n b = log_p (a+b). Tính a/b.
                                                if dang_toan == "[PT_DE2024_39]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_39()                                     
                                                    
                                        #[PT_DE2024_40]. Tìm m để hàm số đồng biến(nghịch biến) trên khoảng K.
                                                if dang_toan == "[PT_DE2024_40]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_40()                                     
                                                    

                                        #[PT_DE2024_41]. Cho f(x) bậc 4 có 3 cực trị, g(x) bậc 2, diện tích. Tính tích phân 
                                                if dang_toan == "[PT_DE2024_41]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_41()                                     
                                                    

                                        #[PT_DE2024_42]. Cho 3 môđun dạng |az+bw|. Tính |z| hoặc |w| hoặc zw.
                                                if dang_toan == "[PT_DE2024_42]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=DeMH2024.PT_DE2024_42()
                                                                                                                                                
                                                #[D11_C1_B1_02]-M1. Đổi số đo từ độ sang radian
                                                if dang_toan == "[D11_C1_B1_02]":                                                                                
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_02()

                                                #[D11_C1_B1_03]-M1. Đổi số đo từ radian sang độ
                                                if dang_toan == "[D11_C1_B1_03]":                                                                                
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_03()

                                                #[D11_C1_B1_04]-M1. Tìm góc có điểm biểu diễn trùng nhau
                                                if dang_toan == "[D11_C1_B1_04]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_04()

                                                #[D11_C1_B1_05]-M2. Tìm góc lượng giác có điểm biểu diễn cho trước
                                                if dang_toan == "[D11_C1_B1_05]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_05()

                                                #[D11_C1_B1_06]-M2. Cho bán kính và góc radian. Tính độ dài của cung.
                                                if dang_toan == "[D11_C1_B1_06]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_06()

                                                #[D11_C1_B1_07]-M2. Cho bán kính và góc độ. Tính độ dài của cung.
                                                if dang_toan == "[D11_C1_B1_07]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_07()

                                                #[D11_C1_B1_08]-TF-M2. Cho góc radian. Xét đúng-sai: Đổi sang độ, Điểm biểu diễn thuộc phần tư, Góc cùng điểm biểu diễn, Đếm số điểm biểu diễn
                                                if dang_toan == "[D11_C1_B1_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B1_08()

                                                #[D11_C1_B1_09]-TF-M2. Cho góc radian. Xét đúng-sai: Cho góc độ. Xét đúng-sai: Đổi sang radian, Điểm biểu diễn thuộc phần tư, Góc cùng điểm biểu diễn, Đếm số điểm biểu diễn
                                                if dang_toan == "[D11_C1_B1_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B1_09()

                                                #[D11_C1_B1_10]-TL-M3. Cho số vòng quay bánh xe sau t1 giây. Tính góc radian sau khi quay trong t2 giây
                                                if dang_toan == "[D11_C1_B1_10]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_10()

                                                #[D11_C1_B1_11]-TL-M3. Cho số vòng quay bánh xe sau t1 giây. Tính quãng đường đi được sau t2 giây
                                                if dang_toan == "[D11_C1_B1_11]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_11()

                                                #[D11_C1_B1_12]-TL-M3. Cho đường kính của bánh trước và bánh sau, vận tốc n vòng/phút. Tính quãng đường xe đi.
                                                if dang_toan == "[D11_C1_B1_12]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B1_12()
                                                

                                                    

                                                ####################### Bài 2: Giá trị lượng giác ######################
                                                #[D11_C1_B2_01]. Tính giá trị đặc biệt của một góc lượng giác.
                                                if dang_toan == "[D11_C1_B2_01]":                                                                                
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B2_01()

                                                #[D11_C1_B2_02]-TL-M2. Cho sinx (hoặc cosx), x thuộc (a;b). Tìm cosx (hoặc sinx)
                                                if dang_toan == "[D11_C1_B2_02]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B2_02()

                                                #[D11_C1_B2_03]-TL-M3. Cho sinx (hoặc cosx), x thuộc (a;b). Tìm tanx (hoặc cotx)
                                                if dang_toan == "[D11_C1_B2_03]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B2_03()

                                                #[D11_C1_B2_04]-TL-M3. Cho tanx (hoặc cotx). Tìm P=(asinx+bcosx)/(csinx+dcosx)
                                                if dang_toan == "[D11_C1_B2_04]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B2_04()

                                                #[D11_C1_B2_05]-TL-M3. Cho tanx (hoặc cotx). Tìm P=(asin^2 x+bsinxcosx)/(csin^2 x+dcos^2 x)
                                                if dang_toan == "[D11_C1_B2_05]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B2_05()

                                                #[D11_C1_B2_06]-TF-M2. Cho góc lượng giác thuộc cung phần tư. Xét Đ-S dấu của sin, cos, tan, cot
                                                if dang_toan == "[D11_C1_B2_06]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B2_06()

                                                #[D11_C1_B2_07]-TF-M2. Cho tanx. Xét Đ-S: cotx, sin^2 x, cos^2x, (asinx+bcosx)/(csinx+dcosx)
                                                if dang_toan == "[D11_C1_B2_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B2_07()

                                                #[D11_C1_B2_08]-TF-M3. Chon sinx (a<x<b). Xét Đ-S: dấu của cosx, cosx, sin(x+kpi/2), P=f(tanx).
                                                if dang_toan == "[D11_C1_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B2_08()

                                                #[D11_C1_B2_09]-TF-M3.  Chon cosx (a<x<b). Xét đúng sai: dấu của sinx, sinx, sin(x+kpi/2), P=f(tanx).
                                                if dang_toan == "[D11_C1_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B2_09()




                                                ####################### Bài 3: Các công thức lượng giác. ######################

                                                #[D11_C1_B3_01]  Cho sina, cosa. Tính sin2a.
                                                if dang_toan == "[D11_C1_B3_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_01()

                                                #[D11_C1_B3_02] Cho sina hoặc cosa. Tính cos2a.
                                                if dang_toan == "[D11_C1_B3_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_02()

                                                #[D11_C1_B3_03]-M2. Tìm khẳng định đúng về công thức đối
                                                if dang_toan == "[D11_C1_B3_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_03()

                                                #[D11_C1_B3_04]-M2. Tìm khẳng định đúng về hai góc bù nhau
                                                if dang_toan == "[D11_C1_B3_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_04()

                                                #[D11_C1_B3_05]-M2. Tìm khẳng định đúng về hai góc phụ nhau
                                                if dang_toan == "[D11_C1_B3_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_05()

                                                #[D11_C1_B3_06]-M2. Tìm khẳng định đúng về hai góc hơn kém pi
                                                if dang_toan == "[D11_C1_B3_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_06()

                                                #[D11_C1_B3_07]-M2. Tìm khẳng định đúng về hai góc hơn liên quan tùy ý
                                                if dang_toan == "[D11_C1_B3_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_07()

                                                #[D11_C1_B3_08]-M2. Tìm khẳng định đúng về công thức nhân đôi
                                                if dang_toan == "[D11_C1_B3_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_08()

                                                #[D11_C1_B3_09]-M2. Tìm khẳng định đúng về công thức nhân đôi
                                                if dang_toan == "[D11_C1_B3_09]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_09()

                                                #[D11_C1_B3_10]-M2. Tìm khẳng định đúng về công thức tổng thành tích
                                                if dang_toan == "[D11_C1_B3_10]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_10()

                                                #[D11_C1_B3_11]-M2. Tìm khẳng định đúng về công thức tích thành tổng
                                                if dang_toan == "[D11_C1_B3_11]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_11()

                                                #[D11_C1_B3_12]-TF-M2. Cho tanx. Xét Đ-S: cotx, cos2x, sin2x, tan2x
                                                if dang_toan == "[D11_C1_B3_12]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B3_12()

                                                #[D11_C1_B3_13]-TF-M2. Cho sinx . Xét Đ-S: cosx, sin2x, sin(x+a), cos(x+b)
                                                if dang_toan == "[D11_C1_B3_13]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B3_13()

                                                #[D11_C1_B3_14]-M2. Cho sinx. Tính sin(x+a) hoặc cos(x+b)
                                                if dang_toan == "[D11_C1_B3_14]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B3_14()

                                                ####################### Bài 4: Hàm số lượng giác và đồ thị. ######################
                                                #[D11_C1_B4_01].Tìm tập xác định hàm số y=a/sinx hoặc y=a/cosx
                                                if dang_toan == "[D11_C1_B4_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_01()

                                                #[D11_C1_B4_02].Tìm tập xác định hàm số y=a/sinmx hoặc y=a/cosmx
                                                if dang_toan == "[D11_C1_B4_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_02()

                                                #[D11_C1_B4_03]. Tìm tập xác định hàm số y=tan(ax+bpi)
                                                if dang_toan == "[D11_C1_B4_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_03()

                                                #[D11_C1_B4_04]. Tìm GTLN,GTNN của hàm số sin, cos
                                                if dang_toan == "[D11_C1_B4_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_04()

                                                #[D11_C1_B4_05]-TF-M2. Cho hàm số y=acos(mx+n)+b. Xét Đ-S: TXĐ, GTLN, GTNN, TGT
                                                if dang_toan == "[D11_C1_B4_05]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B4_05()                                                

                                                #[D11_C1_B4_06]-TF-M2. Cho hàm số y=acos^2(mx+n)+b. Xét Đ-S: TXĐ, GTLN, GTNN, TGT
                                                if dang_toan == "[D11_C1_B4_06]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B4_06()

                                                #[D11_C1_B4_07]-TF-M2. Cho y=acosmx+b hoặc y=asinmx+b. Xét Đ-S: TXĐ, Chẵn-Lẻ, TGT, Cắt trục Oy
                                                if dang_toan == "[D11_C1_B4_07]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D11_C1.ngh_kjg_L11_C1_B4_07()

                                                #[D11_C1_B4_08]. Tìm tập giá trị của HSLG
                                                if dang_toan == "[D11_C1_B4_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_08()

                                                #[D11_C1_B4_09]. Tìm chu kì tuần hoàn của HSLG
                                                if dang_toan == "[D11_C1_B4_09]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B4_09()

                                                ################  Bài 5: Phương trình lượng giác.################
                                                
                                                #[D11_C1_B5_01]. Giải phương trình cosx = m
                                                if dang_toan == "[D11_C1_B5_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_01()

                                                #[D11_C1_B5_02]. Giải phương trình sinx = m
                                                if dang_toan == "[D11_C1_B5_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_02()                                                

                                                #[D11_C1_B5_03]. Giải phương trình tan cơ bản
                                                if dang_toan == "[D11_C1_B5_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_03()

                                                #[D11_C1_B5_04]. Tìm m để phương trình sin, cos có nghiệm
                                                if dang_toan == "[D11_C1_B5_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_04()

                                                #[D11_C1_B5_05]. Giải phương trình cosx = b
                                                if dang_toan == "[D11_C1_B5_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_05()

                                                #[D11_C1_B5_06]. Giải phương trình sinx = m
                                                if dang_toan == "[D11_C1_B5_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_06()

                                                #[D11_C1_B5_07]. Giải phương trình cot(ax) = m
                                                if dang_toan == "[D11_C1_B5_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_07()

                                                #[D11_C1_B5_08]. Giải phương trình cos(ax+m)=cos(bx+n)
                                                if dang_toan == "[D11_C1_B5_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_08()

                                                #[D11_C1_B5_09]. Giải phương trình cos(ax+m)=sin(bx+n)
                                                if dang_toan == "[D11_C1_B5_09]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_09()

                                                #[D11_C1_B5_10]. Giải phương trình sin(ax+m)=sin(bx+n)
                                                if dang_toan == "[D11_C1_B5_10]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_10()

                                                #[D11_C1_B5_11]-M2. Giải phương trình tan(ax+m)=tan(bx+n)
                                                if dang_toan == "[D11_C1_B5_11]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_11()

                                                #[D11_C1_B5_12]-M2. Giải phương trình cot(ax+m)=cot(bx+n)
                                                if dang_toan == "[D11_C1_B5_12]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_12()

                                                #[D11_C1_B5_13]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của sinu=m
                                                if dang_toan == "[D11_C1_B5_13]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_13()

                                                #[D11_C1_B5_14]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của sinu=m
                                                if dang_toan == "[D11_C1_B5_14]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_14()

                                                #[D11_C1_B5_15]-TL-M3. Tìm số nghiệm thuộc khoảng đoạn của tanu=m
                                                if dang_toan == "[D11_C1_B5_15]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D11_C1.ngh_kjg_L11_C1_B5_15()

                                    ######### Toán 11_ Chương 2 _ Dãy số _ Cấp số cộng _ Cấp số nhân ########
                                                #---------------------- Bài 1 - Dãy số -------------------------->
                                                #[D11_C2_B1_01]-M2. Cho dãy số có SHTQ. Hỏi số là số hạng thứ mấy
                                                if dang_toan == "[D11_C2_B1_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B1_01()

                                                #[D11_C2_B1_02]-M2. Cho dãy số. Tìm 3 số đầu tiên
                                                if dang_toan == "[D11_C2_B1_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B1_02()

                                                #[D11_C2_B1_03]-M3.  Cho dãy số dạng truy hồi. Tìm công thức tổng quát
                                                if dang_toan == "[D11_C2_B1_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B1_03()

                                                #[D11_C2_B1_04]-M3.  Cho dãy số có SHQT xét tính tăng giảm, bị chặn
                                                if dang_toan == "[D11_C2_B1_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B1_04()

                                                ####################### Bài 2 - Cấp số cộng ######################
                                                #[DS11_C2_B2_01]. Cho cấp số cộng có u1, d tìm u_k
                                                if dang_toan == "[D11_C2_B2_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_01()

                                                #[D11_C2_B2_02]. Cho cấp số cộng có u1, d . Tính S_n
                                                if dang_toan == "[D11_C2_B2_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_02()

                                                #[D11_C2_B2_03]. Cho cấp số cộng có u_k,u_m. Tìm u1
                                                if dang_toan == "[D11_C2_B2_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_03()

                                                #[D11_C2_B2_04]. Cho cấp số cộng có u_k,u_m. Tìm d
                                                if dang_toan == "[D11_C2_B2_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_04()

                                                #[D11_C2_B2_05]. Cho CSC có u1,d. Xét Đ-S: u_k, số hạng thứ mấy, S_k, u_n
                                                if dang_toan == "[D11_C2_B2_05]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C2.mn8mn_L11_C2_B2_05()

                                                #[D11_C2_B2_06]. Cho CSC có un. Tìm d
                                                if dang_toan == "[D11_C2_B2_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_06()

                                                #[D11_C2_B2_07]. Nhận dạng cấp số cộng từ SHTQ
                                                if dang_toan == "[D11_C2_B2_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_07()

                                                #[D11_C2_B2_08]-M2. Nhận dạng dãy số hữu hạn là 1 cấp số cộng
                                                if dang_toan == "[D11_C2_B2_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B2_08()

                                                #[D11_C2_B2_09]-TF-M3. CSC có u_m,u_n. Xét Đ-S: u_1,d, số hạng thứ mấy, S_k, u_n
                                                if dang_toan == "[D11_C2_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C2.mn8mn_L11_C2_B2_09()

                                                #[D11_C2_B2_10]-TF-M3. CSC có u_1,S_n. Xét Đ-S: u_1,d, u_k, S_k, u_n
                                                if dang_toan == "[D11_C2_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C2.mn8mn_L11_C2_B2_10()

                                                #[D11_C2_B2_11]-TF-M3. Xét đúng-sai từ bài toán thực tế liên quan cấp số cộng
                                                if dang_toan == "[D11_C2_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C2.mn8mn_L11_C2_B2_11()

                                                #[D11_C2_B2_12]-TL-M3. CSC có u_k,u_m. Tính u_n
                                                if dang_toan == "[D11_C2_B2_12]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_12()

                                                #[D11_C2_B2_13]-TL-M3. CSC có u_k,u_m. Tính S_n
                                                if dang_toan == "[D11_C2_B2_13]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_13()

                                                #[D11_C2_B2_14]-TL-M3. CSC có u_m+u_n=S. Tính u_p+u_q
                                                if dang_toan == "[D11_C2_B2_14]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_14()

                                                #[D11_C2_B2_15]-TL-M3. Tính giá của một cái máy sau n năm sử dụng
                                                if dang_toan == "[D11_C2_B2_15]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_15()

                                                #[D11_C2_B2_16]-TL-M3. Bài toán trồng cây theo cấp số cộng
                                                if dang_toan == "[D11_C2_B2_16]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_16()

                                                #[D11_C2_B2_17]-TL-M3. Bài toán tính số vòng hoa thỏa mãn cấp số cộng
                                                if dang_toan == "[D11_C2_B2_17]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_17()

                                                #[D11_C2_B2_18]-TL-M2. Bài toán tính số vật khi xếp chồng và giảm dần theo cấp số cộng
                                                if dang_toan == "[D11_C2_B2_18]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_18()

                                                #[D11_C2_B2_19]-TL-M3. Bài toán tính số hàng khi xếp đồ vật và giảm dần theo cấp số cộng
                                                if dang_toan == "[D11_C2_B2_19]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D11_C2.mn8mn_L11_C2_B2_19()

                                              
                                                 ####################### Bài 3 ######################
                                                #D11_C2_B3_01. Cho cấp số nhân có u1, q. Tìm số hạng thứ k
                                                if dang_toan == "[D11_C2_B3_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_01()

                                                #D11_C2_B3_02. Cho cấp số nhân có u_n, u_(n+1). Tìm q.
                                                if dang_toan == "[D11_C2_B3_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_02()

                                                #D11_C2_B3_03. Cho cấp số nhân có u1, q. Tính tổng vài số hạng.
                                                if dang_toan == "[D11_C2_B3_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_03()

                                                #D11_C2_B3_04. Cho cấp số nhân có u1, q. Tính tổng S_n.
                                                if dang_toan == "[D11_C2_B3_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_04()

                                                #D11_C2_B3_05. (TH) Tính số virut tạo ra sau một số khoảng thời gian.
                                                if dang_toan == "[D11_C2_B3_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_05()

                                                #D11_C2_B3_06. (VDT) Tính số virut tạo ra sau một số khoảng thời gian.
                                                if dang_toan == "[D11_C2_B3_06]": 
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_06()

                                                #[D11_C2_B3_07]-TF-M3. Cho CSN có un,u_n+1. Xét Đ-S: q, u_k, S_k, u_n
                                                if dang_toan == "[D11_C2_B3_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C2.mn8mn_L11_C2_B3_07()

                                                #[D11_C2_B3_08]-M2. Nhận dạng dãy số hữu hạn là 1 cấp số nhân
                                                if dang_toan == "[D11_C2_B3_08]": 
                                                   debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C2.mn8mn_L11_C2_B3_08()

                                                



                                    ######### Toán 11_ Chương 3 _ Giới hạn _ Hàm số liên tục ########
                                                #Bài 1: GIỚI HẠN DÃY SỐ

                                                #[D11_C3_B1_01]. Tính giới hạn phân thức bậc tử = mẫu.
                                                if dang_toan == "[D11_C3_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_01()

                                                #[D11_C3_B1_02]. Tính giới hạn phân thức bậc tử < mẫu.
                                                if dang_toan == "[D11_C3_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_02()

                                                #[D11_C3_B1_03]. Tính giới hạn phân thức can(an^2+bn+c)/(dn+e).
                                                if dang_toan == "[D11_C3_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_03()

                                                #[D11_C3_B1_04]. Tính giới hạn phân thức chứa lũy thừa.
                                                if dang_toan == "[D11_C3_B1_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_04()

                                                #[D11_C3_B1_05]. Tính tổng cấp số nhân lùi vô hạn có u1=1.
                                                if dang_toan == "[D11_C3_B1_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_05()

                                                #[D11_C3_B1_06]. Tính tổng cấp số nhân lùi vô hạn có u1 tùy ý.
                                                if dang_toan == "[D11_C3_B1_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_06()

                                                #[D11_C3_B1_07]. Cho giới hạn của (u_n) và (v_n). Tính giới hạn tổng hiệu.
                                                if dang_toan == "[D11_C3_B1_07]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B1_07()

                                                #Bài 2: GIỚI HẠN HÀM SỐ
                                                #[D11_C3_B2_09]. Cho limf(x) và limg(x). Tính lim f(x)g(x) hoặc lim f(x)/g(x).
                                                if dang_toan == "[D11_C3_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_09()

                                                #[D11_C3_B2_01]. Tính giới hạn tại điểm - thay số trực tiếp
                                                if dang_toan == "[D11_C3_B2_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_01()

                                                #[D11_C3_B2_02]. Tính giới hạn dạng 0/0 - Bậc 2/Bậc 1
                                                if dang_toan == "[D11_C3_B2_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_02()

                                                #[D11_C3_B2_03]. Tính giới hạn dạng 0/0 - Bậc 2/Bậc 2
                                                if dang_toan == "[D11_C3_B2_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_03()

                                                #[D11_C3_B2_04]. Tính giới hạn dạng 0/0 - [Căn(ax+b)-c]/(a[2]x+d)
                                                if dang_toan == "[D11_C3_B2_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_04()

                                                #[D11_C3_B2_05]. Tính giới hạn x-->00 Bậc tử = Bậc mẫu
                                                if dang_toan == "[D11_C3_B2_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_05()

                                                #[D11_C3_B2_06]. Tính giới hạn x-->00 Bậc tử < Bậc mẫu
                                                if dang_toan == "[D11_C3_B2_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_06()

                                                #[D11_C3_B2_07]. Tính giới hạn x-->00: Căn(A)/B
                                                if dang_toan == "[D11_C3_B2_07]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_07()

                                                #[D11_C3_B2_08]. Tính giới hạn x-->00: A/Căn(B)
                                                if dang_toan == "[D11_C3_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_08()

                                                #[D11_C3_B2_10]. Tính giới hạn x-->00: Đa thức
                                                if dang_toan == "[D11_C3_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_10()

                                                #[D11_C3_B2_11]. Tính giới hạn dạng 0/0 - Bậc 1/Bậc 2
                                                if dang_toan == "[D11_C3_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B2_11()

                                                #Bài 3: HÀM SỐ LIÊN TỤC

                                                #[D11_C3_B3_01]. Cho f(x)=căn(ax+b). Xét tính liên tục tại điểm.
                                                if dang_toan == "[D11_C3_B3_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B3_01()

                                                #[D11_C3_B3_02]. Cho f(x)=(ax+b)/(cx+d). Tìm khoảng liên tục.
                                                if dang_toan == "[D11_C3_B3_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B3_02()

                                                #[D11_C3_B3_03]. Cho f(x) có >=,<. Xét tính liên tục tại điểm.
                                                if dang_toan == "[D11_C3_B3_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B3_03()

                                                #[D11_C3_B3_04]. Cho f(x) có phân thức bậc 3. Tìm m để f(x) liên tục tại x_0.
                                                if dang_toan == "[D11_C3_B3_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B3_04()

                                                #[D11_C3_B3_05]. Cho f(x) có phân thức bậc 2. Tìm m để f(x) liên tục tại x_0.
                                                if dang_toan == "[D11_C3_B3_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C3.zz8zz_L11_C3_B3_05()

                                    ######### Toán 11_ Chương 4-QUAN HỆ SONG SONG ########

                                    #Bài 1: ĐƯỜNG THẲNG VÀ MẶT PHẲNG TRONG KHÔNG GIAN
                                                #[D11_C4_B1_01]-M2. Cho tứ diện. Tìm khẳng định đúng về quan hệ điểm và mặt phẳng
                                                if dang_toan == "[D11_C4_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_01()

                                                #[D11_C4_B1_02]-M2. Cho hình chóp đáy là h.b.h. Xét quan hệ giữa điểm - đường thẳng - mặt phẳng
                                                if dang_toan == "[D11_C4_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_02()

                                                #[D11_C4_B1_03]-M2. Cho hình chóp đáy là h.b.h. Xét quan hệ giữa điểm - đường thẳng - mặt phẳng
                                                if dang_toan == "[D11_C4_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_03()

                                                #[D11_C4_B1_04]-M2. Cho tứ diện hoặc hình chóp tứ giác. Tìm giao tuyến giữa các mặt bên
                                                if dang_toan == "[D11_C4_B1_04]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_04()

                                                #[D11_C4_B1_05]-M2. Cho tứ diện. Tìm giao tuyến của các mặt phẳng
                                                if dang_toan == "[D11_C4_B1_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_05()

                                                #[D11_C4_B1_06]-M2. Cho hình chóp đáy là h.b.h.Tìm giao điểm đường thẳng - mặt phẳng
                                                if dang_toan == "[D11_C4_B1_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B1_06()

                                    #Bài 2: HAI ĐƯỜNG THẲNG SONG SONG
                                                #[D11_C4_B2_01]. Cho hình chóp. Xét hai đường thẳng song song - đường trung bình.
                                                if dang_toan == "[D11_C4_B2_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B2_01()

                                    #Bài 3: ĐƯỜNG THẲNG SONG SONG MẶT PHẲNG
                                    #[D11_C4_B3_01]. Cho hình chóp đáy h.b.h. Xét sự song song của một đường thẳng với các mp.
                                                if dang_toan == "[D11_C4_B3_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B3_01()

                                    #[D11_C4_B3_02]. Cho hình chóp đáy hình thang. Xét sự song song của một đường thẳng với các mp.
                                                if dang_toan == "[D11_C4_B3_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B3_02()

                                    #Bài 4: HAI MẶT PHẲNG SONG SONG

                                    #[D11_C4_B4_03]. Cho hình lăng trụ tam giác. Xét sự song song của 2 mp-Nhận biết.
                                                if dang_toan == "[D11_C4_B4_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_03()                                    

                                    #[D11_C4_B4_05]. Tìm khẳng định đúng về hình lăng trụ.
                                                if dang_toan == "[D11_C4_B4_05]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_05()

                                    #[D11_C4_B4_06]. Tìm cặp điểm đối diện của hình hộp.
                                                if dang_toan == "[D11_C4_B4_06]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_06()

                                    #[D11_C4_B4_07]. Cho hình hộp. Xét sự song song của một đường với các đường
                                                if dang_toan == "[D11_C4_B4_07]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_07()
                                                    
                                    #[D11_C4_B4_08]. Cho hình hộp. Xét sự song song của hai đường tùy ý
                                                if dang_toan == "[D11_C4_B4_08]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_08()

                                    #[D11_C4_B4_09]. Cho hình lăng trụ tam giác. Xét sự song song của 2 mp-Thông hiểu.
                                                if dang_toan == "[D11_C4_B4_09]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C4.ghj_7_jkl_L11_C4_B4_09()

                                    ######### Toán 11_ Chương 5-MẪU SỐ LIỆU GHÉP NHÓM ########
                                    #Bài 1: SỐ TRUNG BÌNH VÀ MỐT CỦA MẪU SỐ LIỆU GHÉP NHÓM
                                                #[D11_C5_B1_01]. Cho bảng số liệu khách hàng. Tìm giá trị đại diện.
                                                if dang_toan == "[D11_C5_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B1_01()                                     
                                                    

                                                #[D11_C5_B1_02]. Cho bảng số liệu ghép nhóm ngẫu nhiên. Tìm giá trị đại diện.
                                                if dang_toan == "[D11_C5_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B1_02()                                     
                                                    

                                                #[D11_C5_B1_03]. Cho bảng số liệu ghép nhóm ngẫu nhiên. Tính số trung bình.
                                                if dang_toan == "[D11_C5_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B1_03()                                     
                                                    

                                    #Bài 2: TRUNG VỊ VÀ TỨ PHÂN VỊ VUA MẪU SỐ LIỆU GHÉP NHÓM

                                    #[D11_C5_B2_01]-M3. Tìm số trung bình của mẫu số liệu ghép nhóm ngẫu nhiên.
                                                if dang_toan == "[D11_C5_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B2_01()                                     
                                                    

                                    #[D11_C5_B2_02]-M2. Tìm mốt của mẫu số liệu ghép nhóm ngẫu nhiên.
                                                if dang_toan == "[D11_C5_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B2_02()

                                    #[D11_C5_B2_03]-M2. Tìm tứ phân vị Q1 của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D11_C5_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B2_03()

                                    #[D11_C5_B2_04]-M2. Tìm tứ phân vị Q2 của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D11_C5_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B2_04()

                                    #[D11_C5_B2_05]-M2. Tìm tứ phân vị Q3 của mẫu số liệu ghép nhóm.
                                                if dang_toan == "[D11_C5_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C5.treqw_L11_C5_B2_05()                                  
                                                    

                                
                                    ################ Bài 1: PHÉP TÍNH LŨY THỪA ########################
                                    #[D11_C6_B1_01]. Rút gọn a^m.a^n với m,n là phân số.                      
                                                if dang_toan == "[D11_C6_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_01()                                     
                                                    

                                    #[D11_C6_B1_02]-M2. Rút gọn (a^m.a^n)/a^p với m,n,p là phân số.                     
                                                if dang_toan == "[D11_C6_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_02()                                     
                                                    

                                    #[D11_C6_B1_03]-M1. Biểu diễn căn bậc sang lũy thừa số mũ hữu tỷ.                    
                                                if dang_toan == "[D11_C6_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_03()                                     
                                                    

                                    #[D11_C6_B1_04]-M1. Biểu diễn tích chứa 2 căn thành lũy thừa.                    
                                                if dang_toan == "[D11_C6_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_04()                                     
                                                    

                                    #[D11_C6_B1_05]-M3. Biểu diễn tích chứa 3 căn thành lũy thừa.                    
                                                if dang_toan == "[D11_C6_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_05()                                     
                                                    

                                     #[D11_C6_B1_05]-M3. Biểu diễn tích chứa 3 căn thành lũy thừa.                    
                                                if dang_toan == "[D11_C6_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B1_06()                                     
                                                    
                                    #[D11_C6_B1_07]-TF-M2. Tạo câu đúng-sai: Tính chất lũy thừa.  
                                                if dang_toan == "[D11_C6_B1_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B1_07()                                      


                                    #[D11_C6_B1_08]-TF-M2. Tạo câu đúng-sai: Tính chất của căn bậc n.  
                                                if dang_toan == "[D11_C6_B1_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B1_08()                                      


                                    #Bài 2: PHÉP TÍNH LÔGARIT
                                    #[D11_C6_B2_01]-M1. Tìm khẳng định đúng về log_a (a^m)
                                                if dang_toan == "[D11_C6_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_01()                                     
                                                    

                                    #[D11_C6_B2_02]-M1. Tìm khẳng định đúng về log_a (1/a^m)
                                                if dang_toan == "[D11_C6_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_02()                                     
                                                    

                                    #[D11_C6_B2_03]-M2. Tìm khẳng định đúng về log_can[n](a) (1/a^m)
                                                if dang_toan == "[D11_C6_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_03()                                     
                                                    

                                    #[D11_C6_B2_04]-M1. Tính giá trị biểu thức chứa logarit bằng máy tính
                                                if dang_toan == "[D11_C6_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_04()                                     
                                                    

                                    #[D11_C6_B2_05]-M2. Biễu diễn một logarit theo một logarit khác
                                                if dang_toan == "[D11_C6_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_05()                                     
                                                    

                                    #[D11_C6_B2_06]-M2. Tìm khẳng định đúng của log_a(a^m.b^n)
                                                if dang_toan == "[D11_C6_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_06()                                     
                                                    

                                            #D11_C6_B2_07]-M3. Biểu diễn một logarit theo 2 logarit khác.
                                                if dang_toan == "[D11_C6_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B2_07()                                     
                                                    

                                    

                                    #Bài 3: HÀM SỐ MŨ- HÀM SỐ LOGARIT 

                                                #[D11_C6_B3_01]. Tập xác định hàm số y=(ax+b)^n với n là số nguyên âm.
                                                if dang_toan == "[D11_C6_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_01()                                     
                                                    

                                                #[D11_C6_B3_02]. Tập xác định hàm số y=(ax+b)^n với n là không nguyên.
                                                if dang_toan == "[D11_C6_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_02()                                     
                                                    

                                    #[D11_C6_B3_03]-M2. TXĐ hàm số y=(ax^2+bx+c)^n với n là số nguyên âm.
                                                if dang_toan == "[D11_C6_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_03()                                     
                                                    

                                                #[D11_C6_B3_04]. Tập xác định hàm số y=(ax^2+bx+c)^n với n là không nguyên.
                                                if dang_toan == "[D11_C6_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_04()                                     
                                                    

                                                #[D11_C6_B3_05]. Tập xác định hàm số y=log(ax+b).
                                                if dang_toan == "[D11_C6_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_05()                                     
                                                    

                                                #[D11_C6_B3_06]. Tập xác định hàm số y=log(ax^2+bx+c).
                                                if dang_toan == "[D11_C6_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_06()                                     
                                                    

                                                #[D11_C6_B3_07]. Đồ thị hàm số y=a^x.
                                                if dang_toan == "[D11_C6_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_07()                                     
                                                    

                                                #[D11_C6_B3_08]. Đồ thị hàm số y=log_a(x).
                                                if dang_toan == "[D11_C6_B3_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_08()                                     
                                                    

                                                #[D11_C6_B3_09]-M2. Đồ thị hàm số y=căn(a)^x
                                                if dang_toan == "[D11_C6_B3_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_09()                                     
                                                    

                                                #[D11_C6_B3_10]-M3. Tìm m để log (ax^2 +bx+c) có tập xác định là R.
                                                if dang_toan == "[D11_C6_B3_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_10()                                     
                                                    

                                            #[D11_C6_B3_11]-M3. Tìm m để e^a/căn(ax^2 +bx+c) có tập xác định là R.
                                                if dang_toan == "[D11_C6_B3_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_11()                                     
                                                    

                                            #[D11_C6_B3_12]-TF-M2. Cho y=a^x, a>1. Tạo câu Đ-S.
                                                if dang_toan == "[D11_C6_B3_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_12()                                      


                                            #[D11_C6_B3_13]-TF-M2. Cho y=a^x, 0<a<1. Tạo câu Đ-S.
                                                if dang_toan == "[D11_C6_B3_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_13()                                      


                                            #[D11_C6_B3_14]-TF-M2. Cho y=log_a x, a>1. Tạo câu Đ-S.
                                                if dang_toan == "[D11_C6_B3_14]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_14()                                      


                                            #[D11_C6_B3_15]-TF-M2. Cho y=log_a x, 0<a<1. Tạo câu Đ-S.
                                                if dang_toan == "[D11_C6_B3_15]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_15()                                      


                                            #[D11_C6_B3_16]-M2. Đồ thị hàm số y=a^x, a>1
                                                if dang_toan == "[D11_C6_B3_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_16()                                     
                                                    

                                            #[D11_C6_B3_17]-M2. Đồ thị hàm số y=a^x, 0<a<1
                                                if dang_toan == "[D11_C6_B3_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_17()                                     
                                                    

                                            #[D11_C6_B3_18]-M2. Đồ thị hàm số y=log_a x, a>1
                                                if dang_toan == "[D11_C6_B3_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_18()                                     
                                                    

                                            #[D11_C6_B3_19]-M2. Đồ thị hàm số y=log_a x, 0<a<1
                                                if dang_toan == "[D11_C6_B3_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_19()                                     
                                                    

                                            #[D11_C6_B3_20]-M2. Cho hàm số y=a^x. Xét tính đơn điệu
                                                if dang_toan == "[D11_C6_B3_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_20()                                     
                                                    

                                            #[D11_C6_B3_21]-M2. Cho hàm số y=log_a x. Xét tính đơn điệu
                                                if dang_toan == "[D11_C6_B3_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B3_21()                                     
                                                    
                                            
                                            #[D11_C6_B3_22]-TF-M2. Cho đồ thị y=a^x, 0<a<1. Tạo Đúng-Sai.                                   
                                                if dang_toan == "[D11_C6_B3_22]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_22()                                      


                                            #[D11_C6_B3_23]-TF-M2. Cho đồ thị y=log_a x, a>1. Tạo Đúng-Sai.                                  
                                                if dang_toan == "[D11_C6_B3_23]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_23()                                      


                                            #[D11_C6_B3_24]-TF-M2. Cho đồ thị y=log_a x, 0<a<1. Tạo Đúng-Sai.                                  
                                                if dang_toan == "[D11_C6_B3_24]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_24()                                      


                                            #[D11_C6_B3_25]-TF-M2. Cho đồ thị y=log_a x, 0<a<1. Tạo Đúng-Sai.                                  
                                                if dang_toan == "[D11_C6_B3_25]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B3_25()                                      


                                            #Chương 7 - ĐẠO HÀM
                                            #Bài 1 - Định nghĩa đạo hàm

                                            #[D11_C7_B1_01]-M1. Tìm khẳng định đúng về đạo hàm tại điểm của f(x)
                                                if dang_toan == "[D11_C7_B1_01]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B1_01()                                     
                                                    

                                            #[D11_C7_B1_02]-M2. Cho hàm số đa thức. Biễu diễn công thức đạo hàm tại x_0
                                                if dang_toan == "[D11_C7_B1_02]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B1_02()                                     
                                                    
                                            #Bài 2 - Đạo hàm các hàm số cơ bản

                                            #[D11_C7_B2_01]. Tính đạo hàm của hàm số đa thức.
                                                if dang_toan == "[D11_C7_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_01()                                     
                                                    

                                            #[D11_C7_B2_02]. Tính đạo hàm của hàm số đa thức + a/x
                                                if dang_toan == "[D11_C7_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_02()                                     
                                                    

                                            #[D11_C7_B2_03]. Tính đạo hàm của hàm số đa thức + acăn(x)
                                                if dang_toan == "[D11_C7_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_03()                                     
                                                    

                                            #[D11_C7_B2_04]. Tính đạo hàm của hàm số đa thức + acăn(x)
                                                if dang_toan == "[D11_C7_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_04()                                     
                                                    

                                            #[D11_C7_B2_05]. Tính đạo hàm của hàm số y=asinu
                                                if dang_toan == "[D11_C7_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_05()                                     
                                                    

                                            #[D11_C7_B2_06]. Tính đạo hàm của hàm số y=acosu
                                                if dang_toan == "[D11_C7_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_06()                                     
                                                    

                                            #[D11_C7_B2_07]. Tính đạo hàm của hàm số y=asinx+bcosx
                                                if dang_toan == "[D11_C7_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_07()                                     
                                                    

                                            #[D11_C7_B2_08]. Tính đạo hàm của y=a^x.
                                                if dang_toan == "[D11_C7_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_08()                                     
                                                    

                                                #[D11_C7_B2_09]. Tính đạo hàm của y=a^u.
                                                if dang_toan == "[D11_C7_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_09()                                     
                                                    

                                                #[D11_C7_B3_10]. Tìm đạo hàm của y=m/(ax+b)
                                                if dang_toan == "[D11_C7_B2_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_10()                                     
                                                    

                                                #[D11_C7_B3_11]. Tìm đạo hàm của y=(mx^2+nx+p)/(ax+b)
                                                if dang_toan == "[D11_C7_B2_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_11()                                     
                                                    

                                                #[D11_C7_B3_12]. Tìm đạo hàm của y = đa thức + e^x.
                                                if dang_toan == "[D11_C7_B2_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_12()                                     
                                                    

                                                #[D11_C7_B3_13]. Tìm đạo hàm của y = e^(đa thức).
                                                if dang_toan == "[D11_C7_B2_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_13()                                     
                                                    

                                                #[D11_C7_B3_14]. Tìm đạo hàm của y=atanx + bcotx
                                                if dang_toan == "[D11_C7_B2_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_14()                                     
                                                    

                                                #[D11_C7_B3_15]. Tìm đạo hàm của y=atanu
                                                if dang_toan == "[D11_C7_B2_15]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_15()                                     
                                                    

                                                #[D11_C7_B3_16]. Tìm đạo hàm của y=acotu
                                                if dang_toan == "[D11_C7_B2_16]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_16()                                     
                                                    

                                                #[D11_C7_B3_17]. Tìm đạo hàm y=(ax+b).căn(x)
                                                if dang_toan == "[D11_C7_B2_17]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_17()                                     
                                                    

                                                #[D11_C7_B3_18]. Tìm đạo hàm y=(ax+b).sin(x)
                                                if dang_toan == "[D11_C7_B2_18]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_18()                                     
                                                    

                                                #[D11_C7_B3_19]. Tìm đạo hàm y=(ax+b).cos(x)
                                                if dang_toan == "[D11_C7_B2_19]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_19()                                     
                                                    

                                                #[D11_C7_B3_20]. Tìm đạo hàm y=(ax+b).cos(x)
                                                if dang_toan == "[D11_C7_B2_20]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_20()                                     
                                                    

                                                #[D11_C7_B2_21]-M1. Tìm đạo hàm y=log_a x
                                                if dang_toan == "[D11_C7_B2_21]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_21()                                     
                                                    

                                                #[D11_C7_B2_22]-M2. Tìm đạo hàm y=log_a (đa thức)
                                                if dang_toan == "[D11_C7_B2_22]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_22()                                     
                                                    

                                                #[D11_C7_B2_23]-M2. Tìm đạo hàm y=ln (đa thức)
                                                if dang_toan == "[D11_C7_B2_23]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_23()                                     
                                                    

                                                #[D11_C7_B2_24]-M2. Tìm đạo hàm y=(đa thức).a^x
                                                if dang_toan == "[D11_C7_B2_24]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_24()                                     
                                                    

                                                #[D11_C7_B2_25]-M2. Tìm đạo hàm y=(ax+b)/cos(x) hoặc y=(ax+b)/sin(x)
                                                if dang_toan == "[D11_C7_B2_25]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_25()                                     
                                                    

                                                #[D11_C7_B2_26]-M2. Tìm đạo hàm y=căn(đa thức)
                                                if dang_toan == "[D11_C7_B2_26]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_26()                                     
                                                    

                                                #[D11_C7_B2_27]-M2. Tìm đạo hàm y=căn(sin(ax+b))
                                                if dang_toan == "[D11_C7_B2_27]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_27()                                     
                                                    

                                                #[D11_C7_B2_28]-M2. Tìm đạo hàm y=căn(cos(ax+b))
                                                if dang_toan == "[D11_C7_B2_28]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_28()                                     
                                                    

                                                #[D11_C7_B2_29]-M2. Tìm đạo hàm y=m/u
                                                if dang_toan == "[D11_C7_B2_29]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_29()                                     
                                                    

                                                #[D11_C7_B2_30]-M2. Tìm đạo hàm y=(đa thức)^n
                                                if dang_toan == "[D11_C7_B2_30]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_30()                                     
                                                    

                                                #[D11_C7_B2_31]-M2. Đạo hàm cấp 2 của đa thức
                                                if dang_toan == "[D11_C7_B2_31]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_31()                                     
                                                    

                                                #[D11_C7_B2_32]-M2. Đạo hàm cấp 2 của sin(ax+b)
                                                if dang_toan == "[D11_C7_B2_32]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_32()                                     
                                                    

                                                #[D11_C7_B2_33]-M2. Đạo hàm cấp 2 của cos(ax+b)
                                                if dang_toan == "[D11_C7_B2_33]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_33()                                     
                                                    

                                                #[D11_C7_B2_34]-M2. Đạo hàm cấp 2 của y=(ax+b)/(cx+d)
                                                if dang_toan == "[D11_C7_B2_34]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_34()                                     
                                                    

                                                #[D11_C7_B2_35]-M2. Đạo hàm của y=(ax+b)/(cx+d)
                                                if dang_toan == "[D11_C7_B2_35]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_35()                                     
                                                    

                                                #[D11_C7_B2_36]-M2. Cho hàm số đa thức. Giải phương trình chứa đạo hàm cấp hai.
                                                if dang_toan == "[D11_C7_B2_36]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_36()                                     
                                                    

                                                #[D11_C7_B2_37]-M2. Cho hàm số đa thức.Tính đạo hàm cấp hai tại điểm x_0.
                                                if dang_toan == "[D11_C7_B2_37]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_37()                                     
                                                    

                                                #[D11_C7_B2_38]-M2. Cho hàm số quãng đường.Tính vận tốc tại thời điểm t_0.
                                                if dang_toan == "[D11_C7_B2_38]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_38()                                     
                                                    

                                                #[D11_C7_B2_39]-M2. Cho hàm số quãng đường.Tính gia tốc tại thời điểm t_0.
                                                if dang_toan == "[D11_C7_B2_39]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_39()

                                                #[D11_C7_B2_40]-M2. f(x) bậc 3. Giải phương trình f'(x)=a
                                                if dang_toan == "[D11_C7_B2_40]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_40()

                                                #[D11_C7_B2_41]-M2. Cho f(x) là hàm số bậc 3. Giải BPT f'(x)+a>0 (<0,>=0,<=0)
                                                if dang_toan == "[D11_C7_B2_41]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_41()

                                                #[D11_C7_B2_42]-M2. Cho f(x)=a.căn(x)+ bx +c. Giải BPT f'(x)+a>0 (<0,>=0,<=0)
                                                if dang_toan == "[D11_C7_B2_42]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_42()

                                                #[D11_C7_B2_43]-M2. Cho f(x) là hàm bậc 3. Tìm m để f'(x)>0 (<0) với mọi x.
                                                if dang_toan == "[D11_C7_B2_43]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B2_43()       
                                                
                                                #BÀI 3-PHƯƠNG TRÌNH TIẾP TUYẾN
                                                #[D11_C7_B3_01]-M2. Cho hàm số đa thức. Viết phương trình tiếp tuyến tại điểm (x_0;y_0)

                                                if dang_toan == "[D11_C7_B3_01]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B3_01()                                     
                                                    

                                                #[D11_C7_B3_02]-M2. Cho hàm số đa thức. Viết phương trình tiếp tuyến tại điểm có hoành độ x_0
                                                if dang_toan == "[D11_C7_B3_02]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B3_02()                                     
                                                    

                                                #[D11_C7_B3_03]-M2. Cho hàm số đa thức. Viết phương trình tiếp tuyến tại điểm có tung độ y_0
                                                if dang_toan == "[D11_C7_B3_03]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B3_03()                                     
                                                    

                                                #[D11_C7_B3_04]-M2. Cho hàm số. Tính hệ số góc của tiếp tuyến tại điểm x_0
                                                if dang_toan == "[D11_C7_B3_04]":          
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C7.ui5io_L11_C7_B3_04()                                     
                                                    
                                                #Câu hỏi đúng sai
                                                #[D11_C7_B4_01]-TF-M2. Tạo câu đúng-sai: Đạo hàm của hàm số đa thức  
                                                if dang_toan == "[D11_C7_B4_01]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C7.ui5io_L11_C7_B4_01()                                      

                                                #[D11_C7_B4_02]-TF-M2. Tạo câu đúng-sai: Đạo hàm của hàm số y=(ax+b)/(cx+d)  
                                                if dang_toan == "[D11_C7_B4_02]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C7.ui5io_L11_C7_B4_02()                                      


                                                #[D11_C7_B4_03]-TF-M2. Tạo câu đúng-sai: Đạo hàm của y=msin(ax)+cos(x+b) 
                                                if dang_toan == "[D11_C7_B4_03]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C7.ui5io_L11_C7_B4_03()                                      


                                                #[D11_C7_B4_04]-TF-M2. Tạo câu đúng-sai: Đạo hàm của y=msin(ax)+cos(x+b) 
                                                if dang_toan == "[D11_C7_B4_04]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C7.ui5io_L11_C7_B4_04()                                      


                                                #[D11_C7_B4_05]-TF-M2. Tạo câu đúng-sai: Đạo hàm của y=msin(ax)+cos(x+b) 
                                                if dang_toan == "[D11_C7_B4_05]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C7.ui5io_L11_C7_B4_05()                                      


                                    #BÀI 4 PHƯƠNG TRÌNH MŨ - PHƯƠNG TRÌNH LOGARIT

                                                #[D11_C6_B4_01]. Giải phương trình a^x=b.
                                                if dang_toan == "[D11_C6_B4_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_01()                                     
                                                    

                                                #[D11_C6_B4_02]. Giải phương trình a^(x-m)=b.
                                                if dang_toan == "[D11_C6_B4_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_02()                                     
                                                    

                                                #[D11_C6_B4_03]. Phương trình dạng a^(mx+n)=b^(px+q)
                                                if dang_toan == "[D11_C6_B4_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_03()                                     
                                                    

                                                #[D11_C6_B4_04]. Giải phương trình bậc 2 đối với a^x.
                                                if dang_toan == "[D11_C6_B4_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_04()                                     
                                                    

                                                #[D11_C6_B4_05]. Giải phương trình logarit log_a(x)=b
                                                if dang_toan == "[D11_C6_B4_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_05()                                     
                                                    

                                            #[D11_C6_B4_06]-M2. Giải phương trình log_a(mx+n)=b
                                                if dang_toan == "[D11_C6_B4_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_06()                                     
                                                    

                                            #[D11_C6_B4_07]-M2. Giải phương trình log_m(ax+b)-log_m(cx+d) = e
                                                if dang_toan == "[D11_C6_B4_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_07()                                     
                                                    

                                            #[D11_C6_B4_08]-M2. Giải phương trình log_m(ax+b)+log_m(cx+d) = e
                                                if dang_toan == "[D11_C6_B4_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_08()                                     
                                                    

                                            #[D11_C6_B4_09]-TF-M2. Tạo câu Đ-S: phương trình a^x=b.
                                                if dang_toan == "[D11_C6_B4_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B4_09()                                      


                                            #[D11_C6_B4_10]-TF-M2. Tạo câu Đ-S: log_a x=b.
                                                if dang_toan == "[D11_C6_B4_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B4_10()                                      


                                            #[D11_C6_B4_11]-M3. Giải phương trình m.a^x + n.b^x = p.a^x +q.b^x
                                                if dang_toan == "[D11_C6_B4_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_11()                                     
                                                    

                                            #[D11_C6_B4_12]-M3.  Giải phương trình a^x + a^(x+m) +a^(x-n) = b^x + b^(x+p) + b^(x-q)
                                                if dang_toan == "[D11_C6_B4_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B4_12()                                     
                                                    

                                    #BÀI 5 - BẤT PHƯƠNG TRÌNH MŨ - PHƯƠNG TRÌNH LOGARIT
                                                #[D11_C6_B5_01]. Giải bất phương trình log_a x > b (a>1)
                                                if dang_toan == "[D11_C6_B5_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_01()                                     
                                                    

                                                #[D11_C6_B5_02]. Giải bất phương trình log_a (x-m) > b (a>1)
                                                if dang_toan == "[D11_C6_B5_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_02()                                     
                                                    

                                                #[D11_C6_B5_03]. Giải BPT log_m(ax+b) > log_m(cx+d), (a>1)
                                                if dang_toan == "[D11_C6_B5_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_03()                                     
                                                    

                                                #[D11_C6_B5_04]. Giải bất phương trình a^x > b (a>1)
                                                if dang_toan == "[D11_C6_B5_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_04()                                     
                                                    

                                                #[D11_C6_B5_05]. Giải bất phương trình a^x < b (a>1)
                                                if dang_toan == "[D11_C6_B5_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_05()                                     
                                                    

                                                #[D11_C6_B5_06]. Giải bất phương trình a^x < b (0<a<1)
                                                if dang_toan == "[D11_C6_B5_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_06()                                     
                                                    

                                                #[D11_C6_B5_07]. Giải bất phương trình a^x > b (0<a<1)
                                                if dang_toan == "[D11_C6_B5_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_07()                                     
                                                    

                                                #[D11_C6_B5_08]. Giải BPT m^(ax+b) > m^(cx+d), (m>1)
                                                if dang_toan == "[D11_C6_B5_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_08()                                     
                                                    

                                                #[D11_C6_B5_09]. Giải BPT m^(ax+b) > m^(cx+d), (0<m<1)
                                                if dang_toan == "[D11_C6_B5_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B5_09()                                     
                                                    

                                                #[D11_C6_B5_10]-TF-M2. Tạo câu Đ-S: bất phương trình chứa a^x.
                                                if dang_toan == "[D11_C6_B5_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C6.zz8zz_L11_C6_B5_10()                                      


                                    #BÀI 6 - BÀI TOÁN LÃI SUẤT - TĂNG TRƯỞNG
                                                #[D11_C6_B6_01]. Cho số tiền và lãi suất theo năm. Tính tổng tiền thu được.
                                                if dang_toan == "[D11_C6_B6_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_01()                                     
                                                    

                                                #[D11_C6_B6_02]. Cho số tiền và lãi suất theo tháng. Tính tổng tiền thu được sau n tháng.
                                                if dang_toan == "[D11_C6_B6_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_02()                                     
                                                    

                                                #[D11_C6_B6_03]. Cho số tiền và lãi suất/năm. Tính số năm để thu được khoản tiền nào đó.
                                                if dang_toan == "[D11_C6_B6_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_03()                                     
                                                    

                                                #[D11_C6_B6_04]. Cho số tiền và kì hạn theo quý, lãi suất theo năm. Tính tổng tiền thu được sau n năm.
                                                if dang_toan == "[D11_C6_B6_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_04()                                     
                                                    

                                                #[D11_C6_B6_05]-M2. Cho số dân và tỉ lệ tăng trưởng. Tính số dân sau n năm.
                                                if dang_toan == "[D11_C6_B6_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_05()                                     
                                                    

                                                #[D11_C6_B6_06]-M2. Cho mức tiền lương và tỉ lệ tăng lương. Tính mức lương nhận được sau n năm.
                                                if dang_toan == "[D11_C6_B6_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C6.zz8zz_L11_C6_B6_06()                                     
                                                    

                                    #Toán 11 - Chương 8 - Quan hệ vuông góc
                                    #Bài 1 - Hai đường thẳng vuông góc
                                    #[D11_C8_B1_01]-M1. Cho véctơ pháp tuyến tìm véctơ chỉ phương và ngược lại
                                                if dang_toan == "[D11_C8_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_01()                                    
                                                    

                                    #[D11_C8_B2_02]-M2. S.ABCD: cạnh bên vuông góc đáy h.chữ nhật. Tìm đường vg mặt.
                                                if dang_toan == "[D11_C8_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B1_02()

                                    #[D11_C8_B1_03]-M2. S.ABCD: cạnh v.g đáy h.vuông. Tìm 2 đường vuông góc.
                                                if dang_toan == "[D11_C8_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B1_03()

                                    #[D11_C8_B1_04]-TF-M2. Cho hình lập phương. Tạo câu đúng-sai: góc giữa 2 đường thẳng.
                                                if dang_toan == "[D11_C8_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B1_04()                            
                       
                                                    

                                    #[D11_C8_B2_03]-M2. S.ABCD: cạnh bên vuông góc đáy h.thoi. Tìm đường vg mặt.
                                                if dang_toan == "[D11_C8_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_03()                                     
    
                                                    

                                    #[D11_C8_B2_04]-M2. S.ABC: cạnh bên vuông góc đáy h.tam giác vuông. Tìm đường vg mặt.
                                                if dang_toan == "[D11_C8_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_04()                                     
                                                    

                                    #[D11_C8_B2_05]-M2. S.ABC: cạnh bên vuông góc đáy h.tam giác đều. Tìm đường vg mặt.
                                                if dang_toan == "[D11_C8_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_05()                                     
                                                    

                                    #[D11_C8_B2_06]-M2. S.ABC: cạnh bên vuông góc đáy h.tam giác đều. Tìm đường vg mặt.
                                                if dang_toan == "[D11_C8_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_06()                                     
                                                    

                                    #[D11_C8_B2_07]-M3. S.ABC: cạnh v.g đáy h.tam giác đều. Tìm đường vg đường
                                                if dang_toan == "[D11_C8_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_07()                                     
                                                    

                                    #[D11_C8_B2_08]-TF-M2.  S.ABCD: cạnh v.g đáy h.vuông. Xét tính đúng sai về đường vuông góc đường 
                                                if dang_toan == "[D11_C8_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B2_08()                                       


                                    #[D11_C8_B2_09]-TF-M2.  S.ABCD: cạnh v.g đáy h.vuông. Tạo câu đúng-sai: đường vuông góc mặt.
                                                if dang_toan == "[D11_C8_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B2_09()                                   


                                    #[D11_C8_B2_10]-TF-M2.  S.ABCD: cạnh v.g đáy h.chữ nhật. Tạo câu đúng-sai: đường vuông góc đường.
                                                if dang_toan == "[D11_C8_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B2_10()                                      


                                    #[D11_C8_B2_11]-TF-M2.  S.ABCD: cạnh v.g đáy h.chữ nhật. Tạo câu đúng-sai: đường vuông góc mặt.
                                                if dang_toan == "[D11_C8_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B2_11()                                  

                                    #[D11_C8_B2_12]-TF-M2. S.ABCD: cạnh v.g đáy h.vuông. Tạo câu đúng-sai: lập luận đường vuông góc mặt.
                                                if dang_toan == "[D11_C8_B2_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B2_12()                             


                                    #[D11_C8_B2_13]-M2. S.ABCD: ABCD hình vuông. Tìm hình chiếu của điểm lên mặt phẳng
                                                if dang_toan == "[D11_C8_B2_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_13()                                     
                                                    

                                    #[D11_C8_B2_14]-M2. S.ABCD: ABCD hình vuông. Tìm hình chiếu của điểm lên mặt phẳng
                                                if dang_toan == "[D11_C8_B2_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_14()                                     
                                                    

                                    #[D11_C8_B2_15]-M2. S.ABC: ABC tam giác vuông. Tìm hình chiếu của điểm lên mặt phẳng.
                                                if dang_toan == "[D11_C8_B2_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_15()                                     
                                                    

                                    #[D11_C8_B2_16]-M2. S.ABC: ABC tam giác đều. Tìm hình chiếu của điểm lên mặt phẳng.
                                                if dang_toan == "[D11_C8_B2_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_16()                                     
                                                    

                                    #[D11_C8_B2_17]-M2. S.ABCD: cạnh v.g đáy h.vuông. Tìm hình chiếu của điểm trên mặt nghiêng.
                                                if dang_toan == "[D11_C8_B2_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_17()                                     
                                                    

                                    #[D11_C8_B2_18]-M3. S.ABCD: cạnh v.g đáy h.chữ nhật. Tìm hình chiếu của điểm trên mặt nghiêng.
                                                if dang_toan == "[D11_C8_B2_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_18()                                     
                                                    

                                    #[D11_C8_B2_19]-M1. S.ABCD: ABCD h.vuông. Tìm hình chiếu của đường thẳng trên mặt đáy.
                                                if dang_toan == "[D11_C8_B2_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_19()                                     
                                                    

                                        #[D11_C8_B2_20]-M2. S.ABCD: ABCD h.vuông. Tìm hình chiếu của đường thẳng trên mặt đứng.
                                                if dang_toan == "[D11_C8_B2_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B2_20()                                     
                                                    
                                    #Bài 3: Hai mặt phẳng vuông góc
                                    #[D11_C8_B3_01]-M2. S.ABCD: ABCD h.vuông. Xác định 2 mặt phẳng vuông góc.
                                                if dang_toan == "[D11_C8_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_01()                                     
                                                    

                                    #[D11_C8_B3_02]-M3. S.ABCD: ABCD h.vuông. Xác định 2 mặt phẳng vuông góc.
                                                if dang_toan == "[D11_C8_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_02()                                     
                                                    

                                    #[D11_C8_B3_03]-M2. S.ABCD: ABCD h.chữ nhật. Xác định 2 mặt phẳng vuông góc.
                                                if dang_toan == "[D11_C8_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_03()                                     
                                                    

                                    #[D11_C8_B3_04]-M3. S.ABCD: ABCD h.chữ nhật. Xác định 2 mặt phẳng vuông góc.
                                                if dang_toan == "[D11_C8_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_04()                                     
                                                    

                                    #[D11_C8_B3_05]-M2. S.ABC: ABC vuông tại A. Tìm 2 mpvg.
                                                if dang_toan == "[D11_C8_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_05()                                     
                                                    

                                    #[D11_C8_B3_06]-M2. S.ABCD: ABCD h.chữ nhật. Tìm vị trí góc (mặt nghiêng, mặt đáy).
                                                if dang_toan == "[D11_C8_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_06()                                     
                                                    

                                    #[D11_C8_B3_07]-M2. S.ABCD: ABCD h.vuông. Tìm vị trí góc (mặt nghiêng, mặt đáy).
                                                if dang_toan == "[D11_C8_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_07()                                     
                                                    

                                    #[D11_C8_B3_08]-M2. S.ABC: ABC đều. Tính góc(mặt nghiêng, mặt đáy).
                                                if dang_toan == "[D11_C8_B3_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_08()                                     
                                                    

                                    #[D11_C8_B3_09]-M2. S.ABC: ABC vuông. Tính góc(mặt nghiêng, mặt đáy).
                                                if dang_toan == "[D11_C8_B3_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B3_09()                                     
                                                    
                                    


                                    #Bài 4 - Khoảng cách trong không gian 
                                                #[D11_C8_B6_01]. Cho hình chóp có diện tích đáy và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_01()                                     
                                                    

                                                #[D11_C8_B6_02]. Cho hình chóp có đáy là tam giác đều và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_02()                                     
                                                    

                                                #[D11_C8_B6_03]. Cho hình chóp có đáy là tam giác vuông và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_03()                                     
                                                    

                                                #[D11_C8_B6_04]. Cho hình chóp có đáy là tam giác đều và góc giữa cạnh bên và đáy. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_04()                                     
                                                    

                                                #[D11_C8_B6_05]-M1. Cho hình lăng trụ có diện tích đáy và chiều cao. Tính thể tích
                                                if dang_toan == "[D11_C8_B6_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_05()                                     
                                                    

                                                #[D11_C8_B6_06]-M2. Cho hình lăng trụ có đáy là tam giác đều và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_06()                                     
                                                    

                                                #[D11_C8_B6_07]-M2. Cho hình lăng trụ có đáy là tam giác vuông và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_07()                                     
                                                    

                                                #[D11_C8_B6_08]-M2. Cho hình lăng trụ có đáy là h.vuông,h.c.n và chiều cao. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_08()                                     
                                                    

                                                #[D11_C8_B6_09]-M3. S.ABCD: ABCD h.chữ nhật, có góc giữa 2 đường thẳng. Tính V.
                                                if dang_toan == "[D11_C8_B6_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_09()                                     
                                                    

                                                #[D11_C8_B4_01]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c từ điểm dưới đáy đến mặt đứng.
                                                if dang_toan == "[D11_C8_B4_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_01()                                     
                                                    

                                                #[D11_C8_B4_02]-M3. Tính k.c từ điểm thuộc cạnh bên đến mặt đứng.
                                                if dang_toan == "[D11_C8_B4_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_02()                                     
                                                    

                                                #[D11_C8_B4_03]-M3. S.ABCD: ABCD h.chữ nhật. Tính k.c từ chân đường cao đến mặt nghiêng.
                                                if dang_toan == "[D11_C8_B4_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_03()                                     
                                                    

                                                #[D11_C8_B4_04]-M2. Cho hình lăng trụ đứng tam giác. Tính khoảng cách từ điểm đến mặt đứng
                                                if dang_toan == "[D11_C8_B4_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_04()                                     
                                                    

                                                #[D11_C8_B4_05]-M2. Cho hình lập phương. Tính khoảng cách giữa 2 đường chéo nhau
                                                if dang_toan == "[D11_C8_B4_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_05()                                     
                                                    

                                                #[D11_C8_B4_06]-M2. Cho hình lập phương. Tính khoảng cách giữa 2 đường chéo nhau (dựng mặt song song đường)
                                                if dang_toan == "[D11_C8_B4_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_06()                                     
                                                    

                                                #[D11_C8_B4_07]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c 2 đường chéo nhau-kẻ đường vuông góc chung
                                                if dang_toan == "[D11_C8_B4_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_07()                                     
                                                    

                                                #[D11_C8_B4_08]-M2. S.ABCD: ABCD h.chữ nhật. Tính k.c 2 đường chéo nhau-(tạo mặt song song đường)
                                                if dang_toan == "[D11_C8_B4_08]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_08()

                                                #[D11_C8_B4_09]-M3. Cho hình lăng trụ đứng tam giác. Tính khoảng cách từ điểm đến mặt nghiêng
                                                if dang_toan == "[D11_C8_B4_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B4_09()                                    
                                                    

                                            #BÀI 5-GÓC GIỮA ĐƯỜNG THẲNG VÀ MẶT PHẲNG.
                                            #[D11_C8_B5_01]-M1. S.ABCD: ABCD h.vuông. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_01()                                     
                                                    

                                            #[D11_C8_B5_02]-M2. S.ABCD: ABCD h.vuông. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_02()                                     
                                                    

                                            #[D11_C8_B5_03]-M2. S.ABCD: ABCD h.chữ nhật. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_03()                                     
                                                    

                                            #[D11_C8_B5_04]-M2. S.ABC: ABC tam giác vuông. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_04()                                     
                                                    

                                            #[D11_C8_B5_05]-M2. S.ABC: ABC tam giác vuông. Xác định góc giữa đường thẳng và mặt đứng.
                                                if dang_toan == "[D11_C8_B5_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_05()                                     
                                                    

                                            #[D11_C8_B5_06]-M2. S.ABCD, đáy h.c.n. Tính số đo góc phẳng nhị diện.
                                                if dang_toan == "[D11_C8_B5_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_06()                                     
                                                    

                                            #[D11_C8_B5_07]-M2. S.ABCD, đáy h.c.n. Tính số đo góc phẳng nhị diện.
                                                if dang_toan == "[D11_C8_B5_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_07()                                     
                                                    

                                            #[D11_C8_B5_08]-M2. S.ABCD: ABCD h.vuông. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_08()

                                            #[D11_C8_B5_09]-M2. S.ABCD: ABCD h.vuông. Xác định góc giữa đường thẳng và mặt phẳng.
                                                if dang_toan == "[D11_C8_B5_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B5_09()                                  
                                                    

                                            #[D11_C8_B6_01]-M2. Nhận dạng hình lăng trụ
                                                if dang_toan == "[D11_C8_B6_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_10()                                     
                                                    

                                            #[D11_C8_B6_11]-M1. Cho hình lập phương có độ dài cạnh. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_11()                                     
                                                    

                                            #[D11_C8_B6_12]-M1. Cho hình hộp chữ nhật có độ dài cạnh. Tính thể tích.
                                                if dang_toan == "[D11_C8_B6_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C8.zz8zz_L11_C8_B6_12()                                     
                                                    
                                            #Đúng-Sai
                                            #[D11_C8_B7_01]-TF-M3. S.ABC: đáy tam giác. Tạo câu đúng-sai: Thể tích, Góc đường mặt, Góc mặt mặt, Hai mặt vuông góc
                                                if dang_toan == "[D11_C8_B7_01]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B7_01()                                                                                         
                                
                                            #[D11_C8_B7_02]-TF-M3. S.ABC: đáy tam giác. Tạo câu đúng-sai: Thể tích, Góc đường mặt, Góc mặt mặt, Khoảng cách
                                                if dang_toan == "[D11_C8_B7_02]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D11_C8.zz8zz_L11_C8_B7_02()                                      

                                                            
                                                            

                                    ######### Toán 11 Chương 7 - ĐẠO HÀM ########

                                    ######### Toán 11 Chương 9 - XÁC SUẤT ########
                                            #Bài 1. Biến cố giao và quy tắc nhân xác suất
                                            #[D11_C9_B1_01]-M2. Cho một hộp chứa n quả cầu. Tìm biến cố giao vừa là số chẵn (lẻ), vừa chia hết cho k.
                                                if dang_toan == "[D11_C9_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_01()                                     
                                                    

                                            #[D11_C9_B1_02]-M2. Cho một hộp chứa n tấm thẻ. Tìm biến cố giao: thẻ vừa thuộc phạm vi [a;b], vừa chia hết cho k.
                                                if dang_toan == "[D11_C9_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_02()                                     
                                                    

                                            #[D11_C9_B1_03]-M2. Cho một hộp chứa n tấm thẻ. Tìm biến cố giao: thẻ vừa thuộc phạm vi [a;b], vừa chia hết cho k.
                                                if dang_toan == "[D11_C9_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_03()                                     
                                                    

                                            #[D11_C9_B1_04]-M2. Cho một hộp chứa n tấm thẻ. Tìm biến cố giao: thẻ vừa thuộc phạm vi [a;b], vừa chia hết cho k.
                                                if dang_toan == "[D11_C9_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_04()                                     
                                                    

                                            #[D11_C9_B1_05]-M2. Gieo 2 con xúc xắc. Tìm biến cố giao: mặt k chấm xuất hiện và tổng số chấm =,>,< m.
                                                if dang_toan == "[D11_C9_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_05()                                     
                                                    

                                            #[D11_C9_B1_06]-M2. Gieo 2 con xúc xắc. Tìm biến cố giao: mặt k chấm xuất hiện và tổng số chấm =,>,< m.
                                                if dang_toan == "[D11_C9_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_06()                                     
                                                    

                                            #[D11_C9_B1_07]-M1. Cho xác suất của 2 biến cố độc lập. Tìm xác suất của biến cố giao.
                                                if dang_toan == "[D11_C9_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_07()                                     
                                                    

                                            #[D11_C9_B1_08]-M2. Cho xác suất của 2 biến cố độc lập. Tìm xác suất của biến cố giao chứa biến cố đối.
                                                if dang_toan == "[D11_C9_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_08()                                     
                                                    

                                            #[D11_C9_B1_09]-M2. Hộp chứa các viên bi có 2 màu. Phát biểu biến cố giao.
                                                if dang_toan == "[D11_C9_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_09()                                     
                                                    

                                            #[D11_C9_B1_10]-M2. Cho xác suất bị bệnh của 2 bệnh nhân. Tìm xác suất cả hai bị không bị bệnh.
                                                if dang_toan == "[D11_C9_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_10()                                     
                                                    

                                            #[D11_C9_B1_11]-M2. Cho xác suất bị bệnh của 2 bệnh nhân. Tìm xác suất chỉ một trong hai không bị biến chứng nặng.
                                                if dang_toan == "[D11_C9_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_11()                                     
                                                    

                                            #[D11_C9_B1_12]-M2. Cho xác suất bắn trúng bia của 1 xạ thủ. Tính xác suất cả hai lần bắn trượt.
                                                if dang_toan == "[D11_C9_B1_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_12()                                     
                                                    

                                            #[D11_C9_B1_13]-M2. Cho xác suất bắn trúng bia của 1 xạ thủ. Tính xác suất chỉ một lần bắn trúng.
                                                if dang_toan == "[D11_C9_B1_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_13()                                     
                                                    

                                            #[D11_C9_B1_14]-M2. Cho xác suất bắn trúng bia của 1 xạ thủ. Tính xác suất cả hai lần đều bắn trúng.
                                                if dang_toan == "[D11_C9_B1_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_14()                                     
                                                    

                                            #[D11_C9_B1_15]-M2. Cho xác suất bắn trúng bia của 1 xạ thủ. Tính xác suất cả hai lần đều bắn trúng.
                                                if dang_toan == "[D11_C9_B1_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B1_15()                                     
                                                    

                                            #BÀI 2: BIẾN CỐ HỢP

                                            #[D11_C9_B2_01]-M2. Hộp chứa các viên bi có 2 màu. Phát biểu biến cố hợp.
                                                if dang_toan == "[D11_C9_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_01()                                     
                                                    

                                            #[D11_C9_B2_02]-M2. Hộp chứa các viên bi, lấy ra rồi trả lại. Phát biểu biến cố hợp.
                                                if dang_toan == "[D11_C9_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_02()                                     
                                                    

                                            

                                            #[D11_C9_B2_03]-M2. Gieo một con xúc xắc và 1 đồng xu.Tìm số phần tử của biến cố hợp.
                                                if dang_toan == "[D11_C9_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_03()                                     
                                                    

                                            #[D11_C9_B2_04]-M1. Cho A, B xung khắc và P(A), P(B). Tính P(AUB).
                                                if dang_toan == "[D11_C9_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_04()                                     
                                                    

                                            #[D11_C9_B2_05]-M1. Gieo một con xúc xắc và 1 đồng xu. Tính xác suất của biến cố hợp.
                                                if dang_toan == "[D11_C9_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_05()                                     
                                                    

                                            #[D11_C9_B2_06]-M2. Cho 2 nhóm đồ vật. Tính xác suất để số vật được chọn thuộc cùng 1 nhóm.
                                                if dang_toan == "[D11_C9_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_06()                                     
                                                    

                                            #[D11_C9_B2_07]-M2. Cho 2 nhóm sách tham khảo. Tính xác suất để số sách được chọn thuộc cùng 1 loại.
                                                if dang_toan == "[D11_C9_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_07()                                     
                                                    

                                            #[D11_C9_B2_08]-M2. Cho các viên bi. Tính xác suất để các bi được chọn cùng màu.
                                                if dang_toan == "[D11_C9_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_08()                                     
                                                    

                                            #[D11_C9_B2_09]-M2. Cho các cuốn truyện. Tính xác suất để các cuốn được chọn cùng thể loại.
                                                if dang_toan == "[D11_C9_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D11_C9.zz8zz_L11_C9_B2_09()                                     
                                                    

                                                                                
                                    #Toán 10- Chương 1 - Mệnh đề và tập hợp
                                            #Bài 1 - MỆNH ĐỀ
                                            #[D10_C1_B1_01]-TF-M2. Xét tính Đ-S của mệnh đề, mệnh đề phủ định, mệnh đề chứa biến.
                                                if dang_toan == "[D10_C1_B1_01]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B1_01()

                                            #[D10_C1_B1_02]-TF-M2. Xét Đ-S về số chính phương, số hữu tỉ, số chẵn, lẻ, số chia hết.
                                                if dang_toan == "[D10_C1_B1_02]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B1_02()

                                            #[D10_C1_B1_03]-TF-M2. Cho P(n)=an^2+bn+c. Xét Đ-S: P(x_0) chẵn (lẻ),P(x_1) chia hết cho, P(ta)-P(a), ...
                                                if dang_toan == "[D10_C1_B1_03]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B1_03()

                                            #[D10_C1_B1_04]-M2. Tìm mệnh đề đúng (cảm thán, số là số gì, số chia hết cho, nghiệm của phương trình).
                                                if dang_toan == "[D10_C1_B1_04]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B1_04()

                                            #[D10_C1_B1_05]-M2. Tìm mệnh đề sai (cảm thán, số là số gì, số chia hết cho, nghiệm của phương trình).
                                                if dang_toan == "[D10_C1_B1_05]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B1_05()

                                            #[D10_C1_B1_06]-M2. Tìm mềnh đề phủ định của mệnh đề chứa biến.
                                                if dang_toan == "[D10_C1_B1_06]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B1_06()

                                            #[D10_C1_B1_07]-M2. Tìm mềnh đề phủ định của mệnh đề không chứa biến.
                                                if dang_toan == "[D10_C1_B1_07]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B1_07()

                                            #[D10_C1_B1_08]-M2. Tìm mệnh đề kéo theo đúng.
                                                if dang_toan == "[D10_C1_B1_08]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B1_08()

                                            #Bài 2 - Tập hợp
                                            #[D10_C1_B2_01] Liệt kê các phần tử của tập hợp số nguyên"]

                                                if dang_toan == "[D10_C1_B2_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_01()

                                            #[D10_C1_B2_02] Liệt kê các phần tử là nghiệm của phương trình bậc 2"]
                                                if dang_toan == "[D10_C1_B2_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_02()

                                            #[D10_C1_B2_03] Liệt kê các phần tử là ước của n
                                                if dang_toan == "[D10_C1_B2_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_03()

                                            #[D10_C1_B2_04] Liệt kê các phần tử là bội của n
                                                if dang_toan == "[D10_C1_B2_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_04()

                                            #[D10_C1_B2_05]-M2. Chỉ ra tính chất đặc trưng từ tập hợp liệt kê
                                                if dang_toan == "[D10_C1_B2_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_05()

                                            #[D10_C1_B2_06]-M2. Cho tập hợp dạng đặc trưng. Xét đúng-sai: Số phần tử, Liệt kê, số tập hợp con, xác định tập hợp con
                                                if dang_toan == "[D10_C1_B2_06]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B2_06()

                                            #[D10_C1_B2_07]-M2. Xét đúng sai về số phần tử của: a<x<b, ax^2+bx+c=0, phương trình tích, m<ax+b<n
                                                if dang_toan == "[D10_C1_B2_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B2_07()

                                            #[D10_C1_B2_08]-M2. Cho tập hợp dạng liệt kê. Xét đúng-sai: Số phần tử, Liệt kê, số tập hợp con, xác định tập hợp con
                                                if dang_toan == "[D10_C1_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B2_08()

                                            #[D10_C1_B2_09]-M2. Liệt kê các số là nguyên tố
                                                if dang_toan == "[D10_C1_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_09()

                                            #[D10_C1_B2_10]-M2. Liệt kê các số là số chẵn hoặc số lẻ
                                                if dang_toan == "[D10_C1_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_10()

                                            #[D10_C1_B2_11]-M2. Liệt kê phần tử thỏa mãn điều kiện tùy ý
                                                if dang_toan == "[D10_C1_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B2_11()

                                            #Bài 3 - Các phép toán trên tập hợp

                                            #[D10_C1_B3_01]. Tìm giao của 2 tập hợp dạng liệt kê
                                                if dang_toan == "[D10_C1_B3_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_01()

                                            #[D10_C1_B3_02]. Tìm hợp của 2 tập hợp dạng liệt kê
                                                if dang_toan == "[D10_C1_B3_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_02()

                                            #[D10_C1_B3_03]. Tìm hiệu của 2 tập hợp dạng liệt kê
                                                if dang_toan == "[D10_C1_B3_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_03()

                                            #[D10_C1_B3_04]. Tìm phần bù của 2 tập hợp dạng liệt kê
                                                if dang_toan == "[D10_C1_B3_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_04()

                                            #[D10_C1_B3_05]. Tìm giao của 2 tập hợp dạng tính chất đặc trưng
                                                if dang_toan == "[D10_C1_B3_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_05()

                                            #[D10_C1_B3_06]. Tìm hợp của 2 tập hợp dạng tính chất đặc trưng
                                                if dang_toan == "[D10_C1_B3_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_06()

                                            #[D10_C1_B3_07]. Tìm hợp của 2 tập hợp dạng tính chất đặc trưng
                                                if dang_toan == "[D10_C1_B3_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_07()

                                            #[D10_C1_B3_08]. Tìm số tập hợp con của một tập hợp
                                                if dang_toan == "[D10_C1_B3_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_08()

                                            #[D10_C1_B3_09]. Tìm số tập hợp con gồm k phần tử của một tập hợp
                                                if dang_toan == "[D10_C1_B3_09]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_09()

                                            #[D10_C1_B3_10]-TF-M2. Cho 2 tập hợp dạng liệt kê. Xét Đ-S: Giao, Hợp, Hiệu, Kiểm tra tập con
                                                if dang_toan == "[D10_C1_B3_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B3_10()

                                            #[D10_C1_B3_11]-TF-M2. Cho 2 tập hợp dạng đặc trưng. Xét Đ-S: Giao, Hợp, Hiệu, Kiểm tra tập con
                                                if dang_toan == "[D10_C1_B3_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B3_11()

                                            #[D10_C1_B3_12]-TL-M3. Bài thực tế: Cho n(A), n(B). Tính n(A giao B).
                                                if dang_toan == "[D10_C1_B3_12]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B3_12()

                                            #[D10_C1_B3_13]-TL-M3. Bài thực tế: Cho n(A), n(B), n(A giao B). Tính tổng.
                                                if dang_toan == "[D10_C1_B3_13]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B3_13()

                                            #[D10_C1_B3_14]-TL-M3. Bài thực tế: Cho n(A), n(B), n(A giao B), n_khong(AB). Tính tổng.
                                                if dang_toan == "[D10_C1_B3_14]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B3_14()

                                            #[D10_C1_B3_15]-TF-M3. Bài thực tế: Cho n(A), n(B), n(A giao B), n_khôngAB. Xét Đ-S: Chỉ A, Chỉ B, A hoặc B, Tổng số.
                                                if dang_toan == "[D10_C1_B3_15]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B3_15()

                                            #[D10_C1_B3_16]. Cho kí hiệu khoảng,đoạn. Mô tả tính chất đặc trưng
                                                if dang_toan == "[D10_C1_B3_16]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B3_16()

                                            #[D10_C1_B3_17]-TF-M2. Cho A-đặc trưng, B-liệt kê. Xét Đ-S: Liệt kê A, Số tập hợp con, phép toán, A con X con B
                                                if dang_toan == "[D10_C1_B3_17]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B3_17()

                                            #[D10_C1_B3_18]-TL-M4. Bài toán thực tế liên quan 3 tập hợp (biết ít nhất 1 môn)
                                                if dang_toan == "[D10_C1_B3_18]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B3_18()

                                            #[D10_C1_B3_19]-TL-M4. Bài toán thực tế liên quan 3 tập hợp (biết đúng 1 môn)
                                                if dang_toan == "[D10_C1_B3_19]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B3_19()

                                            #Bài 4 - Các tập hợp con của R 
                                            #[D10_C1_B4_01]. Cho kí hiệu khoảng,đoạn. Mô tả tính chất đặc trưng
                                                if dang_toan == "[D10_C1_B4_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_01()

                                            #[D10_C1_B4_02]. Cho tính chất đặc trưng. Tìm kí hiệu khoảng đoạn
                                                if dang_toan == "[D10_C1_B4_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_02()

                                            #[D10_C1_B4_03]. Tìm giao các khoảng, đoạn
                                                if dang_toan == "[D10_C1_B4_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_03()

                                            #[D10_C1_B4_04]. Tìm hợp các khoảng, đoạn
                                                if dang_toan == "[D10_C1_B4_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_04()

                                            #[D10_C1_B4_05]. Tìm hiệu các khoảng, đoạn
                                                if dang_toan == "[D10_C1_B4_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_05()

                                            #[D10_C1_B4_06]. Tìm hiệu các khoảng, đoạn
                                                if dang_toan == "[D10_C1_B4_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_06()

                                            #[D10_C1_B4_07]-TF-M2. Cho A, B là các khoảng đoạn. Xét Đ-S: A giao B, A hợp B, A\B, C_R(A) 
                                                if dang_toan == "[D10_C1_B4_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B4_07()

                                            #[D10_C1_B4_08]-TL-M2. Cho A, B là các khoảng đoạn. Tìm m để A giao B có đúng 1 phần tử
                                                if dang_toan == "[D10_C1_B4_08]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_08()

                                            #[D10_C1_B4_09]-TL-M3. Cho A, B là các khoảng đoạn. Tìm m để A giao B khác rỗng.
                                                if dang_toan == "[D10_C1_B4_09]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_09()

                                            #[D10_C1_B4_10]-TL-M3. Cho A, B là các khoảng đoạn. Tìm m để A giao B bằng rỗng.
                                                if dang_toan == "[D10_C1_B4_10]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_10()

                                            #[D10_C1_B4_11]-TL-M3. Cho A, B là các khoảng đoạn. Tìm m để A hợp B =A.
                                                if dang_toan == "[D10_C1_B4_11]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_11()

                                            #[D10_C1_B4_12]-TL-M3. Cho A, B là các khoảng đoạn. Tìm m để A con B.
                                                if dang_toan == "[D10_C1_B4_12]": 
                                                        debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_12()

                                            #[D10_C1_B4_13]-TF-M2.  Xét đúng-sai về kí hiệu khoảng, đoạn, nữa khoảng.
                                                if dang_toan == "[D10_C1_B4_13]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B4_13()

                                            #[D10_C1_B4_14]-M3. Tìm kí hiệu khoảng, đoạn ứng với tập hợp chứa điều kiện.
                                                if dang_toan == "[D10_C1_B4_14]":
                                                        debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_14()

                                            #[D10_C1_B4_15]-TF-M3. Xét Đ-S: Kí hiệu, phép toán, Giao với N,Z,  tìm m
                                                if dang_toan == "[D10_C1_B4_15]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C1.mjulk_L10_C1_B4_15()

                                            #[D10_C1_B4_16]-M2. Tìm giao (hoặc hợp, hiệu, phần bù) các tập con của R
                                                if dang_toan == "[D10_C1_B4_16]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C1.mjulk_L10_C1_B4_16()

                                            #[D10_C1_B4_17]-M2. Bài toán tìm m để thỏa mãn phép toán liên quan khoảng đoạn 
                                                if dang_toan == "[D10_C1_B4_17]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_17()

                                            #[D10_C1_B4_18]-TL-M2. Tìm số phần tử là số nguyên của một phép toán khoảng đoạn
                                                if dang_toan == "[D10_C1_B4_18]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C1.mjulk_L10_C1_B4_18()

                                    #TOÁN 10-CHƯƠNG 2 - BẤT PHƯƠNG TRÌNH VÀ HỆ BẤT PHƯƠNG TRÌNH BẬC NHẤT 2 ẨN
                                    # Bài 1: BẤT PHƯƠNG TRÌNH BẬC NHẤT 2 ẨN
                                                #[D10_C2_B1_01] Tìm cặp số là nghiệm của BPT ax+by+c>0,>=0.
                                                if dang_toan == "[D10_C2_B1_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_01()

                                                #[D10_C2_B1_02] Tìm cặp số là nghiệm của BPT ax+by+c<0,<=0.
                                                if dang_toan == "[D10_C2_B1_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_02()

                                                #[D10_C2_B1_03]-M1. Nhận dạng bất phương trình bậc nhất 2 ẩn
                                                if dang_toan == "[D10_C2_B1_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_03()

                                                #[D10_C2_B1_04]-M2. Xác định miền nghiệm là nữa mặt phẳng chứa điểm nào
                                                if dang_toan == "[D10_C2_B1_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_04()

                                                #[D10_C2_B1_05]-M2. Xác định miền nghiệm của ax+by+c >(<) ex+fy+g  là nữa mặt phẳng chứa điểm nào
                                                if dang_toan == "[D10_C2_B1_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_05()

                                                #[D10_C2_B1_06]-M2. Xác định miền nghiệm của ax+by+c >(<) ex+fy+g  là nữa mặt phẳng chứa điểm nào
                                                if dang_toan == "[D10_C2_B1_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_06()

                                                #[D10_C2_B1_07]-M2. Xác định miền nghiệm của ax+by+c >(<) ex+fy+g  là nữa mặt phẳng chứa điểm nào
                                                if dang_toan == "[D10_C2_B1_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_07()

                                                #[D10_C2_B1_08]-TF-M2. Cho BPT ax+by+c>0 (<0). Xét Đ-S: nghiệm, không là nghiệm, nghiệm chứa bờ, miền nghiệm
                                                if dang_toan == "[D10_C2_B1_08]":           
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C2.bch_12_L10_C2_B1_08()

                                                #[D10_C2_B1_09]-M2. Lập BPT bậc nhất 2 ẩn từ bài toán thực tế
                                                if dang_toan == "[D10_C2_B1_09]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B1_09()

                                                #Bài 2
                                                #[D10_C2_B2_01]-M2. Cho Hệ BPT ax+by+c>0 (<0). Tìm cặp số là nghiệm của hệ.
                                                if dang_toan == "[D10_C2_B2_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B2_01()

                                                #[D10_C2_B2_02]-M2. Cho Hệ BPT ax+by+c>0 (<0). Tìm cặp số là không là nghiệm của hệ.
                                                if dang_toan == "[D10_C2_B2_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B2_02()

                                                #[D10_C2_B2_03]-M2. Cho Hệ 3 BPT ax+by+c>0 (<0). Tìm cặp số là không là nghiệm của hệ.
                                                if dang_toan == "[D10_C2_B2_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C2.bch_12_L10_C2_B2_03()

                                                #[D10_C2_B2_04]-TF-M2. Cho hệ 2 BPT. Xét đúng sai về nghiệm của hệ.
                                                if dang_toan == "[D10_C2_B2_04]":           
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C2.bch_12_L10_C2_B2_04()



                                    #Toán 10- Chương 3 - Hàm số và đồ thị
                                                ####################### Bài 1: Hàm số và đồ thị ######################
                                               
                                                #D10_C3_B1_01. Tính giá trị của hàm số tại một điểm
                                                if dang_toan == "[D10_C3_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_01()

                                                #D10_C3_B1_02. Tính giá trị của hàm số có nhiều biểu thức tại một điểm
                                                if dang_toan == "[D10_C3_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_02()

                                                #[D10_C3_B1_03]. Tìm tập xác định hàm số y =(ax+b)/(cx+d)
                                                if dang_toan == "[D10_C3_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_03()                                     
                                                    

                                                #[D10_C3_B1_04]. Tìm tập xác định hàm số y = căn(ax+b)
                                                if dang_toan == "[D10_C3_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_04()                                     
                                                    

                                                #[D10_C3_B1_05]. Tìm tập xác định y = căn (A) + B/C
                                                if dang_toan == "[D10_C3_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_05()                                     
                                                    

                                                #[D10_C3_B1_06]. Tìm tập giá trị dựa vào hình vẽ đồ thị
                                                if dang_toan == "[D10_C3_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_06()                                     
                                                    

                                            #[D10_C3_B1_07]-M2. Tìm tập xác định y =căn(ax+b) + căn(cx+d).
                                                if dang_toan == "[D10_C3_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_07()                                     
                                                    

                                            #[D10_C3_B1_08]-M2.Tìm tập xác định y =A/(ax^2+bx+c).
                                                if dang_toan == "[D10_C3_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B1_08()                                     
                                                    

                                                ####################### Bài 1: Hàm số và đồ thị ######################
                                                #[D10_C3_B2_01]. Tìm tọa độ đỉnh của đồ thị hàm số bậc 2
                                                if dang_toan == "[D10_C3_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_01()                                     
                                                    

                                                #[D10_C3_B2_02].Tìm trục đối xứng của đồ thị hàm số bậc 2
                                                if dang_toan == "[D10_C3_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_02()                                     
                                                    

                                                #[D10_C3_B2_03]. Cho hàm số bậc 2. Tìm khoảng biến thiên.
                                                if dang_toan == "[D10_C3_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_03()                                     
                                                    

                                                #[D10_C3_B2_04]. Cho BBT hàm số bậc 2. Tìm khoảng biến thiên.
                                                if dang_toan == "[D10_C3_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_04()                                     
                                                    

                                                #[D10_C3_B2_05]. Cho bảng biến thiên. Tìm hàm số bậc 2.
                                                if dang_toan == "[D10_C3_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_05()                                     
                                                    

                                                #[D10_C3_B2_06]. Cho đồ thị. Tìm hàm số bậc 2.
                                                if dang_toan == "[D10_C3_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_06()                                     
                                                    

                                                #[D10_C3_B2_07]. Tìm hàm số bậc 2 có đồ thị đi qua một điểm.
                                                if dang_toan == "[D10_C3_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_07()                                     
                                                    

                                                #[D10_C3_B2_08]. Tìm hàm số bậc 2 có đồ thị đi qua hai điểm.
                                                if dang_toan == "[D10_C3_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_08()                                     
                                                    

                                                #[D10_C3_B2_09]. Tìm hàm số bậc 2 có đồ thị đi qua điểm và trục đối xứng.
                                                if dang_toan == "[D10_C3_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_09()                                     
                                                    

                                                #[D10_C3_B2_10]. Cho đồ thị hàm số bậc 2. Xét dấu các hệ số a,b,c.
                                                if dang_toan == "[D10_C3_B2_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C3.npl_mk_L10_C3_B2_10()                                     
                                                    

                                                #[D10_C3_B2_11]-TF-M2. Cho hàm số bậc hai. Tạo Đúng-Sai. 
                                                if dang_toan == "[D10_C3_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C3.npl_mk_L10_C3_B2_11()                                     


                                                #[D10_C3_B2_12]-TF-M2. Cho hàm số bậc hai. Tạo Đúng-Sai về đồ thị. 
                                                if dang_toan == "[D10_C3_B2_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C3.npl_mk_L10_C3_B2_12()                                      


                                                #[D10_C3_B2_13]-TF-M2. Cho BBT hàm số bậc hai. Tạo Đúng-Sai. 
                                                if dang_toan == "[D10_C3_B2_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C3.npl_mk_L10_C3_B2_13()                                      


                                            #CHƯƠNG 4 - HỆ THỨC LƯỢNG TRONG TAM GIÁC 
                                            #BÀI 1- GIÁ TRỊ LƯỢNG GIÁC CỦA MỘT GÓC TỪ 0 ĐẾN 180

                                                #[D10_C4_B1_01]. Tìm dấu của một giá trị lượng giác từ 0 đến 180 độ.
                                                if dang_toan == "[D10_C4_B1_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B1_01() 

                                                #[D10_C4_B1_02]. Cho giá trị lượng giác, tìm góc.
                                                if dang_toan == "[D10_C4_B1_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B1_02()

                                            #BÀI 2 - ĐỊNH LÍ COSIN VÀ ĐỊNH LÍ SIN

                                                #[D10_C4_B2_01]. Tính diện tích biết 2 cạnh và góc xen giữa
                                                if dang_toan == "[D10_C4_B2_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_01()

                                                #[D10_C4_B2_02]. Tính diện tích biết 2 cạnh và góc xen giữa
                                                if dang_toan == "[D10_C4_B2_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_02()

                                               #[D10_C4_B2_03]. Tính số đo một góc biết độ dài 3 cạnh
                                                if dang_toan == "[D10_C4_B2_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_03()

                                                #[D10_C4_B2_04]. Cho tam giác có một cạnh và góc đối diện. Tính bán kính đường tròn ngoại tiếp.
                                                if dang_toan == "[D10_C4_B2_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_04()

                                                #[D10_C4_B2_05]. Cho tam giác có độ dài 3 cạnh. Tính diện tích tam giác.
                                                if dang_toan == "[D10_C4_B2_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_05()

                                                #[D10_C4_B2_06]-TF-M2. Cho a,b và góc C. Xét Đ-S: cạnh C, cos A, góc B, diện tích
                                                if dang_toan == "[D10_C4_B2_06]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_06()

                                                #[D10_C4_B2_07]-TF-M2. Cho AC,BC và góc C. Xét Đ-S: cạnh AB, cos A, góc B, diện tích
                                                if dang_toan == "[D10_C4_B2_07]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_07()

                                                #[D10_C4_B2_08]-TF-M2. Cho a,b c. Xét Đ-S: p, S, r, R
                                                if dang_toan == "[D10_C4_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_08()

                                                #[D10_C4_B2_09]-TF-M2. Cho a,b,c. Xét Đ-S: cosA, B, S, sinC, R, r
                                                if dang_toan == "[D10_C4_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_09()

                                                #[D10_C4_B2_10]-TF-M2. Cho AB,AC,BC. Xét Đ-S: cosA, B, S, sinC, R, r
                                                if dang_toan == "[D10_C4_B2_10]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_10()

                                                #[D10_C4_B2_11]-TF-M2. Cho A,B và cạnh a. Xét Đ-S: góc C, cạnh a, cạnh b, diện tích
                                                if dang_toan == "[D10_C4_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C4.yy3yy_L10_C4_B2_11()

                                                #[D10_C4_B2_12]-M2. Cho A,B và cạnh a. Tính cạnh b hoặc c
                                                if dang_toan == "[D10_C4_B2_12]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C4.yy3yy_L10_C4_B2_12()

                                                #[D10_C4_B2_13]-TL-M3. Cho b,c và góc C. Tính cạnh a                                                
                                                if dang_toan == "[D10_C4_B2_13]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C4.yy3yy_L10_C4_B2_13()

                                         

                                            #CHƯƠNG 5 - VÉCTƠ 
                                            #Bài 1 - Véctơ
                                                #[D10_C5_B4_01].  Cho hai véctơ có độ dài và góc. Tính tích vô hướng
                                                if dang_toan == "[D10_C5_B1_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_01()

                                                #[D10_C5_B4_02]. Cho hình dạng bình hành. Tìm khẳng định sai về vectơ bằng nhau, độ dài bằng nhau.
                                                if dang_toan == "[D10_C5_B1_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_02()

                                                #[D10_C5_B1_03]-M2. Cho hình dạng bình hành. Tìm khẳng định đúng về vectơ hướng, phương.
                                                if dang_toan == "[D10_C5_B1_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_03()

                                                #[D10_C5_B1_04]-M2. Cho trung điểm của đoạn. Tìm khẳng định đúng về hướng, phương, véctơ bằng nhau, độ dài bằng nhau
                                                if dang_toan == "[D10_C5_B1_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_04()

                                                #[D10_C5_B1_05]-M2. Cho 3 điểm. Tìm khẳng định về hướng và phương.
                                                if dang_toan == "[D10_C5_B1_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_05()

                                                #[D10_C5_B1_06]-M2. Cho hình vuông (chữ nhật). Tính độ dài vectơ cạnh.
                                                if dang_toan == "[D10_C5_B1_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_06()

                                                #[D10_C5_B1_07]-M2. Cho hình vuông (chữ nhật). Tính độ dài vectơ đường chéo.
                                                if dang_toan == "[D10_C5_B1_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B1_07()

                                            #Bài 2 - Tổng hiệu hai vectơ
                                                #[D10_C5_B2_01]-M2. Cho trung điểm của đoạn. Tìm khẳng định tổng hiệu các vectơ.
                                                if dang_toan == "[D10_C5_B2_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_01()

                                                #[D10_C5_B2_02]-M2. Cho các điểm. Tính tổng hiệu các vectơ
                                                if dang_toan == "[D10_C5_B2_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_02()

                                                #[D10_C5_B2_03]-M2. Tìm đẳng thức đúng liên quan đến 3 điểm (phép cộng trừ)
                                                if dang_toan == "[D10_C5_B2_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_03()

                                                #[D10_C5_B2_04]-M2. Cho tứ giác. Tìm khẳng định đúng về quy tắc cộng trừ.
                                                if dang_toan == "[D10_C5_B2_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_04()

                                                #[D10_C5_B2_05]-M2. Cho hình bình hành. Tìm khẳng định sai về phép toán vectơ
                                                if dang_toan == "[D10_C5_B2_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_05()

                                                #[D10_C5_B2_06]-M2. Cho hình bình hành. Tính tổng hiệu các vectơ
                                                if dang_toan == "[D10_C5_B2_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_06()

                                                #[D10_C5_B2_07]-M2. Cho các điểm. Tính tổng các vectơ
                                                if dang_toan == "[D10_C5_B2_07]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_07()

                                                #[D10_C5_B2_08]-M2. Nhận dạng quy tắc hình bình hành
                                                if dang_toan == "[D10_C5_B2_08]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B2_08()

                                                #[D10_C5_B2_09]-TF-M2. Cho h.b.h. Xét Đ-S:quy tắc h.b.h, tổng-hiệu các vectơ.
                                                if dang_toan == "[D10_C5_B2_09]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C5.y7y7u_L10_C5_B2_09()

                                                #[D10_C5_B2_10]-TL-M3. Cho hai lực hợp góc. Tính độ lớn của tổng lực.
                                                if dang_toan == "[D10_C5_B2_10]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C5.y7y7u_L10_C5_B2_10()

                                                #[D10_C5_B2_11]-TF-M3.Cho h.c.n. Xét Đ-S:hướng-phương-bằng, quy tắc h.b.h, tổng-hiệu, độ dài các vectơ.
                                                if dang_toan == "[D10_C5_B2_11]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C5.y7y7u_L10_C5_B2_11()

                                                #[D10_C5_B2_12]-TL-M3. Cho tam giác vuông. Tính độ dài của vectơ có điểm thỏa mãn đẳng thức vectơ.
                                                if dang_toan == "[D10_C5_B2_12]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C5.y7y7u_L10_C5_B2_12()

                                                #[D10_C5_B2_13]-TL-M3. Cho h.c.n. Tính độ dài tổng-hiệu, độ dài các vectơ.
                                                if dang_toan == "[D10_C5_B2_13]": 
                                                    debai_word,loigiai_word,latex_tuluan,dap_an=D10_C5.y7y7u_L10_C5_B2_13()                            

                                            #Bài 3 - Tích vectơ với một số
                                                #[D10_C5_B3_01]-M2. Cho tứ giác. Tìm khẳng định đúng về quy tắc cộng trừ.
                                                if dang_toan == "[D10_C5_B3_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B3_01()

                                                #[D10_C5_B3_02]-M2. Cho tam giác. Tìm đẳng thức đúng về vectơ
                                                if dang_toan == "[D10_C5_B3_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B3_02()

                                                #[D10_C5_B3_03]-M2. Cho MA=kMB. Tìm đẳng thức đúng về vectơ
                                                if dang_toan == "[D10_C5_B3_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B3_03()

                                                #[D10_C5_B3_04]-M2. Cho hình bình hành. Tìm đẳng thức đúng về vectơ
                                                if dang_toan == "[D10_C5_B3_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B3_04()

                                                #[D10_C5_B3_05]-TF-M3. Cho hình bình hành. Xét đúng sai các đẳng thức vectơ 
                                                if dang_toan == "[D10_C5_B3_05]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C5.y7y7u_L10_C5_B3_05()

                                            #Bài 4 - Tích vô hướng của hai vectơ 
                                                #[D10_C5_B4_01].  Cho hai véctơ có độ dài và góc. Tính tích vô hướng
                                                if dang_toan == "[D10_C5_B4_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B4_01()                                
                                                    

                                                #[D10_C5_B4_02].  Cho hai véctơ có độ dài và tích vô hướng. Tìm góc
                                                if dang_toan == "[D10_C5_B4_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B4_02()                                     
                                                    

                                                #[D10_C5_B4_03].  Cho hai véctơ có độ dài và góc. Tính |vta+vtb|
                                                if dang_toan == "[D10_C5_B4_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B4_03()                                     
                                                    

                                                #[D10_C5_B4_04].  Cho hai véctơ có độ dài và góc. Tính |vta-vtb|
                                                if dang_toan == "[D10_C5_B4_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C5.y7y7u_L10_C5_B4_04()

                                                #[D10_C5_B4_05]-TF-M2. Cho hình vuông. Xét Đ-S: các phép toán về vectơ, tích vô hướng
                                                if dang_toan == "[D10_C5_B4_05]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D10_C5.y7y7u_L10_C5_B4_05()

                                                #[D10_C5_B4_06]-TF-M2. Cho tứ giác. Xét Đ-S: các phép toán về vectơ, tích vô hướng
                                                if dang_toan == "[D10_C5_B4_06]": 
                                                        debai_word,debai_latex,loigiai_word,dap_an=D10_C5.y7y7u_L10_C5_B4_06()                                 
                                                    

                                    #Bài 5 - Tọa độ vectơ

                                            #[D10_CX_B0_01]. Cho hai điểm. Tìm tọa độ vectơ 
                                                if dang_toan == "[D10_CX_B0_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_01()                                     
                                                    

                                            #[D10_CX_B0_02]. Cho hai véctơ. Tìm tọa độ vectơ tổng.
                                                if dang_toan == "[D10_CX_B0_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_02()                                     
                                                    

                                            #[D10_CX_B0_03]. Cho hai véctơ. Tìm tọa độ vectơ hiệu.
                                                if dang_toan == "[D10_CX_B0_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_03()                                     
                                                    

                                            #[D10_CX_B0_04]. Cho hai véctơ tính tích vô hướng.
                                                if dang_toan == "[D10_CX_B0_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_04()                                     
                                                    

                                            #[D10_CX_B0_05]. Cho véctơ. Tính độ dài vectơ.
                                                if dang_toan == "[D10_CX_B0_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_05()                                     
                                                    

                                             #[D10_CX_B0_06]. Cho 2 điểm. Tính độ dài.
                                                if dang_toan == "[D10_CX_B0_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_06()                                     
                                                    

                                            #[D10_CX_B0_07]-M1. Cho hai điểm. Tìm tọa độ trung điểm.
                                                if dang_toan == "[D10_CX_B0_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_07()                                     
                                                    

                                            #[D10_CX_B0_08]-M1. Cho tam giác. Tìm tọa độ trọng tâm.
                                                if dang_toan == "[D10_CX_B0_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_08()                                     
                                                    

                                            #[D10_CX_B0_09]-M2. Cho hai điểm A,B. Tìm tọa độ C để AC nhận B làm trung điểm.
                                                if dang_toan == "[D10_CX_B0_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_09()                                     
                                                    

                                            #[D10_CX_B0_10]-M2. Cho ba điểm A,B,G. Tìm tọa độ C để ABC nhận G làm trọng tâm.
                                                if dang_toan == "[D10_CX_B0_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_10()                                     
                                                    

                                            #[D10_CX_B0_11]-M2. Cho hai véctơ. Tìm m để 2 vectơ cùng phương.
                                                if dang_toan == "[D10_CX_B0_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_11()                                     
                                                    

                                            #[D10_CX_B0_12]-M2. Cho hai véctơ. Tìm m để 2 vectơ vuông góc.
                                                if dang_toan == "[D10_CX_B0_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_12()                                     
                                                    

                                            #[D10_CX_B0_13]-M2. Tìm A thuộc Ox cách B một khoảng cho trước.
                                                if dang_toan == "[D10_CX_B0_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B0_13()                                     
                                                    

                                    #Chương 6- Thống kê
                                            #Bài 1: Số gần đúng
                                            #[D10_C6_B1_01]. Cho dãy số liệu. Tính số trung bình.
                                                if dang_toan == "[D10_C6_B1_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B1_01()

                                            #Bài 3:Các số đặc trưng đo xu thế trung tâm của mẫu số liệu
                                            #[D10_C6_B3_01]. Cho dãy số liệu. Tính số trung bình.
                                                if dang_toan == "[D10_C6_B3_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_01()

                                            #[D10_C6_B3_02]. Cho dãy số liệu. Tính số trung vị.
                                                if dang_toan == "[D10_C6_B3_02]":    
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_02()

                                            #[D10_C6_B3_03]. Cho dãy số liệu. Tính mốt.
                                                if dang_toan == "[D10_C6_B3_03]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_03()

                                            #[D10_C6_B4_04]. Cho dãy số liệu. Tính tứ phân vị thứ nhất
                                                if dang_toan == "[D10_C6_B3_04]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_04()

                                            #[D10_C6_B3_05]. Cho dãy số liệu. Tính tứ phân vị thứ hai
                                                if dang_toan == "[D10_C6_B3_05]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_05()

                                            #[D10_C6_B3_06]. Cho dãy số liệu. Tính tứ phân vị thứ ba
                                                if dang_toan == "[D10_C6_B3_06]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B3_06()

                                            #Bài 4
                                            #[D10_C6_B4_01]. Cho dãy số liệu. Tính độ lệch chuẩn.
                                                if dang_toan == "[D10_C6_B4_01]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B4_01()

                                            #[D10_C6_B4_02]. Cho dãy số liệu. Tính phương sai.
                                                if dang_toan == "[D10_C6_B4_02]":           
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C6.tktk_L10_C6_B4_02()

                                            
                                    #Toán 10- Chương 7- Bất phương trình bậc 2 một ẩn
                                            #Bài 1: Dấu của tam thức bậc 2
                                            #[D10_C7_B1_01]-M1. Cho biểu thức bậc 2 vô nghiệm. Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_01()                                    
                                                    

                                            #[D10_C7_B1_02]-M1. Cho biểu thức bậc 2 luôn nghiệm kép. Tìm khẳng định đúng về dấu.
                                                
                                                if dang_toan == "[D10_C7_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_02()                                     
                                                    

                                            #[D10_C7_B1_03]-M2. Cho biểu thức bậc 2 có 2 nghiệm. Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_03()                                     
                                                    

                                            #[D10_C7_B1_04]-M1. Tìm m để biểu thức là bậc 2.
                                                                                       
                                                if dang_toan == "[D10_C7_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_04()                                     
                                                    

                                            #[D10_C7_B1_05]-M1. Xác định biểu thức nào là tam thức bậc 2.
                                                if dang_toan == "[D10_C7_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_05()                                     
                                                    

                                            #[D10_C7_B1_06]-M1.  Cho đồ thị bậc 2 luôn âm hoặc luôn dương. Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_06()                                     
                                                    

                                            #[D10_C7_B1_07]-M1. Cho đồ thị bậc 2 có nghiệm kép. Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_07()                                     
                                                    

                                            #[D10_C7_B1_08]-M1. Cho đồ thị bậc 2 có 2 nghiệm . Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_08()                                     
                                                    

                                            #[D10_C7_B1_09]-M1. Cho BXD bậc 2 có 2 nghiệm . Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_09()                                     
                                                    

                                            #[D10_C7_B1_10]-M1. Cho BXD bậc 2 có 1 nghiệm . Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_10()                                     
                                                    

                                            #[D10_C7_B1_11]-M1. Cho BXD bậc 2 có 2 nghiệm . Tìm khẳng định đúng về dấu.
                                                if dang_toan == "[D10_C7_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_11()                                     
                                                    

                                            #[D10_C7_B1_12]-TF-M2. Cho biểu thức bậc 2 có 2 nghiệm. Tạo câu đúng-sai.
                                                if dang_toan == "[D10_C7_B1_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C7.aaa_pry_L10_C7_B1_12()                                      


                                            #[D10_C7_B1_13]-TF-M2. Cho bảng xét dấu hai nghiệm. Tạo câu Đ-S.
                                                if dang_toan == "[D10_C7_B1_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C7.aaa_pry_L10_C7_B1_13()                                      


                                            #[D10_C7_B1_14]-M2. Cho bxd bậc 2 có 2 nghiệm. Tìm hàm số.
                                                if dang_toan == "[D10_C7_B1_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_14()                                     
                                                    

                                            #[D10_C7_B1_15]-M2. Cho bxd bậc 2 có nghiệm kép. Tìm hàm số.
                                                if dang_toan == "[D10_C7_B1_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B1_15()                                     
                                                    



                                            #Bài 2 - Giải bất phương trình bậc 2:
                                            #[D10_C7_B2_01]-M1. Giải bất phương trình bậc 2, tam thức vô nghiệm.
                                                if dang_toan == "[D10_C7_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_01()                                     
                                                    

                                            #[D10_C7_B2_01]-M1. Giải bất phương trình bậc 2, tam thức nghiệm kép.
                                                if dang_toan == "[D10_C7_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_02()                                     
                                                    

                                            #[D10_C7_B2_03]-M1. Giải bất phương trình bậc 2, tam thức 2 nghiệm.
                                                if dang_toan == "[D10_C7_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_03()                                     
                                                    

                                            #[D10_C7_B2_04]-M1. Giải BPT có 2 vế là bậc 2
                                                if dang_toan == "[D10_C7_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_04()                                     
                                                    

                                            #[D10_C7_B2_05]-M3. Tìm các giá trị của tham số m để ax^2 +bx+c>0 (<0)  với mọi x.
                                                if dang_toan == "[D10_C7_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_05()                                     
                                                    

                                            #[D10_C7_B2_06]-M3. Tìm các giá trị của tham số m để ax^2 +bx+c>=0 (<=0)  với mọi x.
                                                if dang_toan == "[D10_C7_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_06()                                     
                                                    

                                            #[D10_C7_B2_07]-M3. Tìm các giá trị của tham số m để ax^2 +bx+c>=0 (<=0) vô nghiệm.
                                                if dang_toan == "[D10_C7_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_07()                                     
                                                    

                                            #[D10_C7_B2_08]-M3. Tìm các giá trị của tham số m để ax^2 +bx+c>0 (<0) vô nghiệm.
                                                if dang_toan == "[D10_C7_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_08()                                     
                                                    

                                            #[D10_C7_B2_09]-M2. Tìm tập xác định hàm số y=căn(ax^2+bx+c)
                                                if dang_toan == "[D10_C7_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B2_09()                                     
                                                    

                                            #Bài 3: Phương trình quy về bậc 2
                                            #[D10_C7_B3_01]-M2. Giải PT căn(ax^2 + bx^2 + c)=căn(dx^2 + ex + f)

                                                if dang_toan == "[D10_C7_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B3_01()                                     
                                                    

                                            #[D10_C7_B3_02]-M2. Giải PT căn(ax^2 + bx^2 + c)=dx+e
                                                if dang_toan == "[D10_C7_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C7.aaa_pry_L10_C7_B3_02()                                     
                                                    
                                    #Toán 11 - Chương 8 - Đại số tổ hợp
                                    #Bài 1: Quy tắc cộng và Quy tắc nhân
                                    #[D10_C8_B1_01]-M1. Cho 2 nhóm đồ vật. Tìm số cách chọn 1 vật.
                                                if dang_toan == "[D10_C8_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_01()                                     
                                                    

                                    #[D10_C8_B1_02]-M1. Cho số lượng học sinh. Chọn 1 bạn để giữ chức vụ.
                                                if dang_toan == "[D10_C8_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_02()                                     
                                                    

                                    #[D10_C8_B1_03]-M1. Chọn một thức uống từ các loại nước.
                                                if dang_toan == "[D10_C8_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_03()                                     
                                                    

                                    #[D10_C8_B1_04]-M1. Chọn một địa điểm đi du lịch.
                                                if dang_toan == "[D10_C8_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_04()                                     
                                                    
                                    #9.1.2 Quy tắc nhân

                                    #[D10_C8_B1_05]-M1. Cho 2 nhóm đồ vật. Tìm số cách chọn 2 vật.
                                                if dang_toan == "[D10_C8_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_05()                                     
                                                    

                                    #[D10_C8_B1_06]-M1. Cho số lượng học sinh. Chọn 2 bạn để giữ chức vụ.
                                                if dang_toan == "[D10_C8_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_06()                                     
                                                    

                                    #[D10_C8_B1_07]-M1. Tìm số đường đi từ A đến B và từ B đến C.
                                                if dang_toan == "[D10_C8_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_07()                                     
                                                    

                                    #[D10_C8_B1_08]-M1. Từ n chữ số lập được bao nhiêu số có k chữ số.
                                                if dang_toan == "[D10_C8_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_08()                                     
                                                    

                                    #[D10_C8_B1_09]-M1. Từ n chữ số lập được bao nhiêu số có k chữ số.
                                                if dang_toan == "[D10_C8_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_09()                                     
                                                    

                                    #[D10_C8_B1_10]-M1. Từ n chữ số lập được bao nhiêu số có k chữ số.
                                                if dang_toan == "[D10_C8_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_10()                                     
                                                    

                                    #[D10_C8_B1_11]-M3. Chọn 2 quả cầu khác màu và tổng là chẵn hoặc lẻ.
                                                if dang_toan == "[D10_C8_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B1_11()                                     
                                                    
                                    #8.2.1 Hoán vị

                                    #[D10_C8_B2_01]-M1. Số cách xếp n bạn vào một hàng ngang(dọc).
                                                if dang_toan == "[D10_C8_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_01()                                     
                                                    

                                    #[D10_C8_B2_02]-M1. Tìm số cách chọn k vật từ n nhóm đồ vật.
                                                if dang_toan == "[D10_C8_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_02()                                     
                                                    

                                    #[D10_C8_B2_03]-M1. Tìm số cách chọn k người từ n người và sắp xếp chức vụ trong lớp.
                                                if dang_toan == "[D10_C8_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_03()                                     
                                                    

                                    #[D10_C8_B2_04]-M2. Tìm số cách chọn k người từ n người và sắp xếp chức vụ trong lớp.
                                                if dang_toan == "[D10_C8_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_04()                                     
                                                    

                                    #[D10_C8_B2_05]-M2. Từ n chữ số lập được bao nhiêu số có k chữ số khác nhau.
                                                if dang_toan == "[D10_C8_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_05()                                     
                                                    

                                    #[D10_C8_B2_06]-M2. Lập số có k chữ số khác nhau từ n chữ số (chứa số 0).
                                                if dang_toan == "[D10_C8_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_06()                                     
                                                    

                                    #[D10_C8_B2_07]-M2. Xếp 3 nhóm đồ vật vào một hàng.
                                                if dang_toan == "[D10_C8_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_07()                                     
                                                    
                                    #Tổ hợp
                                    #[D10_C8_B2_08]-TF-M2. Tạo câu đúng-sai: Chọn 2 nhóm đồ vật.
                                                if dang_toan == "[D10_C8_B2_08]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_08()


                                    #[D10_C8_B2_09]-M1. Chọn k đối tượng tự n đối tượng.
                                                if dang_toan == "[D10_C8_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_09()                                     
                                                    

                                    #[D10_C8_B2_10]-M2. Chọn k đối tượng tự từ 2 nhóm đối tượng.
                                                if dang_toan == "[D10_C8_B2_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_10()                                     
                                                    

                                    #[D10_C8_B2_11]-M2. Chọn k đối tượng tự từ 2 nhóm đối tượng thỏa mãn điều kiện.
                                                if dang_toan == "[D10_C8_B2_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_11()                                     
                                                    

                                    #[D10_C8_B2_12]-M2. Xếp k học sinh vào 1 hàng, có 2 bạn đứng ở 2 đầu.
                                                if dang_toan == "[D10_C8_B2_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_12()                                     
                                                    

                                    #[D10_C8_B2_13]-M3. Xếp k học sinh vào 1 hàng, có 2 bạn đứng cạnh nhau hoặc không cạnh nhau.
                                                if dang_toan == "[D10_C8_B2_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_13()                                     
                                                    

                                    #[D10_C8_B2_14]-M2. Số đoạn thẳng tạo bởi n điểm
                                                if dang_toan == "[D10_C8_B2_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_14()                                     
                                                    

                                    #[D10_C8_B2_15]-M2. Số tam giác tạo bởi n điểm
                                                if dang_toan == "[D10_C8_B2_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_15()                                     
                                                    

                                    #[D10_C8_B2_16]-M1. Số tập hợp con gồm k phần tử của n phần tử
                                                if dang_toan == "[D10_C8_B2_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_16()                                     
                                                    

                                    #[D10_C8_B2_17]-M3. Số tập hợp con gồm k phần tử của n phần tử
                                                if dang_toan == "[D10_C8_B2_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_17()                                     
                                                    

                                    #[D10_C8_B2_18]-M2. Số giao điểm tối đa của n đường thẳng phân biệt
                                                if dang_toan == "[D10_C8_B2_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_18()                                     
                                                    

                                    #[D10_C8_B2_19]-M2. Số hình chữ nhật tạo từ n đường song song và m đường vuông góc
                                                if dang_toan == "[D10_C8_B2_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_19()                                     
                                                    

                                    #[D10_C8_B2_20]-M2. Số giao điểm tối đa của n đường tròn
                                                if dang_toan == "[D10_C8_B2_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_20()                                     
                                                    

                                    #[D10_C8_B2_21]-M2. Số giao điểm tối đa của m đường thẳng và n đường tròn.
                                                if dang_toan == "[D10_C8_B2_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B2_21()                                     
                                                    

                                    #[D10_C8_B2_22]-TF-M2. Tạo câu đúng-sai: Chọn 2 cuốn sách.
                                                if dang_toan == "[D10_C8_B2_22]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_22()                                       

                                    #[D10_C8_B2_23]-TF-M2. Tạo câu đúng-sai: Chọn 2 cuốn truyện.
                                                if dang_toan == "[D10_C8_B2_23]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_23()

                                    #[D10_C8_B2_22]-TF-M2. Tạo câu đúng-sai: Chọn 2 cuốn sách.
                                                if dang_toan == "[D10_C8_B2_22]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_22()                                       


                                    #[D10_C8_B2_24]-TF-M2. Tạo câu đúng-sai: Chọn 2 bức tranh.
                                                if dang_toan == "[D10_C8_B2_24]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_24()                                       


                                    #[D10_C8_B2_25]-TF-M2. Tạo câu đúng-sai: Chọn 2 viên bi.
                                                if dang_toan == "[D10_C8_B2_25]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_25()                                       
                                                    

                                    #[D10_C8_B2_26]-TF-M2. Tạo câu đúng-sai: Chọn 2 người.
                                                if dang_toan == "[D10_C8_B2_26]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C8.mcn__L10_C8_B2_26()                                       
                                                    

                                    #[D10_C8_B3_01]-M1. Khai triển (x+a)^n
                                                if dang_toan == "[D10_C8_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_01()                                     
                                                    

                                    #[D10_C8_B3_02]-M2. Khai triển (a-x)^n
                                                if dang_toan == "[D10_C8_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_02()                                     
                                                    

                                    #[D10_C8_B3_03]-M2.  Khai triển (ax+b)^n
                                                if dang_toan == "[D10_C8_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_03()                                     
                                                    
                                    
                                    #[D10_C8_B3_04]-M2.  Tìm hệ số của x^k trong khai triển của (ax+b)^n
                                                if dang_toan == "[D10_C8_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_04()                                     
                                                    

                                    #[D10_C8_B3_05]-M2.  Tìm số hạng chứa x^k trong khai triển của (ax+b)^n
                                                if dang_toan == "[D10_C8_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_05()                                     
                                                    

                                    #[D10_C8_B3_06]-M2.  Tìm hệ số của x^k trong khai triển của (ax+b)^n, n=4,5
                                                if dang_toan == "[D10_C8_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_06()                                     
                                                    
                                    #[D10_C8_B3_07]-M2.  Tìm số hạng chứa x^k trong khai triển của (ax+b)^n, n=4,5
                                                if dang_toan == "[D10_C8_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_07()

                                    #[D10_C8_B3_08]-M2.  Tìm số hạng thứ k trong khai triển của (ax+b)^n số mũ giảm dần
                                                if dang_toan == "[D10_C8_B3_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C8.mcn__L10_C8_B3_08()                                  
                                                    
                                    #Toán 10 - Chương 9 - Xác suất
                                    #Bài 1. Không gian mẫu và biến cố

                                    #[D10_C9_B1_01]-M1. Chọn ngẫu nhiên 1 vật từ 2 nhóm đồ vật. Tìm số phần tử không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_01()                                     
                                                    

                                    #[D10_C9_B1_02]-M1. Cho ngẫu nhiên 2 vật từ 2 nhóm đồ vật. Tìm số phần tử của không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_02()                                     
                                                    

                                    #[D10_C9_B1_03]-M2. Chọn k quả cầu khác màu. Tìm số phần tử của không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_03()                                     
                                                    

                                    #[D10_C9_B1_04]-M2. Chọn k con xúc xắc. Tìm số phần tử của không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_04()                                     
                                                    

                                        #[D10_C9_B1_05]-M2. Gieo k đồng xu. Tìm số phần tử của không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_05()                                     
                                                    

                                        #[D10_C9_B1_06]-M2. Chọn k vật từ n vật. Tính số phần tử của không gian mẫu
                                                if dang_toan == "[D10_C9_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_06()                                     
                                                    

                                        #[D10_C9_B1_07]-M2. Chọn k đối tượng từ 2 nhóm đối tượng. Tính số phần tử của không gian mẫu.
                                                if dang_toan == "[D10_C9_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B1_07()

                                        #[D10_C9_B1_08]-M2. Chọn k vật từ 2 nhóm. Xét đúng-sai: không gian mẫu, biến cố k vật được chọn, cả 2 loại được chọn, ít nhất 1 được chọn.
                                                if dang_toan == "[D10_C9_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_08()

                                        #[D10_C9_B1_09]-M2. Gieo con xúc sắc 2 lần: không gian mẫu, biến cố liên quan tổng, tích số chấm 2 lần gieo.
                                                if dang_toan == "[D10_C9_B1_09]":     
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_09()
                                        
                                        #[D10_C9_B1_10]-TF-M3. Chọn k vật từ 3 nhóm. Xét Đ-S: không gian mẫu, biến cố ít nhất một, có đúng m được chọn, chỉ một loại được chọn.
                                                if dang_toan == "[D10_C9_B1_10]":
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_10()

                                        #[D10_C9_B1_11]-TF-M3. Có n tấm thẻ, lấy ngẫu nhiên k thẻ. Đúng-Sai: không gian mẫu, các thẻ đều chẵn (lẻ), ít nhất một thẻ chia hết cho m, m thẻ chẵn (lẻ) được chọn.
                                                if dang_toan == "[D10_C9_B1_11]":
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_11()       
                                                    
                                        #Bài 2- Xác suất

                                        #[D10_C9_B2_01]-M2. Chọn n vật từ 2 nhóm. Tính xác suất k vật được chọn.
                                                if dang_toan == "[D10_C9_B2_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_01()                                     
                                                    

                                        #[D10_C9_B2_02]-M3. Chọn n vật từ 2 nhóm. Tính xác suất k vật được chọn.
                                                if dang_toan == "[D10_C9_B2_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_02()                                     
                                                    

                                        #[D10_C9_B2_03]-M3. Cho một quả cầu từ n quả cầu được đánh số. Tính xác suất quả là số chẵn (lẻ).
                                                if dang_toan == "[D10_C9_B2_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_03()                                     
                                                    
                                        
                                        #[D10_C9_B2_04]-M3. Cho một quả cầu từ n quả cầu được đánh số. Tính xác suất quả là số chẵn (lẻ) và chia hết cho k.
                                                if dang_toan == "[D10_C9_B2_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_04()                                     
                                                    

                                        #[D10_C9_B2_05]-M3. Cho một hộp chứa n tấm thẻ đánh số. Tính xác suất thẻ được chọn ghi số thuộc [a;b].
                                                if dang_toan == "[D10_C9_B2_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_05()                                     
                                                    

                                        #[D10_C9_B2_06]-M2. Gieo 2 con xúc xắc. Tính xác suất để  i + j>k.
                                                if dang_toan == "[D10_C9_B2_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_06()                                     
                                                    

                                        #[D10_C9_B2_07]-M2. Gieo 2 con xúc xắc. Tính xác suất để  i + j<k.
                                                if dang_toan == "[D10_C9_B2_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_07()                                     
                                                    

                                        #[D10_C9_B2_08]-M2. Gieo 2 con xúc xắc. Tính xác suất để  i + j=k.
                                                if dang_toan == "[D10_C9_B2_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_08()                                     
                                                    

                                        #[D10_C9_B2_09]-M2. Gieo 2 con xúc xắc. Tìm biến cố giao: i + j = k và i.j <(>) m.
                                                if dang_toan == "[D10_C9_B2_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C9.mjulk_L10_C9_B2_09()
                                        #ĐÚNGG-SAI
                                        #[D10_C9_B1_12]-TF-M3. Chọn 2 nhóm đồ vật. Đúng-sai: không gian mẫu, xác suất
                                                if dang_toan == "[D10_C9_B1_12]":
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_12()

                                        #Gieo một con xúc sắc 2 lần. Đúng-Sai: không gian mẫu, xác suất.
                                                if dang_toan == "[D10_C9_B1_13]":
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_13()

                                        #Chọn k vật từ 3 nhóm. Xét Đ-S: không gian mẫu, xác suất.
                                                if dang_toan == "[D10_C9_B1_14]":
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C9.mjulk_L10_C9_B1_14()
                                    #Toán 10 - Chương 10 - Phương pháp tọa độ trong mặt phẳng
                                    #Bài 1. Phương trình đường thẳng
                                    #[D10_CX_B1_01]-M1. Cho véctơ pháp tuyến tìm véctơ chỉ phương và ngược lại
                                                if dang_toan == "[D10_CX_B1_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_01()                                     
                                                    

                                    #[D10_CX_B1_02]-M2. Đọc véctơ pháp tuyến(véctơ chỉ phương) từ phương trình tổng quát
                                                if dang_toan == "[D10_CX_B1_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_02()                                     
                                                    

                                    #[D10_CX_B1_03]-M2. Đọc véctơ pháp tuyến(véctơ chỉ phương) từ phương trình tham số
                                                if dang_toan == "[D10_CX_B1_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_03()                                     
                                                    
                                    #10.1.3. Phương trình tổng quát 

                                    #[D10_CX_B1_04]-M2. Cho tọa độ điểm và véctơ pháp tuyến, viết phương trình tổng quát
                                                if dang_toan == "[D10_CX_B1_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_04()                                     
                                                    

                                    #[D10_CX_B1_05]-M2. Cho tọa độ điểm và véctơ chỉ phương, viết phương trình tổng quát
                                                if dang_toan == "[D10_CX_B1_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_05()                                     
                                                    

                                    #[D10_CX_B1_06]-M2. Viết phương trình tổng quát đường thẳng qua điểm song song với đường thẳng.
                                                if dang_toan == "[D10_CX_B1_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_06()                                     
                                                    

                                    #[D10_CX_B1_07]-M2. Viết phương trình tổng quát đường thẳng qua điểm vuông góc với đường thẳng.
                                                if dang_toan == "[D10_CX_B1_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_07()                                     
                                                    



                                    #[D10_CX_B1_08]-M2. Viết phương trình tổng quát đường thẳng qua 2 điểm
                                                if dang_toan == "[D10_CX_B1_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_08()                                     
                                                    

                                    #[D10_CX_B1_09]-M2. Viết phương trình tổng quát đường cao trong tam giác
                                                if dang_toan == "[D10_CX_B1_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_09()                                     
                                                    

                                    #[D10_CX_B1_10]-M2. Viết phương trình tổng quát đường trung trực trong tam giác
                                                if dang_toan == "[D10_CX_B1_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_10()                                     
                                                    
                                    
                                    #[D10_CX_B1_11]-M2. Viết phương trình tổng quát đường trung tuyến trong tam giác
                                                if dang_toan == "[D10_CX_B1_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_11()                                     
                                                    

                                    #[D10_CX_B1_12]-M1. Lập PTTS của d qua điểm và có VT chỉ phương
                                                if dang_toan == "[D10_CX_B1_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_12()                                     
                                                    

                                    #[D10_CX_B1_13]-M1. Lập PTTS của d qua điểm và có VT pháp tuyến
                                                if dang_toan == "[D10_CX_B1_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_13()                                     
                                                    

                                    #[D10_CX_B1_14]-M2. Lập PTTS của d qua 2 điểm 
                                                if dang_toan == "[D10_CX_B1_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_14()                                     
                                                    

                                    #[D10_CX_B1_15]-M2. Lập PTTS của d qua 2 điểm 
                                                if dang_toan == "[D10_CX_B1_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_15()                                     
                                                    

                                    #[D10_CX_B1_16]-M2. Lập PTTQ của d từ PTTS. 
                                                if dang_toan == "[D10_CX_B1_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_16()                                     
                                                    

                                    #[D10_CX_B1_17]-M1. Tìm điểm thuộc d biết PTTS.
                                                if dang_toan == "[D10_CX_B1_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_17()                                     
                                                    

                                            #[D10_CX_B1_18]-M1. Tìm điểm thuộc d biết PTTQ.
                                                if dang_toan == "[D10_CX_B1_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_18()                                     
                                                    

                                            #[D10_CX_B1_19]-M2. Cho PTTQ của 2 đường thẳng, xét vị trí tương đối.
                                                if dang_toan == "[D10_CX_B1_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_19()                                     
                                                    

                                            #[D10_CX_B1_20]-M2. Cho PTTQ của 2 đường thẳng, tính cosin của góc.
                                                if dang_toan == "[D10_CX_B1_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_20()                                     
                                                    

                                            #[D10_CX_B1_21]-M2. Cho điểm và PTTQ của đường thẳng, tính khoảng cách.
                                                if dang_toan == "[D10_CX_B1_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_21()                                     
                                                    

                                            #[D10_CX_B1_22]-M2. Cho điểm và PTTS của đường thẳng, tính khoảng cách.
                                                if dang_toan == "[D10_CX_B1_22]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_22()                                     
                                                    

                                            #[D10_CX_B1_23]-M2. Cho PTTS của 2 đường thẳng, tính cosin của góc.
                                                if dang_toan == "[D10_CX_B1_23]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_23()                                     
                                                    

                                            #[D10_CX_B1_24]-M2. Cho PTTS của 2 đường thẳng, tính cosin của góc.
                                                if dang_toan == "[D10_CX_B1_24]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_24()                                     
                                                    

                                            #[D10_CX_B1_25]-M2. Cho PTTS của 2 đường thẳng, xét vị trí tương đối.
                                                if dang_toan == "[D10_CX_B1_25]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_25()                                     
                                                    

                                            #[D10_CX_B1_26]-TF-M2. Tạo câu đúng-sai: Cho PTTQ của 1 đường thẳng, xét đúng-sai về VTCP,VTPT, điểm thuộc đường thẳng.    
                                                if dang_toan == "[D10_CX_B1_26]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B1_26()                                      

                                            #[D10_CX_B1_27]-TF-M2. Tạo câu đúng-sai: Cho PTTS của 1 đường thẳng, xét đúng-sai về VTCP,VTPT, điểm thuộc đường thẳng.   
                                                if dang_toan == "[D10_CX_B1_27]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B1_27()                                      

                                            #[D10_CX_B1_28]-TF-M2. Tạo câu đúng-sai: Tạo câu đúng-sai: Cho PTTS của 1 đường thẳng, xét đúng-sai về VTCP, VTPT, PTTS, PTTQ.  
                                                if dang_toan == "[D10_CX_B1_28]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B1_28()                                      


                                            #[D10_CX_B1_29]-TF-M2. Tạo câu đúng-sai: Cho PTTQ của 1 đường thẳng, xét đúng-sai về vị trí tương đối.   
                                                if dang_toan == "[D10_CX_B1_29]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B1_29()                                      

                                            #[D10_CX_B1_30]-TF-M2. Cho 1 điểm và PTTQ của 1 đường thẳng, xét đúng-sai về vị trí tương đối, khoảng cách.
                                                if dang_toan == "[D10_CX_B1_30]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B1_30()                                      


                                            #[D10_CX_B1_31]-M1. Viết phương trình mặt cầu có tâm và bán kính
                                                if dang_toan == "[D10_CX_B1_31]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B1_31()                                     
                                                    

                                            #Bài 3 - Phương trình đường tròn
                                            #[D10_CX_B3_01]-M1. Viết phương trình mặt cầu có tâm và bán kính
                                                if dang_toan == "[D10_CX_B3_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_01()                                     
                                                    

                                            #[D10_CX_B3_02]-M1. Đọc tọa độ tâm từ phương trình mặt cầu
                                                if dang_toan == "[D10_CX_B3_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_02()                                     
                                                    

                                            #[D10_CX_B3_03]-M1. Đọc bán kính từ phương trình mặt cầu thu gọn
                                                if dang_toan == "[D10_CX_B3_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_03()                                     
                                                    

                                            #[D10_CX_B3_04]-M1. Đọc tọa độ tâm từ phương trình mặt cầu khai triển
                                                if dang_toan == "[D10_CX_B3_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_04()                                     
                                                    

                                            #[D10_CX_B3_05]-M1. Đọc bán kính từ phương trình mặt cầu khai triển
                                                if dang_toan == "[D10_CX_B3_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_05()                                     
                                                    

                                            #[D10_CX_B3_06]-M2. Viết phương trình mặt cầu có tâm và đường kính
                                                if dang_toan == "[D10_CX_B3_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_06()                                     
                                                    

                                            #[D10_CX_B3_07]-M3. Viết phương trình mặt cầu có đường kính AB
                                                if dang_toan == "[D10_CX_B3_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_07()                                     
                                                    

                                            #[D10_CX_B3_08]-M3. Viết phương trình đường tròn đi qua 3 điểm
                                                if dang_toan == "[D10_CX_B3_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_08()                                     
                                                    

                                            #[D10_CX_B3_09]-M2. Viết tiếp tuyến của đường tròn tại điểm (x_0;y_0)
                                                if dang_toan == "[D10_CX_B3_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_09()                                     
                                                    

                                            #[D10_CX_B3_10]-M3. Viết tiếp tuyến của đường tròn song song với đường thẳng
                                                if dang_toan == "[D10_CX_B3_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_10()                                     
                                                    

                                            #[D10_CX_B3_11]-M3. Viết tiếp tuyến của đường tròn vuông góc với đường thẳng
                                                if dang_toan == "[D10_CX_B3_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_11()                                     
                                                    

                                            #[D10_CX_B3_12]-TF-M2. Cho phương trình đường tròn dạng thu gọn. Tạo Đúng-Sai
                                                if dang_toan == "[D10_CX_B3_12]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B3_12()

                                            #[D10_CX_B3_13]-TF-M2. Cho phương trình đường tròn dạng khai triển. Tạo Đúng-Sai
                                                if dang_toan == "[D10_CX_B3_13]": 
                                                    debai_word,debai_latex,loigiai_word,dap_an=D10_C10.gghik_L10_CX_B3_13()


                                            #[D10_CX_B3_14]-M2. Viết phương trình đường tròn có tâm và đi qua điểm
                                                if dang_toan == "[D10_CX_B3_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_14()                                     
                                                    

                                            #[D10_CX_B3_15]-M2. Viết phương trình đường tròn có tâm và tiếp xúc với đường thẳng
                                                if dang_toan == "[D10_CX_B3_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_15()                                     
                                                

                                                #[D10_CX_B3_16]-M2. Viết phương trình đường tròn có tâm và tiếp xúc với Ox(Oy)
                                                if dang_toan == "[D10_CX_B3_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_16()                                     
                                                    

                                                #[D10_CX_B3_17]-M2. Xét vị trí tương đối giữa điểm và đường tròn có phương trình thu gọn
                                                if dang_toan == "[D10_CX_B3_17]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B3_17()                                     
                                                    

                                                #Bài 4: PHƯƠNG TRÌNH ĐƯỜNG ELIP
                                                #[D10_CX_B4_01]-M2. Cho phương trình Elip. Đọc trục lớn.
                                                if dang_toan == "[D10_CX_B4_01]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_01()                                     
                                                    

                                            #[D10_CX_B4_02]-M2. Cho phương trình Elip. Đọc trục lớn.
                                                if dang_toan == "[D10_CX_B4_02]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_02()                                     
                                                    

                                            #[D10_CX_B4_03]-M2. Cho phương trình Elip. Đọc tiêu cự.
                                                if dang_toan == "[D10_CX_B4_03]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_03()                                     
                                                    

                                            #[D10_CX_B4_04]-M2. Cho phương trình Elip.Đọc tiêu điểm.
                                                if dang_toan == "[D10_CX_B4_04]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_04()                                     
                                                    

                                            #[D10_CX_B4_05]-M2. Cho phương trình Elip.Đọc tọa độ đỉnh.
                                                if dang_toan == "[D10_CX_B4_05]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_05()                                     
                                                    

                                            #[D10_CX_B4_06]-M2. Cho trục lớn, trục nhỏ. Viết phương trình Elip.
                                                if dang_toan == "[D10_CX_B4_06]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_06()                                     
                                                    

                                            #[D10_CX_B4_07]-M2. Cho trục lớn, tiêu cự. Viết phương trình Elip.
                                                if dang_toan == "[D10_CX_B4_07]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_07()                                     
                                                    

                                            #[D10_CX_B4_08]-M2. Cho trục nhỏ, tiêu cự. Viết phương trình Elip.
                                                if dang_toan == "[D10_CX_B4_08]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_08()                                     
                                                    

                                            #[D10_CX_B4_09]-M2. Cho đỉnh thuộc trục lớn, trục nhỏ. Viết phương trình Elip.
                                                if dang_toan == "[D10_CX_B4_09]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_09()                                     
                                                    

                                            #[D10_CX_B4_10]-M2. Cho đỉnh thuộc trục lớn, tiêu điểm. Viết phương trình Elip.
                                                if dang_toan == "[D10_CX_B4_10]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_10()                                     
                                                    

                                            #[D10_CX_B4_11]-M2. Cho phương trình Hypebol. Tìm độ dài trục thực.
                                                if dang_toan == "[D10_CX_B4_11]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_11()                                     
                                                    

                                            #[D10_CX_B4_12]-M2. Cho phương trình Hypebol. Tìm độ dài trục ảo.
                                                if dang_toan == "[D10_CX_B4_12]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_12()                                     
                                                    

                                            #[D10_CX_B4_13]-M2. Cho phương trình Hypebol. Tìm tiêu cự.
                                                if dang_toan == "[D10_CX_B4_13]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_13()                                     
                                                    

                                            #[D10_CX_B4_14]-M2. Cho phương trình Hypebol. Đọc tiêu điểm.
                                                if dang_toan == "[D10_CX_B4_14]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_14()                                     
                                                    

                                            #[D10_CX_B4_15]-M2. Cho phương trình Hypebol. Đọc tọa độ đỉnh.
                                                if dang_toan == "[D10_CX_B4_15]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_15()                                     
                                                    

                                            #[D10_CX_B4_16]-M2.  Cho trục thực, trục ảo. Viết phương trình Hypebol.
                                                if dang_toan == "[D10_CX_B4_16]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_16()                                     
                                                    

                                                #[D10_CX_B4_17]-M2.  Cho trục thực, tiêu cự. Viết phương trình Hypebol.
                                                if dang_toan == "[D10_CX_B4_17]":
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_17()                                     
                                                    

                                                #[D10_CX_B4_18]-M2.  Cho trục ảo, tiêu cự. Viết phương trình Hypebol.
                                                if dang_toan == "[D10_CX_B4_18]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_18()                                     
                                                    

                                                #[D10_CX_B4_19]-M2. Cho đỉnh thuộc trục thực, tiêu điểm. Viết phương trình Hypebol
                                                if dang_toan == "[D10_CX_B4_19]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_19()                                     
                                                    

                                                #[D10_CX_B4_20]-M2. Cho parabol. Tìm tham số tiêu.
                                                if dang_toan == "[D10_CX_B4_20]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_20()                                     
                                                    

                                                #[D10_CX_B4_21]-M2. Cho parabol. Tìm đường chuẩn.
                                                if dang_toan == "[D10_CX_B4_21]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_21()                                     
                                                    

                                                #[D10_CX_B4_22]-M2. Cho parabol. Tìm tiêu điểm.
                                                if dang_toan == "[D10_CX_B4_22]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_22()                                     
                                                    

                                                #[D10_CX_B4_23]-M2. Viết phương trình parabol có tham số tiêu.
                                                if dang_toan == "[D10_CX_B4_23]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_23()                                     
                                                    

                                            #[D10_CX_B4_24]-M2. Viết phương trình parabol có tiêu điểm.
                                                if dang_toan == "[D10_CX_B4_24]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_24()                                     
                                                    

                                            #[D10_CX_B4_25]-M2. Viết phương trình parabol có đường chuẩn.
                                                if dang_toan == "[D10_CX_B4_25]":                                        
                                                    debai_word,debai_latex,loigiai_word,phuongan,latex_tuluan,loigiai_traloingan,dap_an=D10_C10.gghik_L10_CX_B4_24()                                     
                                                                                         
                                                                               
                                            #Xử lí xuất câu hỏi                                                
                                                if loai_cau=="TN":
                                                    if self.combo_taode.currentText() in ["Tạo đề Word - Equation",  "Tạo đề Word - MathType"]:
                                                        list_tracnghiem.append(f'{debai_word}\n{phuongan}\n')
                                                        list_tracnghiem_HDG.append(f'{debai_word}\n{phuongan}\n{loigiai_word}\n')
                                                        list_dapan_TN.append(f'{dap_an}')                                                       
                                                        
                                                    else:                                                       
                                                        list_tracnghiem.append(debai_latex)
                                                        list_tracnghiem_HDG.append(debai_latex)
                                                        list_dapan_TN.append('')

                                                if loai_cau=="Đ-S":
                                                    if self.combo_taode.currentText() in ["Tạo đề Latex - PDF", "Tạo code Latex"]:                                                          
                                                        list_dungsai.append(debai_latex)
                                                        list_dungsai_HDG.append(debai_latex)
                                                        list_dapan_TF.append('')
                                                    else:                                                           
                                                        list_dungsai.append(f'{debai_word}\n')
                                                        list_dungsai_HDG.append(f'{debai_word}\n{loigiai_word}\n')
                                                        list_dapan_TF.append(f'{dap_an}')
                                                    
                                           
                                                if loai_cau=="TL":
                                                    if self.combo_taode.currentText() in ["Tạo đề Latex - PDF", "Tạo code Latex"]:                                                        
                                                        list_tuluan.append(latex_tuluan)
                                                        list_tuluan_HDG.append(latex_tuluan)
                                                        list_dapan_TL.append('')
                                                    else:                                                                                                       
                                                        list_tuluan.append(f'{debai_word}\n')                                                    
                                                        list_tuluan_HDG.append(f'{debai_word}\n{loigiai_word}\n')
                                                        list_dapan_TL.append(f'{dap_an}')

                                                                                                                            
 
#end
#############################################################################################################################

                                                #Cập nhật số câu đã xử lí
                                                if dong_dangtoan[0]=="[":
                                                        socau_daxuli += 1
                                                        self.label_socau_daxuli.setText(f"Đã tạo xong câu {socau_daxuli}/{sum_toanbo} của đề số {j+1}.")
                                    
                                    self.progress_bar.setValue(i)                                    
                                    self.sleep(1)
                                self.progress_bar.setValue((j+1)*20)

                                #Trộn câu hỏi trong từng nhóm
                                if self.checkbox_shuffle_dangtoan.isChecked():

                                    if len(list_tracnghiem)>0:                           
                                    
                                        combined = list(zip(list_tracnghiem, list_tracnghiem_HDG, list_dapan_TN)) 
                                        random.shuffle(combined)
                                        # Tách lại thành các danh sách
                                        a_shuffled, b_shuffled, TN_shuffled= zip(*combined)
                                        list_tracnghiem = list(a_shuffled)
                                        list_tracnghiem_HDG = list(b_shuffled)
                                        list_dapan_TN=list(TN_shuffled)
                                        

                                    if len(list_dungsai)>0:

                                        combined = list(zip(list_dungsai, list_dungsai_HDG, list_dapan_TF)) 
                                        random.shuffle(combined)
                                        # Tách lại thành hai danh sách a và b
                                        a_shuffled, b_shuffled,TF_shuffled = zip(*combined)
                                        list_dungsai = list(a_shuffled)
                                        list_dungsai_HDG = list(b_shuffled)
                                        list_dapan_TF=list(TF_shuffled)

                                    if len(list_tuluan)>0:

                                        combined = list(zip(list_tuluan, list_tuluan_HDG, list_dapan_TL)) 
                                        random.shuffle(combined)
                                        # Tách lại thành hai danh sách a và b
                                        a_shuffled, b_shuffled, TL_shuffled = zip(*combined)
                                        list_tuluan = list(a_shuffled)
                                        list_tuluan_HDG = list(b_shuffled)
                                        list_dapan_TL=list(TL_shuffled)                            
     
                                                                              
                                if self.combo_taode.currentText() == "Tạo đề Word - Equation":
                                    if len(list_tracnghiem)>0:                                        
                                        list_noi_dung+=f"{{\\bf PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.}}\n"
                                        list_noi_dung_HDG+=f"{{\\bf PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.}}\n" 
                                        for chi_so in range(len(list_tracnghiem)):
                                            list_noi_dung+=f'{{\\bf Câu {chi_so+1}.}} {str(list_tracnghiem[chi_so])}\n'
                                            list_noi_dung_HDG+=f'{{\\bf Câu {chi_so+1}.}} {str(list_tracnghiem_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TN:
                                                list_dapan_word.append(phan_tu)

                                    if len(list_dungsai)>0:
                                        list_noi_dung+=f'\n {{\\bf PHẦN II. Câu trắc nghiệm đúng sai.}}\n'
                                        list_noi_dung_HDG+=f'\n {{\\bf PHẦN II. Câu trắc nghiệm đúng sai.}}\n' 
                                        for chi_so in range(len(list_dungsai)):
                                            list_noi_dung+=f'Câu {chi_so+1}. {str(list_dungsai[chi_so])}\n'
                                            list_noi_dung_HDG+=f'Câu {chi_so+1}. {str(list_dungsai_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TF:
                                                list_dapan_word.append(phan_tu)

                                    if len(list_tuluan)>0:
                                        list_noi_dung+=f'\n {{\\bf PHẦN III. Câu trắc nghiệm trả lời ngắn.}}\n'
                                        list_noi_dung_HDG+=f'\n {{\\bf PHẦN III. Câu trắc nghiệm trả lời ngắn.}}\n'
                                        for chi_so in range(len(list_tuluan)):
                                            list_noi_dung+=f'{{\\bf Câu {chi_so+1}.}} {str(list_tuluan[chi_so])}\n'
                                            list_noi_dung_HDG+=f'{{\\bf Câu {chi_so+1}.}} {str(list_tuluan_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TL:
                                                list_dapan_word.append(phan_tu)

                                    list_noi_dung+=f"{{\\bf -----HẾT-----}} \n"
                                    list_noi_dung_HDG+=f"{{\\bf -----HẾT-----}} \n"
                                    list_noi_dung=list_noi_dung.replace(f"\n",f"\n\n")
                                    list_noi_dung_HDG=list_noi_dung_HDG.replace(f"\n",f"\n\n")

                                    self.export_msword(name_thu_muc, name_de, list_noi_dung, list_noi_dung_HDG)                                    
                                    

                                if self.combo_taode.currentText() == "Tạo đề Word - MathType":
                                    if len(list_tracnghiem)>0:                                        
                                        list_noi_dung+=f"PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.\n"
                                        list_noi_dung_HDG+=f"PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.\n" 
                                        for chi_so in range(len(list_tracnghiem)):
                                            list_noi_dung+=f'Câu {chi_so+1}. {str(list_tracnghiem[chi_so])}\n'
                                            list_noi_dung_HDG+=f'Câu {chi_so+1}. {str(list_tracnghiem_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TN:
                                                list_dapan_word.append(phan_tu)

                                    if len(list_dungsai)>0:
                                        list_noi_dung+=f'\n PHẦN II. Câu trắc nghiệm đúng sai.\n'
                                        list_noi_dung_HDG+=f'\n PHẦN II. Câu trắc nghiệm đúng sai.\n' 
                                        for chi_so in range(len(list_dungsai)):
                                            list_noi_dung+=f'Câu {chi_so+1}. {str(list_dungsai[chi_so])}\n'
                                            list_noi_dung_HDG+=f'Câu {chi_so+1}. {str(list_dungsai_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TF:
                                                list_dapan_word.append(phan_tu)

                                    if len(list_tuluan)>0:
                                        list_noi_dung+=f'\n PHẦN III. Câu trắc nghiệm trả lời ngắn.\n'
                                        list_noi_dung_HDG+=f'\n PHẦN III. Câu trắc nghiệm trả lời ngắn.\n'
                                        for chi_so in range(len(list_tuluan)):
                                            list_noi_dung+=f' Câu {chi_so+1}. {str(list_tuluan[chi_so])}\n'
                                            list_noi_dung_HDG+=f'Câu {chi_so+1}. {str(list_tuluan_HDG[chi_so])}\n'

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TL:
                                                list_dapan_word.append(phan_tu)

                                    list_noi_dung+=f" -----HẾT-----\n"
                                    list_noi_dung_HDG+=f"-----HẾT-----\n"
                                    list_noi_dung=list_noi_dung.replace(f"\n",f"\n\n")
                                    list_noi_dung_HDG=list_noi_dung_HDG.replace(f"\n",f"\n\n")

                                    list_tonghop+=f"{list_noi_dung}"
                                    list_tonghop_HDG+=f"{list_noi_dung_HDG}"
                                    self.text_taode.append(list_noi_dung)
                                    self.text_taode_HDG.append(list_noi_dung_HDG)                                    
                                    self.export_msword_latex(name_thu_muc, name_de)

                                if self.combo_taode.currentText() == "Tạo đề Latex - PDF":
                                    if len(list_tracnghiem)>0:
                                        ghep_tracnghiem='\n'.join(list_tracnghiem)
                                        list_noi_dung+=(f"{{\\bf PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.}}\n"
                                        f'\\setcounter{{ex}}{{0}}\n'
                                        f"\\Opensolutionfile{{ans}}[ans/ans{name_de}-1]\n"                                        
                                        f"{ghep_tracnghiem}\n"
                                        f"\\Closesolutionfile{{ans}}\n")

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TN:
                                                list_dapan_word.append(phan_tu)

                                    if len(list_dungsai)>0:
                                        ghep_dungsai='\n'.join(list_dungsai)
                                        list_noi_dung+=(f'{{\\bf PHẦN II. Câu trắc nghiệm đúng sai.}}\n'
                                        f'\\setcounter{{ex}}{{0}}\n'
                                        f"\\Opensolutionfile{{ans}}[ans/ans{name_de}-2]\n"                                        
                                        f'{ghep_dungsai}\n'
                                        f"\\Closesolutionfile{{ans}}\n")

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TF:
                                                list_dapan_word.append(phan_tu)
                                        

                                    if len(list_tuluan)>0:
                                        ghep_tuluan='\n'.join(list_tuluan)
                                        list_noi_dung+=(f'{{\\bf PHẦN III. Câu trắc nghiệm trả lời ngắn.}}\n'
                                        f'\\setcounter{{ex}}{{0}}\n'
                                        f"\\Opensolutionfile{{ans}}[ans/ans{name_de}-3]\n"                                        
                                        f'{ghep_tuluan}\n'
                                        f"\\Closesolutionfile{{ans}}\n")

                                        #Lấy danh sách đáp án word:
                                        for phan_tu in list_dapan_TL:
                                                list_dapan_word.append(phan_tu)
                                    

                                    list_noi_dung+=(f"\n \\begin{{center}}\n-----HẾT-----\n\\end{{center}}\n"
                                    f"\n %\\newpage \n"
                                    f"%\\begin{{center}}\n"
                                    f"%{{\\bf BẢNG ĐÁP ÁN MÃ ĐỀ {ma_de} }}\n"
                                    f"%\\end{{center}}\n")

                                    code_bang_dap_an+=(f"\\begin{{center}}\n{{\\bf BẢNG ĐÁP ÁN MÃ ĐỀ {ma_de} }}\n\\end{{center}}\n")

                                    if len(list_tracnghiem)>0:
                                        list_noi_dung+=(
                                    f"%{{\\bf Phần 1 }}\n"
                                    f"% \\inputansbox{{6}}{{ans{name_de}-1}}\n")

                                        code_bang_dap_an+=(f"{{\\bf Phần 1 }}\n"
                                    f"\\inputansbox{{6}}{{ans{name_de}-1}}\n")

                                    if len(list_dungsai)>0:
                                        list_noi_dung+=(f"%{{\\bf Phần 2 }}\n"
                                    f"% \\inputansbox{{2}}{{ans{name_de}-2}}\n")

                                        code_bang_dap_an+=(f"{{\\bf Phần 2 }}\n"
                                    f"\\inputansbox{{2}}{{ans{name_de}-2}}\n")

                                    if len(list_tuluan)>0:

                                        list_noi_dung+=(f"%{{\\bf Phần 3 }}\n"
                                    f"% \\inputansbox{{6}}{{ans{name_de}-3}}\n"
                                    f"\\newpage \n")

                                        code_bang_dap_an+=(f"{{\\bf Phần 3 }}\n"
                                    f"\\inputansbox{{6}}{{ans{name_de}-3}}\n")

                                    list_tonghop+=f"{list_noi_dung}"                                

                                                                       
                                    self.text_taode.append(list_noi_dung)
                                    self.open_latex_pdf(name_thu_muc, name_de, list_noi_dung)

                                if self.combo_taode.currentText() == "Tạo code Latex":
                                    if len(list_tracnghiem)>0:
                                        ghep_tracnghiem='\n'.join(list_tracnghiem)
                                        list_noi_dung+= f"{ghep_tracnghiem}"

                                    if len(list_dungsai)>0:
                                        ghep_dungsai='\n'.join(list_dungsai)
                                        list_noi_dung+=f'{ghep_dungsai}'                  

                                    if len(list_tuluan)>0:
                                        ghep_tuluan='\n'.join(list_tuluan)
                                        list_noi_dung+=f'{ghep_tuluan}'                                       

                                    list_tonghop+=f"{list_noi_dung}"

                                    self.text_taode.append(list_noi_dung)

                            #end

                            self.timer.stop()
                            self.progress_bar.setValue(100)
                            self.label_dangxuli.setText("")
                            self.label_nhapmade.setText("")                           



                            #Mở thư mục chứa file
                            if self.combo_taode.currentText() == "Tạo đề Latex - PDF":                           
                                
                                # current_directory = os.path.dirname(os.path.abspath(__file__))
                                # doc_folder_path = os.path.join(current_directory, 'LATEX')
                                # new_folder_path = os.path.join(doc_folder_path, name_thu_muc)                               
                                self.tao_bang_dap_an_latex(name_thu_muc,code_bang_dap_an)

                                #self.tao_tnmaker_word(name_thu_muc, list_ma_de, list_dapan_word)                            

                                self.tao_tnmaker_latex(name_thu_muc, list_ma_de,len(list_tracnghiem),len(list_dungsai),len(list_tuluan))
                                name_thu_muc=name_thu_muc.replace("/","\\")
                                subprocess.Popen(['explorer', name_thu_muc])

                            if self.combo_taode.currentText() in ["Tạo đề Word - Equation", "Tạo đề Word - MathType"]:                                
                                self.tao_tnmaker_word(name_thu_muc, list_ma_de, list_dapan_word)                   
                               
                                # current_directory = os.path.dirname(os.path.abspath(__file__))
                                # doc_folder_path = os.path.join(current_directory, 'DOC')
                                # new_folder_path = os.path.join(doc_folder_path, name_thu_muc)
                                name_thu_muc=name_thu_muc.replace("/","\\")
                                subprocess.Popen(['explorer', name_thu_muc])

                            if self.combo_taode.currentText() == "Tạo code Latex":
                                self.text_taode.clear()
                                self.text_taode.append(list_tonghop)
                                pyperclip.copy(list_tonghop)
                                              
                           
                        else:
                            show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo bản quyền', 'Thầy(cô) cần đăng kí bản quyền để sử dụng chức năng này.\nHãy ủng hộ tác giả để tác giả có động lực hỗ trợ cho thầy (cô) nhé.\nVui lòng vào tab Bản Quyền, copy Mã máy gửi cho tác giả để được cung cấp key sử dụng.')
                            show_msg_box.exec_()

                except Exception as e:
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                        show_msg_box.exec_()
                return

        def updateProgressBar(self):
                # Cập nhật giá trị của progress bar
                current_value = self.progress_bar.value()
                self.progress_bar.setValue(current_value + 1)
                return

        def sleep(self, seconds):
                loop = QEventLoop()
                QTimer.singleShot(int(seconds * 1000), loop.quit)
                loop.exec_()
                return

        #Copy đề đã được tạo
        def copy_noidung(self):
                text = self.text_taode.toPlainText()
                pyperclip.copy(text)
                show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo bản quyền', 'Đã sao chép nội dung đề bài vào bộ nhớ.')
                show_msg_box.exec_()
                return

        #Xuất file word đề đã được tạo
        def export_msword(self, name_thu_muc, name_de, list_noi_dung, list_noi_dung_HDG):
                try:                                 
                                        
                    new_folder_path =  name_thu_muc 

                    # Tạo file đề bài
                    file_path = os.path.join(new_folder_path, f'De_{name_de}.tex')
                    with open(file_path, 'w',encoding="utf-8") as file:
                            file.write(list_noi_dung)
                    
                    os.chdir(os.path.dirname(file_path))
                    
                    latex_command = f"pandoc De_{name_de}.tex -o De_{name_de}.docx"
                    subprocess.run(latex_command, shell=True)

                    file_path = os.path.join(new_folder_path, f'De_{name_de}_loigiai.tex')
                    with open(file_path, 'w',encoding="utf-8") as file:
                            file.write(list_noi_dung_HDG)
                    
                    os.chdir(os.path.dirname(file_path))
                    
                    latex_command = f"pandoc De_{name_de}_loigiai.tex -o De_{name_de}_loigiai.docx"
                    subprocess.run(latex_command, shell=True)

                    # Tìm tất cả các file .docx trong thư mục
                    tex_files = glob.glob(os.path.join(new_folder_path, '*.tex'))
                    for file_path in tex_files:
                        os.remove(file_path)

                    folder_hinh = my_module.get_folder_hinh()
                    file_path = os.path.join(new_folder_path, f'De_{name_de}.docx')             
                    doc = Document(file_path)
                    my_module.find_and_insert_image(doc)               
                    doc.save(file_path)
                    
                    file_path = os.path.join(new_folder_path, f'De_{name_de}_loigiai.docx')             
                    doc = Document(file_path)                                        
                    
                    my_module.find_and_insert_image(doc)               
                    doc.save(file_path)

                    name_matran=f"{new_folder_path}\\Matran.xlsx"
                    self.save_matran(name_matran)

                    #Xuất file excel chứa bảng đáp án

               
                    
                except Exception as e:
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                        show_msg_box.exec_()
                return

        #Xuất file word đề đã được tạo
        def export_msword_latex(self, name_thu_muc, name_de):
            try:
                #Tạo thư mục với tên là name,
                current_directory = os.path.dirname(os.path.abspath(__file__))
                doc_folder_path = os.path.join(current_directory, 'DOC')
                new_folder_path = os.path.join(doc_folder_path, name_thu_muc)
                if not os.path.exists(new_folder_path):
                        os.makedirs(new_folder_path)         
               
                doc = Document()                
                table = doc.add_table(rows=1, cols=2)                
                table.style = 'Table Grid'
                
                cell = table.cell(0, 0)
                cell.text = f"{self.ten_sogd.toPlainText()} \n {self.ten_truong.toPlainText()}"
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)
                
                cell = table.cell(0, 1)
                cell.text = f"{self.ten_kythi.toPlainText()} \n  Môn học: {self.ten_monthi.toPlainText()} \n Thời gian làm bài: {self.ten_thoigian.toPlainText()} phút \n Mã đề: {name_de}"
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)

                paragraph="Họ tên HS:............................................................................\t Số báo danh:......................"
                doc.add_paragraph(paragraph) 

                #Thêm nội dung câu hỏi
                noidung = self.text_taode.toPlainText()        
                paragraphs = noidung.split('\n')
                for paragraph in paragraphs:
                    doc.add_paragraph(paragraph)                            

                
                file_path = os.path.join(new_folder_path , f"De_{name_de}.docx")
                doc.save(file_path)
                
                folder_hinh = my_module.get_folder_hinh()        
                
                # Tìm kiếm cụm từ và chèn ảnh
                doc = Document(file_path)
                my_module.find_and_insert_image(doc)

                
                doc.save(file_path)
                #Xuất file lời giải
                doc = Document()
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'

                cell = table.cell(0, 0)
                cell.text = f"{self.ten_sogd.toPlainText()} \n {self.ten_truong.toPlainText()}"
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)

                cell = table.cell(0, 1)
                cell.text = f"{self.ten_kythi.toPlainText()} \n  Môn học: {self.ten_monthi.toPlainText()} \n Thời gian làm bài: {self.ten_thoigian.toPlainText()} phút \n Mã đề: {name_de}"
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(12)

                noidung = self.text_taode_HDG.toPlainText()        
                paragraphs = noidung.split('\n')
                for paragraph in paragraphs:
                    doc.add_paragraph(paragraph)                          

                file_path = os.path.join(new_folder_path , f"De_{name_de}_loigiai.docx")
                doc.save(file_path)

                folder_hinh = my_module.get_folder_hinh()  
                doc = Document(file_path)
                my_module.find_and_insert_image(doc)
                doc.save(file_path)
                name_matran=f"{new_folder_path}\\Matran.xlsx"
                self.save_matran(name_matran)
            except Exception as e:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                    show_msg_box.exec_()
            return

        #Save file word đề đã được tạo
        def save_file(self):
            try:
   
                    # Tạo một đối tượng Document
                    doc = Document()
                    noidung = self.text_taode.toPlainText()        
                    paragraphs = noidung.split('\n')
                    for paragraph in paragraphs:
                        doc.add_paragraph(paragraph) 
                     #Lấy thư mục để lưu                    

                    options = QtWidgets.QFileDialog.Options()
                    options |= QtWidgets.QFileDialog.DontUseNativeDialog
                    file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self.tab_taode, "Lưu đề bài", "", "Text Files (*.docx);;All Files (*)", options=options) 
                    file_path=f"{file_path}.docx"           
        
                    doc.save(file_path)
              
                    folder_hinh = my_module.get_folder_hinh()        
                    
                    # Tìm kiếm cụm từ và chèn ảnh
                    doc = Document(file_path)
                    my_module.find_and_insert_image(doc)

                    # Lưu tài liệu Word đã chỉnh sửa
                    doc.save(file_path)       

                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', f'Đã xuất câu hỏi và lưu tại {file_path}!')
                    show_msg_box.exec_()
            except Exception as e:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                    show_msg_box.exec_()
            return

        
        #Tao ma tran
        def get_chi_muc(self, item, column):
                if item.checkState(0) == Qt.CheckState.Checked:
                    name = item.text(0)  # Lấy nội dung cột 0 (tên)
                    if name[0]=="[":
                        row_count = self.tableWidget.rowCount()
                        self.tableWidget.setRowCount(row_count + 1)
                        self.tableWidget.setItem(row_count, 0, QTableWidgetItem(f"{name}"))

                        #Tạo số câu là 1
                        item = QTableWidgetItem(f"1")
                        item.setTextAlignment(Qt.AlignCenter)
                        self.tableWidget.setItem(row_count, 3, item)
                

                        #Lấy loại câu và mức độ của câu trăc nghiệm
                        if name[14]=="-":
                            item = QTableWidgetItem(f"TN")
                            item.setTextAlignment(Qt.AlignCenter)
                            self.tableWidget.setItem(row_count, 1, item)                        
                            if name[16]=="1":
                                item = QTableWidgetItem(f"NB")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)                                                      
                            if name[16]=="2":
                                item = QTableWidgetItem(f"TH")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)                            
                            if name[16]=="3":
                                item = QTableWidgetItem(f"VDT")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)                           
                            if name[16]=="4":
                                item = QTableWidgetItem(f"VDC")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)                      
                                

                        #Lấy loại câu và mức độ của câu đúng sai
                        if name[15:17]=="TF":
                            item = QTableWidgetItem(f"Đ-S")
                            item.setTextAlignment(Qt.AlignCenter)
                            self.tableWidget.setItem(row_count, 1, item)
                            
                            if name[19]=="1":
                                item = QTableWidgetItem(f"NB")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="2":
                                item = QTableWidgetItem(f"TH")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="3":
                                item = QTableWidgetItem(f"VDT")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="4":
                                item = QTableWidgetItem(f"VDC")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            

                        #Lấy loại câu và mức độ của câu đúng sai
                        if name[15:17]=="TL":
                            item = QTableWidgetItem(f"TL")
                            item.setTextAlignment(Qt.AlignCenter)
                            self.tableWidget.setItem(row_count, 1, item)
                            
                            if name[19]=="1":
                                item = QTableWidgetItem(f"NB")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="2":
                                item = QTableWidgetItem(f"TH")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="3":
                                item = QTableWidgetItem(f"VDT")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)
                            if name[19]=="4":
                                item = QTableWidgetItem(f"VDC")
                                item.setTextAlignment(Qt.AlignCenter)
                                self.tableWidget.setItem(row_count, 2, item)


                    for row in range(self.tableWidget.rowCount()):
                        item = self.tableWidget.item(row, 0)
                        item.setBackground(QtGui.QColor(251, 243, 221))               
                    self.thongke()                 
                return
                            

        #Hiện tooltip khi rê chuột đến chỉ mục
        def showToolTip(self, item, column):
                # Hiển thị tooltip khi rê chuột lên một mục
                tooltip = item.toolTip(column)
                if tooltip:
                    QToolTip.showText(tooltip)
                    #QToolTip.showText(self.mapToGlobal(item.rect().center()), tooltip, self)
                return 



        def clear_dangtoan(self):                
                self.tableWidget.setRowCount(0)
                self.thongke()
                return               



        def btn_tron_dangtoan_click(self):
                # Tách các dòng văn bản thành một danh sách
                lines = self.text_dangtoan.toPlainText().split('\n')
                
                # Xáo trộn ngẫu nhiên danh sách
                random.shuffle(lines)
                
                # Ghép lại các dòng đã xáo trộn thành một chuỗi mới
                shuffled_text = '\n'.join(lines)
                
                # Hiển thị chuỗi đã xáo trộn trong QTextEdit
                self.text_dangtoan.setPlainText(shuffled_text)
                return 

        #Luu ma tran
        def luu_matran(self):
            try:
                options = QtWidgets.QFileDialog.Options()
                options |= QtWidgets.QFileDialog.DontUseNativeDialog
                fileName, _ = QtWidgets.QFileDialog.getSaveFileName(self.tab_taode, "Lưu ma trận", "", "Text Files (*.xlsx);;All Files (*)", options=options)
                fileName = f"{fileName}.xlsx"
                if fileName:
                    # Tạo mới một workbook
                    wb = Workbook()
                    ws = wb.active
                    ws.column_dimensions['A'].width = 80
                    ws.cell(row=1,column=1, value=f"Dạng Toán")
                    ws.cell(row=1,column=2, value=f"Loại Câu")
                    ws.cell(row=1,column=3, value=f"Mức Độ")
                    ws.cell(row=1,column=4, value=f"Số Câu") 

                    #Điền đáp án cho từng dòng
                    row_count = self.tableWidget.rowCount()
                    for i in range(row_count):
                        dang_toan= self.tableWidget.item(i, 0).text()
                        loai_cau= self.tableWidget.item(i, 1).text()
                        muc_do = self.tableWidget.item(i, 2).text()
                        so_cau = int(self.tableWidget.item(i, 3).text())
                        if all([dang_toan, loai_cau, muc_do, so_cau]):   
                            ws.cell(row=i+2, column=1, value=dang_toan)
                            ws.cell(row=i+2, column=2, value=loai_cau)
                            ws.cell(row=i+2, column=3, value=muc_do)
                            ws.cell(row=i+2, column=4, value=so_cau)                                  
                    wb.save(f"{fileName}")
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', f'Đã lưu ma trận tại {fileName}')
                    show_msg_box.exec_()
            except Exception as e:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                    show_msg_box.exec_()
            return

        def save_matran(self,fileName):
            fileName = f"{fileName}"
            if fileName:
                # Tạo mới một workbook
                wb = Workbook()
                ws = wb.active
                ws.column_dimensions['A'].width = 80
                ws.cell(row=1,column=1, value=f"Dạng Toán")
                ws.cell(row=1,column=2, value=f"Loại Câu")
                ws.cell(row=1,column=3, value=f"Mức Độ")
                ws.cell(row=1,column=4, value=f"Số Câu")
                # Lấy số hàng và số cột của tableWidget
                num_rows = self.tableWidget.rowCount()
                num_cols = self.tableWidget.columnCount()

                # Lặp qua từng hàng và cột của tableWidget để lưu dữ liệu vào file Excel
                for row in range(num_rows):
                    for col in range(num_cols):
                        # Lấy dữ liệu từ tableWidget
                        cell_value = self.tableWidget.item(row, col).text() if self.tableWidget.item(row, col) else ""
                        # Ghi dữ liệu vào cell tương ứng trong worksheet
                        ws.cell(row=row + 1, column=col + 1, value=cell_value)

                wb.save(f"{fileName}")
            return

        def load_matran(self):
            try:
                options = QtWidgets.QFileDialog.Options()
                options |= QtWidgets.QFileDialog.DontUseNativeDialog
                fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self.tab_taode, "Chọn file chứa dạng toán", "", "Excel Files (*.xlsx);;All Files (*)", options=options)

                if fileName:                    
                    wb = load_workbook(fileName)
                    sheet = wb['Sheet']
                    # Tìm số dòng cuối cùng có dữ liệu trong cột A
                    rows_count = 0                    
                    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=1):
                        cell_value = row[0].value
                        # Kiểm tra ô có giá trị không rỗng
                        if cell_value is not None:
                            rows_count += 1

                    self.tableWidget.setRowCount(rows_count-1)  
                    
                    for i in range(1,rows_count):
                        cell_value = sheet.cell(row=i+1, column=1).value
                        dang_toan = sheet.cell(row=i+1, column=1).value                            
                        item = QTableWidgetItem(dang_toan)                            
                        self.tableWidget.setItem(i-1, 0, item)

                        loai_cau = sheet.cell(row=i+1, column=2).value
                        item = QTableWidgetItem(loai_cau)
                        item.setTextAlignment(Qt.AlignCenter)
                        self.tableWidget.setItem(i-1, 1, item)

                        muc_do = sheet.cell(row=i+1, column=3).value
                        item = QTableWidgetItem(muc_do)
                        item.setTextAlignment(Qt.AlignCenter)
                        self.tableWidget.setItem(i-1, 2, item)

                        so_cau = str(sheet.cell(row=i+1, column=4).value)
                        item = QTableWidgetItem(so_cau)
                        item.setTextAlignment(Qt.AlignCenter)
                        self.tableWidget.setItem(i-1, 3, item)
            except Exception as e:
                show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                show_msg_box.exec_()


            return 
        #vedothi
        def btn_vedothi_click(self):
                try:
                        self.text_hamso.clear()
                        self.label_showpic.clear()
                        x = symbols('x')
                        #Vẽ đồ thị bậc nhất
                        if self.combo_dothihamso.currentText() == "Hàm số bậc nhất":
                            # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax+b cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok:                
                                # Xử lý dữ liệu nhập vào
                                
                                if text.count(",") == 1:
                                    a,b=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    f=a*x+b                                                      
                                    self.text_hamso.append(f" y = {latex(f)} ")    
                                  
                                    code_dothi= my_module.codelatex_dothi_bac_1(a,b)                    
                                    my_module.pdftoimage(code_dothi)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)
                        #self.label_tab_bang_so_lieu_getbang.setPixmap(pixmap)
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show() 
                                    self.tab_dothi_textcode.setText(code_dothi)

                        if self.combo_dothihamso.currentText() == "Hàm số bậc 2":
                             # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^2+bx+c cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok:                
                                # Xử lý dữ liệu nhập vào                                
                                
                                if text.count(",") == 2:
                                    a,b,c=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c) 
                                    f=a*x**2+b*x+c                    
                                    self.text_hamso.append(f" y = {latex(f)} ")           
                                  
                                    #Vẽ hình
                                    x = -b/(2*a)
                                    y = (4*a*c-b**2)/(4*a)                                                                                                         
                                                                                                
                                    code_dothi = my_module.codelatex_dothi_bac_2(a,b,c)                                                   
                                    my_module.pdftoimage(code_dothi)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)
                        #self.label_tab_bang_so_lieu_getbang.setPixmap(pixmap)
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show()
                                    self.tab_dothi_textcode.setText(code_dothi)
                        if self.combo_dothihamso.currentText() == "Hàm số bậc ba":
                             # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^3+bx^2+cx+d cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok and text!= "":
                                # Xử lý dữ liệu nhập vào
                                
                                if text.count(",") == 3:
                                    a,b,c,d=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    f=a*x**3+b*x**2+c*x+d                    
                                    self.text_hamso.append(f" y = {latex(f)} ")          
                                  
                                    #Vẽ hình         
                                    code_dothi= my_module.codelatex_dothi_bac_3(a,b,c,d)        
                                    my_module.pdftoimage(code_dothi)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show()
                                    self.tab_dothi_textcode.setText(code_dothi)
                                
                        if self.combo_dothihamso.currentText() == "Hàm số bậc bốn trùng phương":
                             # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^4+bx^2+c cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok and text!= "":
                                # Xử lý dữ liệu nhập vào
                                if text.count(",") == 2:
                                    a,b,c=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)                                    
                                    f=a*x**4+b*x**2+c                    
                                    self.text_hamso.append(f" y = {latex(f)} ")                      
                                  
                                    #Vẽ hình         
                                    code_dothi= my_module.codelatex_dothi_bac_4(a,b,c)        
                                    my_module.pdftoimage(code_dothi)  

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"   

                                    pixmap = QPixmap(fname)
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show()
                                    self.tab_dothi_textcode.setText(code_dothi)
                                    
                        if self.combo_dothihamso.currentText() == "Hàm số bậc 1/bậc 1":
                             # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=(ax+b)/(cx+d) cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok and text!= "":
                                # Xử lý dữ liệu nhập vào                                
                                if text.count(",") == 3:
                                    a,b,c,d=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    f=(a*x+b)/(c*x+d)                    
                                    self.text_hamso.append(f" y = {latex(f)} ")            
                                  
                                    #Vẽ hình       

                                    code_dothi= my_module.codelatex_dothi_phanthuc_bac1(a,b,c,d)        
                                    my_module.pdftoimage(code_dothi)
                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"
                                    pixmap = QPixmap(fname)              
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show()
                                    self.tab_dothi_textcode.setText(code_dothi)

                        if self.combo_dothihamso.currentText() == "Hàm số bậc 2/bậc 1":
                             # Hiển thị hộp thoại nhập liệu
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=(ax^2+bx+c)/(dx+e) cách nhau bằng dấu ','")
                            #Kiểm tra xem người dùng đã nhấp vào OK hay không
                            if ok and text!="":
                                # Xử lý dữ liệu nhập vào
                                if text.count(",") == 4:
                                    a,b,c,d,e=text.split(',')                                
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    e=eval(e)
                                    f=(a*x**2+b*x+c)/(d*x+e)                    
                                    self.text_hamso.append(f" y = {latex(f)} ")      
                                    code_dothi= my_module.code_dothi_phanthuc_bac2(a,b,c,d,e)        
                                    my_module.pdftoimage(code_dothi)
                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"
                                    pixmap = QPixmap(fname)              
                                    new_width = 500  # Đặt chiều rộng mới
                                    new_height = 500  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                                    self.label_showpic.setPixmap(scaled_pixmap)
                                    self.label_showpic.show()
                                    self.tab_dothi_textcode.setText(code_dothi)
                except Exception as e:
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                        show_msg_box.exec_()
                return
                                              
        def btn_copydothi_click(self):
                folder_path = my_module.get_folder_hinh()       
                image=QImage(f"{folder_path}\\hinhve.png")
                QApplication.clipboard().setImage(image)
                return

        #Code biên dịch vẽ lại đồ thị
        def tab_dothi_velaidothi(self):
                code = self.tab_dothi_textcode.toPlainText()
                if code != "":
                    file_name=my_module.pdftoimage(code)

                    #Hiện thị hình ảnh
                    folder_path = my_module.get_folder_hinh()
                    fname = f"{folder_path}\\hinhve.png"
                    pixmap = QPixmap(fname)
                    new_width = 500  # Đặt chiều rộng mới
                    new_height = 500  # Đặt chiều cao mới
                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                    self.label_showpic.setPixmap(scaled_pixmap)
                    self.label_showpic.show()

                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Công việc đã hoàn thành!')
                    show_msg_box.exec_()
                else:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Không tìm thấy code để biên dịch!')
                    show_msg_box.exec_()
                return
        #Code biên dịch vẽ lại hình
        def tab_vehinh_velaihinh(self):
                code = self.tab_vehinh_textcode.toPlainText()
                if code != "":
                    file_name=my_module.pdftoimage(code)

                    #Hiện thị hình ảnh
                    folder_path = my_module.get_folder_hinh()
                    fname = f"{folder_path}\\hinhve.png"
                    pixmap = QPixmap(fname)
                    new_width = 500  # Đặt chiều rộng mới
                    new_height = 500  # Đặt chiều cao mới
                    scaled_pixmap = pixmap.scaled(new_width, new_height)      
                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                    self.label_showpic_vehinh.show()

                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Công việc đã hoàn thành!')
                    show_msg_box.exec_()
                else:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Không tìm thấy code để biên dịch!')
                    show_msg_box.exec_()
                return                     
                
                  
        def btn_copyimage_tab_run_latex_click(self):
                folder_path = my_module.get_folder_hinh()
                image=QImage(f"{folder_path}\\hinhve.png")
                QApplication.clipboard().setImage(image)
                return

        #Code mẫu số liệu ghép nhóm
        #Hàm xử lí tạo bảng số liệu
        def tao_bang_so_lieu(self):
                ten_nhom=self.tab_bang_so_lieu_tennhom.toPlainText()
                so_nhom =self.tab_bang_so_lieu_sonhom.value() + 1
                ten_tan_so=self.tab_bang_so_lieu_tentanso.toPlainText()
                so_bat_dau =int(self.tab_bang_so_lieu_giatribatdau.toPlainText())
                khoang_cach =float(self.tab_bang_so_lieu_khoangcach.toPlainText())
                tan_so_min =int(self.tab_bang_so_lieu_tansomin.toPlainText())
                tan_so_max =int(self.tab_bang_so_lieu_tansomax.toPlainText())
                #Tạo mẫu số liệu kết quả trả về là: gia_tri,khoang_gt,list_tan_so, tan_so
                gia_tri,list_khoang_gia_tri,list_tan_so,tan_so=my_module.tao_mau_ghep_nhom(so_nhom,so_bat_dau,khoang_cach,tan_so_min,tan_so_max)

                #Tạo code vẽ bảng:codelatex_bang_ghep_nhom(ten_nhom,list_khoang_gia_tri,ten_tan_so,list_tan_so):
                code = my_module.codelatex_bang_ghep_nhom(ten_nhom,list_khoang_gia_tri,ten_tan_so,list_tan_so)
                file_name=my_module.pdftoimage(code)

                #Hiện thị hình ảnh
                folder_path = my_module.get_folder_hinh()
                fname = f"{folder_path}\\hinhve.png"
                pixmap = QPixmap(fname)
                new_width = 600  # Đặt chiều rộng mới
                new_height = 80  # Đặt chiều cao mới
                scaled_pixmap = pixmap.scaled(new_width, new_height)
                #self.label_tab_bang_so_lieu_getbang.setPixmap(pixmap)
                self.label_tab_bang_so_lieu_getbang.setPixmap(scaled_pixmap)
                self.label_tab_bang_so_lieu_getbang.show()

                #Tạo code latex
                self.tab_bang_so_lieu_code.setText(code)
                return


        def tao_bang_so_lieu_runcode(self):
                #Tạo code vẽ bảng:codelatex_bang_ghep_nhom(ten_nhom,list_khoang_gia_tri,ten_tan_so,list_tan_so):
                code = self.tab_bang_so_lieu_code.toPlainText()
                if code != "":
                    file_name=my_module.pdftoimage(code)

                    #Hiện thị hình ảnh
                    folder_path = my_module.get_folder_hinh()
                    fname = f"{folder_path}\\hinhve.png"
                    pixmap = QPixmap(fname)
                    new_width = 600  # Đặt chiều rộng mới
                    new_height = 80  # Đặt chiều cao mới
                    scaled_pixmap = pixmap.scaled(new_width, new_height)
                    #self.label_tab_bang_so_lieu_getbang.setPixmap(pixmap)
                    self.label_tab_bang_so_lieu_getbang.setPixmap(scaled_pixmap)
                    self.label_tab_bang_so_lieu_getbang.show()

                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Công việc đã hoàn thành!')
                    show_msg_box.exec_()
                else:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Không tìm thấy code để biên dịch!')
                    show_msg_box.exec_()
                return

        def copy_bang_so_lieu(self):
                folder_path = my_module.get_folder_hinh()
                image=QImage(f"{folder_path}\\hinhve.png")
                QApplication.clipboard().setImage(image)
                return

        def copy_ma_may(self):
                text = self.tab_ban_quyen_ma_may.toPlainText()
                pyperclip.copy(text)
                show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo', 'Đã copy mã máy vào clipboard. Vui lòng gửi mã này cho tác giả để được cung cấp key sử dụng.')
                show_msg_box.exec_()
                return

        def dang_ki_ban_quyen(self):
                key = self.tab_ban_quyen_ma_dangki.toPlainText()
                if license.check_banquyen(key):            
                    license.create_registry_key(key)
                    thoi_gian = f"Thời gian sử dụng còn lại là {license.so_ngay_banquyen()} ngày."
                    self.label_thoigian.setText(thoi_gian)        
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Đăng kí thành công', 'Đăng kí bản quyền thành công. Chúc thầy cô soạn giảng nhanh hơn với iMath.')
                    show_msg_box.exec_()
                else:
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Đăng kí thất bại', 'Key cung cấp không hợp lệ. Thầy cô vui lòng liên hệ tác gia để được hỗ trợ.')
                    show_msg_box.exec_()
                return

        #vebbt
        def btn_veBBT_click(self):
                try:
                        self.text_hamso_BBT.clear()
                        self.label_showpic_BBT.clear()
                        x = symbols('x')
                        #Vẽ bảng biến thiên bậc 2

                        if self.combo_dothihamso_BBT.currentText() == "Hàm số bậc 2":
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^2+bx+c cách nhau bằng dấu ','")
                            if ok and text!= "":
                                if text.count(",") == 2:
                                    a,b,c=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    f=a*x**2+b*x+c                    
                                    self.text_hamso_BBT.append(f" y = {latex(f)} ")         
                                    code_BBT= my_module.codelatex_bbt_bac2(a,b,c)                    
                                    my_module.pdftoimage(code_BBT)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_BBT.setPixmap(scaled_pixmap)
                                    self.label_showpic_BBT.show()
                                    self.tab_bbt_textcode.setText(code_BBT)

                        if self.combo_dothihamso_BBT.currentText() == "Hàm số bậc ba":
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^3+bx^2+cx+d cách nhau bằng dấu ','")
                            if ok and text!= "":
                                if text.count(",") == 3:
                                    a,b,c,d=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    f=a*x**3+b*x**2+c*x+d                    
                                    self.text_hamso_BBT.append(f" y = {latex(f)} ")

                                    #Vẽ BBT                         
                                    code_BBT= my_module.codelatex_bbt_bac3(a,b,c,d)                    
                                    my_module.pdftoimage(code_BBT)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_BBT.setPixmap(scaled_pixmap)
                                    self.label_showpic_BBT.show()
                                    self.tab_bbt_textcode.setText(code_BBT)
                        
                        if self.combo_dothihamso_BBT.currentText() == "Hàm số bậc bốn trùng phương":                            
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=ax^4+bx^2+c cách nhau bằng dấu ','")
                            if ok and text!= "":
                                if text.count(",") == 2:
                                    a,b,c =text.split(',')                                        
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    f=a*x**4+b*x**2+c                    
                                    self.text_hamso.append(f" y = {latex(f)} ")                                  
                                    code_BBT=my_module.codelatex_bbt_bac4(a,b,c)                                                                 
                                    my_module.pdftoimage(code_BBT)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_BBT.setPixmap(scaled_pixmap)
                                    self.label_showpic_BBT.show()
                                    self.tab_bbt_textcode.setText(code_BBT)
                        

                        if self.combo_dothihamso_BBT.currentText() == "Hàm số bậc 1/bậc 1":
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=(ax+b)/(cx+d) cách nhau bằng dấu ','")
                            if ok and text!= "":
                                if text.count(",") ==3:
                                    a,b,c,d=text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    f=(a*x+b)/(c*x+d)
                                    self.text_hamso_BBT.append(f" y = {latex(f)} ")                         
                                    code_BBT= my_module.codelatex_bbt_phanthucbac1(a,b,c,d)
                                    my_module.pdftoimage(code_BBT)

                                    folder_path = my_module.get_folder_hinh()
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_BBT.setPixmap(scaled_pixmap)
                                    self.label_showpic_BBT.show()
                                    self.tab_bbt_textcode.setText(code_BBT)

                        if self.combo_dothihamso_BBT.currentText() == "Hàm số bậc 2/bậc 1":
                            text, ok = QInputDialog.getText(None, "Nhập hệ số", "Nhập các hệ số của hàm số y=(ax^2+bx+c)/(dx+e) cách nhau bằng dấu ','")
                            if ok and text!= "":
                                if text.count(",") ==4:
                                    a,b,c,d,e =text.split(',')
                                    a=eval(a)
                                    b=eval(b)
                                    c=eval(c)
                                    d=eval(d)
                                    e=eval(e)
                                    f=(a*x**2+b*x+c)/(d*x+e)
                                    self.text_hamso_BBT.append(f" y = {latex(f)} ")                             
                                    code_BBT= my_module.codelatex_bbt_phanthucbac2(a,b,c,d,e)
                                    my_module.pdftoimage(code_BBT)
                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 600  # Đặt chiều rộng mới
                                    new_height = 300  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_BBT.setPixmap(scaled_pixmap)
                                    self.label_showpic_BBT.show()     
                 
                                    self.tab_bbt_textcode.setText(code_BBT)
                except Exception as e:
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                        show_msg_box.exec_() 
                return

        #Code vẽ bảng biến thiên
        def btn_vehinh_click(self):
                try:
                        self.text_hinhve.clear()
                        self.label_showpic_vehinh.clear()                      
                        
                        if self.combo_loaihinh.currentText() == "Hình chóp tam giác cạnh vuông góc":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C")
                            if ok and text!= "":
                                if text.count(",") == 3:
                                    s,a,b,c=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c}")         
                                    code= my_module.ve_hinhchop_tamgiac_canhvg(s,a,b,c)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp tam giác mặt vuông góc":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 3:
                                    s,a,b,c=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c}")         
                                    code= my_module.ve_hinhchop_tamgiac_matvg(s,a,b,c)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp tam giác đều":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 3:
                                    s,a,b,c=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c}")         
                                    code= my_module.ve_hinhchop_tamgiac_deu(s,a,b,c)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình tứ diện":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 3:
                                    s,a,b,c=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c}")         
                                    code= my_module.ve_hinh_tu_dien(s,a,b,c)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp tứ giác":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 4:
                                    s,a,b,c,d=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c},{d}")         
                                    code= my_module.ve_hinhchop_tugiac(s,a,b,c,d)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp đáy h.b.h cạnh vuông góc":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n  Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 4:
                                    s,a,b,c,d=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c},{d}")         
                                    code= my_module.ve_hinhchop_hbh_canhvg(s,a,b,c,d)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp đáy h.b.h mặt vuông góc":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n  Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 4:
                                    s,a,b,c,d=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c},{d}")         
                                    code= my_module.ve_hinhchop_hbh_matvg(s,a,b,c,d)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình chóp tứ giác đều":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh.\n Ví dụ: S,A,B,C,D")
                            if ok and text!= "":
                                if text.count(",") == 4:
                                    s,a,b,c,d=text.split(',')                                                    
                                    self.text_hinhve.append(f"{s},{a},{b},{c},{d}")         
                                    code= my_module.ve_hinhchop_tugiacdeu(s,a,b,c,d)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình hộp chữ nhật":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh cách nhau bằng dấu ','.\n Ví dụ: A,B,C,D,E,F,G,H")
                            if ok and text!= "":
                                if text.count(",") == 7:
                                    a,b,c,d,e,f,g,h=text.split(',')                                                    
                                    self.text_hinhve.append(f"{a},{b},{c},{d},{e},{f},{g},{h}")         
                                    code= my_module.codelatex_hinh_hop(a,b,c,d,e,f,g,h)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình lập phương":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh cách nhau bằng dấu ','. \n Ví dụ: A,B,C,D,E,F,G,H")
                            if ok and text!= "":
                                if text.count(",") == 7:
                                    a,b,c,d,e,f,g,h=text.split(',')                                                    
                                    self.text_hinhve.append(f"{a},{b},{c},{d},{e},{f},{g},{h}")         
                                    code= my_module.ve_hinh_lapphuong(a,b,c,d,e,f,g,h)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình lăng trụ xiên tam giác":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh cách nhau bằng dấu ','. \n Ví dụ: A,B,C,A',B',C'")
                            if ok and text!= "":
                                if text.count(",") == 5:
                                    a,b,c,d,e,f=text.split(',')                                                    
                                    self.text_hinhve.append(f"{a},{b},{c},{d},{e},{f}")         
                                    code= my_module.codelatex_hinh_langtruxien_tamgiac(a,b,c,d,e,f)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)

                        if self.combo_loaihinh.currentText() == "Hình lăng trụ đứng tam giác":
                            text, ok = QInputDialog.getText(None, "Nhập tên các đỉnh", "Nhập tên các đỉnh cách nhau bằng dấu ','. \n Ví dụ: A,B,C,A',B',C'")
                            if ok and text!= "":
                                if text.count(",") == 5:
                                    a,b,c,d,e,f=text.split(',')                                                    
                                    self.text_hinhve.append(f"{a},{b},{c},{d},{e},{f}")         
                                    code= my_module.ve_hinh_langtrudung_tamgiac(a,b,c,d,e,f)                    
                                    my_module.pdftoimage(code)      

                                    folder_path = my_module.get_folder_hinh()    
                                    fname = f"{folder_path}\\hinhve.png"

                                    pixmap = QPixmap(fname)
                                    new_width = 400  # Đặt chiều rộng mới
                                    new_height = 250  # Đặt chiều cao mới
                                    scaled_pixmap = pixmap.scaled(new_width, new_height)

                                    self.label_showpic_vehinh.setPixmap(scaled_pixmap)
                                    self.label_showpic_vehinh.show()
                                    self.tab_vehinh_textcode.setText(code)                        

                        
                except Exception as e:
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Thông báo lỗi', f'Lỗi {str(e)}!')
                        show_msg_box.exec_() 
                return 

        def btn_copyBBT_click(self):
                folder_path = my_module.get_folder_hinh()       
                image=QImage(f"{folder_path}\\hinhve.png")
                QApplication.clipboard().setImage(image)
                return

        def btn_copyhinh_click(self):
                folder_path = my_module.get_folder_hinh()       
                image=QImage(f"{folder_path}\\hinhve.png")
                QApplication.clipboard().setImage(image)
                return

        def uncheck_all_items(self):
                # Lặp qua tất cả các item trong QTreeWidget và bỏ chọn chúng
                top_level_items = self.treeWidget.invisibleRootItem().childCount()
                for i in range(top_level_items):
                    top_level_item = self.treeWidget.invisibleRootItem().child(i)
                    top_level_item.setCheckState(0, Qt.CheckState.Qt.Unchecked)
                return

        #Hàm biên dịch latex
        def btn_moitruong_latex_click(self):
                file_path=f"{my_module.get_folder_latex()}\\test_code.pdf"
                code=my_module.moi_truong_latex("")
                self.text_run_latex.setText(code)
                return

        def btn_run_latex_click(self):
                code  = self.text_run_latex.toPlainText()      
                my_module.run_latex(code) 
                folder_path = my_module.get_folder_latex()
                file_path=f"{my_module.get_folder_latex()}\\test_code.pdf"
                subprocess.Popen(["start", "", file_path], shell=True)
                return

        #open latex và pdf
        def open_latex_pdf(self,name_thu_muc, name_de,list_noi_dung):
                
                code  = list_noi_dung                
                code_debai = my_module.moi_truong_latex(code)
                code_debai=code_debai.replace(r"{\tenso}{}", f"{{\\tenso}}{{{self.ten_sogd.toPlainText()}}}")                
                code_debai=code_debai.replace(r"{\tentruong}{}", f"{{\\tentruong}}{{{self.ten_truong.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\tenkythi}{}", f"{{\\tenkythi}}{{{self.ten_kythi.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\tenmonthi}{}", f"{{\\tenmonthi}}{{Môn học: {self.ten_monthi.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\thoigian}{}", f"{{\\thoigian}}{{{self.ten_thoigian.toPlainText()}}}")               

                code_loigiai=my_module.moi_truong_latex_loigiai(code)
                code_loigiai=code_loigiai.replace(r"{\tenso}{}", f"{{\\tenso}}{{{self.ten_sogd.toPlainText()}}}")
                #code_debai=code_debai.replace(f"&",f"\\&")
                code_loigiai=code_loigiai.replace(r"{\tentruong}{}", f"{{\\tentruong}}{{{self.ten_truong.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\tenkythi}{}", f"{{\\tenkythi}}{{{self.ten_kythi.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\tenmonthi}{}", f"{{\\tenmonthi}}{{Môn thi: {self.ten_monthi.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\thoigian}{}", f"{{\\thoigian}}{{{self.ten_thoigian.toPlainText()}}}")        

                
                current_directory = os.path.dirname(os.path.abspath(__file__))
                latex_folder_path = os.path.join(current_directory, 'LATEX')

                # new_folder_path=os.path.join(latex_folder_path, name_thu_muc)
                # goc_foler_path=os.path.join(latex_folder_path, name_thu_muc)
                new_folder_path=name_thu_muc
                goc_foler_path=name_thu_muc                

                #Tạo thư mục chứa các mã đề, thư mục DAP AN
                new_folder_path=os.path.join(new_folder_path, f"De_{name_de}")
                ans_folder_path=os.path.join(new_folder_path, 'ans')
                dapan_folder_path=os.path.join(goc_foler_path, f"DAP AN")
                ans_dapan_folder_path=os.path.join(dapan_folder_path, f"ans")

                if not os.path.exists(new_folder_path):
                        os.makedirs(new_folder_path)  
                if not os.path.exists(ans_folder_path):
                        os.makedirs(ans_folder_path)
                if not os.path.exists(dapan_folder_path):
                        os.makedirs(dapan_folder_path)
                if not os.path.exists(ans_dapan_folder_path):
                        os.makedirs(ans_dapan_folder_path)

                #Copy file ex_test.sty
                source_file_path = f"{latex_folder_path}\\ex_test.sty"
                if os.path.exists(source_file_path):                             
                        destination_file_path = os.path.join(new_folder_path, os.path.basename(source_file_path))                                
                        shutil.copy2(source_file_path, destination_file_path)
                        shutil.copy2(source_file_path, ans_dapan_folder_path)

                # Tạo file đề bài
                file_path = os.path.join(new_folder_path, f'De_{name_de}.tex')               
                with open(file_path, 'w',encoding="utf-8") as file:
                        file.write(code_debai)
                
                os.chdir(os.path.dirname(file_path))
                
                latex_command = f"pdflatex -jobname=De_{name_de} De_{name_de}.tex"
                subprocess.run(latex_command, shell=True)
  
                file_path_pdf=os.path.join(new_folder_path, f'De_{name_de}.pdf')
                subprocess.Popen(["start", "", file_path_pdf], shell=True)
                
                #Copy file đề PDF ra thư mục gốc                
                if os.path.exists(f"{file_path_pdf}"):                         
                    shutil.copy2(f"{file_path_pdf}", goc_foler_path)

                file_path = os.path.join(new_folder_path, f'De_{name_de}_loigiai.tex')               
                with open(file_path, 'w',encoding="utf-8") as file:
                        file.write(code_loigiai)                
                os.chdir(os.path.dirname(file_path))              
                
                latex_command = f"pdflatex -jobname=De_{name_de}_loigiai De_{name_de}_loigiai.tex"
                subprocess.run(latex_command, shell=True)

                file_path_pdf=os.path.join(new_folder_path, f'De_{name_de}_loigiai.pdf')

                dapan_path_pdf=os.path.join(new_folder_path, f'De_{name_de}_loigiai.pdf')             
                if os.path.exists(f"{dapan_path_pdf}"):                         
                    shutil.copy2(f"{dapan_path_pdf}", dapan_folder_path)

                #Copy file ans vào thư mục DAPAN/ans                
                ans_file_path = f"{new_folder_path}\\ans\\ans{name_de}-1.tex"
                if os.path.exists(f"{ans_file_path}"): 
                    shutil.copy2(f"{ans_file_path}", ans_dapan_folder_path)

                ans_file_path = f"{new_folder_path}\\ans\\ans{name_de}-2.tex"
                if os.path.exists(f"{ans_file_path}"):
                    shutil.copy2(f"{ans_file_path}", ans_dapan_folder_path)

                ans_file_path = f"{new_folder_path}\\ans\\ans{name_de}-3.tex"
                if os.path.exists(f"{ans_file_path}"):
                    shutil.copy2(f"{ans_file_path}", ans_dapan_folder_path)

                name_matran=f"{goc_foler_path}\\Matran.xlsx"
                self.save_matran(name_matran)   
                return

        def tao_bang_dap_an_latex(self,name_thu_muc, code_bang_dap_an):            
                current_directory = os.path.dirname(os.path.abspath(__file__))
                latex_folder_path = os.path.join(current_directory, 'LATEX')                
                goc_foler_path=os.path.join(latex_folder_path, name_thu_muc)
                goc_foler_path= name_thu_muc

                dapan_folder_path=os.path.join(goc_foler_path, f"DAP AN")
                ans_dapan_folder_path=os.path.join(dapan_folder_path, f"ans")

                #Tạo file để in bảng đáp án
                code_tex_dap_an = my_module.moi_truong_latex(code_bang_dap_an)
                file_path = os.path.join(ans_dapan_folder_path, f'ans.tex')               
                with open(file_path, 'w',encoding="utf-8") as file:
                        file.write(code_tex_dap_an)

                os.chdir(os.path.dirname(file_path))           
                
                latex_command = f"pdflatex -jobname=BANG_DAP_AN ans.tex"
                subprocess.run(latex_command, shell=True)
                #Copy file bảng đáp án ra thư mục gốc
                file_ans_pdf = os.path.join(ans_dapan_folder_path, f'BANG_DAP_AN.pdf')                          
                if os.path.exists(f"{file_ans_pdf }"):                         
                    shutil.copy2(f"{file_ans_pdf }", goc_foler_path)

                return

        #Check Updates
        def check_update(self):
                #Lấy version trên desktop
                try:
                        desktop_version_path=f"{my_module.get_folder_update()}\\imathversion.txt"
                        f=open(desktop_version_path)
                        desktop_version=f.readline()

                        #Lấy version trên web                       
                        url = "http://c3trandainghia.daklak.edu.vn/imath.html"                        
                        response = requests.get(url)
                        # Kiểm tra xem yêu cầu có thành công hay không (status code 200 là thành công)
                        if response.status_code == 200:
                            # Sử dụng BeautifulSoup để phân tích cú pháp HTML của trang web
                            soup = BeautifulSoup(response.text, 'html.parser')

                            # Tìm và in ra giá trị của thẻ có id là 'ngaythi'
                            ngay_thi_element = soup.find('div', {'id': 'content-post'})
                            
                            if ngay_thi_element:
                                web_version = ngay_thi_element.text.strip()
                                                                              
                            else:
                                print("Không tìm thấy thông tin phiên bản trên server.")
                        else:
                            print("Yêu cầu không thành công. Status code:", response.status_code)
                        
                        if desktop_version != web_version:
                                link_update=f"https://www.mediafire.com/file/pqve39k9k5gq6zg/iMath_Update.rar/file"
                                thong_bao=f"iMath đã có bản cập nhật mới ver {web_version}.\n Thầy cô tải và cài tại zalo nhóm."

                                show_msg_box = ShowMessageBox(QMessageBox.Information, f'Cập nhật phiên bản {web_version}', thong_bao)
                                show_msg_box.exec_()                                
                                #webbrowser.open_new_tab(link_update)
                except Exception as e:
                        return


        def btn_check_update_click(self):
                try:
                        #Lấy version trên desktop
                        desktop_version_path=f"{my_module.get_folder_update()}\\imathversion.txt"
                        f=open(desktop_version_path)
                        desktop_version=f.readline()
                        
                        url = "http://c3trandainghia.daklak.edu.vn/imath-2.html"
                        
                        response = requests.get(url)
                        
                        if response.status_code == 200:                            
                            soup = BeautifulSoup(response.text, 'html.parser')                           
                            ngay_thi_element = soup.find('div', {'id': 'content-post'})
                            
                            if ngay_thi_element:
                                web_version = ngay_thi_element.text.strip()
                                                        
                            else:
                                print("Không tìm thấy thông tin phiên bản trên server.")
                        else:
                            print("Yêu cầu không thành công. Status code:", response.status_code)
                            
                        if desktop_version != web_version:
                                link_update=f"https://www.mediafire.com/file/pqve39k9k5gq6zg/iMath_Update.rar/file"
                                thong_bao=f"iMath đã có bản cập nhật mới ver {web_version}.\n Thầy cô tải về tại zalo nhóm. Trân trọng."
                                        
                                show_msg_box = ShowMessageBox(QMessageBox.Information, f'Cập nhật phiên bản {web_version}', thong_bao)
                                show_msg_box.exec_()                                
                                #webbrowser.open_new_tab(link_update)
                        else:
                                thong_bao=f"Thầy cô đang sử dụng iMath phiên bản mới nhất. Trân trọng."                            
                                show_msg_box = ShowMessageBox(QMessageBox.Information, 'Cập nhật phiên bản mới', thong_bao)
                                show_msg_box.exec_()
                except Exception as e:
                        thong_bao=f"Không tìm thấy kết nối mạng để tra cứu bản cập nhật.\n Thầy (cô) vui lòng kiểm tra lại kết nối mạng."                            
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Cập nhật phiên bản mới', thong_bao)
                        show_msg_box.exec_()
                return

        def btn_version_update_click(self):
                try:
                    current_directory = os.path.dirname(os.path.abspath(__file__))
                    update_folder_path = os.path.join(current_directory, 'UPDATE\\update list.txt')                    
                    subprocess.Popen(['explorer', update_folder_path])
                        
                except Exception as e:
                        thong_bao=f"Không tìm thấy file."                            
                        show_msg_box = ShowMessageBox(QMessageBox.Information, 'Lịch sử cập nhật', thong_bao)
                        show_msg_box.exec_()
                return

        #Chọn các thư mục con khi thư mục mẹ được chọn        
        # def handleItemChanged(self, item, column):
        #         if item.flags() & Qt.ItemFlag.ItemIsUserCheckable:
        #             # If the item is user checkable, update its children
        #             self.updateChildren(item)
        #         return

        #Chọn ngẫu nhiên mục con khi chọn mục mẹ
        # def handleItemChanged(self, item, column):
        #     if self.checkbox_tree_random is check:                
        #         text, ok = QInputDialog.getText(None, "Nhập số lượng dạng toán", "Nhập số lượng dạng toán ngẫu nhiên cần chọn")
                
        #         if ok:                
                

        #         if text.count(",") == 1:
        #     if item.flags() & Qt.ItemFlag.ItemIsUserCheckable:
        #         # Kiểm tra xem mục mẹ có được chọn hay không
        #         if item.checkState(0) == Qt.Checked:
        #             # Lấy danh sách tất cả các mục con
        #             child_count = item.childCount()
                    
        #             if child_count > 0:
        #                 # Lựa chọn ngẫu nhiên 5 mục con từ danh sách các mục con
        #                 list_child_count=[i for i in range(child_count)]
        #                 selected_children = random.sample(list_child_count, 3)
        #                 for i in range(child_count):
        #                     child = item.child(i)
        #                     # Nếu mục con nằm trong danh sách chọn ngẫu nhiên, đánh dấu mục đó là được chọn
        #                     if i in selected_children:
        #                         child.setCheckState(0, Qt.Checked)
        #                     else:
        #                         child.setCheckState(0, Qt.Unchecked)
        #         else:
        #             # Nếu mục mẹ không được chọn, bỏ chọn tất cả các mục con
        #             for i in range(item.childCount()):
        #                 child = item.child(i)
        #                 child.setCheckState(0, Qt.Unchecked)
        #     return

        def handleItemChanged(self, item, column):
                if  self.soluong_dangtoan.toPlainText()!="":
                           
                    n=int(self.soluong_dangtoan.toPlainText())
                    if n>0:
                        if item.flags() & Qt.ItemFlag.ItemIsUserCheckable:                            
                        # Kiểm tra xem mục mẹ có được chọn hay không
                                if item.checkState(0) == Qt.Checked:
                                    # Lấy tất cả các mục con từ tất cả các cấp
                                    all_children = []
                                    self.collectAllChildren(item, all_children)

                                    # Chọn ngẫu nhiên 5 mục từ tất cả các mục con
                                    selected_children = random.sample(all_children, min(n, len(all_children)))                             
                                    

                                    # Đánh dấu 5 mục con ngẫu nhiên là được chọn
                                    for child in all_children:
                                        if child in selected_children:
                                            child.setCheckState(0, Qt.Checked)
                                        else:
                                            child.setCheckState(0, Qt.Unchecked)
                                    self.uncheckAllChildren(item)
                                else:
                                    # Nếu mục mẹ không được chọn, bỏ chọn tất cả các mục con
                                    self.uncheckAllChildren(item)
                else:
                    if item.flags() & Qt.ItemFlag.ItemIsUserCheckable:        
                        self.updateChildren(item)

                return

        def collectAllChildren(self, parent_item, all_children):
            """
            Đệ quy để thu thập tất cả các mục con từ mọi cấp của parent_item.
            """
            for i in range(parent_item.childCount()):
                child = parent_item.child(i)
                all_children.append(child)
                # Gọi đệ quy để thu thập các mục con của mục con
                self.collectAllChildren(child, all_children)

        def uncheckAllChildren(self, parent_item):
            """
            Bỏ chọn tất cả các mục con của parent_item.
            """
            for i in range(parent_item.childCount()):
                child = parent_item.child(i)
                child.setCheckState(0, Qt.Unchecked)
                # Gọi đệ quy để bỏ chọn các mục con của mục con
                self.uncheckAllChildren(child)



        def updateChildren(self, parentItem):
                for row in range(parentItem.childCount()):
                    childItem = parentItem.child(row)
                    childItem.setCheckState(0, parentItem.checkState(0))
                return

        #Lấy thông tin tiêu đề
        def open_form_tieude(self):
                self.form_tieude = Form_tieude(self)
                self.form_tieude.show()
                return
        def update_label(self, text):
                self.label_title_chung.setText(text)
                return    
                           

        #Hàm tạo tiêu đề nhóm
        def tao_title_nhom(self,index):
            if self.combo_title_nhom.currentText() == "Nhóm trắc nghiệm":
                text_tieude = "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn."
            if self.combo_title_nhom.currentText() == "Nhóm đúng sai":
                text_tieude= "PHẦN II. Câu trắc nghiệm đúng sai."
            if self.combo_title_nhom.currentText() == "Nhóm trả lời ngắn":
                text_tieude="PHẦN III. Câu trắc nghiệm trả lời ngắn."
            row_count = self.tableWidget.rowCount()
            self.tableWidget.setRowCount(row_count + 1)
            
            item = QTableWidgetItem(f"{text_tieude}")
            item.setTextAlignment(Qt.AlignCenter)
            self.tableWidget.setItem(row_count, 0, item)

            item = QTableWidgetItem("x")
            item.setTextAlignment(Qt.AlignCenter)
            self.tableWidget.setItem(row_count, 1, item)

            item = QTableWidgetItem("x")
            item.setTextAlignment(Qt.AlignCenter)
            self.tableWidget.setItem(row_count, 2, item)
            t=0
            item = QTableWidgetItem(f"0")
            item.setTextAlignment(Qt.AlignCenter)
            self.tableWidget.setItem(row_count, 3, item)                
            return
        #Hàm thiết lập tiêu đề
        def thietlap_tieude(self):
            self.tab_main.setCurrentIndex(1)
            return

        #Lưu thông tin đề thi
        def save_thongtin_dethi(self):
            license.create_reg_thongtin("So_gd", self.ten_sogd.toPlainText())
            license.create_reg_thongtin("Truong_THPT", self.ten_truong.toPlainText())
            license.create_reg_thongtin("ten_kythi", "ĐỀ ÔN TẬP")
            license.create_reg_thongtin("ten_monthi", self.ten_monthi.toPlainText())
            return

        #Tạo đáp án tnmaker cho latex
        def tao_tnmaker_latex(self, name_thu_muc, list_ma_de, socau_TN, socau_DS, socau_TLN):
            current_directory = os.path.dirname(os.path.abspath(__file__))
            latex_folder_path = os.path.join(current_directory, 'LATEX')                
            goc_foler_path=os.path.join(latex_folder_path, name_thu_muc)             
            dapan_folder_path=os.path.join(goc_foler_path, f"DAP AN")
            ans_dapan_folder_path=os.path.join(dapan_folder_path, f"ans")

            #Tạo file để in bảng đáp án
            so_made=len(list_ma_de)              
 
            wb = Workbook()
            ws = wb.active            
            ws.cell(row=1,column=1, value=f"Đề \\ Câu")           

            #Điền đáp án cho từng dòng            
            for i in range(so_made):
                #Điền tên mã đề
                ws.cell(row=1, column=i+2, value=list_ma_de[i])
                
                #Xử lí đáp án trắc nghiệm
                if socau_TN>0:
                        data = []  
                
                        file_path = os.path.join(ans_dapan_folder_path, f'ans{list_ma_de[i]}-1.tex')           
                
                        with open(file_path, "r") as file:
                            for line in file:                                
                                line = line.strip()
                                if len(line) == 1:
                                    data.append(line)                        
                                  
                        for j in range(1, len(data)+1):
                                ws.cell(row=j+1, column=i+2, value=data[j-1])

                #Xử lí đáp án đúng-sai
                if socau_DS>0:
                        data = []                
                        file_path = os.path.join(ans_dapan_folder_path, f'ans{list_ma_de[i]}-2.tex')         
                
                        with open(file_path, "r") as file:
                            for line in file:                                
                                line = line.strip()
                                if len(line) == 36:
                                        dapan=""
                                        if line[5]=="T":
                                                dapan+=f"Đ"
                                        if line[5]=="F":
                                                dapan+=f"S"
                                        if line[14]=="T":
                                                dapan+=f"Đ"
                                        if line[14]=="F":
                                                dapan+=f"S"
                                        if line[23]=="T":
                                                dapan+=f"Đ"
                                        if line[23]=="F":
                                                dapan+=f"S"
                                        if line[32]=="T":
                                                dapan+=f"Đ"
                                        if line[32]=="F":
                                                dapan+=f"S"                                
                                        data.append(dapan)
                        
                        t=1       
                        for j in range(socau_TN+1, socau_TN+socau_DS+1):                                
                                ws.cell(row=j+1, column=i+2, value=data[t-1])
                                t+=1

                #Xử lí đáp án tra loi ngan
                if socau_TLN>0:
                        data = []                
                        file_path = os.path.join(ans_dapan_folder_path, f'ans{list_ma_de[i]}-3.tex')         
                
                        with open(file_path, "r") as file:                                
                                for line in file:                                
                                        line = line.strip()
                                        if "Solution" not in line:
                                                dapan=line.replace("$","").replace("{","").replace("}","")
                                                data.append(dapan)
                        
                        t=1       
                        for j in range(socau_TN+socau_DS+1, socau_TN+socau_DS+socau_TLN+1):                                
                                ws.cell(row=j+1, column=i+2, value=data[t-1])
                                t+=1              

            #Điền số câu
            so_cau=socau_TN+socau_DS+socau_TLN
            for i in range(1, so_cau+1):
                ws.cell(row=i+1, column=1, value=i)            
            wb.save(f"{goc_foler_path}\\QR_TNMaker_SmartTest.xlsx")
            return

        #Tạo đáp án tnmaker cho word
        def tao_tnmaker_word(self, name_thu_muc, list_ma_de, list_dapan_word):
            current_directory = os.path.dirname(os.path.abspath(__file__))
            doc_folder_path = os.path.join(current_directory, 'DOC')                
            goc_foler_path=os.path.join(doc_folder_path, name_thu_muc)           

            #Tạo file để in bảng đáp án
            so_made=len(list_ma_de)
            so_cau=int(len(list_dapan_word)/so_made)
                    
 
            wb = Workbook()
            ws = wb.active            
            ws.cell(row=1,column=1, value=f"Đề \\ Câu")

            t=0
            #Điền đáp án cho từng dòng            
            for i in range(so_made):
                #Điền tên mã đề               
                ws.cell(row=1, column=i+2, value=list_ma_de[i])                      
           
                for j in range(1, so_cau+1):
                        dap_an=list_dapan_word[t]                        
                        ws.cell(row=j+1, column=i+2, value=dap_an)
                        t+=1
            #Điền số câu
            for i in range(1, so_cau+1):
                ws.cell(row=i+1, column=1, value=i)            
            wb.save(f"{goc_foler_path}\\QR_TNMaker_SmartTest.xlsx")
            return

           
        def ghep_de_pdf(self,name_thu_muc, list_noi_dung):
                
                code  = list_noi_dung
                
                code_debai = my_module.moi_truong_latex(code)
                code_debai=code_debai.replace(r"{\tenso}{}", f"{{\\tenso}}{{{self.ten_sogd.toPlainText()}}}")
                code_debai=code_debai.replace(f"&",f"\\&")
                code_debai=code_debai.replace(r"{\tentruong}{}", f"{{\\tentruong}}{{{self.ten_truong.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\tenkythi}{}", f"{{\\tenkythi}}{{{self.ten_kythi.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\tenmonthi}{}", f"{{\\tenmonthi}}{{Môn học: {self.ten_monthi.toPlainText()}}}")
                code_debai=code_debai.replace(r"{\thoigian}{}", f"{{\\thoigian}}{{{self.ten_thoigian.toPlainText()}}}")               

                code_loigiai=my_module.moi_truong_latex_loigiai(code)
                code_loigiai=code_loigiai.replace(r"{\tenso}{}", f"{{\\tenso}}{{{self.ten_sogd.toPlainText()}}}")
                code_debai=code_debai.replace(f"&",f"\\&")
                code_loigiai=code_loigiai.replace(r"{\tentruong}{}", f"{{\\tentruong}}{{{self.ten_truong.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\tenkythi}{}", f"{{\\tenkythi}}{{{self.ten_kythi.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\tenmonthi}{}", f"{{\\tenmonthi}}{{Môn thi: {self.ten_monthi.toPlainText()}}}")
                code_loigiai=code_loigiai.replace(r"{\thoigian}{}", f"{{\\thoigian}}{{{self.ten_thoigian.toPlainText()}}}")             

                #Tạo thư mục với tên là name, Tạo thư mục ans
                current_directory = os.path.dirname(os.path.abspath(__file__))
                latex_folder_path = os.path.join(current_directory, 'LATEX')

                new_folder_path=os.path.join(latex_folder_path, name_thu_muc)
                tonghop_foler_path=os.path.join(new_folder_path, f"De_tonghop")

                if not os.path.exists(tonghop_folder_path):
                    os.makedirs(tonghop_folder_path) 

                #Copy file ex_test.sty
                source_file_path = f"{latex_folder_path}\\ex_test.sty"
                if os.path.exists(source_file_path):                             
                    destination_file_path = os.path.join(new_folder_path, os.path.basename(source_file_path))                                
                    shutil.copy2(source_file_path, destination_file_path)                       

                # Tạo file đề bài
                file_path = os.path.join(tonghop_folder_path, f'De_tonghop.tex')               
                with open(file_path, 'w',encoding="utf-8") as file:
                        file.write(code_debai)

                #Di chuyển đến thư mục chứa file tex
                os.chdir(os.path.dirname(file_path))

                #Biên dịch tệp LaTeX thành PDF
                latex_command = f"pdflatex -jobname=De_tonghop De_tonghop.tex"
                subprocess.run(latex_command, shell=True)

                latex_command = f"pdflatex -jobname=De_tonghop De_tonghop.tex"
                subprocess.run(latex_command, shell=True)
  
                file_path_pdf=os.path.join(new_folder_path, f'De_tonghop.pdf')                
                
                #Copy file đề PDF ra thư mục gốc                
                if os.path.exists(f"{file_path_pdf}"):                         
                    shutil.copy2(f"{file_path_pdf}", new_foler_path)
                 
                return

        def nhapmade(self):
            text, ok = QInputDialog.getText(None, "Nhập mã đề", "Nhập danh sách mã đề cách nhau bằng dấu ',' :")
            #Kiểm tra xem người dùng đã nhấp vào OK hay không
            if ok:              
                self.label_nhapmade.setPlainText(text)
                if text.count(',')+1 != self.spin_soluong_de.value():
                    thong_bao=f"Số lượng mã đề đã nhập không bằng số lượng đề cần tạo"                            
                    show_msg_box = ShowMessageBox(QMessageBox.Information, 'Cảnh báo', thong_bao)
                    show_msg_box.exec_()
            else:
                thong_bao=f"Bạn chưa nhập danh sách mã đề. Chương trình sẽ tạo mã đề ngãu nhiên."                            
                show_msg_box = ShowMessageBox(QMessageBox.Information, 'Cảnh báo', thong_bao)
                show_msg_box.exec_()
            return

        def btn_xoa_dong_click(self):
            selected_row = self.tableWidget.currentRow()
            if selected_row >= 0:
                self.tableWidget.removeRow(selected_row)
            self.thongke()
            return

        #Code thống kê số lượng
        def thongke(self):            
            #Đếm số lượng câu trắc nghiệm
            socau_tracnghiem_NB=0
            socau_tracnghiem_TH=0
            socau_tracnghiem_VDT=0
            socau_tracnghiem_VDC=0

            socau_dungsai_NB=0
            socau_dungsai_TH=0
            socau_dungsai_VDT=0
            socau_dungsai_VDC=0

            socau_tuluan_NB=0
            socau_tuluan_TH=0
            socau_tuluan_VDT=0
            socau_tuluan_VDC=0
            row_count = self.tableWidget.rowCount()
            for i in range(row_count):
                loai_cau= self.tableWidget.item(i, 1)
                muc_do = self.tableWidget.item(i, 2)
                so_cau = self.tableWidget.item(i, 3)
                if loai_cau and muc_do and so_cau:
                    so_cau=so_cau.text()       
                    if loai_cau.text()=="TN" and muc_do.text()=="NB":
                        socau_tracnghiem_NB+=int(so_cau)
                    if loai_cau.text()=="TN" and muc_do.text()=="TH":
                        socau_tracnghiem_TH+=int(so_cau)
                    if loai_cau.text()=="TN" and muc_do.text()=="VDT":
                        socau_tracnghiem_VDT+=int(so_cau)
                    if loai_cau.text()=="TN" and muc_do.text()=="VDC":
                        socau_tracnghiem_VDC+=int(so_cau)

                    if loai_cau.text()=="Đ-S" and muc_do.text()=="NB":
                        socau_dungsai_NB+=int(so_cau)
                    if loai_cau.text()=="Đ-S" and muc_do.text()=="TH":
                        socau_dungsai_TH+=int(so_cau)
                    if loai_cau.text()=="Đ-S" and muc_do.text()=="VDT":
                        socau_dungsai_VDT+=int(so_cau)
                    if loai_cau.text()=="Đ-S" and muc_do.text()=="VDC":
                        socau_dungsai_VDC+=int(so_cau)

                    if loai_cau.text()=="TL" and muc_do.text()=="NB":
                        socau_tuluan_NB+=int(so_cau)
                    if loai_cau.text()=="TL" and muc_do.text()=="TH":
                        socau_tuluan_TH+=int(so_cau)
                    if loai_cau.text()=="TL" and muc_do.text()=="VDT":
                        socau_tuluan_VDT+=int(so_cau)
                    if loai_cau.text()=="TL" and muc_do.text()=="VDC":
                        socau_tuluan_VDC+=int(so_cau)

            #Ghi thống kê số câu trắc nghiệm
            item = QTableWidgetItem(f"{socau_tracnghiem_NB}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 0, item)

            item = QTableWidgetItem(f"{socau_tracnghiem_TH}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 3, item)

            item = QTableWidgetItem(f"{socau_tracnghiem_VDT}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 6, item)

            item = QTableWidgetItem(f"{socau_tracnghiem_VDC}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 9, item)
            

            #Ghi thống kê số câu đúng sai
            item = QTableWidgetItem(f"{socau_dungsai_NB}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 1, item)

            item = QTableWidgetItem(f"{socau_dungsai_TH}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 4, item)

            item = QTableWidgetItem(f"{socau_dungsai_VDT}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 7, item)

            item = QTableWidgetItem(f"{socau_dungsai_VDC}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 10, item)                    

            #Ghi thống kê số câu tự luận
            item = QTableWidgetItem(f"{socau_tuluan_NB}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 2, item)

            item = QTableWidgetItem(f"{socau_tuluan_TH}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 5, item)

            item = QTableWidgetItem(f"{socau_tuluan_VDT}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 7, item)

            item = QTableWidgetItem(f"{socau_tuluan_VDC}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(2, 11, item)

            #Tính tổng số câu của mỗi loại mức độ
            sum_NB=socau_tracnghiem_NB+socau_dungsai_NB+socau_tuluan_NB
            sum_TH=socau_tracnghiem_TH+socau_dungsai_TH+socau_tuluan_TH
            sum_VDT=socau_tracnghiem_VDT+socau_dungsai_VDT+socau_tuluan_VDT
            sum_VDC=socau_tracnghiem_VDC+ socau_dungsai_VDC+socau_tuluan_VDC

            item = QTableWidgetItem(f"{sum_NB}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(3, 0, item)

            item = QTableWidgetItem(f"{sum_TH}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(3, 3, item)

            item = QTableWidgetItem(f"{sum_VDT}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(3, 6, item)

            item = QTableWidgetItem(f"{sum_VDC}")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(3, 9, item)

            # Tính tổng số câu trắc nghiệm, đúng-sai, trả lời ngắn và tổng số câu
            sum_TN=socau_tracnghiem_NB+socau_tracnghiem_TH+socau_tracnghiem_VDT+socau_tracnghiem_VDC
            sum_DS=socau_dungsai_NB+socau_dungsai_TH+socau_dungsai_VDT+socau_dungsai_VDC
            sum_TL=socau_tuluan_NB+socau_tuluan_TH+socau_tuluan_VDT+socau_tuluan_VDC
            sum_toanbo=sum_TN + sum_DS + sum_TL

            item = QTableWidgetItem(f"Trắc nghiệm: {sum_TN} câu")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(4, 0, item)
            self.table_thongke.setSpan(4, 0, 1, 4)

            item = QTableWidgetItem(f"Đúng-Sai: {sum_DS} câu")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(4, 4, item)
            self.table_thongke.setSpan(4, 4, 1, 4)

            item = QTableWidgetItem(f"Trả lời ngắn: {sum_TL} câu")
            item.setTextAlignment(Qt.AlignCenter)
            self.table_thongke.setItem(4, 8, item)
            self.table_thongke.setSpan(4, 8, 1, 4)

            item = QTableWidgetItem(f"Tổng cộng: {sum_toanbo} câu")
            item.setTextAlignment(Qt.AlignCenter)
            item.setBackground(QtGui.QColor(251, 243, 221))
            self.table_thongke.setItem(5, 0, item)
            self.table_thongke.setSpan(5, 0, 1, 12)
            return sum_toanbo

        def chuyen_tuluan(self):
            selected_row = self.tableWidget.currentRow()
            if selected_row >= 0:
                item = QTableWidgetItem(f"TL")
                item.setTextAlignment(Qt.AlignCenter)
                self.tableWidget.setItem(selected_row, 1, item)
            self.thongke()
            return

        def move_row_up(self):
            current_row = self.tableWidget.currentRow()
            if current_row > 0:  # Đảm bảo rằng dòng hiện tại không phải là dòng đầu tiên
                self.swap_rows(current_row, current_row - 1)
                self.tableWidget.setCurrentCell(current_row - 1, self.tableWidget.currentColumn())
            return

        def move_row_down(self):
            current_row = self.tableWidget.currentRow()
            if current_row < self.tableWidget.rowCount() - 1:  # Đảm bảo rằng dòng hiện tại không phải là dòng cuối cùng
                self.swap_rows(current_row, current_row + 1)
                self.tableWidget.setCurrentCell(current_row + 1, self.tableWidget.currentColumn())
            return
            
        def swap_rows(self, row1, row2):
            for column in range(self.tableWidget.columnCount()):
                item1 = self.tableWidget.takeItem(row1, column)
                item2 = self.tableWidget.takeItem(row2, column)
                self.tableWidget.setItem(row1, column, item2)
                self.tableWidget.setItem(row2, column, item1)

        def hd_taode_video(self):
            QDesktopServices.openUrl(QUrl('https://youtu.be/GKglF9yjmeg'))

        def hd_taode_pdf(self):            
            current_directory = os.path.dirname(os.path.abspath(__file__))
            help_folder_path = os.path.join(current_directory, 'HELP')
            file_path = os.path.join(help_folder_path, "iMath-guide.pdf")
            subprocess.Popen(['explorer', file_path])

        def show_tool_Word2PPT(self):            
            current_directory = os.path.dirname(os.path.abspath(__file__))
            help_folder_path = os.path.join(current_directory, 'HELP')
            file_path = os.path.join(help_folder_path, "show-tool-Word-to-PPT.pdf")
            subprocess.Popen(['explorer', file_path])
            
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)    
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec())


